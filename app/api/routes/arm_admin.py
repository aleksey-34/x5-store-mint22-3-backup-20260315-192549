from __future__ import annotations

import io
import json
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import tempfile
from collections import defaultdict
from dataclasses import asdict
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from html import escape
from pathlib import Path, PurePosixPath
from urllib.parse import quote

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse, Response
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.db.session import get_db
from app.models.document import Document
from app.models.document_content import DocumentContent
from app.models.journal_entry import JournalEntry
from app.models.work_schedule import WorkSchedule
from app.schemas.arm_admin import (
    ArmAssistRequest,
    ArmAssistResponse,
    ArmActionResponse,
    ArmChecklistItem,
    ArmDashboardResponse,
    ArmObjectProfileResponse,
    ArmObjectProfileUpdateRequest,
    ArmPprImportRequest,
    ArmEmployeeChecklistGenerateRequest,
    ArmEmployeeChecklistGenerateResponse,
    ArmEmployeeChecklistItem,
    ArmEmployeeChecklistOverviewResponse,
    ArmEmployeeChecklistResponse,
    ArmEmployeeCatalogItem,
    ArmEmployeeCatalogResponse,
    ArmEmployeeOverviewAction,
    ArmEmployeeOverviewEmployee,
    ArmEmployeeOverviewGroup,
    ArmFileReadResponse,
    ArmFileWriteRequest,
    ArmFsEntry,
    ArmFsTreeResponse,
    ArmMetrics,
    ArmScanCaptureRequest,
    ArmScanIngestItem,
    ArmScanIngestRequest,
    ArmScanIngestResponse,
    ArmScannerDevice,
    ArmScannerDevicesResponse,
    ArmSpeechTranscribeResponse,
    ArmProfessionOption,
    ArmTodoItem,
    ArmTodoResponse,
)
from app.services.local_llm import (
    LocalLLMConnectionError,
    LocalLLMRequestError,
    check_local_llm_available,
    generate_with_local_llm_profile,
)
from app.services.office_export import (
    OfficeExportDependencyError,
    VALID_EXPORT_CLASSIFICATIONS,
    build_office_pack_zip,
    export_orders_docx_bundle,
    export_registers_xlsx,
)
from app.services.scan_archiver import ingest_inbox
from app.services.scan_classifier import classify_scan_candidate

try:
    import speech_recognition as sr
except Exception:  # noqa: BLE001
    sr = None

try:
    import markdown as md
except Exception:  # noqa: BLE001
    md = None

try:
    from pypdf import PdfReader
except Exception:  # noqa: BLE001
    PdfReader = None

router = APIRouter(prefix="/arm", tags=["arm"])

DOC_STATUS_LABELS: dict[str, str] = {
    "approved": "Утверждено",
    "new": "Вновь созданные",
    "fix": "Исправить",
}


def _extract_order_info_from_rel_path(rel_path: str) -> tuple[str, str]:
    safe = (rel_path or "").strip().replace("\\", "/")
    if not safe:
        return "", ""

    stem = PurePosixPath(safe).stem
    match = re.search(r"ORDER[_-]?([0-9]{1,3})(?:[_-](.*))?", stem, flags=re.IGNORECASE)
    if not match:
        return "", ""

    order_no = (match.group(1) or "").strip()
    if order_no.isdigit():
        order_no = order_no.zfill(2)

    title_tail = (match.group(2) or "").strip()
    title_tail = re.sub(r"[_-]?v[0-9]+$", "", title_tail, flags=re.IGNORECASE).strip(" _-")
    title_tail = re.sub(r"[_-]+", " ", title_tail)
    title_tail = re.sub(r"\s+", " ", title_tail).strip()
    if title_tail:
        title_tail = title_tail[0].upper() + title_tail[1:]

    return order_no, title_tail


def _extract_order_header_info(root: Path, rel_path: str) -> tuple[str, str]:
    if not rel_path:
        return "", ""
    suffix = PurePosixPath(rel_path).suffix.lower()
    if suffix not in {".md", ".txt", ".html"}:
        return _extract_order_info_from_rel_path(rel_path)

    try:
        target = _resolve_safe_path(root=root, rel_path=rel_path)
        if not target.exists() or not target.is_file():
            return _extract_order_info_from_rel_path(rel_path)
        text = target.read_text(encoding="utf-8", errors="ignore")
    except Exception:  # noqa: BLE001
        return _extract_order_info_from_rel_path(rel_path)

    order_no = ""
    order_title = ""

    no_match = re.search(r"Приказ\s*[№N]\s*([0-9]{1,3}(?:\s*/\s*[А-ЯA-ZЁ]{2,6}-\d{4})?)", text, flags=re.IGNORECASE)
    if no_match:
        order_no = re.sub(r"\s+", "", no_match.group(1).replace("№", "")).strip()
        if order_no.isdigit():
            order_no = order_no.zfill(2)

    quoted_match = re.search(r"[«\"]\s*О\s+([^»\"\n]+)[»\"]", text, flags=re.IGNORECASE)
    if quoted_match:
        order_title = f"О {quoted_match.group(1).strip()}"
    else:
        h2_match = re.search(r"^##\s+(О\s+.+)$", text, flags=re.IGNORECASE | re.MULTILINE)
        if h2_match:
            order_title = h2_match.group(1).strip()

    fallback_no, fallback_title = _extract_order_info_from_rel_path(rel_path)
    if not order_no:
        order_no = fallback_no
    if not order_title:
        order_title = fallback_title

    return order_no, order_title


@dataclass(frozen=True)
class ChecklistRule:
    code: str
    title: str
    folder: str
    pattern: str | tuple[str, ...]
    required_min: int


@dataclass(frozen=True)
class EmployeeChecklistRule:
    code: str
    title: str
    folder: str
    patterns: tuple[str, ...]
    required_count: int
    guidance: str
    scope: str = "employee"


EMPLOYEE_TB_BASE_RULES: tuple[EmployeeChecklistRule, ...] = (
    EmployeeChecklistRule(
        code="E01_PROFILE",
        title="Карточка сотрудника (employee_profile)",
        folder=".",
        patterns=("employee_profile.txt",),
        required_count=1,
        guidance="Заполните профиль сотрудника: ФИО, должность, разряд и дата рождения.",
    ),
    EmployeeChecklistRule(
        code="E02_IDENTITY_SCAN",
        title="Скан удостоверения/личного документа",
        folder="01_identity_and_contract",
        patterns=("*.pdf", "*.jpg", "*.jpeg", "*.png", "*.tif", "*.tiff"),
        required_count=1,
        guidance="Добавьте как минимум один валидный скан (PDF/JPG/PNG/TIFF).",
    ),
    EmployeeChecklistRule(
        code="E03_ADMISSION_ORDERS",
        title="Приказы/допуски по сотруднику",
        folder="02_admission_orders",
        patterns=("*.md", "*.pdf", "*.jpg", "*.jpeg", "*.png"),
        required_count=1,
        guidance="Добавьте приказ или подтверждение допуска по сотруднику.",
    ),
    EmployeeChecklistRule(
        code="E04_BRIEFINGS",
        title="Инструктажи и обучение",
        folder="03_briefings_and_training",
        patterns=("*.md", "*.pdf", "*.jpg", "*.jpeg", "*.png"),
        required_count=1,
        guidance="Нужен документ о прохождении инструктажа/обучения.",
    ),
    EmployeeChecklistRule(
        code="E05_ATTESTATION",
        title="Аттестация и удостоверения",
        folder="04_attestation_and_certificates",
        patterns=("*.md", "*.pdf", "*.jpg", "*.jpeg", "*.png"),
        required_count=1,
        guidance="Нужен документ аттестации/удостоверения по профессии.",
    ),
    EmployeeChecklistRule(
        code="E06_PPE",
        title="Выдача СИЗ",
        folder="05_ppe_issue",
        patterns=("*.md", "*.pdf", "*.jpg", "*.jpeg", "*.png"),
        required_count=1,
        guidance="Добавьте ведомость/акт выдачи СИЗ.",
    ),
    EmployeeChecklistRule(
        code="E07_PERMITS",
        title="Наряды и допуски",
        folder="06_permits_and_work_admission",
        patterns=("*.md", "*.pdf", "*.jpg", "*.jpeg", "*.png"),
        required_count=1,
        guidance="Добавьте действующий наряд-допуск или документ допуска к работам.",
    ),
    EmployeeChecklistRule(
        code="E08_MEDICAL",
        title="Медосмотр и первая помощь",
        folder="07_medical_and_first_aid",
        patterns=("*.md", "*.pdf", "*.jpg", "*.jpeg", "*.png"),
        required_count=1,
        guidance="Добавьте подтверждение медосмотра/допуска по здоровью.",
    ),
)


PROJECT_TB_ORDER_RULES: tuple[EmployeeChecklistRule, ...] = (
    EmployeeChecklistRule(
        code="P01_ORDER_REGISTER",
        title="Реестр приказов по охране труда",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_REGISTER*",),
        required_count=1,
        guidance="Ведите единый реестр приказов и подписей ознакомления.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P11_ORDER_PS",
        title="Приказ N11: ответственные лица за безопасные работы с ПС",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_11_*",),
        required_count=1,
        guidance="Формат валиден по вводным генподрядчика.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P12_ORDER_PERMIT",
        title="Приказ N12: ответственные лица за выдачу нарядов-допусков",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_12_*", "*PERMIT_12_*"),
        required_count=1,
        guidance="Формат валиден по вводным генподрядчика.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P13_ORDER_HEIGHT",
        title="Приказ N13: ответственные лица и меры безопасности на высоте",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_13_*",),
        required_count=1,
        guidance="Формат валиден по вводным генподрядчика.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P14_ORDER_FIRE",
        title="Приказ N14: ответственные лица за пожарную безопасность",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_14_*",),
        required_count=1,
        guidance="Формат валиден по вводным генподрядчика.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P15_ORDER_LOADING",
        title="Приказ N15: ответственные лица за погрузочно-разгрузочные работы",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_15_*",),
        required_count=1,
        guidance="Формат валиден по вводным генподрядчика.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P16_ORDER_PRESSURE",
        title="Приказ N16: ответственные лица за сосуды под давлением",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_16_*",),
        required_count=1,
        guidance="Формат валиден по вводным генподрядчика.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P17_ORDER_CLOSE_SHIFT",
        title="Приказ N17: закрытие помещений по окончании смены",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_17_*",),
        required_count=1,
        guidance="Обычно нужен для ответственных лиц смены/руководителей.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P18_ORDER_INTERNSHIP",
        title="Приказ N18: о стажировке вновь принятых рабочих",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_18_*",),
        required_count=1,
        guidance="Формат валиден по вводным генподрядчика.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P19_ORDER_AFTER_INTERNSHIP",
        title="Приказ N19: допуск к самостоятельной работе после стажировки",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_19_*",),
        required_count=1,
        guidance="Формат валиден по вводным генподрядчика.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P20_ORDER_HEATING",
        title="Приказ N20: ответственные лица по электропрогреву бетона в зимних условиях",
        folder="01_orders_and_appointments",
        patterns=("*ORDER_20_*",),
        required_count=1,
        guidance="Критично для электротехнических/зимних работ.",
        scope="project",
    ),
    EmployeeChecklistRule(
        code="P21_ADMISSION_LETTER",
        title="Письмо-допуск работников и техники в адрес генподрядчика",
        folder="00_incoming_requests",
        patterns=("*LETTER_ADMISSION*", "*допуск*работников*техники*"),
        required_count=1,
        guidance="Используйте фирменный бланк и приложите списки работников/техники.",
        scope="project",
    ),
)


PROFESSION_PROJECT_RULES: dict[str, set[str]] = {
    "default": {
        "P01_ORDER_REGISTER",
        "P11_ORDER_PS",
        "P12_ORDER_PERMIT",
        "P13_ORDER_HEIGHT",
        "P14_ORDER_FIRE",
        "P15_ORDER_LOADING",
        "P16_ORDER_PRESSURE",
        "P18_ORDER_INTERNSHIP",
        "P19_ORDER_AFTER_INTERNSHIP",
        "P21_ADMISSION_LETTER",
    },
    "electric": {
        "P01_ORDER_REGISTER",
        "P11_ORDER_PS",
        "P12_ORDER_PERMIT",
        "P13_ORDER_HEIGHT",
        "P14_ORDER_FIRE",
        "P15_ORDER_LOADING",
        "P16_ORDER_PRESSURE",
        "P18_ORDER_INTERNSHIP",
        "P19_ORDER_AFTER_INTERNSHIP",
        "P20_ORDER_HEATING",
        "P21_ADMISSION_LETTER",
    },
    "supervisor": {
        "P01_ORDER_REGISTER",
        "P11_ORDER_PS",
        "P12_ORDER_PERMIT",
        "P13_ORDER_HEIGHT",
        "P14_ORDER_FIRE",
        "P15_ORDER_LOADING",
        "P16_ORDER_PRESSURE",
        "P17_ORDER_CLOSE_SHIFT",
        "P18_ORDER_INTERNSHIP",
        "P19_ORDER_AFTER_INTERNSHIP",
        "P20_ORDER_HEATING",
        "P21_ADMISSION_LETTER",
    },
    "itr": {
        "P01_ORDER_REGISTER",
        "P11_ORDER_PS",
        "P12_ORDER_PERMIT",
        "P13_ORDER_HEIGHT",
        "P14_ORDER_FIRE",
        "P15_ORDER_LOADING",
        "P16_ORDER_PRESSURE",
        "P18_ORDER_INTERNSHIP",
        "P19_ORDER_AFTER_INTERNSHIP",
        "P21_ADMISSION_LETTER",
    },
}

PROFESSION_LABELS: dict[str, str] = {
    "default": "Общий персонал",
    "electric": "Электротехнический персонал",
    "supervisor": "Руководители и прорабы",
    "itr": "ИТР (инженеры, геодезисты)",
    "custom": "Пользовательская роль",
}


@dataclass(frozen=True)
class EmployeeCatalogRow:
    employee_root: Path
    employee_rel_path: str
    employee_id: str | None
    employee_name: str
    position: str | None
    profession: str
    profession_group: str


@dataclass(frozen=True)
class PeriodicDocRule:
    code: str
    title: str
    folder: str
    pattern: str
    period_days: int
    details: str


@dataclass(frozen=True)
class PeriodicDocStatus:
    rule: PeriodicDocRule
    latest_file: str | None
    latest_mtime: datetime | None
    days_since_update: int | None
    is_due: bool
    action_path: str


CHECKLIST_RULES: tuple[ChecklistRule, ...] = (
    ChecklistRule("1.1.1", "Приказ: уполномоченный представитель", "01_orders_and_appointments", "*ORDER_01_*", 1),
    ChecklistRule("1.1.2", "Приказ: допуск к СМР", "01_orders_and_appointments", "*ORDER_02_*", 1),
    ChecklistRule("1.1.3", "Приказы: ответственные по направлениям", "01_orders_and_appointments", "*ORDER_03_*", 1),
    ChecklistRule("1.1.4", "Приказ: стропальщики", "01_orders_and_appointments", "*ORDER_04_*", 1),
    ChecklistRule("1.1.5", "Приказ: электрохозяйство", "01_orders_and_appointments", "*ORDER_05_*", 1),
    ChecklistRule("1.1.6", "Приказ: стажировка рабочих", "01_orders_and_appointments", "*ORDER_06_*", 1),
    ChecklistRule("1.1.7", "Приказ: допуск по профессиям", "01_orders_and_appointments", "*ORDER_07_*", 1),
    ChecklistRule("1.1.8", "Приказы: распределение ответственности по прорабам", "01_orders_and_appointments", "*ORDER_09_*", 1),
    ChecklistRule("1.1.9", "Приказы: ТБ по прорабам", "01_orders_and_appointments", "*ORDER_10_*", 1),
    ChecklistRule("1.1.10", "Приказы: ответственные лица по ПС", "01_orders_and_appointments", "*ORDER_11_*", 1),
    ChecklistRule("1.3.3", "Наряд-допуск на опасные работы", "01_orders_and_appointments", "*PERMIT_12_*", 1),
    ChecklistRule("1.4", "ППР", "05_execution_docs/ppr", "*", 1),
    ChecklistRule("1.5", "ППРв", "05_execution_docs/pprv_work_at_height", "*", 1),
    ChecklistRule("1.6", "Акт-допуск на производство СМР", "05_execution_docs/admission_acts", "*", 1),
    ChecklistRule("3.1", "Журналы производства", "04_journals/production", "*", 1),
    ChecklistRule("3.2", "Журналы ОТ/ПБ", "04_journals/labor_safety", "*", 1),
    ChecklistRule(
        "4",
        "Сканы удостоверений и протоколов",
        "02_personnel/employees",
        ("*.pdf", "*.jpg", "*.jpeg", "*.png", "*.tif", "*.tiff"),
        4,
    ),
    ChecklistRule("5", "Нормативная база", "06_normative_base", "*", 1),
)

PERIODIC_DOC_RULES: tuple[PeriodicDocRule, ...] = (
    PeriodicDocRule(
        code="PD01",
        title="Наряд-допуск на высоте",
        folder="03_hse_and_fire_safety/permits/наряды_допуски",
        pattern="*НД*высот*.*",
        period_days=14,
        details="Обновлять не реже 1 раза в 14 дней и перед новым фронтом работ.",
    ),
)

PROJECT_METADATA_DEFAULT_NAME = "20260310_PROJECT_METADATA_v01.md"
PPR_CONTEXT_REL_PATH = "06_normative_base/ppr_imports/ppr_context_columns.md"

PROJECT_ROOT = Path(__file__).resolve().parents[3]
SCANNER_COMMAND_TIMEOUT_SEC = 240
MAX_TEXT_PREVIEW_BYTES = 1_500_000
UPLOAD_ALLOWED_EXTENSIONS = {".docx", ".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".xlsx"}
UPLOAD_MAX_BYTES = 50 * 1024 * 1024  # 50 MB
TEXT_PREVIEW_EXTENSIONS = {
    ".md",
    ".txt",
    ".csv",
    ".json",
    ".yml",
    ".yaml",
    ".py",
    ".log",
    ".ini",
}
INLINE_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tif", ".tiff"}
COMPETITOR_RESEARCH_HTML = PROJECT_ROOT / "docs" / "document_control" / "arm_competitors_research.html"
COMPANY_REQUISITES_CARD_PATH = PROJECT_ROOT / "docs" / "templates" / "organization" / "company_requisites_card.md"
ORDER_FORMAT_TEMPLATES_DIR = PROJECT_ROOT / "docs" / "templates" / "orders"

ORDER_DRAFT_TEMPLATE_MAP: dict[str, str] = {
    "P01_ORDER_REGISTER": "order_01_register_template.md",
    "P11_ORDER_PS": "order_11_ps_responsible_template.md",
    "P12_ORDER_PERMIT": "order_12_permit_issuer_template.md",
    "P13_ORDER_HEIGHT": "order_13_height_works_template.md",
    "P14_ORDER_FIRE": "order_14_fire_safety_template.md",
    "P15_ORDER_LOADING": "order_15_loading_unloading_template.md",
    "P16_ORDER_PRESSURE": "order_16_pressure_vessels_template.md",
    "P17_ORDER_CLOSE_SHIFT": "order_17_close_shift_template.md",
    "P18_ORDER_INTERNSHIP": "order_18_internship_template.md",
    "P19_ORDER_AFTER_INTERNSHIP": "order_19_independent_work_admission_template.md",
    "P20_ORDER_HEATING": "order_20_concrete_heating_template.md",
    "P21_ADMISSION_LETTER": "letter_admission_workers_equipment_template.md",
}

ORDER_DRAFT_FILE_NAME_MAP: dict[str, str] = {
    "P01_ORDER_REGISTER": "ORDER_01_реестр_приказов_черновик.md",
    "P11_ORDER_PS": "ORDER_11_ответственные_лица_по_пс_черновик.md",
    "P12_ORDER_PERMIT": "ORDER_12_выдача_наряда_допуска_черновик.md",
    "P13_ORDER_HEIGHT": "ORDER_13_работы_на_высоте_черновик.md",
    "P14_ORDER_FIRE": "ORDER_14_ответственные_лица_пожарная_безопасность_черновик.md",
    "P15_ORDER_LOADING": "ORDER_15_ответственные_погрузочно_разгрузочные_черновик.md",
    "P16_ORDER_PRESSURE": "ORDER_16_ответственные_сосуды_под_давлением_черновик.md",
    "P17_ORDER_CLOSE_SHIFT": "ORDER_17_закрытие_смены_черновик.md",
    "P18_ORDER_INTERNSHIP": "ORDER_18_стажировка_черновик.md",
    "P19_ORDER_AFTER_INTERNSHIP": "ORDER_19_допуск_к_самостоятельной_черновик.md",
    "P20_ORDER_HEATING": "ORDER_20_ответственные_лица_электропрогрев_бетона_черновик.md",
    "P21_ADMISSION_LETTER": "LETTER_ADMISSION_допуск_персонала_и_техники.md",
}

ORDER_REGISTER_TITLE_BY_NUMBER: dict[str, str] = {
    "01": "О назначении уполномоченного представителя",
    "02": "О допуске к СМР",
    "03": "О назначении ответственного лица за геодезический контроль",
    "04": "Об утверждении положения по охране труда",
    "05": "Об утверждении вводного инструктажа и программы вводного инструктажа",
    "06": "Об утверждении программы первичного инструктажа по охране труда",
    "07": "Об утверждении программы противопожарных инструктажей",
    "08": "О назначении ответственного лица по охране труда",
    "09": "Об утверждении перечней работ по наряду-допуску",
    "10": "О назначении ответственного лица по электробезопасности",
    "11": "О назначении ответственного лица за безопасные работы с ПС",
    "12": "О назначении ответственного лица за выдачу наряда-допуска",
    "13": "О назначении ответственного лица при работах на высоте",
    "14": "О назначении ответственного лица за пожарную безопасность",
    "15": "О назначении ответственных лиц за погрузочно-разгрузочные работы",
    "16": "О назначении ответственного лица за безопасную эксплуатацию сосудов под давлением",
    "17": "О назначении ответственного лица по закрытию помещений по окончании смены",
    "18": "О стажировке вновь принятых рабочих",
    "19": "О допуске к самостоятельной работе после стажировки",
    "20": "О назначении ответственных лиц по электропрогреву бетона в зимних условиях",
}

MANUAL_REVIEW_TARGET_BY_DOC_TYPE: dict[str, str] = {
    "order": "01_orders_and_appointments",
    "hidden_work_act": "05_execution_docs/hidden_work_acts",
    "employee_passport": "02_personnel/employees",
    "unknown": "10_scan_inbox/manual_review",
}

SCANNER_DOC_TYPES: dict[str, str] = {
    "ORDER": "Приказ",
    "AWR": "Акт выполненных работ",
    "PASSPORT": "Удостоверение/протокол",
    "INVOICE": "Счет",
    "UPD": "УПД",
    "TTN": "ТТН",
    "ACT": "Акт",
    "OTHER": "Другое",
}
SCANNER_DOC_TYPES_REQUIRING_EMPLOYEE: set[str] = {"PASSPORT"}
SCANNER_PROFILE_SETTINGS: dict[int, dict[str, object]] = {
    1: {
        "label": "Профиль 1: документы + печати + OCR",
        "dpi": 300,
        "grayscale": True,
    },
    2: {
        "label": "Профиль 2: максимально читаемо",
        "dpi": 400,
        "grayscale": True,
    },
    3: {
        "label": "Профиль 3: профиль 2 + цвет",
        "dpi": 400,
        "grayscale": False,
    },
}

MANUAL_REVIEW_TARGET_BY_SCAN_TYPE: dict[str, str] = {
    "ORDER": "01_orders_and_appointments",
    "AWR": "05_execution_docs/work_reports",
    "PASSPORT": "02_personnel/employees",
    "INVOICE": "08_outgoing_submissions/бухгалтерия/счета",
    "UPD": "08_outgoing_submissions/бухгалтерия/упд",
    "TTN": "08_outgoing_submissions/логистика/ттн",
    "ACT": "05_execution_docs/admission_acts",
    "OTHER": "10_scan_inbox/manual_review",
}

MAINTENANCE_STATIC_FOLDERS: tuple[str, ...] = (
    "00_incoming_requests",
    "01_orders_and_appointments",
    "01_orders_and_appointments/заявки",
    "02_personnel/00_registries",
    "02_personnel/employees",
    "02_personnel/табели",
    "03_hse_and_fire_safety/instructions",
    "03_hse_and_fire_safety/briefings",
    "03_hse_and_fire_safety/incidents_and_microtrauma",
    "03_hse_and_fire_safety/ppe_and_equipment_checks",
    "03_hse_and_fire_safety/permits",
    "04_journals/production",
    "04_journals/labor_safety",
    "05_execution_docs/ppr",
    "05_execution_docs/pprv_work_at_height",
    "05_execution_docs/admission_acts",
    "05_execution_docs/hidden_work_acts",
    "05_execution_docs/work_reports",
    "06_normative_base",
    "07_monthly_control",
    "08_outgoing_submissions",
    "09_archive",
    "09_archive/scan_bundles",
    "10_scan_inbox",
    "10_scan_inbox/manual_review",
)

MAINTENANCE_RESET_TARGETS: tuple[str, ...] = (
    "00_incoming_requests",
    "01_orders_and_appointments/drafts_from_assistant",
    "01_orders_and_appointments/drafts_from_checklist",
    "01_orders_and_appointments/print_pdf_ready",
    "01_orders_and_appointments/print_office",
    "04_journals/production",
    "04_journals/labor_safety",
    "05_execution_docs/work_reports",
    "10_scan_inbox",
)

DEFAULT_ORDER_CONTEXT: dict[str, str] = {
    "ORG_FULL_NAME": "Индивидуальный предприниматель Исмагилов Вадим Шакирович",
    "ORG_SHORT_NAME": "ИП Исмагилов В.Ш.",
    "ORG_INN": "743003400568",
    "ORG_OGRNIP": "319745600096672",
    "ORG_ADDRESS": "454030, Россия, Челябинская область, г. Челябинск, ул. Скульптора Головницкого, д. 166, кв. 3",
    "PROJECT_OBJECT_NAME": "Логистический парк, г. Уфа, Этап 2",
    "PROJECT_ADDRESS": "Республика Башкортостан, город Уфа, Октябрьский район. Этап 2",
    "PROJECT_CODE": "10-2/07-2025ПР-1-КЖ3",
    "ORDER_CITY": "г. Уфа",
    "LEADER_NAME": "Исмагилов Вадим Шакирович",
    "LEADER_POSITION": "Руководитель",
    "LEADER_SIGNATURE": "Исмагилов В.Ш.",
    "RESPONSIBLE_PERSON_2": "Якупов Расим Рафаилович",
    "RESPONSIBLE_POSITION_2": "начальник участка (прораб)",
    "ELECTRICIAN_NAME": "Мельников",
    "ELECTRICIAN_POSITION": "электрик-монтажник",
    "LETTER_TARGET": "Генеральному подрядчику",
    "GENCONTRACTOR_NAME": "ООО «АНТТЕК»",
}

DEFAULT_ORDER_DATE = "01.03.2026"
ORDER_DATE_PATTERN = re.compile(r"^\d{2}\.\d{2}\.\d{4}$")
PRINT_SANITIZE_PATH_PATTERN = re.compile(r"(?:\b\d{2}_personnel|02_personnel)/employees/[\w\-./а-яА-ЯёЁ]+", flags=re.IGNORECASE)
PRINT_SANITIZE_INLINE_ID_PATTERN = re.compile(r"\s*\([^\)]*\bID\b[^\)]*\)", flags=re.IGNORECASE)
PRINT_SANITIZE_MANUAL_NOTE_PATTERN = re.compile(r"\s*\(\s*ручное\s+ведение\s*\)", flags=re.IGNORECASE)
PRINT_SANITIZE_APPENDIX_HEADING_PATTERN = re.compile(r"^\s*#{2,6}\s+Приложение\b", flags=re.IGNORECASE)
PRINT_SANITIZE_TABLE_ROW_PATTERN = re.compile(r"^\s*\|.*\|\s*$")

COMPANY_CARD_KEY_MAP: dict[str, str] = {
    "Полное наименование": "ORG_FULL_NAME",
    "Краткое наименование": "ORG_SHORT_NAME",
    "ИНН": "ORG_INN",
    "ОГРНИП": "ORG_OGRNIP",
    "Юридический адрес": "ORG_ADDRESS",
    "Адрес объекта": "PROJECT_ADDRESS",
    "Город приказа": "ORDER_CITY",
    "Шифр проекта": "PROJECT_CODE",
    "Должность руководителя": "LEADER_POSITION",
    "Руководитель (ФИО)": "LEADER_NAME",
}

EMPLOYEE_REQUIRED_FOLDERS: tuple[str, ...] = (
    "01_identity_and_contract",
    "02_admission_orders",
    "03_briefings_and_training",
    "04_attestation_and_certificates",
    "05_ppe_issue",
    "06_permits_and_work_admission",
    "07_medical_and_first_aid",
    "07_templates_to_print",
)

PROJECT_RULES_BY_CODE: dict[str, EmployeeChecklistRule] = {
    rule.code: rule for rule in PROJECT_TB_ORDER_RULES
}


def resolve_object_root() -> Path:
    return Path(settings.object_root).resolve()


def _to_rel_path(root: Path, target: Path) -> str:
    return str(target.relative_to(root)).replace("\\", "/")


def _resolve_safe_path(root: Path, rel_path: str) -> Path:
    cleaned = (rel_path or "").strip().replace("\\", "/")
    target = (root / cleaned).resolve()

    try:
        target.relative_to(root)
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Путь выходит за пределы объекта") from exc

    return target


def _guess_media_type(path: Path) -> str:
    media_type, _ = mimetypes.guess_type(path.name)
    return media_type or "application/octet-stream"


def _read_docx_preview(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Для предпросмотра DOCX установите python-docx",
        ) from exc

    try:
        document = DocxDocument(str(path))
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Не удалось прочитать DOCX-файл: {exc}",
        ) from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    content = "\n".join(lines).strip()
    if not content:
        content = "Документ DOCX не содержит извлекаемого текста для предпросмотра."

    return content


def _analyze_xlsx_brief(path: Path) -> dict[str, object]:
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Для аналитики XLSX установите openpyxl",
        ) from exc

    workbook = load_workbook(filename=str(path), data_only=True, read_only=True)
    sheets: list[dict[str, object]] = []
    for sheet in workbook.worksheets:
        max_row = int(sheet.max_row or 0)
        max_col = int(sheet.max_column or 0)
        headers: list[str] = []
        if max_row > 0:
            first_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), ())
            headers = [str(cell).strip() for cell in first_row if cell is not None and str(cell).strip()][:10]
        sheets.append(
            {
                "title": sheet.title,
                "rows": max_row,
                "cols": max_col,
                "headers": headers,
            }
        )

    return {
        "file_name": path.name,
        "sheets": sheets,
    }


def _scan_profile_settings(profile: int | None) -> dict[str, object]:
    safe_profile = int(profile or 1)
    return SCANNER_PROFILE_SETTINGS.get(safe_profile, SCANNER_PROFILE_SETTINGS[1]).copy()


def _detect_recompress_profile_for_file(path: Path) -> int:
    rel = str(path).lower().replace("\\", "/")
    if any(token in rel for token in ("счет", "invoice", "упд", "upd", "печат", "подпис", "stamp", "signature")):
        return 3
    return 2


def _recompress_image_file(path: Path, profile_id: int) -> tuple[bool, int, int, int]:
    try:
        from PIL import Image
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Для сжатия сканов установите Pillow: pip install pillow",
        ) from exc

    settings = _scan_profile_settings(profile_id)
    grayscale = bool(settings.get("grayscale", True))
    dpi = int(settings.get("dpi", 300))

    before_size = path.stat().st_size
    suffix = path.suffix.lower()

    try:
        with Image.open(path) as image:
            src = image
            if grayscale:
                src = image.convert("L")
            elif image.mode not in {"RGB", "RGBA"}:
                src = image.convert("RGB")

            save_kwargs: dict[str, object] = {}
            if suffix in {".jpg", ".jpeg"}:
                save_kwargs.update({"quality": 80, "optimize": True, "progressive": True, "dpi": (dpi, dpi)})
            elif suffix in {".tif", ".tiff"}:
                save_kwargs.update({"compression": "tiff_lzw", "dpi": (dpi, dpi)})
            elif suffix == ".png":
                save_kwargs.update({"optimize": True, "compress_level": 9})
            else:
                save_kwargs.update({"dpi": (dpi, dpi)})

            src.save(path, **save_kwargs)
    except HTTPException:
        raise
    except Exception:
        return False, before_size, before_size, profile_id

    after_size = path.stat().st_size
    return True, before_size, after_size, profile_id


def _read_text_preview_content(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx_preview(path), "docx-extracted"

    try:
        return path.read_text(encoding="utf-8"), "utf-8"
    except UnicodeDecodeError:
        return path.read_text(encoding="cp1251", errors="ignore"), "cp1251"


def _sanitize_markdown_for_print(markdown_text: str) -> str:
    normalized = markdown_text.replace("\r\n", "\n").replace("\r", "\n")
    lines: list[str] = []
    skip_appendix_section = False
    strip_first_table_column = False

    def _remove_first_markdown_column(raw_line: str) -> str:
        parts = raw_line.split("|")
        if len(parts) < 4:
            return raw_line
        if parts[0].strip() != "" or parts[-1].strip() != "":
            return raw_line
        del parts[1]
        return "|".join(parts)

    for raw_line in normalized.split("\n"):
        line = re.sub(r"^\s*#\s+##\s+", "## ", raw_line)

        if skip_appendix_section:
            if re.match(r"^\s*#{1,6}\s+", line):
                skip_appendix_section = False
            else:
                continue

        if PRINT_SANITIZE_APPENDIX_HEADING_PATTERN.match(line):
            skip_appendix_section = True
            continue

        line = PRINT_SANITIZE_PATH_PATTERN.sub("по сканам", line)

        line = PRINT_SANITIZE_INLINE_ID_PATTERN.sub("", line)
        line = re.sub(r"\bID\s*\d{3,}(?:\s*[-–]\s*\d{3,})?(?:\s*,\s*\d{3,})*\b", "", line, flags=re.IGNORECASE)
        line = PRINT_SANITIZE_MANUAL_NOTE_PATTERN.sub("", line)
        line = re.sub(r"\s{2,}", " ", line).rstrip()

        if PRINT_SANITIZE_TABLE_ROW_PATTERN.match(line):
            header_row = line.strip().lower()
            if re.match(r"^\|\s*id\s*\|", header_row):
                strip_first_table_column = True

            if strip_first_table_column:
                line = _remove_first_markdown_column(line)
        else:
            strip_first_table_column = False

        lines.append(line)

    return "\n".join(lines)


def _split_markdown_table_cells(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def _is_markdown_table_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", (cell or "").strip()) is not None for cell in cells)


def _render_markdown_fallback_html(markdown_text: str) -> str:
    lines = markdown_text.splitlines()
    html: list[str] = []
    i = 0

    def _render_table(start_index: int) -> tuple[str, int]:
        rows: list[list[str]] = []
        index = start_index
        while index < len(lines):
            raw = lines[index].rstrip()
            stripped = raw.strip()
            if not (stripped.startswith("|") and stripped.endswith("|")):
                break
            cells = _split_markdown_table_cells(stripped)
            if not _is_markdown_table_separator_row(cells):
                rows.append(cells)
            index += 1

        if not rows:
            return "", start_index

        col_count = max(len(row) for row in rows)
        table_parts = ["<table>"]
        table_parts.append("<thead><tr>")
        for c in range(col_count):
            value = rows[0][c] if c < len(rows[0]) else ""
            table_parts.append(f"<th>{escape(value)}</th>")
        table_parts.append("</tr></thead>")

        if len(rows) > 1:
            table_parts.append("<tbody>")
            for row in rows[1:]:
                table_parts.append("<tr>")
                for c in range(col_count):
                    value = row[c] if c < len(row) else ""
                    table_parts.append(f"<td>{escape(value)}</td>")
                table_parts.append("</tr>")
            table_parts.append("</tbody>")

        table_parts.append("</table>")
        return "".join(table_parts), index

    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.strip()

        if not line:
            i += 1
            continue

        # Keep simple inline HTML blocks from templates.
        if line.startswith("<") and line.endswith(">"):
            html.append(line)
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_html, next_index = _render_table(i)
            if table_html:
                html.append(table_html)
                i = next_index
                continue

        if line.startswith("### "):
            html.append(f"<h3>{escape(line[4:].strip())}</h3>")
            i += 1
            continue
        if line.startswith("## "):
            html.append(f"<h2>{escape(line[3:].strip())}</h2>")
            i += 1
            continue
        if line.startswith("# "):
            html.append(f"<h1>{escape(line[2:].strip())}</h1>")
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", line)
        if bullet:
            items: list[str] = [escape(bullet.group(1).strip())]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                nxt_match = re.match(r"^[-*]\s+(.*)$", nxt)
                if not nxt_match:
                    break
                items.append(escape(nxt_match.group(1).strip()))
                i += 1
            html.append("<ul>" + "".join(f"<li>{item}</li>" for item in items) + "</ul>")
            continue

        ordered = re.match(r"^\d+\.\s+(.*)$", line)
        if ordered:
            items = [escape(ordered.group(1).strip())]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                nxt_match = re.match(r"^\d+\.\s+(.*)$", nxt)
                if not nxt_match:
                    break
                items.append(escape(nxt_match.group(1).strip()))
                i += 1
            html.append("<ol>" + "".join(f"<li>{item}</li>" for item in items) + "</ol>")
            continue

        if line.startswith("/") and line.endswith("/"):
            html.append(f"<p class=\"fill-note\">{escape(line)}</p>")
            i += 1
            continue

        if "Исх." in line and "№" in line and "г." in line:
            parts = re.split(r"(Исх\.?\s*№)", line, maxsplit=1)
            if len(parts) == 3:
                left = escape(parts[0].strip())
                right = escape(f"{parts[1]}{parts[2]}".strip())
                html.append(f"<div class=\"meta-row\"><span>{left}</span><span>{right}</span></div>")
                i += 1
                continue

        html.append(f"<p>{escape(line)}</p>")
        i += 1

    return "".join(html)


def _is_admission_letter_document(text: str) -> bool:
    return "прошу разрешить допуск работников и техники" in (text or "").lower()



def _is_timesheet_document(text: str) -> bool:
    lowered = (text or "").lower()
    return "табель учета рабочего времени" in lowered or "табель учёта рабочего времени" in lowered



def _render_markdown_for_print(markdown_text: str) -> str:
    cleaned = _sanitize_markdown_for_print(markdown_text)
    classes = ["doc-body"]
    if _is_admission_letter_document(cleaned):
        classes.append("admission-letter")
    if _is_timesheet_document(cleaned):
        classes.append("timesheet")

    rendered = ""
    if md is not None:
        rendered = md.markdown(
            cleaned,
            extensions=["tables", "fenced_code", "sane_lists", "nl2br"],
        )

    has_table_in_source = "|" in cleaned
    has_table_in_rendered = "<table" in rendered.lower() if rendered else False
    if not rendered or (has_table_in_source and not has_table_in_rendered):
        rendered = _render_markdown_fallback_html(cleaned)

    return f"<article class=\"{' '.join(classes)}\">{rendered}</article>"


def _resolve_chrome_executable() -> Path | None:
    candidates = [
        Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
        Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"),
        Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
        Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
    ]

    chrome_from_path = shutil.which("chrome")
    edge_from_path = shutil.which("msedge")
    if chrome_from_path:
        candidates.append(Path(chrome_from_path))
    if edge_from_path:
        candidates.append(Path(edge_from_path))

    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate

    return None


def _build_pdf_print_html_document(*, title: str, body_html: str, is_permit: bool, is_timesheet: bool = False) -> str:
    if is_timesheet:
        page_size = "A4 landscape"
        page_margin = "6mm 7mm"
        body_font = "9.2pt"
        body_line = "1.16"
        heading_margin = "0.12em 0 0.08em"
        h1_size = "15pt"
        h2_size = "11.6pt"
        h3_size = "10.2pt"
        p_margin = "0.08em 0"
        table_margin = "0.18em 0 0.30em"
        cell_padding = "2px"
        cell_font = "7.2pt"
        cell_line = "1.05"
        hr_margin = "4px 0"
        list_margin = "0.10em 0 0.18em 1.0em"
    else:
        page_size = "A4"
        page_margin = "8mm 8mm" if is_permit else "10mm 10mm"
        body_font = "10.3pt" if is_permit else "11pt"
        body_line = "1.18" if is_permit else "1.26"
        heading_margin = "0.30em 0 0.15em" if is_permit else "0.42em 0 0.22em"
        h1_size = "15pt" if is_permit else "17pt"
        h2_size = "12.8pt" if is_permit else "14pt"
        h3_size = "11.2pt" if is_permit else "12pt"
        p_margin = "0.12em 0" if is_permit else "0.2em 0"
        table_margin = "0.20em 0 0.35em" if is_permit else "0.35em 0 0.55em"
        cell_padding = "3px" if is_permit else "4px"
        cell_font = "9.3pt" if is_permit else "10pt"
        cell_line = "1.12" if is_permit else "1.2"
        hr_margin = "5px 0" if is_permit else "7px 0"
        list_margin = "0.15em 0 0.22em 1.0em" if is_permit else "0.2em 0 0.35em 1.1em"

    return f"""<!doctype html>
<html lang=\"ru\">
<head>
  <meta charset=\"utf-8\" />
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
  <title>{escape(title)}</title>
  <style>
    @page {{
      size: {page_size};
      margin: {page_margin};
    }}
    body {{
      font-family: \"Times New Roman\", serif;
      font-size: {body_font};
      line-height: {body_line};
      margin: 0;
      color: #111;
    }}
    h1, h2, h3 {{
      margin: {heading_margin};
      break-after: avoid-page;
    }}
    h1 {{ font-size: {h1_size}; }}
    h2 {{ font-size: {h2_size}; }}
    h3 {{ font-size: {h3_size}; }}
    p {{ margin: {p_margin}; }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: {table_margin};
      table-layout: fixed;
      word-break: break-word;
      page-break-inside: auto;
    }}
    th, td {{
      border: 1px solid #222;
      padding: {cell_padding};
      vertical-align: top;
      text-align: left;
      font-size: {cell_font};
      line-height: {cell_line};
      overflow-wrap: anywhere;
    }}
    tr {{ page-break-inside: avoid; page-break-after: auto; }}
    hr {{ border: none; border-top: 1px solid #444; margin: {hr_margin}; }}
    ul, ol {{ margin: {list_margin}; }}
    li {{ margin: 0.08em 0; }}
        .doc-body.admission-letter {{ font-size: 12pt; line-height: 1.28; }}
        .doc-body.admission-letter p {{ margin: 0 0 0.18em; }}
        .doc-body.admission-letter .meta-row {{ display: grid; grid-template-columns: 110px 1fr 180px; align-items: end; column-gap: 18px; margin: 0.20em 0 0.35em; }}
        .doc-body.admission-letter .meta-row .number {{ justify-self: center; white-space: nowrap; }}
        .doc-body.admission-letter .meta-row .date {{ justify-self: end; white-space: nowrap; }}
        .doc-body.admission-letter .fill-line {{ border-bottom: 1px solid #222; min-height: 1.15em; padding: 0 4px 1px; text-align: center; margin: 0 auto 0.04em; }}
        .doc-body.admission-letter .org-line {{ width: 74%; }}
        .doc-body.admission-letter .work-line {{ width: 76%; }}
        .doc-body.admission-letter .fill-note {{ text-align: center; margin: 0.02em 0 0.16em; }}
        .doc-body.admission-letter .signature-line {{ margin: 0.10em 0 0.35em; }}
        .doc-body.admission-letter table {{ margin: 0.32em 0 0.62em; }}
        .doc-body.admission-letter th, .doc-body.admission-letter td {{ font-size: 10pt; padding: 3px 5px; overflow-wrap: normal; word-break: normal; }}
        .doc-body.admission-letter table:nth-of-type(1) th:nth-child(1), .doc-body.admission-letter table:nth-of-type(1) td:nth-child(1), .doc-body.admission-letter table:nth-of-type(2) th:nth-child(1), .doc-body.admission-letter table:nth-of-type(2) td:nth-child(1) {{ width: 8%; }}
        .doc-body.admission-letter table:nth-of-type(1) th:nth-child(2), .doc-body.admission-letter table:nth-of-type(1) td:nth-child(2), .doc-body.admission-letter table:nth-of-type(2) th:nth-child(2), .doc-body.admission-letter table:nth-of-type(2) td:nth-child(2) {{ width: 41%; }}
        .doc-body.admission-letter table:nth-of-type(1) th:nth-child(3), .doc-body.admission-letter table:nth-of-type(1) td:nth-child(3), .doc-body.admission-letter table:nth-of-type(2) th:nth-child(3), .doc-body.admission-letter table:nth-of-type(2) td:nth-child(3) {{ width: 31%; }}
        .doc-body.admission-letter table:nth-of-type(1) th:nth-child(4), .doc-body.admission-letter table:nth-of-type(1) td:nth-child(4), .doc-body.admission-letter table:nth-of-type(2) th:nth-child(4), .doc-body.admission-letter table:nth-of-type(2) td:nth-child(4) {{ width: 20%; }}
        .doc-body.timesheet {{ font-size: 9pt; line-height: 1.16; }}
        .doc-body.timesheet h1 {{ margin: 0.08em 0 0.04em; font-size: 15pt; letter-spacing: 0.02em; }}
        .doc-body.timesheet h2 {{ margin: 0.10em 0 0.08em; font-size: 11.5pt; }}
        .doc-body.timesheet p {{ margin: 0 0 0.14em; }}
        .doc-body.timesheet .timesheet-meta {{ display: grid; grid-template-columns: 1.35fr 1fr 1fr; gap: 10px; margin: 0.10em 0 0.22em; }}
        .doc-body.timesheet .timesheet-meta-item {{ border-bottom: 1px solid #222; padding: 0 3px 2px; text-align: center; }}
        .doc-body.timesheet .timesheet-meta-label {{ display: block; margin-top: 2px; font-size: 7.2pt; color: #444; }}
        .doc-body.timesheet .timesheet-note {{ margin: 0.08em 0 0.18em; font-size: 7.4pt; }}
        .doc-body.timesheet table.timesheet-grid {{ margin: 0.14em 0 0.28em; table-layout: fixed; }}
        .doc-body.timesheet table.timesheet-grid th, .doc-body.timesheet table.timesheet-grid td {{ padding: 2px 1px; font-size: 7pt; line-height: 1.05; text-align: center; vertical-align: middle; overflow-wrap: normal; word-break: normal; }}
        .doc-body.timesheet table.timesheet-grid thead th {{ white-space: nowrap; }}
        .doc-body.timesheet table.timesheet-grid th.employee, .doc-body.timesheet table.timesheet-grid td.employee {{ text-align: left; padding-left: 4px; white-space: nowrap; }}
        .doc-body.timesheet table.timesheet-grid .weekday {{ font-size: 6.4pt; letter-spacing: 0.02em; }}
        .doc-body.timesheet table.timesheet-grid .day-cell {{ height: 19px; }}
        .doc-body.timesheet table.timesheet-grid .weekend {{ background: #eef2f7; }}
        .doc-body.timesheet table.timesheet-grid .summary {{ background: #f8fafc; font-weight: 700; }}
  </style>
</head>
<body>
{body_html}
</body>
</html>
"""


def _render_markdown_to_pdf_bytes(target: Path) -> bytes:
    chrome_exe = _resolve_chrome_executable()
    if chrome_exe is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Chrome/Edge не найден для серверной печати Markdown в PDF",
        )

    markdown_text, _ = _read_text_preview_content(target)
    body_html = _render_markdown_for_print(markdown_text)
    is_timesheet = _is_timesheet_document(markdown_text)
    html = _build_pdf_print_html_document(
        title=target.stem,
        body_html=body_html,
        is_permit="PERMIT" in target.stem.upper(),
        is_timesheet=is_timesheet,
    )

    with tempfile.TemporaryDirectory(prefix="arm_md_print_") as temp_dir:
        temp_root = Path(temp_dir)
        html_path = temp_root / "input.html"
        pdf_path = temp_root / "output.pdf"
        html_path.write_text(html, encoding="utf-8")

        cmd = [
            str(chrome_exe),
            "--headless=new",
            "--disable-gpu",
            "--no-pdf-header-footer",
            "--print-to-pdf-no-header",
            f"--print-to-pdf={pdf_path.resolve()}",
            str(html_path.resolve()),
        ]
        completed = subprocess.run(cmd, capture_output=True, text=True)
        if completed.returncode != 0 or not pdf_path.exists():
            detail = (completed.stderr or completed.stdout or "Неизвестная ошибка Chrome/Edge").strip()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Не удалось сформировать PDF из Markdown: {detail}",
            )

        return pdf_path.read_bytes()


def _resolve_print_ready_pdf_for_markdown(root: Path, target: Path) -> Path | None:
    if target.suffix.lower() != ".md":
        return None

    # If a sibling PDF exists next to the markdown file, prefer it for preview/print.
    # This avoids browser print headers/footers from HTML fallback mode.
    sibling_pdf = target.with_suffix(".pdf").resolve()
    try:
        sibling_pdf.relative_to(root)
    except ValueError:
        sibling_pdf = None
    if sibling_pdf and sibling_pdf.exists() and sibling_pdf.is_file():
        return sibling_pdf

    try:
        rel_target = target.relative_to(root)
    except ValueError:
        return None

    if not rel_target.parts or rel_target.parts[0] != "01_orders_and_appointments":
        return None

    for folder_name in ("print_pdf_ready", "print_pdf"):
        candidate = (root / "01_orders_and_appointments" / folder_name / f"{target.stem}.pdf").resolve()
        try:
            candidate.relative_to(root)
        except ValueError:
            continue
        if candidate.exists() and candidate.is_file():
            return candidate

    return None


def _build_print_preview_html(*, file_rel_path: str, file_name: str, body_html: str, auto_print: bool, is_timesheet: bool = False) -> str:
    auto_print_script = (
        "window.addEventListener('load', () => { setTimeout(() => printClean(), 350); });"
        if auto_print
        else ""
    )
    page_size = "A4 landscape" if is_timesheet else "A4"
    print_panel_margin = "6mm 7mm" if is_timesheet else "9mm 11mm"
    wrap_class = "wrap timesheet-preview" if is_timesheet else "wrap"
    panel_class = "panel timesheet-panel" if is_timesheet else "panel"

    return f"""
<!doctype html>
<html lang=\"ru\">
<head>
    <meta charset=\"utf-8\" />
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
    <title>{escape(file_name)}</title>
    <style id=\"printPreviewStyles\">
        :root {{ color-scheme: light; }}
        *, *::before, *::after {{ box-sizing: border-box; }}
        body {{ margin: 0; font-family: \"Segoe UI\", Tahoma, sans-serif; color: #0f172a; background: #f3f4f6; }}
        .wrap {{ max-width: 1060px; margin: 12px auto; padding: 0 12px 16px; }}
        .wrap.timesheet-preview {{ max-width: 1520px; }}
        .panel {{ border: 1px solid #d1d5db; border-radius: 10px; background: #fff; padding: 12px; }}
        .panel.timesheet-panel {{ overflow-x: auto; }}
        .meta {{ color: #475569; font-size: 13px; margin-bottom: 8px; }}
        .actions {{ display: flex; gap: 8px; flex-wrap: wrap; margin-bottom: 10px; }}
        .btn {{ border: none; border-radius: 10px; background: #334155; color: #fff; padding: 8px 12px; cursor: pointer; font-size: 14px; text-decoration: none; }}
        .btn.primary {{ background: #0f766e; }}
        pre {{ margin: 0; white-space: pre-wrap; word-break: break-word; font-family: Consolas, \"Courier New\", monospace; font-size: 13px; }}
        .doc-body {{ font-family: \"Times New Roman\", serif; font-size: 12pt; line-height: 1.28; color: #111; }}
        .doc-body h1 {{ margin: 0.24em 0 0.14em; font-size: 17pt; text-align: center; }}
        .doc-body h2 {{ margin: 0.24em 0 0.14em; font-size: 14pt; text-align: center; }}
        .doc-body h3 {{ margin: 0.24em 0 0.14em; font-size: 12pt; }}
        .doc-body p {{ margin: 0 0 0.36em; }}
        .doc-body ul, .doc-body ol {{ margin: 0.2em 0 0.42em 1.2em; }}
        .doc-body li {{ margin: 0.07em 0; }}
        .doc-body table {{ width: 100%; max-width: 100%; border-collapse: collapse; margin: 0.35em 0 0.55em; table-layout: fixed; word-break: break-word; }}
        .doc-body th, .doc-body td {{ border: 1px solid #222; padding: 4px 6px; vertical-align: top; text-align: left; overflow-wrap: anywhere; }}
        .doc-body th {{ font-weight: 700; }}
        .doc-body hr {{ border: none; border-top: 1px solid #333; margin: 0.35em 0; }}
        .doc-body.admission-letter {{ font-size: 12pt; line-height: 1.28; }}
        .doc-body.admission-letter p {{ margin: 0 0 0.18em; }}
        .doc-body.admission-letter .meta-row {{ display: grid; grid-template-columns: 110px 1fr 180px; align-items: end; column-gap: 18px; margin: 0.20em 0 0.35em; }}
        .doc-body.admission-letter .meta-row .number {{ justify-self: center; white-space: nowrap; }}
        .doc-body.admission-letter .meta-row .date {{ justify-self: end; white-space: nowrap; }}
        .doc-body.admission-letter .fill-line {{ border-bottom: 1px solid #222; min-height: 1.15em; padding: 0 4px 1px; text-align: center; margin: 0 auto 0.04em; }}
        .doc-body.admission-letter .org-line {{ width: 74%; }}
        .doc-body.admission-letter .work-line {{ width: 76%; }}
        .doc-body.admission-letter .fill-note {{ text-align: center; margin: 0.02em 0 0.16em; }}
        .doc-body.admission-letter .signature-line {{ margin: 0.10em 0 0.35em; }}
        .doc-body.admission-letter table {{ margin: 0.32em 0 0.62em; }}
        .doc-body.admission-letter th, .doc-body.admission-letter td {{ font-size: 10pt; padding: 3px 5px; overflow-wrap: normal; word-break: normal; }}
        .doc-body.admission-letter table:nth-of-type(1) th:nth-child(1), .doc-body.admission-letter table:nth-of-type(1) td:nth-child(1), .doc-body.admission-letter table:nth-of-type(2) th:nth-child(1), .doc-body.admission-letter table:nth-of-type(2) td:nth-child(1) {{ width: 8%; }}
        .doc-body.admission-letter table:nth-of-type(1) th:nth-child(2), .doc-body.admission-letter table:nth-of-type(1) td:nth-child(2), .doc-body.admission-letter table:nth-of-type(2) th:nth-child(2), .doc-body.admission-letter table:nth-of-type(2) td:nth-child(2) {{ width: 41%; }}
        .doc-body.admission-letter table:nth-of-type(1) th:nth-child(3), .doc-body.admission-letter table:nth-of-type(1) td:nth-child(3), .doc-body.admission-letter table:nth-of-type(2) th:nth-child(3), .doc-body.admission-letter table:nth-of-type(2) td:nth-child(3) {{ width: 31%; }}
        .doc-body.admission-letter table:nth-of-type(1) th:nth-child(4), .doc-body.admission-letter table:nth-of-type(1) td:nth-child(4), .doc-body.admission-letter table:nth-of-type(2) th:nth-child(4), .doc-body.admission-letter table:nth-of-type(2) td:nth-child(4) {{ width: 20%; }}
        .doc-body.timesheet {{ font-size: 9pt; line-height: 1.16; }}
        .doc-body.timesheet h1 {{ margin: 0.08em 0 0.04em; font-size: 15pt; letter-spacing: 0.02em; }}
        .doc-body.timesheet h2 {{ margin: 0.10em 0 0.08em; font-size: 11.5pt; }}
        .doc-body.timesheet p {{ margin: 0 0 0.14em; }}
        .doc-body.timesheet .timesheet-meta {{ display: grid; grid-template-columns: 1.35fr 1fr 1fr; gap: 10px; margin: 0.10em 0 0.22em; }}
        .doc-body.timesheet .timesheet-meta-item {{ border-bottom: 1px solid #222; padding: 0 3px 2px; text-align: center; }}
        .doc-body.timesheet .timesheet-meta-label {{ display: block; margin-top: 2px; font-size: 7.2pt; color: #444; }}
        .doc-body.timesheet .timesheet-note {{ margin: 0.08em 0 0.18em; font-size: 7.4pt; }}
        .doc-body.timesheet table.timesheet-grid {{ margin: 0.14em 0 0.28em; table-layout: fixed; }}
        .doc-body.timesheet table.timesheet-grid th, .doc-body.timesheet table.timesheet-grid td {{ padding: 2px 1px; font-size: 7pt; line-height: 1.05; text-align: center; vertical-align: middle; overflow-wrap: normal; word-break: normal; }}
        .doc-body.timesheet table.timesheet-grid thead th {{ white-space: nowrap; }}
        .doc-body.timesheet table.timesheet-grid th.employee, .doc-body.timesheet table.timesheet-grid td.employee {{ text-align: left; padding-left: 4px; white-space: nowrap; }}
        .doc-body.timesheet table.timesheet-grid .weekday {{ font-size: 6.4pt; letter-spacing: 0.02em; }}
        .doc-body.timesheet table.timesheet-grid .day-cell {{ height: 19px; }}
        .doc-body.timesheet table.timesheet-grid .weekend {{ background: #eef2f7; }}
        .doc-body.timesheet table.timesheet-grid .summary {{ background: #f8fafc; font-weight: 700; }}
        .preview-frame {{ width: 100%; min-height: 78vh; border: 1px solid #d1d5db; border-radius: 10px; }}
        .preview-image {{ width: 100%; height: auto; border: 1px solid #d1d5db; border-radius: 10px; background: #fff; }}
        @page {{ size: {page_size}; margin: 0; }}
        @media print {{
            .actions {{ display: none !important; }}
            html, body {{ margin: 0 !important; padding: 0 !important; background: #fff; }}
            .wrap {{ max-width: none; margin: 0; padding: 0; }}
            .panel {{ border: none; border-radius: 0; margin: {print_panel_margin}; padding: 0; }}
            .meta {{ display: none; }}
            .doc-body {{ font-size: 11.5pt; line-height: 1.24; }}
            .doc-body table {{ font-size: 10.5pt; }}
            .doc-body a {{ color: #111; text-decoration: none; }}
            .doc-body.admission-letter {{ font-size: 12pt; line-height: 1.24; }}
            .doc-body.admission-letter .meta-row {{ display: grid; grid-template-columns: 110px 1fr 180px; align-items: end; column-gap: 18px; }}
            .doc-body.admission-letter .meta-row .number {{ justify-self: center; white-space: nowrap; }}
            .doc-body.admission-letter .meta-row .date {{ justify-self: end; white-space: nowrap; }}
            .doc-body.admission-letter .fill-line {{ border-bottom: 1px solid #222; min-height: 1.12em; padding: 0 4px 1px; text-align: center; margin: 0 auto 0.04em; }}
            .doc-body.admission-letter .org-line {{ width: 74%; }}
            .doc-body.admission-letter .work-line {{ width: 76%; }}
            .doc-body.admission-letter th, .doc-body.admission-letter td {{ font-size: 9.8pt; padding: 3px 4px; overflow-wrap: normal; word-break: normal; }}
            .doc-body.timesheet {{ font-size: 9pt; line-height: 1.16; }}
            .doc-body.timesheet table.timesheet-grid th, .doc-body.timesheet table.timesheet-grid td {{ font-size: 7pt; padding: 2px 1px; }}
            .doc-body.timesheet table.timesheet-grid thead th {{ white-space: nowrap; }}
            .doc-body.timesheet table.timesheet-grid th.employee, .doc-body.timesheet table.timesheet-grid td.employee {{ white-space: nowrap; }}
        }}
    </style>
</head>
<body>
    <div class=\"{wrap_class}\">
        <div class=\"{panel_class}\">
            <div class=\"meta\">Файл: {escape(file_rel_path)}</div>
            <div class=\"actions\">
                <button class=\"btn primary\" type=\"button\" onclick=\"printClean()\">Печать</button>
                <a class=\"btn\" href=\"/arm/fs/download?rel_path={quote(file_rel_path, safe='')}\" target=\"_blank\" rel=\"noopener\">Скачать оригинал</a>
                <a class=\"btn\" href=\"/arm/dashboard\" target=\"_blank\" rel=\"noopener\">Вернуться в АРМ</a>
            </div>
            {body_html}
        </div>
    </div>
    <script>
        function printClean() {{
            const panel = document.querySelector('.panel');
            if (!panel) {{
                window.print();
                return;
            }}

            const printHost = panel.cloneNode(true);
            const actions = printHost.querySelector('.actions');
            const meta = printHost.querySelector('.meta');
            if (actions) {{
                actions.remove();
            }}
            if (meta) {{
                meta.remove();
            }}

            const styleNode = document.getElementById('printPreviewStyles');
            const styleText = styleNode ? styleNode.textContent : '';
            const popup = window.open('about:blank', '_blank', 'noopener,noreferrer');
            if (!popup) {{
                window.print();
                return;
            }}

            popup.document.open();
            popup.document.write(`<!doctype html><html lang=\"ru\"><head><meta charset=\"utf-8\" /><title></title><style>${{styleText || ''}}</style></head><body><div class=\"wrap\"><div class=\"panel\">${{printHost.innerHTML}}</div></div></body></html>`);
            popup.document.close();
            popup.focus();

            window.setTimeout(() => {{
                popup.print();
                window.setTimeout(() => {{ popup.close(); }}, 220);
            }}, 120);
        }}
    </script>
    <script>{auto_print_script}</script>
</body>
</html>
"""


def _resolve_employee_root(root: Path, employee_rel_path: str) -> Path:
    target = _resolve_safe_path(root=root, rel_path=employee_rel_path)
    if not target.exists() or not target.is_dir():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Папка сотрудника не найдена")

    employees_root = (root / "02_personnel" / "employees").resolve()
    try:
        target.relative_to(employees_root)
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Путь сотрудника должен быть внутри 02_personnel/employees",
        ) from exc

    return target


def _read_employee_profile(employee_root: Path) -> dict[str, str]:
    profile_path = employee_root / "employee_profile.txt"
    if not profile_path.exists() or not profile_path.is_file():
        return {}

    profile: dict[str, str] = {}
    for raw_line in profile_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw_line.strip()
        if not line or ":" not in line:
            continue
        key, value = line.split(":", 1)
        profile[key.strip().lower()] = value.strip()
    return profile


def _employee_display_name(employee_root: Path, profile: dict[str, str]) -> str:
    last_name = profile.get("last_name")
    first_name = profile.get("first_name")
    middle_name = profile.get("middle_name")
    full_name = " ".join(part for part in [last_name, first_name, middle_name] if part)
    if full_name:
        return full_name

    folder_name = employee_root.name
    if "_" in folder_name:
        _, slug = folder_name.split("_", 1)
        return slug.replace("_", " ").strip().title()
    return folder_name


def _infer_profession_key(profession: str) -> str:
    lowered = (profession or "").strip().lower()
    if any(token in lowered for token in ("руковод", "прораб", "начальник")):
        return "supervisor"
    if any(token in lowered for token in ("инженер", "пто", "геодез", "мастер", "техник")):
        return "itr"
    # Welders are handled as a separate non-electric profile in this workflow.
    if any(token in lowered for token in ("свар", "газосвар", "электрогазосвар", "электросвар")):
        return "default"
    if any(token in lowered for token in ("элект", "электро")):
        return "electric"
    return "default"


def _suggest_manual_review_target(predicted_doc_type: str) -> str:
    normalized = (predicted_doc_type or "").strip().lower()
    return MANUAL_REVIEW_TARGET_BY_DOC_TYPE.get(normalized, MANUAL_REVIEW_TARGET_BY_DOC_TYPE["unknown"])


def _extract_scan_subject_tag(file_name: str) -> str:
    name = (file_name or "").strip()
    match = re.match(r"^\d{8}__([A-Z0-9_]+)__([^\.]+)", name)
    if not match:
        return ""
    return (match.group(2) or "").strip().lower()


def _detect_scan_doc_type_from_name(file_name: str) -> str:
    name = (file_name or "").strip()
    match = re.match(r"^\d{8}__([A-Z0-9_]+)__", name)
    if not match:
        return "OTHER"
    scan_doc_type = (match.group(1) or "").upper().strip()
    if scan_doc_type not in SCANNER_DOC_TYPES:
        return "OTHER"
    return scan_doc_type


def _suggest_manual_review_target_from_scan_name(file_name: str) -> str:
    scan_doc_type = _detect_scan_doc_type_from_name(file_name)
    return MANUAL_REVIEW_TARGET_BY_SCAN_TYPE.get(scan_doc_type, MANUAL_REVIEW_TARGET_BY_SCAN_TYPE["OTHER"])


def _collect_rule_matches(target_folder: Path, patterns: tuple[str, ...]) -> list[Path]:
    if not target_folder.exists() or not target_folder.is_dir():
        return []

    seen: set[Path] = set()
    matched: list[Path] = []
    for pattern in patterns:
        glob_pattern = pattern if pattern.startswith("**/") else f"**/{pattern}"
        for item in target_folder.glob(glob_pattern):
            if not item.is_file():
                continue
            resolved = item.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)
            matched.append(resolved)

    return sorted(matched, key=lambda p: p.name.lower())


def _normalize_match_text(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").lower().replace("ё", "е")).strip()


def _rule_tag_keywords(rule: EmployeeChecklistRule) -> tuple[str, ...]:
    by_code: dict[str, tuple[str, ...]] = {
        "E02_IDENTITY_SCAN": ("паспорт", "личност", "договор", "удостовер"),
        "E03_ADMISSION_ORDERS": ("приказ", "допуск", "назнач"),
        "E04_BRIEFINGS": ("инструктаж", "обучен", "проверк знаний"),
        "E05_ATTESTATION": ("аттеста", "удостовер", "протокол"),
        "E06_PPE": ("сиз", "выдач", "каск", "жилет"),
        "E07_PERMITS": ("наряд", "допуск", "опасн"),
        "P11_ORDER_PS": ("подъем", "кран", "пс"),
        "P12_ORDER_PERMIT": ("наряд", "допуск", "повышенн"),
        "P13_ORDER_HEIGHT": ("высот", "на высоте"),
        "P14_ORDER_FIRE": ("пожар", "огнев"),
        "P15_ORDER_LOADING": ("погруз", "разгруз"),
        "P16_ORDER_PRESSURE": ("давлен", "сосуд", "баллон"),
        "P19_ORDER_AFTER_INTERNSHIP": ("самостоятельн", "стажиров"),
        "P21_ADMISSION_LETTER": ("заявк", "допуск", "персонал", "техник"),
    }
    fallback = _normalize_match_text(rule.title)
    return by_code.get(rule.code, tuple(token for token in fallback.split(" ") if len(token) >= 5)[:4])


def _employee_name_tokens(profile: dict[str, str], employee_name: str) -> list[str]:
    tokens = [
        _normalize_match_text(profile.get("last_name", "")),
        _normalize_match_text(profile.get("first_name", "")),
        _normalize_match_text(profile.get("middle_name", "")),
    ]
    compact = [token for token in tokens if token]
    if compact:
        return compact

    parsed = [_normalize_match_text(part) for part in re.split(r"\s+", employee_name or "") if part.strip()]
    return [part for part in parsed if part]


def _doc_text_for_matching(path: Path) -> str:
    suffix = path.suffix.lower()
    content = ""
    if suffix in {".md", ".txt", ".csv", ".json", ".yml", ".yaml", ".ini", ".log"}:
        content = path.read_text(encoding="utf-8", errors="ignore")
    elif suffix == ".docx":
        try:
            content = _read_docx_preview(path)
        except Exception:  # noqa: BLE001
            content = ""
    elif suffix == ".pdf" and PdfReader is not None:
        try:
            reader = PdfReader(str(path))
            chunks: list[str] = []
            for page in reader.pages[:20]:
                text = page.extract_text() or ""
                if text:
                    chunks.append(text)
            content = "\n".join(chunks)
        except Exception:  # noqa: BLE001
            content = ""
    return _normalize_match_text(f"{path.name} {content}")


def _employee_mentioned_in_text(text: str, *, employee_tokens: list[str], employee_id: str | None) -> bool:
    if not text:
        return False

    if len(employee_tokens) >= 2:
        last_name = employee_tokens[0]
        first_name = employee_tokens[1]
        if last_name and (last_name in text) and (first_name in text or f"{first_name[:1]}." in text):
            return True
    elif employee_tokens:
        if employee_tokens[0] in text:
            return True

    if employee_id and len(employee_id) >= 2:
        # Restrict numeric ID matches to explicit ID/tab-number contexts.
        id_re = re.compile(rf"(?:id|таб(?:ельн(?:ый|ого)?)?|номер|сотрудник|№|#)\s*[:=]?\s*0*{re.escape(employee_id)}\b")
        if id_re.search(text):
            return True

    return False


def _collect_related_employee_docs(
    *,
    root: Path,
    rule: EmployeeChecklistRule,
    employee_tokens: list[str],
    employee_id: str | None,
) -> list[Path]:
    # Search only in project-level document areas to map participation links.
    source_folders = [
        "01_orders_and_appointments",
        "03_hse_and_fire_safety",
        "04_journals",
        "05_execution_docs",
        "08_outgoing_submissions",
        "00_incoming_requests",
    ]
    source_ext = {".md", ".txt", ".docx", ".pdf"}
    keywords = _rule_tag_keywords(rule)

    related: list[Path] = []
    seen: set[Path] = set()
    for rel_folder in source_folders:
        folder = _resolve_safe_path(root=root, rel_path=rel_folder)
        if not folder.exists() or not folder.is_dir():
            continue

        for item in folder.rglob("*"):
            if not item.is_file() or item.suffix.lower() not in source_ext:
                continue

            resolved = item.resolve()
            if resolved in seen:
                continue

            text = _doc_text_for_matching(resolved)
            if not _employee_mentioned_in_text(text, employee_tokens=employee_tokens, employee_id=employee_id):
                continue

            if keywords and not any(keyword in text for keyword in keywords):
                continue

            seen.add(resolved)
            related.append(resolved)
            if len(related) >= 12:
                return related

    return related


def _iter_employee_tb_rules(profession: str) -> list[EmployeeChecklistRule]:
    profession_key = _infer_profession_key(profession)
    required_project_codes = PROFESSION_PROJECT_RULES.get(profession_key) or PROFESSION_PROJECT_RULES["default"]

    rules = list(EMPLOYEE_TB_BASE_RULES)
    rules.extend(rule for rule in PROJECT_TB_ORDER_RULES if rule.code in required_project_codes)
    return rules


def _profession_label(key: str) -> str:
    return PROFESSION_LABELS.get(key, PROFESSION_LABELS["custom"])


def _sample_profession_for_group(key: str) -> str:
    if key == "electric":
        return "электромонтажник"
    if key == "supervisor":
        return "прораб"
    if key == "itr":
        return "инженер ПТО"
    return "общий персонал"


def _parse_employee_id_from_folder(folder_name: str) -> str | None:
    match = re.match(r"^(\d{2,10})(?:[_-].*)?$", (folder_name or "").strip())
    return match.group(1) if match else None


def _iter_employee_catalog_rows(root: Path) -> list[EmployeeCatalogRow]:
    employees_root = (root / "02_personnel" / "employees").resolve()
    if not employees_root.exists() or not employees_root.is_dir():
        return []

    rows: list[EmployeeCatalogRow] = []
    for folder in sorted(employees_root.iterdir(), key=lambda p: p.name.lower()):
        if not folder.is_dir():
            continue

        profile = _read_employee_profile(folder)
        position = (profile.get("position") or "").strip() or None
        profession = position or "общий персонал"
        profession_group = _infer_profession_key(profession)

        employee_id = (profile.get("employee_id") or "").strip() or _parse_employee_id_from_folder(folder.name)
        employee_name = _employee_display_name(folder, profile)

        rows.append(
            EmployeeCatalogRow(
                employee_root=folder,
                employee_rel_path=_to_rel_path(root, folder),
                employee_id=employee_id,
                employee_name=employee_name,
                position=position,
                profession=profession,
                profession_group=profession_group,
            )
        )

    return rows


def _matches_profession_filter(row: EmployeeCatalogRow, profession_filter: str | None) -> bool:
    normalized = (profession_filter or "").strip().lower()
    if not normalized or normalized == "all":
        return True

    if normalized in PROFESSION_LABELS:
        return row.profession_group == normalized

    return (
        normalized in (row.profession or "").lower()
        or normalized in (row.position or "").lower()
    )


def _build_employee_catalog(root: Path, profession_filter: str | None = None) -> ArmEmployeeCatalogResponse:
    rows = [row for row in _iter_employee_catalog_rows(root) if _matches_profession_filter(row, profession_filter)]

    return ArmEmployeeCatalogResponse(
        total=len(rows),
        items=[
            ArmEmployeeCatalogItem(
                employee_rel_path=row.employee_rel_path,
                employee_id=row.employee_id,
                employee_name=row.employee_name,
                position=row.position,
                profession_group=row.profession_group,
                profession_label=_profession_label(row.profession_group),
            )
            for row in rows
        ],
        profession_options=[
            ArmProfessionOption(key="default", label=PROFESSION_LABELS["default"]),
            ArmProfessionOption(key="electric", label=PROFESSION_LABELS["electric"]),
            ArmProfessionOption(key="supervisor", label=PROFESSION_LABELS["supervisor"]),
            ArmProfessionOption(key="itr", label=PROFESSION_LABELS["itr"]),
        ],
    )


def _build_employee_checklist_overview(
    *,
    root: Path,
    profession_filter: str | None,
) -> ArmEmployeeChecklistOverviewResponse:
    rows = [row for row in _iter_employee_catalog_rows(root) if _matches_profession_filter(row, profession_filter)]

    grouped_rows: dict[str, list[EmployeeCatalogRow]] = defaultdict(list)
    for row in rows:
        grouped_rows[row.profession_group].append(row)

    groups: list[ArmEmployeeOverviewGroup] = []
    for profession_group, group_rows in sorted(grouped_rows.items(), key=lambda item: _profession_label(item[0])):
        missing_counter: dict[str, int] = defaultdict(int)
        employee_items: list[ArmEmployeeOverviewEmployee] = []

        for row in group_rows:
            analysis = _build_employee_checklist_data(
                root=root,
                employee_root=row.employee_root,
                profession=row.profession,
            )
            top_missing_codes = [item.code for item in analysis.items if not item.ready][:8]
            for item in analysis.items:
                if not item.ready:
                    missing_counter[item.code] += 1

            employee_items.append(
                ArmEmployeeOverviewEmployee(
                    employee_rel_path=analysis.employee_rel_path,
                    employee_id=analysis.employee_id,
                    employee_name=analysis.employee_name or row.employee_name,
                    position=analysis.profile_position,
                    profession=analysis.profession,
                    total_required=analysis.total_required,
                    ready_count=analysis.ready_count,
                    missing_count=analysis.missing_count,
                    progress_percent=analysis.progress_percent,
                    top_missing_codes=top_missing_codes,
                )
            )

        # Keep full list of activities visible for type-level planning.
        sample_profession = _sample_profession_for_group(profession_group)
        group_rules = _iter_employee_tb_rules(sample_profession)
        activities = [
            ArmEmployeeOverviewAction(
                code=rule.code,
                title=rule.title,
                guidance=rule.guidance,
                scope=rule.scope,
                missing_employees=missing_counter.get(rule.code, 0),
            )
            for rule in group_rules
        ]

        employees_total = len(employee_items)
        ready_employees = sum(1 for employee in employee_items if employee.missing_count == 0)
        avg_progress = round(
            (sum(employee.progress_percent for employee in employee_items) / employees_total) if employees_total else 0.0,
            1,
        )

        employee_items.sort(key=lambda row: (-row.missing_count, row.employee_name.lower()))

        groups.append(
            ArmEmployeeOverviewGroup(
                profession_group=profession_group,  # type: ignore[arg-type]
                profession_label=_profession_label(profession_group),
                employees_total=employees_total,
                ready_employees=ready_employees,
                average_progress_percent=avg_progress,
                missing_actions=activities,
                employees=employee_items,
            )
        )

    return ArmEmployeeChecklistOverviewResponse(
        generated_at=datetime.now(tz=timezone.utc),
        profession_filter=(profession_filter or "").strip() or None,
        groups=groups,
    )


def _build_employee_checklist_data(
    *,
    root: Path,
    employee_root: Path,
    profession: str | None,
) -> ArmEmployeeChecklistResponse:
    profile = _read_employee_profile(employee_root)
    employee_name = _employee_display_name(employee_root, profile)
    employee_id = profile.get("employee_id")
    employee_tokens = _employee_name_tokens(profile, employee_name)
    effective_profession = (profession or profile.get("position") or "общий персонал").strip()
    rules = _iter_employee_tb_rules(effective_profession)

    items: list[ArmEmployeeChecklistItem] = []
    for rule in rules:
        if rule.scope == "project":
            target_folder = _resolve_safe_path(root=root, rel_path=rule.folder)
        else:
            target_folder = employee_root if rule.folder == "." else (employee_root / rule.folder).resolve()

        matches = _collect_rule_matches(target_folder=target_folder, patterns=rule.patterns)
        related = _collect_related_employee_docs(
            root=root,
            rule=rule,
            employee_tokens=employee_tokens,
            employee_id=employee_id,
        )
        found_count = len(matches)
        ready = found_count >= rule.required_count

        items.append(
            ArmEmployeeChecklistItem(
                code=rule.code,
                title=rule.title,
                folder_rel_path=_to_rel_path(root, target_folder),
                expected_patterns=list(rule.patterns),
                required_count=rule.required_count,
                found_count=found_count,
                ready=ready,
                found_files=[_to_rel_path(root, match) for match in matches[:8]],
                related_count=len(related),
                related_files=[_to_rel_path(root, match) for match in related[:8]],
                guidance=rule.guidance,
            )
        )

    total_required = len(items)
    ready_count = sum(1 for item in items if item.ready)
    missing_count = total_required - ready_count
    progress_percent = round((ready_count / total_required) * 100.0 if total_required else 0.0, 1)

    return ArmEmployeeChecklistResponse(
        employee_rel_path=_to_rel_path(root, employee_root),
        employee_id=employee_id,
        employee_name=employee_name,
        profile_position=profile.get("position"),
        profession=effective_profession,
        total_required=total_required,
        ready_count=ready_count,
        missing_count=missing_count,
        progress_percent=progress_percent,
        items=items,
    )


def _build_draft_content(
    *,
    root: Path,
    employee_root: Path,
    rule: EmployeeChecklistRule,
    employee_rel_path: str,
    employee_name: str,
    employee_id: str | None,
    profession: str,
    context_overrides: dict[str, str] | None = None,
) -> str:
    template_text = _load_order_template_text(rule.code)
    if template_text:
        order_date_override = None
        if context_overrides:
            order_date_override = (
                context_overrides.get("ORDER_DATE")
                or context_overrides.get("ISSUE_DATE")
            )

        context = _build_order_template_context(
            root=root,
            employee_root=employee_root,
            rule=rule,
            employee_rel_path=employee_rel_path,
            employee_name=employee_name,
            employee_id=employee_id,
            profession=profession,
            order_date_override=order_date_override,
        )
        if context_overrides:
            context.update({key: value for key, value in context_overrides.items() if value is not None})
        return _render_order_template(template_text, context)

    generated_at = datetime.now(tz=timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    return (
        f"# Черновик: {rule.title}\n\n"
        f"- Код чеклиста: {rule.code}\n"
        f"- Сотрудник: {employee_name}\n"
        f"- Код сотрудника: {employee_id or '-'}\n"
        f"- Профессия/роль: {profession}\n"
        f"- Папка сотрудника: {employee_rel_path}\n"
        f"- Сгенерировано: {generated_at}\n\n"
        "## Что нужно заполнить\n"
        f"- {rule.guidance}\n"
        "- Проверьте соответствие формату документов, принятому генподрядчиком.\n"
        "- После заполнения переименуйте файл в боевой формат объекта.\n"
    )


def _load_order_template_text(rule_code: str) -> str | None:
    template_name = ORDER_DRAFT_TEMPLATE_MAP.get(rule_code)
    if not template_name:
        return None

    template_path = ORDER_FORMAT_TEMPLATES_DIR / template_name
    if not template_path.exists() or not template_path.is_file():
        return None

    try:
        return template_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return template_path.read_text(encoding="cp1251", errors="ignore")


def _load_company_context() -> dict[str, str]:
    context = dict(DEFAULT_ORDER_CONTEXT)
    if not COMPANY_REQUISITES_CARD_PATH.exists() or not COMPANY_REQUISITES_CARD_PATH.is_file():
        return context

    for raw_line in COMPANY_REQUISITES_CARD_PATH.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw_line.strip()
        if not line or ":" not in line:
            continue
        key, value = line.split(":", 1)
        mapped_key = COMPANY_CARD_KEY_MAP.get(key.strip())
        if mapped_key and value.strip():
            context[mapped_key] = value.strip()

    return context


def _short_signature(full_name: str) -> str:
    clean = " ".join(part for part in (full_name or "").strip().split() if part)
    if not clean:
        return "________________"

    parts = clean.split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    if len(parts) == 2:
        return f"{parts[0]} {parts[1][0]}."
    return clean


def _extract_rule_number(code: str) -> str:
    match = re.search(r"(\d+)", code or "")
    return match.group(1) if match else "__"


def _collect_team_members(root: Path, employee_root: Path, profile: dict[str, str]) -> list[tuple[str, str]]:
    employees_root = (root / "02_personnel" / "employees").resolve()
    desired_team = (profile.get("team") or "").strip().lower()

    if not desired_team:
        fallback_name = _employee_display_name(employee_root, profile)
        fallback_position = (profile.get("position") or "Сотрудник").strip()
        return [(fallback_name, fallback_position)]

    rows: list[tuple[str, str]] = []
    if employees_root.exists() and employees_root.is_dir():
        for folder in sorted(employees_root.iterdir(), key=lambda p: p.name.lower()):
            if not folder.is_dir():
                continue

            current_profile = _read_employee_profile(folder)
            if desired_team:
                if (current_profile.get("team") or "").strip().lower() != desired_team:
                    continue

            worker_name = _employee_display_name(folder, current_profile)
            worker_position = (current_profile.get("position") or "Сотрудник").strip()
            rows.append((worker_name, worker_position))

    if rows:
        return rows

    fallback_name = _employee_display_name(employee_root, profile)
    fallback_position = (profile.get("position") or "Сотрудник").strip()
    return [(fallback_name, fallback_position)]


def _workers_table_markdown(rows: list[tuple[str, str]]) -> str:
    lines = ["| ФИО | Должность |", "|---|---|"]
    for name, position in rows:
        safe_name = (name or "Сотрудник").replace("|", "/")
        safe_position = (position or "Сотрудник").replace("|", "/")
        lines.append(f"| {safe_name} | {safe_position} |")
    return "\n".join(lines)


def _workers_list_table_markdown(rows: list[tuple[str, str, str]]) -> str:
    table = [
        "| № п/п | Фамилия, имя, отчество | Профессия, должность | Дата рождения (число, месяц, год) |",
        "|---|---|---|---|",
    ]

    safe_rows = rows or [("________________", "________________", "________________")]
    index = 1
    for name, position, birth_date in safe_rows[:20]:
        table.append(
            "| "
            f"{index} | {_sanitize_markdown_cell(name)} | {_sanitize_markdown_cell(position)} | {_sanitize_markdown_cell(birth_date)} |"
        )
        index += 1

    while index <= max(6, len(safe_rows)):
        table.append(f"| {index} |  |  |  |")
        index += 1

    return "\n".join(table)


def _equipment_list_table_markdown(rows: list[tuple[str, str, str]]) -> str:
    table = [
        "| № п/п | Наименование техники | Марка, модель | Гос. номер |",
        "|---|---|---|---|",
    ]

    safe_rows = rows or [("Автомобиль", "________________", "________________")]
    index = 1
    for equipment_name, model, plate in safe_rows[:20]:
        table.append(
            "| "
            f"{index} | {_sanitize_markdown_cell(equipment_name)} | {_sanitize_markdown_cell(model)} | {_sanitize_markdown_cell(plate)} |"
        )
        index += 1

    while index <= max(5, len(safe_rows)):
        table.append(f"| {index} |  |  |  |")
        index += 1

    return "\n".join(table)


def _workers_bullets(rows: list[tuple[str, str]]) -> str:
    return "\n".join(f"- {name}, {position}" for name, position in rows)


def _sanitize_markdown_cell(value: str) -> str:
    return (value or "").replace("|", "/").replace("\n", " ").strip()


def _workers_ack_table_markdown(rows: list[tuple[str, str]], order_date: str) -> str:
    lines = [
        '<table style="width: 100%; border-collapse: collapse; margin-top: 0.5em;">',
        '  <colgroup>',
        '    <col style="width: 6%;" />',
        '    <col style="width: 44%;" />',
        '    <col style="width: 25%;" />',
        '    <col style="width: 25%;" />',
        '  </colgroup>',
        '  <tr>',
        '    <th style="border: 1px solid black; padding: 4px 2px; text-align: center;">№ п/п</th>',
        '    <th style="border: 1px solid black; padding: 4px 6px; text-align: center;">Ф.И.О. работника, профессия, должность</th>',
        '    <th style="border: 1px solid black; padding: 4px 6px; text-align: center;">Дата ознакомления</th>',
        '    <th style="border: 1px solid black; padding: 4px 6px; text-align: center;">Личная подпись работника</th>',
        '  </tr>',
    ]

    if not rows:
        rows = [("________________", "________________")]

    index = 1
    for name, position in rows[:18]:
        full = _sanitize_markdown_cell(f"{name}, {position}")
        lines += [
            '  <tr>',
            f'    <td style="border: 1px solid black; padding: 4px 2px; text-align: center;">{index}.</td>',
            f'    <td style="border: 1px solid black; padding: 4px 6px;">{full}</td>',
            '    <td style="border: 1px solid black; padding: 4px 6px;"></td>',
            '    <td style="border: 1px solid black; padding: 4px 6px;"></td>',
            '  </tr>',
        ]
        index += 1

    # Keep several blank rows for manual signatures.
    target_rows = max(4, len(rows))
    while index <= target_rows:
        lines += [
            '  <tr>',
            f'    <td style="border: 1px solid black; padding: 4px 2px; text-align: center;">{index}.</td>',
            '    <td style="border: 1px solid black; padding: 4px 6px;"></td>',
            '    <td style="border: 1px solid black; padding: 4px 6px;"></td>',
            '    <td style="border: 1px solid black; padding: 4px 6px;"></td>',
            '  </tr>',
        ]
        index += 1

    lines.append('</table>')
    return "\n".join(lines)


def _order_register_rows_markdown(order_year: str, order_date: str) -> str:
    rows = [
        "| № п/п | Наименование приказа | Дата | Примечание |",
        "|---|---|---|---|",
    ]

    for number in range(1, 31):
        number_code = f"{number:02d}"
        title = ORDER_REGISTER_TITLE_BY_NUMBER.get(number_code, "")
        order_code = f"{number_code}/{order_year}"
        date_cell = order_date if title else ""
        rows.append(f"| {order_code} | {_sanitize_markdown_cell(title)} | {date_cell} |  |")

    return "\n".join(rows)


def _render_order_template(template_text: str, context: dict[str, str]) -> str:
    rendered = template_text
    for key, value in context.items():
        rendered = rendered.replace(f"{{{{{key}}}}}", value)

    # Keep placeholders visible if context is missing.
    return re.sub(r"\{\{[A-Z0-9_]+\}\}", "________________", rendered)


def _parse_order_date_or_default(order_date_raw: str | None) -> datetime:
    candidate = (order_date_raw or "").strip()
    if not candidate:
        candidate = DEFAULT_ORDER_DATE

    if not ORDER_DATE_PATTERN.fullmatch(candidate):
        candidate = DEFAULT_ORDER_DATE

    try:
        return datetime.strptime(candidate, "%d.%m.%Y")
    except ValueError:
        return datetime.strptime(DEFAULT_ORDER_DATE, "%d.%m.%Y")


def _build_order_template_context(
    *,
    root: Path,
    employee_root: Path,
    rule: EmployeeChecklistRule,
    employee_rel_path: str,
    employee_name: str,
    employee_id: str | None,
    profession: str,
    order_date_override: str | None = None,
) -> dict[str, str]:
    profile = _read_employee_profile(employee_root)
    company = _load_company_context()

    order_date_dt = _parse_order_date_or_default(order_date_override)
    order_date = order_date_dt.strftime("%d.%m.%Y")
    order_date_with_suffix = f"{order_date}г."
    order_year = order_date_dt.strftime("%Y")
    permit_end = (order_date_dt + timedelta(days=13)).strftime("%d.%m.%Y")

    rule_number = _extract_rule_number(rule.code)
    responsible_name = employee_name or _employee_display_name(employee_root, profile)
    responsible_position = (profession or profile.get("position") or "Ответственное лицо").strip()

    workers = _collect_team_members(root=root, employee_root=employee_root, profile=profile)
    workers_table = _workers_table_markdown(workers)
    workers_bullets = _workers_bullets(workers)
    workers_ack_table = _workers_ack_table_markdown(workers, order_date)
    register_rows = _order_register_rows_markdown(order_year, order_date)

    team_name = (profile.get("team") or "Строительно-монтажная бригада").strip()
    project_address = company.get("PROJECT_ADDRESS") or DEFAULT_ORDER_CONTEXT["PROJECT_ADDRESS"]
    leader_name = company.get("LEADER_NAME") or DEFAULT_ORDER_CONTEXT["LEADER_NAME"]
    leader_position = company.get("LEADER_POSITION") or DEFAULT_ORDER_CONTEXT["LEADER_POSITION"]

    context = {
        **company,
        "ORDER_NUMBER": rule_number,
        "ORDER_YEAR": order_year,
        "ORDER_DATE": order_date,
        "ORDER_DATE_WITH_SUFFIX": order_date_with_suffix,
        "ISSUE_DATE": order_date,
        "ISSUE_CITY": company.get("ORDER_CITY", "г. Уфа"),
        "PROJECT_OBJECT_NAME": company.get("PROJECT_OBJECT_NAME", DEFAULT_ORDER_CONTEXT["PROJECT_OBJECT_NAME"]),
        "PROJECT_ADDRESS": project_address,
        "PROJECT_CODE": company.get("PROJECT_CODE", DEFAULT_ORDER_CONTEXT["PROJECT_CODE"]),
        "ORG_REGISTRATION_CITY": company.get("ORG_REGISTRATION_CITY", company.get("ORDER_CITY", "г. Уфа")),
        "RESPONSIBLE_PERSON": responsible_name,
        "RESPONSIBLE_POSITION": responsible_position,
        "RESPONSIBLE_SIGNATURE": _short_signature(responsible_name),
        "RESPONSIBLE_PERSON_2": company.get("RESPONSIBLE_PERSON_2", "________________"),
        "RESPONSIBLE_POSITION_2": company.get("RESPONSIBLE_POSITION_2", "________________"),
        "ELECTRICIAN_NAME": company.get("ELECTRICIAN_NAME", DEFAULT_ORDER_CONTEXT["ELECTRICIAN_NAME"]),
        "ELECTRICIAN_POSITION": company.get("ELECTRICIAN_POSITION", DEFAULT_ORDER_CONTEXT["ELECTRICIAN_POSITION"]),
        "LEADER_NAME": leader_name,
        "LEADER_POSITION": leader_position,
        "LEADER_SIGNATURE": _short_signature(leader_name),
        "PREPARED_BY_NAME": responsible_name,
        "PREPARED_BY_POSITION": (profile.get("position") or "Начальник участка").strip(),
        "PREPARED_BY_SIGNATURE": _short_signature(responsible_name),
        "EMPLOYEE_NAME": responsible_name,
        "EMPLOYEE_POSITION": responsible_position,
        "EMPLOYEE_ID": employee_id or "-",
        "EMPLOYEE_REL_PATH": employee_rel_path,
        "TEAM_NAME": team_name,
        "WORKERS_TABLE": workers_table,
        "WORKERS_BULLETS": workers_bullets,
        "WORKERS_ACK_TABLE": workers_ack_table,
        "ORDER_REGISTER_ROWS": register_rows,
        "PERMIT_START_DATE": order_date,
        "PERMIT_END_DATE": permit_end,
        "GUIDANCE": rule.guidance,
        "GENERATED_AT_UTC": datetime.now(tz=timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC"),
        "EQUIPMENT_BULLETS": "- Автокран\n- Электросварочный аппарат\n- Такелажные приспособления",
        "LETTER_NUMBER": f"{rule_number}/{order_year}",
        "ADDRESSEE_ORG": company.get("ADDRESSEE_ORG", "ООО «Новый Элемент»"),
        "ADDRESSEE_PERSON": company.get("ADDRESSEE_PERSON", "Паньковой Т.Н."),
        "ADDRESSEE_ADDRESS": company.get("ADDRESSEE_ADDRESS", "г. Уфа"),
        "WORK_PERIOD": f"с {order_date} по __.__.____",
        "WORKS_NAME": company.get("WORKS_NAME", "строительно-монтажных работ"),
        "WORKERS_LIST_TABLE": _workers_list_table_markdown(
            [(name, pos, "________________") for name, pos in workers]
        ),
        "EQUIPMENT_LIST_TABLE": _equipment_list_table_markdown(
            [("Легковой автомобиль", "________________", "________________")]
        ),
    }

    return context


def _slugify_filename(value: str) -> str:
    normalized = "".join(char.lower() if char.isalnum() else "_" for char in value)
    while "__" in normalized:
        normalized = normalized.replace("__", "_")
    normalized = normalized.strip("_")
    return normalized or "document"


def _clear_directory_contents(target: Path) -> int:
    if not target.exists() or not target.is_dir():
        return 0

    removed = 0
    for item in list(target.iterdir()):
        if item.is_dir():
            shutil.rmtree(item)
        else:
            item.unlink()
        removed += 1
    return removed


def _ensure_object_structure_for_maintenance(root: Path) -> None:
    root.mkdir(parents=True, exist_ok=True)
    for rel_path in MAINTENANCE_STATIC_FOLDERS:
        (root / rel_path).mkdir(parents=True, exist_ok=True)


def _pick_maintenance_seed_employee(root: Path) -> EmployeeCatalogRow | None:
    rows = _iter_employee_catalog_rows(root)
    if not rows:
        return None

    for row in rows:
        if row.profession_group == "supervisor":
            return row
    return rows[0]


def _rebuild_project_order_drafts_from_seed_employee(
    *,
    root: Path,
    seed_employee: EmployeeCatalogRow,
    overwrite: bool,
    order_date: str | None = None,
) -> int:
    employee_root = _resolve_employee_root(root=root, employee_rel_path=seed_employee.employee_rel_path)
    analysis = _build_employee_checklist_data(root=root, employee_root=employee_root, profession=None)

    project_rules = sorted(
        [rule for rule in _iter_employee_tb_rules(analysis.profession) if rule.scope == "project"],
        key=lambda rule: rule.code,
    )
    if not project_rules:
        return 0

    project_output_dir = (root / "01_orders_and_appointments" / "drafts_from_checklist").resolve()
    project_output_dir.mkdir(parents=True, exist_ok=True)

    employee_name = analysis.employee_name or employee_root.name
    created_files = 0
    context_overrides = (
        {
            "ORDER_DATE": order_date,
            "ISSUE_DATE": order_date,
        }
        if order_date
        else None
    )

    for rule in project_rules:
        file_name = ORDER_DRAFT_FILE_NAME_MAP.get(rule.code) or f"{rule.code}_{_slugify_filename(rule.title)}_draft.md"
        draft_path = project_output_dir / file_name

        if draft_path.exists() and not overwrite:
            continue

        draft_content = _build_draft_content(
            root=root,
            employee_root=employee_root,
            rule=rule,
            employee_rel_path=analysis.employee_rel_path,
            employee_name=employee_name,
            employee_id=analysis.employee_id,
            profession=analysis.profession,
            context_overrides=context_overrides,
        )
        draft_path.write_text(draft_content, encoding="utf-8")
        created_files += 1

    return created_files


def _run_scanner_command(args: list[str]) -> subprocess.CompletedProcess[str]:
    script_path = PROJECT_ROOT / "scripts" / "scanner_control.py"
    cmd = [sys.executable, str(script_path), *args]
    try:
        return subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            check=False,
            timeout=SCANNER_COMMAND_TIMEOUT_SEC,
        )
    except subprocess.TimeoutExpired as exc:
        return subprocess.CompletedProcess(
            args=cmd,
            returncode=124,
            stdout=exc.stdout or "",
            stderr=(
                f"Команда сканера превысила таймаут {SCANNER_COMMAND_TIMEOUT_SEC} сек. "
                "Проверьте питание сканера и подключение по USB/WIA."
            ),
        )


def _parse_scanner_list_stdout(stdout: str) -> list[ArmScannerDevice]:
    devices: list[ArmScannerDevice] = []
    current_index: int | None = None

    for raw_line in stdout.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("[") and "]" in line:
            try:
                left, right = line.split("]", 1)
                idx = int(left.strip("[] "))
                name = right.strip()
                devices.append(ArmScannerDevice(index=idx, name=name, device_id=None))
                current_index = idx
            except Exception:
                continue
            continue

        if current_index is not None and devices:
            if line.startswith("\\") or line.lower().startswith("usb"):
                devices[-1].device_id = line

    return devices


def _humanize_scanner_error(stderr: str, stdout: str) -> str:
    raw = (stderr or stdout or "").strip()
    if not raw:
        return "Команда сканирования завершилась с ошибкой"

    lowered = raw.lower()
    if "employee_id is required for passport" in lowered:
        return "Для типа «Удостоверение/протокол» заполните код сотрудника и повторите сканирование."
    if "no scanner" in lowered or "сканер не найден" in lowered:
        return "Сканер не найден. Проверьте USB/WIA-подключение и обновите список устройств."
    if "таймаут" in lowered or "timeout" in lowered:
        return f"Сканер не ответил вовремя (таймаут {SCANNER_COMMAND_TIMEOUT_SEC} сек). Повторите попытку."

    lines = [line.strip() for line in raw.splitlines() if line.strip()]
    if lines:
        last_line = lines[-1]
        if last_line.lower().startswith("error:"):
            last_line = last_line[6:].strip()
        return f"Ошибка сканирования: {last_line}"

    return raw


def _count_files(root: Path, folder: str, pattern: str | tuple[str, ...]) -> int:
    target = root / folder
    if not target.exists() or not target.is_dir():
        return 0

    patterns = (pattern,) if isinstance(pattern, str) else tuple(pattern)
    if patterns == ("*",):
        return sum(1 for item in target.rglob("*") if item.is_file())

    seen: set[Path] = set()
    for current_pattern in patterns:
        if current_pattern == "*":
            iterator = target.rglob("*")
        else:
            iterator = target.rglob(current_pattern)

        for item in iterator:
            if not item.is_file():
                continue
            seen.add(item.resolve())

    return len(seen)


def _list_files(root: Path, folder: str, pattern: str | tuple[str, ...], limit: int = 8) -> list[str]:
    target = root / folder
    if not target.exists() or not target.is_dir():
        return []

    patterns = (pattern,) if isinstance(pattern, str) else tuple(pattern)
    seen: set[Path] = set()
    for current_pattern in patterns:
        iterator = target.rglob("*") if current_pattern == "*" else target.rglob(current_pattern)
        for item in iterator:
            if item.is_file():
                seen.add(item.resolve())

    result = [_to_rel_path(root, path) for path in sorted(seen, key=lambda item: str(item).lower())]
    return result[:limit]


def _build_checklist(root: Path) -> list[ArmChecklistItem]:
    items: list[ArmChecklistItem] = []
    for rule in CHECKLIST_RULES:
        found = _count_files(root=root, folder=rule.folder, pattern=rule.pattern)
        found_files = _list_files(root=root, folder=rule.folder, pattern=rule.pattern)
        items.append(
            ArmChecklistItem(
                code=rule.code,
                title=rule.title,
                location=str((root / rule.folder).as_posix()),
                required_min=rule.required_min,
                found=found,
                ready=found >= rule.required_min,
                found_files=found_files,
            )
        )
    return items


def _collect_metrics(db: Session, root: Path) -> ArmMetrics:
    db_documents_total = int(db.execute(select(func.count(Document.id))).scalar_one())
    db_journal_entries_total = int(db.execute(select(func.count(JournalEntry.id))).scalar_one())
    db_schedules_total = int(db.execute(select(func.count(WorkSchedule.id))).scalar_one())

    orders_md_total = _count_files(root=root, folder="01_orders_and_appointments", pattern="*.md")
    orders_pdf_ready_total = _count_files(
        root=root,
        folder="01_orders_and_appointments/print_pdf_ready",
        pattern="*.pdf",
    )
    journals_production_total = _count_files(root=root, folder="04_journals/production", pattern="*")
    journals_labor_safety_total = _count_files(root=root, folder="04_journals/labor_safety", pattern="*")

    scan_root = root / "10_scan_inbox"
    scan_inbox_pending_total = 0
    if scan_root.exists():
        scan_inbox_pending_total = sum(1 for item in scan_root.iterdir() if item.is_file())

    scan_manual_review_total = _count_files(
        root=root,
        folder="10_scan_inbox/manual_review",
        pattern="*",
    )

    return ArmMetrics(
        db_documents_total=db_documents_total,
        db_journal_entries_total=db_journal_entries_total,
        db_schedules_total=db_schedules_total,
        orders_md_total=orders_md_total,
        orders_pdf_ready_total=orders_pdf_ready_total,
        journals_production_total=journals_production_total,
        journals_labor_safety_total=journals_labor_safety_total,
        scan_inbox_pending_total=scan_inbox_pending_total,
        scan_manual_review_total=scan_manual_review_total,
    )


def _build_todos(
    checklist: list[ArmChecklistItem],
    metrics: ArmMetrics,
    local_llm_reachable: bool,
) -> list[ArmTodoItem]:
    todos: list[ArmTodoItem] = []
    root = resolve_object_root()

    for item in checklist:
        if item.ready:
            continue

        action_path: str | None = None
        try:
            action_path = _to_rel_path(root, Path(item.location).resolve())
        except Exception:  # noqa: BLE001
            action_path = None

        todos.append(
            ArmTodoItem(
                priority="high" if item.code in {"3.1", "3.2", "4"} else "medium",
                title=f"Закрыть позицию {item.code}: {item.title}",
                details=f"Найдено {item.found} из минимум {item.required_min}",
                action_path=action_path,
            )
        )

    if metrics.scan_manual_review_total > 0:
        todos.append(
            ArmTodoItem(
                priority="high",
                title="Разобрать папку ручного разбора после сканирования",
                details=f"В очереди {metrics.scan_manual_review_total} файлов",
                action_path="10_scan_inbox/manual_review",
            )
        )

    if metrics.orders_pdf_ready_total == 0 and metrics.orders_md_total > 0:
        todos.append(
            ArmTodoItem(
                priority="medium",
                title="Собрать пакет PDF для печати",
                details="Запустить scripts/build_and_open_print_pack.ps1",
                action_path="01_orders_and_appointments",
            )
        )

    if not local_llm_reachable:
        todos.append(
            ArmTodoItem(
                priority="medium",
                title="Проверить доступность локальной LLM",
                details="Проверить Ollama: /local-llm/status",
            )
        )

    for status_row in _collect_periodic_doc_statuses(root=root):
        if status_row.is_due:
            details = status_row.rule.details
            if status_row.days_since_update is None:
                details = f"Документ не найден. {details}"
            else:
                details = (
                    f"Последнее обновление {status_row.days_since_update} дн. назад, "
                    f"норма {status_row.rule.period_days} дн. {details}"
                )
            todos.append(
                ArmTodoItem(
                    priority="high",
                    title=f"Периодический документ: {status_row.rule.title}",
                    details=details,
                    action_path=status_row.action_path,
                )
            )

    # Keep daily list compact for foreman workflow.
    return todos[:12]


def _collect_periodic_doc_statuses(root: Path) -> list[PeriodicDocStatus]:
    now = datetime.now(tz=timezone.utc)
    metadata = _read_project_metadata(root)
    start_date_raw = (metadata.get("start_date") or "").strip()
    start_date_dt: datetime | None = None
    if ORDER_DATE_PATTERN.fullmatch(start_date_raw):
        try:
            start_date_dt = datetime.strptime(start_date_raw, "%d.%m.%Y").replace(tzinfo=timezone.utc)
        except ValueError:
            start_date_dt = None

    rows: list[PeriodicDocStatus] = []

    for rule in PERIODIC_DOC_RULES:
        folder = _resolve_safe_path(root=root, rel_path=rule.folder)
        latest_file: str | None = None
        latest_mtime: datetime | None = None

        if folder.exists() and folder.is_dir():
            candidates = [p for p in folder.glob(rule.pattern) if p.is_file()]
            if candidates:
                latest = max(candidates, key=lambda p: p.stat().st_mtime)
                latest_file = _to_rel_path(root, latest)
                latest_mtime = datetime.fromtimestamp(latest.stat().st_mtime, tz=timezone.utc)

        days_since_update: int | None = None
        if latest_mtime is not None:
            days_since_update = max(0, (now - latest_mtime).days)

        if latest_mtime is None:
            if start_date_dt is not None:
                days_since_update = max(0, (now - start_date_dt).days)
                is_due = days_since_update >= rule.period_days
            else:
                is_due = False
        else:
            is_due = days_since_update is not None and days_since_update >= rule.period_days

        rows.append(
            PeriodicDocStatus(
                rule=rule,
                latest_file=latest_file,
                latest_mtime=latest_mtime,
                days_since_update=days_since_update,
                is_due=is_due,
                action_path=rule.folder,
            )
        )

    return rows


def _project_metadata_path(root: Path) -> Path:
    incoming = (root / "00_incoming_requests").resolve()
    matches = sorted(incoming.glob("*PROJECT_METADATA*.md"))
    if matches:
        return matches[0].resolve()
    return (incoming / PROJECT_METADATA_DEFAULT_NAME).resolve()


def _read_project_metadata(root: Path) -> dict[str, str]:
    path = _project_metadata_path(root)
    if not path.exists() or not path.is_file():
        return {
            "object_name": "",
            "project_code": "",
            "organization": "",
            "work_stage": "",
            "start_date": "",
        }

    text = path.read_text(encoding="utf-8", errors="ignore")
    meta = {
        "object_name": "",
        "project_code": "",
        "organization": "",
        "work_stage": "",
        "start_date": "",
    }
    patterns = {
        "object_name": r"^Объект:\s*(.+)$",
        "project_code": r"^Шифр проекта:\s*(.+)$",
        "organization": r"^Организация:\s*(.+)$",
        "work_stage": r"^Этап работ:\s*(.+)$",
        "start_date": r"^Дата начала работ:\s*(.+)$",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, text, flags=re.MULTILINE)
        if match:
            meta[key] = match.group(1).strip()

    if not meta["work_stage"]:
        object_line = meta["object_name"]
        stage_match = re.search(r"(Этап\s*\d+)", object_line, flags=re.IGNORECASE)
        if stage_match:
            meta["work_stage"] = stage_match.group(1).strip()

    return meta


def _write_project_metadata(root: Path, payload: ArmObjectProfileUpdateRequest) -> Path:
    path = _project_metadata_path(root)
    path.parent.mkdir(parents=True, exist_ok=True)

    current = _read_project_metadata(root)
    object_name = (payload.object_name if payload.object_name is not None else current.get("object_name") or "").strip()
    project_code = (payload.project_code if payload.project_code is not None else current.get("project_code") or "").strip()
    organization = (payload.organization if payload.organization is not None else current.get("organization") or "").strip()
    work_stage = (payload.work_stage if payload.work_stage is not None else current.get("work_stage") or "").strip()
    start_date = (payload.start_date if payload.start_date is not None else current.get("start_date") or "").strip()

    lines = [
        "# Метаданные проекта",
        "",
        f"Дата фиксации: {datetime.now(tz=timezone.utc).strftime('%d.%m.%Y')}",
        "",
        f"Объект: {object_name}",
        "",
        f"Этап работ: {work_stage}",
        "",
        f"Дата начала работ: {start_date}",
        "",
        f"Шифр проекта: {project_code}",
        "",
        f"Организация: {organization}",
        "",
        "Назначение:",
        "",
        "- использовать шифр проекта в приказах, исполнительной документации и реестрах;",
        "- указывать в сопроводительных письмах и исходящих пакетах;",
        "- использовать дату начала работ для периодических документов и ТУДУ.",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def _extract_source_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt", ".csv", ".json", ".yml", ".yaml", ".ini", ".log"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        return _read_docx_preview(path)
    if suffix == ".pdf" and PdfReader is not None:
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages[:40])
    raise HTTPException(status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE, detail="Неподдерживаемый формат для импорта ППР")


def _extract_ppr_context_markdown(text: str) -> str:
    compact = re.sub(r"\s+", " ", text).strip()

    def find(pattern: str) -> str:
        match = re.search(pattern, compact, flags=re.IGNORECASE)
        return match.group(1).strip() if match else ""

    object_name = find(r"«?Логистический парк[^»]*?Этап\s*\d+")
    project_code = find(r"(\d{1,2}-\d+/\d{2}-\d{4}(?:/?ПР-1-[А-ЯA-Z0-9]+)?)")
    stage = find(r"(Этап\s*\d+)")
    work_scope = find(r"Проект производства работ разработан на\s+([^\.]+\.)") or find(r"Конструкции железобетонные\.\s*([^\.]+)")

    section_titles = []
    for match in re.finditer(r"(?:(?:^|\s))(\d+(?:\.\d+)*)\.\s+([А-ЯA-Z][^.]{3,140})", text, flags=re.MULTILINE):
        code = match.group(1).strip()
        title = re.sub(r"\s+", " ", match.group(2)).strip(" .")
        if len(title) >= 4:
            section_titles.append(f"- {code}. {title}")
        if len(section_titles) >= 18:
            break

    safety_keywords = []
    for phrase in [
        "монтаж колонн",
        "геодезические работы",
        "погрузочно-разгрузочные работы",
        "охрана труда",
        "пожарная безопасность",
        "стропальщики",
        "машиниста крана",
    ]:
        if phrase in compact.lower():
            safety_keywords.append(f"- {phrase}")

    lines = [
        "# Контекст ППР для автозаполнения",
        "",
        f"Источник: импорт ППР",
        f"Объект: {object_name}",
        f"Этап: {stage}",
        f"Шифр проекта: {project_code}",
        f"Вид работ: {work_scope}",
        "",
        "## Разделы ППР",
        *section_titles,
        "",
        "## Темы и работы для заполнения документов",
        *safety_keywords,
    ]
    return "\n".join(line for line in lines if line is not None).strip() + "\n"


def _build_dashboard_payload(db: Session) -> ArmDashboardResponse:
    root = resolve_object_root()
    checklist = _build_checklist(root=root)
    metrics = _collect_metrics(db=db, root=root)
    local_llm_reachable, local_llm_version = check_local_llm_available()

    checklist_total = len(checklist)
    checklist_ready = sum(1 for item in checklist if item.ready)
    checklist_progress_percent = round(
        (checklist_ready / checklist_total) * 100.0 if checklist_total else 0.0,
        1,
    )
    top_gaps = [item.title for item in checklist if not item.ready][:6]

    return ArmDashboardResponse(
        generated_at=datetime.now(tz=timezone.utc),
        object_root=str(root.as_posix()),
        checklist_total=checklist_total,
        checklist_ready=checklist_ready,
        checklist_progress_percent=checklist_progress_percent,
        local_llm_reachable=local_llm_reachable,
        local_llm_version=local_llm_version,
        top_gaps=top_gaps,
        metrics=metrics,
        checklist=checklist,
    )


def _collect_export_data(db: Session) -> tuple[list[Document], list[JournalEntry], list[WorkSchedule]]:
    documents = db.execute(select(Document).order_by(Document.created_at.desc())).scalars().all()
    journal_entries = db.execute(select(JournalEntry).order_by(JournalEntry.created_at.desc())).scalars().all()
    schedules = db.execute(select(WorkSchedule).order_by(WorkSchedule.planned_start.asc())).scalars().all()
    return list(documents), list(journal_entries), list(schedules)


def _build_gap_fix_hint(item: ArmChecklistItem) -> str:
    fix_map = {
        "1.4": "Добавьте и утвердите ППР в папке 05_execution_docs/ppr (минимум 1 файл).",
        "1.5": "Добавьте ППРв для работ на высоте в 05_execution_docs/pprv_work_at_height.",
        "1.6": "Подготовьте акт-допуск на производство СМР в 05_execution_docs/admission_acts.",
        "3.1": "Заполните журналы производства (минимум 6 файлов) в 04_journals/production.",
        "3.2": "Заполните журналы ОТ/ПБ (минимум 9 файлов) в 04_journals/labor_safety.",
        "4": "Дозагрузите сканы удостоверений/протоколов сотрудников в 02_personnel/employees.",
        "5": "Добавьте действующие нормативные документы в 06_normative_base.",
    }
    return fix_map.get(
        item.code,
        f"Доведите комплект до минимума: сейчас {item.found} из {item.required_min}.",
    )


def _build_arm_context(payload: ArmDashboardResponse, todos: ArmTodoResponse) -> str:
    todo_lines = "\n".join(f"- [{item.priority}] {item.title}" for item in todos.items)
    gaps_lines = "\n".join(f"- {item}" for item in payload.top_gaps)

    return (
        "Контекст АРМ объекта:\n"
        f"- Путь объекта: {payload.object_root}\n"
        f"- Комплектность по чек-листу: {payload.checklist_ready}/{payload.checklist_total} "
        f"({payload.checklist_progress_percent}%)\n"
        f"- Документов в БД: {payload.metrics.db_documents_total}\n"
        f"- Записей журнала в БД: {payload.metrics.db_journal_entries_total}\n"
        f"- Графиков в БД: {payload.metrics.db_schedules_total}\n"
        f"- PDF готово: {payload.metrics.orders_pdf_ready_total}\n"
        f"- Входящие сканы (очередь): {payload.metrics.scan_inbox_pending_total}\n"
        f"- Ручной разбор: {payload.metrics.scan_manual_review_total}\n"
        "Основные пробелы:\n"
        f"{gaps_lines or '- Нет критичных пробелов'}\n"
        "Задачи на сегодня:\n"
        f"{todo_lines or '- Список задач пуст'}\n"
        "Доступные разделы на фронте: Структура и действия, ТБ-чеклист сотрудника, Предпросмотр/правка/печать, "
        "Сканирование и OCR, Экспорт и печать офисных документов, Ассистент. "
        "Не предлагай разделы и кнопки, которых нет в этом списке."
    )


def _normalize_export_classification(classification: str | None) -> str:
    selected = (classification or "all").strip().lower()
    if selected not in VALID_EXPORT_CLASSIFICATIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                "Недопустимая классификация экспорта. "
                f"Разрешено: {', '.join(VALID_EXPORT_CLASSIFICATIONS)}"
            ),
        )
    return selected


def _list_order_markdown_paths(root: Path) -> list[str]:
    source = root / "01_orders_and_appointments"
    files: list[Path] = []
    files.extend(sorted(source.glob("*.md")))

    drafts = source / "drafts_from_checklist"
    if drafts.exists() and drafts.is_dir():
        files.extend(sorted(drafts.glob("*.md")))

    employees_root = root / "02_personnel" / "employees"
    if employees_root.exists() and employees_root.is_dir():
        files.extend(sorted(employees_root.glob("*/07_templates_to_print/*.md")))

    unique: list[str] = []
    seen: set[str] = set()
    for file in files:
        rel = _to_rel_path(root, file)
        if rel in seen:
            continue
        seen.add(rel)
        unique.append(rel)

    return sorted(unique)


def _classify_manual_review_rows(root: Path) -> list[dict[str, object]]:
    manual = root / "10_scan_inbox" / "manual_review"
    if not manual.exists() or not manual.is_dir():
        return []

    rows: list[dict[str, object]] = []
    for file in sorted(manual.iterdir()):
        if not file.is_file():
            continue
        if file.suffix.lower() not in {".pdf", ".jpg", ".jpeg", ".png", ".tif", ".tiff"}:
            continue

        sidecar = file.with_suffix(f"{file.suffix}.ocr.txt")
        ocr_text = sidecar.read_text(encoding="utf-8") if sidecar.exists() else None
        prediction = classify_scan_candidate(filename=file.name, ocr_text=ocr_text)
        rows.append(
            {
                "rel_path": _to_rel_path(root, file),
                "predicted_doc_type": prediction.predicted_doc_type,
                "confidence": prediction.confidence,
            }
        )

    return rows


def _normalize_employee_id(raw_value: str | None) -> str | None:
    if not raw_value:
        return None
    digits = "".join(ch for ch in str(raw_value) if ch.isdigit())
    if not digits:
        return None
    if len(digits) < 3:
        return digits.zfill(3)
    return digits


def _extract_employee_id_from_text(raw_text: str) -> str | None:
    patterns = (
        r"(?:id|айди|таб(?:ельный)?(?:\s*номер)?|сотрудник(?:а)?\s*(?:№|номер))\s*[:=]?\s*(\d{1,10})",
        r"сотрудник(?:а)?\s*#\s*(\d{1,10})",
    )
    for pattern in patterns:
        match = re.search(pattern, raw_text, flags=re.IGNORECASE)
        if match:
            return match.group(1)
    return None


def _extract_employee_name_from_text(raw_text: str) -> str | None:
    quoted_match = re.search(
        r"(?:\bсотрудник(?:а|у|ом|е|и|ов|ам|ами|ах)?\b|\bфио\b)\s*[:=]?\s*[\"«“](.+?)[\"»”]",
        raw_text,
        flags=re.IGNORECASE,
    )
    if quoted_match:
        return " ".join(quoted_match.group(1).split()).strip()

    plain_match = re.search(
        r"(?:\bсотрудник(?:а|у|ом|е|и|ов|ам|ами|ах)?\b|\bфио\b)\s*[:=]?\s*([A-Za-zА-Яа-яЁё-]+(?:\s+[A-Za-zА-Яа-яЁё-]+){1,2})",
        raw_text,
        flags=re.IGNORECASE,
    )
    if plain_match:
        return " ".join(plain_match.group(1).split()).strip()

    return None


def _infer_equipment_name(model_text: str) -> str:
    lowered = (model_text or "").lower()
    if "кран" in lowered:
        return "Автокран"
    if any(token in lowered for token in ("xcmg", "zoomlion", "зулион", "с/у", "погруз", "экскават", "манипулятор")):
        return "Строительная техника"
    if any(token in lowered for token in ("kia", "skoda", "toyota", "hyundai", "renault", "nissan", "ford", "авто", "автомоб")):
        return "Легковой автомобиль"
    return "Техника"


def _extract_employee_transport_entries(raw_text: str) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    seen: set[str] = set()
    plate_pattern = re.compile(
        r"([АВЕКМНОРСТУХABEKMHOPCTYX]\s*\d{3}\s*[АВЕКМНОРСТУХABEKMHOPCTYX]{2}\s*\d{2,3})",
        flags=re.IGNORECASE,
    )

    for raw_line in raw_text.splitlines():
        line = re.sub(r"^[\-\*•\s]+", "", raw_line.strip())
        if not line or "," not in line:
            continue

        lowered_line = line.lower()
        if lowered_line.startswith(("привет", "добав", "удал", "сдел", "отред", "и ")):
            continue

        name_raw, details_raw = line.split(",", 1)
        name_hint = " ".join(name_raw.split()).strip(" .;:")
        details = " ".join(details_raw.split()).strip(" .;:")
        if not name_hint or not details:
            continue

        plate_match = plate_pattern.search(details)
        plate = re.sub(r"\s+", "", plate_match.group(1)).upper() if plate_match else ""
        model_text = details
        if plate_match:
            model_text = (details[: plate_match.start()] + details[plate_match.end() :]).strip(" ,.;:")
        model_text = re.sub(r"\s+", " ", model_text).strip(" ,.;:")

        if not plate and not any(
            token in details.lower() for token in ("kia", "skoda", "xcmg", "кран", "зулион", "zoomlion", "техник", "авто", "машин")
        ):
            continue

        parts = [part for part in name_hint.split() if part]
        full_name = " ".join(parts) if len(parts) >= 2 else ""
        last_name_hint = parts[0] if parts else ""
        is_driver_line = len(parts) >= 2 and any(
            token in details.lower() for token in ("кран", "xcmg", "зулион", "zoomlion", "с/у", "авто", "машин")
        )
        position = "Водитель" if is_driver_line else ""
        equipment_name = _infer_equipment_name(model_text)

        dedupe_key = f"{name_hint.lower()}::{model_text.lower()}::{plate.lower()}"
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)

        entries.append(
            {
                "employee_id": "",
                "full_name": full_name,
                "last_name_hint": last_name_hint,
                "position": position,
                "position_explicit": "1" if position else "0",
                "birth_date": "",
                "vehicle_model": model_text,
                "vehicle_plate": plate,
                "equipment_name": equipment_name,
            }
        )

    return entries


def _extract_employee_position_from_text(raw_text: str) -> str | None:
    match = re.search(
        r"(?:должност(?:ь|и)?|позици(?:я|и)|роль)\s*[:=]?\s*([^\n,.;]+)",
        raw_text,
        flags=re.IGNORECASE,
    )
    if match:
        value = " ".join(match.group(1).split()).strip()
        if value:
            return value

    lowered = raw_text.lower()
    if "прораб" in lowered:
        return "прораб"
    if "электромонтаж" in lowered:
        return "электромонтажник"
    if "электрик" in lowered:
        return "электрик"
    if "свар" in lowered:
        return "сварщик"
    return None


def _extract_vehicle_plate_from_text(raw_text: str) -> str | None:
    match = re.search(
        r"\b([АВЕКМНОРСТУХABEKMHOPCTYX]\s*\d{3}\s*[АВЕКМНОРСТУХABEKMHOPCTYX]{2}\s*\d{2,3})\b",
        raw_text,
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    return re.sub(r"\s+", "", match.group(1)).upper()


def _extract_employee_list_entries(raw_text: str) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    seen_keys: set[str] = set()

    line_pattern = re.compile(
        r"^([A-Za-zА-Яа-яЁё-]+(?:\s+[A-Za-zА-Яа-яЁё-]+){2})\s+(\d{2}\.\d{2}\.\d{4})\s+(.+)$"
    )

    for raw_line in raw_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        line = re.sub(r"^[\-\*•\s]+", "", line).strip()
        match = line_pattern.match(line)
        if not match:
            continue

        full_name = " ".join(match.group(1).split()).strip()
        birth_date = match.group(2).strip()
        position = " ".join(match.group(3).split()).strip(" .,;")
        if not full_name or not position:
            continue

        dedupe_key = f"{full_name.lower()}::{position.lower()}::{birth_date}"
        if dedupe_key in seen_keys:
            continue
        seen_keys.add(dedupe_key)

        entries.append(
            {
                "full_name": full_name,
                "birth_date": birth_date,
                "position": position,
            }
        )

    return entries


def _employee_root_by_id(root: Path, employee_id: str | None) -> Path | None:
    if not employee_id:
        return None

    employees_root = (root / "02_personnel" / "employees").resolve()
    if not employees_root.exists() or not employees_root.is_dir():
        return None

    pattern = re.compile(rf"^{re.escape(employee_id)}(?:_|$)")
    for folder in sorted(employees_root.iterdir(), key=lambda p: p.name.lower()):
        if folder.is_dir() and pattern.match(folder.name):
            return folder
    return None


def _employee_root_by_name(root: Path, full_name: str | None) -> Path | None:
    if not full_name:
        return None

    target_slug = _slugify_filename(full_name)
    if not target_slug:
        return None

    employees_root = (root / "02_personnel" / "employees").resolve()
    if not employees_root.exists() or not employees_root.is_dir():
        return None

    for folder in sorted(employees_root.iterdir(), key=lambda p: p.name.lower()):
        if not folder.is_dir():
            continue
        _, _, tail = folder.name.partition("_")
        folder_slug = tail if tail else folder.name
        if folder_slug == target_slug:
            return folder
    return None


def _employee_folder_numeric_id(folder: Path, profile: dict[str, str]) -> int:
    profile_id = _normalize_employee_id(profile.get("employee_id"))
    if profile_id and profile_id.isdigit():
        return int(profile_id)

    match = re.match(r"^(\d{1,10})(?:_|$)", folder.name)
    if match:
        return int(match.group(1))
    return -1


def _employee_root_by_last_name(root: Path, last_name: str | None) -> Path | None:
    if not last_name:
        return None

    target = (last_name or "").strip().lower()
    target_slug = _slugify_filename(last_name)
    if not target_slug:
        return None

    employees_root = (root / "02_personnel" / "employees").resolve()
    if not employees_root.exists() or not employees_root.is_dir():
        return None

    candidates: list[tuple[int, int, Path]] = []
    for folder in sorted(employees_root.iterdir(), key=lambda p: p.name.lower()):
        if not folder.is_dir():
            continue
        profile = _read_employee_profile(folder)
        profile_last_name = (profile.get("last_name") or "").strip().lower()
        _, _, tail = folder.name.partition("_")
        folder_slug = tail if tail else folder.name

        match_rank = -1
        if profile_last_name == target:
            match_rank = 3
        elif folder_slug == target_slug or folder_slug.startswith(target_slug + "_"):
            match_rank = 2
        elif folder_slug.endswith("_" + target_slug) or ("_" + target_slug + "_") in folder_slug:
            match_rank = 1

        if match_rank >= 0:
            candidates.append((match_rank, _employee_folder_numeric_id(folder, profile), folder))

    if not candidates:
        return None

    candidates.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return candidates[0][2]


def _archive_employee_admission_requests(root: Path, employee_root: Path) -> list[Path]:
    permits_dir = (employee_root / "06_permits_and_work_admission").resolve()
    if not permits_dir.exists() or not permits_dir.is_dir():
        return []

    candidates = [
        item
        for item in sorted(permits_dir.iterdir(), key=lambda p: p.name.lower())
        if item.is_file() and ("letter_admission" in item.name.lower() or "допуск" in item.name.lower())
    ]
    if not candidates:
        return []

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    archive_dir = (
        root
        / "09_archive"
        / "removed_admission_requests"
        / f"{stamp}_{_slugify_filename(employee_root.name)}"
    ).resolve()
    archive_dir.mkdir(parents=True, exist_ok=True)

    archived_paths: list[Path] = []
    for item in candidates:
        target = archive_dir / item.name
        shutil.move(str(item), str(target))
        archived_paths.append(target)

    return archived_paths


def _archive_object_admission_requests_by_keyword(root: Path, keyword: str) -> list[Path]:
    token = (keyword or "").strip().lower()
    if not token:
        return []

    source_dirs = [
        (root / "00_incoming_requests").resolve(),
        (root / "01_orders_and_appointments" / "заявки").resolve(),
    ]

    candidates: list[Path] = []
    seen: set[Path] = set()
    for source_dir in source_dirs:
        if not source_dir.exists() or not source_dir.is_dir():
            continue
        for item in sorted(source_dir.glob("**/*LETTER_ADMISSION*"), key=lambda p: p.name.lower()):
            if not item.is_file():
                continue
            if token not in item.name.lower():
                continue
            resolved = item.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)
            candidates.append(resolved)

    if not candidates:
        return []

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    archive_dir = (
        root
        / "09_archive"
        / "removed_admission_requests"
        / f"{stamp}_object_{_slugify_filename(token)}"
    ).resolve()
    archive_dir.mkdir(parents=True, exist_ok=True)

    archived_paths: list[Path] = []
    for item in candidates:
        target = archive_dir / item.name
        suffix = 2
        while target.exists():
            target = archive_dir / f"{target.stem}_v{suffix:02d}{target.suffix}"
            suffix += 1
        shutil.move(str(item), str(target))
        archived_paths.append(target)

    return archived_paths


def _create_equipment_registry_file(root: Path, equipment_rows: list[dict[str, str]]) -> Path | None:
    if not equipment_rows:
        return None

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    folder = (root / "00_incoming_requests" / "equipment_from_assistant" / stamp).resolve()
    folder.mkdir(parents=True, exist_ok=True)

    file_path = folder / "equipment_registry.md"
    lines = [
        "# Реестр техники (сформировано ассистентом)",
        "",
        "| № | Водитель/ответственный | Тип техники | Марка/модель | Госномер |",
        "|---|---|---|---|---|",
    ]

    for index, row in enumerate(equipment_rows, start=1):
        lines.append(
            "| "
            f"{index} | {_sanitize_markdown_cell(row.get('owner_name') or '-')} | "
            f"{_sanitize_markdown_cell(row.get('equipment_name') or 'Техника')} | "
            f"{_sanitize_markdown_cell(row.get('vehicle_model') or '________________')} | "
            f"{_sanitize_markdown_cell(row.get('vehicle_plate') or '________________')} |"
        )

    file_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    return file_path


def _next_employee_id(root: Path) -> str:
    employees_root = (root / "02_personnel" / "employees").resolve()
    if not employees_root.exists() or not employees_root.is_dir():
        return "001"

    max_id = 0
    for folder in employees_root.iterdir():
        if not folder.is_dir():
            continue
        match = re.match(r"^(\d{1,10})(?:_|$)", folder.name)
        if not match:
            continue
        max_id = max(max_id, int(match.group(1)))

    return str(max_id + 1).zfill(3)


def _split_person_name(full_name: str) -> tuple[str, str, str]:
    parts = [part.strip() for part in full_name.split() if part.strip()]
    if not parts:
        return "Неизвестно", "", ""

    normalized = [part[:1].upper() + part[1:].lower() if part else "" for part in parts]
    last_name = normalized[0]
    first_name = normalized[1] if len(normalized) > 1 else ""
    middle_name = normalized[2] if len(normalized) > 2 else ""
    return last_name, first_name, middle_name


def _ensure_employee_workspace(
    *,
    root: Path,
    employee_id: str,
    full_name: str,
    position: str,
    employee_root: Path | None = None,
) -> tuple[Path, bool, bool]:
    employees_root = (root / "02_personnel" / "employees").resolve()
    employees_root.mkdir(parents=True, exist_ok=True)

    normalized_id = _normalize_employee_id(employee_id) or _next_employee_id(root)
    safe_name = _slugify_filename(full_name)
    folder_name = f"{normalized_id}_{safe_name}" if safe_name else normalized_id

    target_root = employee_root if employee_root is not None else (employees_root / folder_name)
    created = not target_root.exists()
    target_root.mkdir(parents=True, exist_ok=True)

    for section in EMPLOYEE_REQUIRED_FOLDERS:
        (target_root / section).mkdir(parents=True, exist_ok=True)

    profile_path = target_root / "employee_profile.txt"
    last_name, first_name, middle_name = _split_person_name(full_name)

    profile_lines = [
        f"employee_id: {normalized_id}",
        f"last_name: {last_name}",
        f"first_name: {first_name}",
        f"middle_name: {middle_name}",
        f"position: {position or 'Сотрудник'}",
    ]
    next_content = "\n".join(profile_lines).strip() + "\n"
    prev_content = profile_path.read_text(encoding="utf-8", errors="ignore") if profile_path.exists() else ""
    profile_updated = prev_content != next_content

    if profile_updated:
        profile_path.write_text(next_content, encoding="utf-8")

    return target_root, created, profile_updated


def _detect_project_order_code(question_lower: str) -> str:
    by_number = {
        "11": "P11_ORDER_PS",
        "12": "P12_ORDER_PERMIT",
        "13": "P13_ORDER_HEIGHT",
        "14": "P14_ORDER_FIRE",
        "15": "P15_ORDER_LOADING",
        "16": "P16_ORDER_PRESSURE",
        "17": "P17_ORDER_CLOSE_SHIFT",
        "18": "P18_ORDER_INTERNSHIP",
        "19": "P19_ORDER_AFTER_INTERNSHIP",
        "20": "P20_ORDER_HEATING",
    }

    match = re.search(r"приказ(?:а|у)?\s*(?:№|n)?\s*(\d{1,2})\b", question_lower)
    if match and match.group(1) in by_number:
        return by_number[match.group(1)]

    if "пожар" in question_lower:
        return "P14_ORDER_FIRE"
    if "высот" in question_lower:
        return "P13_ORDER_HEIGHT"
    if "стажиров" in question_lower:
        return "P18_ORDER_INTERNSHIP"
    if "электропрогрев" in question_lower:
        return "P20_ORDER_HEATING"
    if "допуск" in question_lower or "наряд" in question_lower:
        return "P12_ORDER_PERMIT"

    return "P12_ORDER_PERMIT"


def _create_assistant_draft(
    *,
    root: Path,
    employee_root: Path,
    rule_code: str,
    employee_name: str,
    employee_id: str | None,
    profession: str,
    context_overrides: dict[str, str] | None = None,
) -> Path:
    rule = PROJECT_RULES_BY_CODE.get(rule_code)
    if rule is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Неизвестный код шаблона: {rule_code}")

    file_name = ORDER_DRAFT_FILE_NAME_MAP.get(rule.code) or f"{rule.code}_{_slugify_filename(rule.title)}_draft.md"
    stem = Path(file_name).stem
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    target_dir = (
        (root / "01_orders_and_appointments" / "заявки").resolve()
        if rule.code == "P21_ADMISSION_LETTER"
        else (root / "01_orders_and_appointments" / "drafts_from_assistant").resolve()
    )
    target_dir.mkdir(parents=True, exist_ok=True)

    draft_path = target_dir / f"{stem}_{stamp}.md"
    suffix = 2
    while draft_path.exists():
        draft_path = target_dir / f"{stem}_{stamp}_v{suffix:02d}.md"
        suffix += 1

    employee_rel_path = _to_rel_path(root, employee_root)
    draft_content = _build_draft_content(
        root=root,
        employee_root=employee_root,
        rule=rule,
        employee_rel_path=employee_rel_path,
        employee_name=employee_name,
        employee_id=employee_id,
        profession=profession,
        context_overrides=context_overrides,
    )
    draft_path.write_text(draft_content, encoding="utf-8")
    return draft_path


def _execute_employee_documents_command(raw_text: str, root: Path) -> str | None:
    question_lower = raw_text.lower()
    employee_list_entries = _extract_employee_list_entries(raw_text)
    transport_entries = _extract_employee_transport_entries(raw_text)

    wants_employee = any(
        token in question_lower
        for token in (
            "добавь сотрудника",
            "добавь сотрудников",
            "добавить сотрудника",
            "добавить сотрудников",
            "создай сотрудника",
            "создай сотрудников",
            "новый сотрудник",
            "новые сотрудники",
            "заведи сотрудника",
            "заведи сотрудников",
            "оформи сотрудника",
            "оформи сотрудников",
        )
    ) or bool(employee_list_entries) or bool(transport_entries)
    wants_order = "приказ" in question_lower and any(
        token in question_lower
        for token in ("сделай", "создай", "сформируй", "подготовь", "оформи", "выполни", "нужен", "нужно")
    )
    wants_vehicle_pass = (
        ("пропуск" in question_lower or "письмо-допуск" in question_lower or "заявк" in question_lower)
        and any(token in question_lower for token in ("авто", "автомоб", "машин", "техник"))
    )
    wants_shared_pass = wants_vehicle_pass and (
        "общую заявк" in question_lower
        or "одну общую" in question_lower
        or "общий пропуск" in question_lower
        or "один общий" in question_lower
    )
    wants_equipment_folder = "папк" in question_lower and "техник" in question_lower
    wants_delete_bikbulatov_pass = (
        "удал" in question_lower and "заявк" in question_lower and "бикбулатов" in question_lower and "пропуск" in question_lower
    )

    if not (wants_employee or wants_order or wants_vehicle_pass or wants_delete_bikbulatov_pass):
        return None

    employee_id = _normalize_employee_id(_extract_employee_id_from_text(raw_text))
    employee_name = _extract_employee_name_from_text(raw_text)
    position = _extract_employee_position_from_text(raw_text) or "Сотрудник"

    placeholder_tokens = {"такого", "такой", "то", "этого", "этот", "сотрудника"}
    if employee_name:
        name_tokens = {token.lower() for token in employee_name.split()}
        if name_tokens.intersection(placeholder_tokens):
            employee_name = None

    employee_requests: list[dict[str, str]] = []
    if employee_list_entries:
        employee_requests.extend(
            {
                "employee_id": "",
                "full_name": item.get("full_name", ""),
                "position": item.get("position", "") or position,
                "position_explicit": "1" if item.get("position") else "0",
                "birth_date": item.get("birth_date", ""),
                "last_name_hint": "",
                "vehicle_model": "",
                "vehicle_plate": "",
                "equipment_name": "",
            }
            for item in employee_list_entries
        )
    if transport_entries:
        employee_requests.extend(transport_entries)

    if not employee_requests and (wants_employee or employee_id or employee_name):
        employee_requests.append(
            {
                "employee_id": employee_id or "",
                "full_name": employee_name or "",
                "position": position,
                "position_explicit": "1" if _extract_employee_position_from_text(raw_text) else "0",
                "birth_date": "",
                "last_name_hint": "",
                "vehicle_model": "",
                "vehicle_plate": "",
                "equipment_name": "",
            }
        )

    has_employee_ref = any(
        (request.get("employee_id") or "").strip()
        or (request.get("full_name") or "").strip()
        or (request.get("last_name_hint") or "").strip()
        for request in employee_requests
    )
    if (wants_order or wants_vehicle_pass) and not has_employee_ref and not wants_delete_bikbulatov_pass:
        return None

    processed_employees: list[dict[str, object]] = []
    equipment_rows: list[dict[str, str]] = []

    for request in employee_requests:
        requested_id = _normalize_employee_id(request.get("employee_id"))
        requested_name = (request.get("full_name") or "").strip()
        requested_last_name = (request.get("last_name_hint") or "").strip()
        requested_position = (request.get("position") or position or "Сотрудник").strip()
        position_explicit = (request.get("position_explicit") or "") == "1"
        requested_birth_date = (request.get("birth_date") or "").strip()
        requested_vehicle_model = (request.get("vehicle_model") or "").strip()
        requested_vehicle_plate = (request.get("vehicle_plate") or "").strip()
        requested_equipment_name = (request.get("equipment_name") or "").strip()

        if not requested_id and not requested_name and not requested_last_name:
            continue

        employee_root = (
            _employee_root_by_id(root, requested_id)
            or _employee_root_by_name(root, requested_name)
            or _employee_root_by_last_name(root, requested_last_name)
        )
        created_employee = False
        profile_updated = False

        if employee_root is None:
            if not requested_name and requested_last_name:
                requested_name = requested_last_name
            if not requested_name:
                continue

            requested_id = requested_id or _next_employee_id(root)
            employee_root, created_employee, profile_updated = _ensure_employee_workspace(
                root=root,
                employee_id=requested_id,
                full_name=requested_name,
                position=requested_position,
                employee_root=None,
            )
        elif wants_employee:
            current_profile = _read_employee_profile(employee_root)
            base_name = requested_name or _employee_display_name(employee_root, current_profile)
            inferred_id = requested_id
            if not inferred_id:
                match = re.match(r"^(\d{1,10})(?:_|$)", employee_root.name)
                inferred_id = _normalize_employee_id(match.group(1) if match else None) or _next_employee_id(root)

            existing_position = (current_profile.get("position") or "Сотрудник").strip()
            effective_position = requested_position if position_explicit else existing_position

            employee_root, _created_unused, profile_updated = _ensure_employee_workspace(
                root=root,
                employee_id=inferred_id,
                full_name=base_name,
                position=effective_position,
                employee_root=employee_root,
            )

        profile = _read_employee_profile(employee_root)
        employee_rel_path = _to_rel_path(root, employee_root)
        employee_name_effective = _employee_display_name(employee_root, profile)
        employee_id_effective = profile.get("employee_id") or requested_id
        profession_effective = (profile.get("position") or requested_position or "Сотрудник").strip()
        birth_date_effective = requested_birth_date or (profile.get("birth_date") or "").strip()

        if created_employee:
            status_text = "создана новая карточка сотрудника"
        elif profile_updated:
            status_text = "профиль сотрудника обновлен"
        else:
            status_text = "использована существующая карточка"

        processed_employees.append(
            {
                "employee_root": employee_root,
                "employee_name": employee_name_effective,
                "employee_id": employee_id_effective or "",
                "profession": profession_effective,
                "employee_rel_path": employee_rel_path,
                "birth_date": birth_date_effective,
                "status": status_text,
            }
        )

        if requested_vehicle_model or requested_vehicle_plate:
            equipment_rows.append(
                {
                    "owner_name": employee_name_effective,
                    "equipment_name": requested_equipment_name or _infer_equipment_name(requested_vehicle_model),
                    "vehicle_model": requested_vehicle_model or "________________",
                    "vehicle_plate": requested_vehicle_plate or "________________",
                }
            )

    if not processed_employees:
        if wants_employee or wants_delete_bikbulatov_pass:
            return (
                "Чтобы выполнить команду, укажите ФИО сотрудника.\n"
                "Пример: «добавь сотрудника Иванов Иван Иванович id 007 должность электромонтажник "
                "и сделай приказ 12 и заявку на пропуск авто А123АА102»."
            )
        return None

    archived_bikbulatov_requests: list[Path] = []
    if wants_delete_bikbulatov_pass:
        bikbulatov_root = _employee_root_by_last_name(root, "Бикбулатов")
        if bikbulatov_root is not None:
            archived_bikbulatov_requests = _archive_employee_admission_requests(root, bikbulatov_root)
        archived_bikbulatov_requests.extend(_archive_object_admission_requests_by_keyword(root, "бикбулатов"))

    primary_employee = processed_employees[0] if (wants_shared_pass or len(processed_employees) > 1) else processed_employees[-1]
    employee_root = primary_employee["employee_root"]
    employee_name_effective = str(primary_employee["employee_name"])
    employee_id_effective = str(primary_employee["employee_id"])
    profession = str(primary_employee["profession"])

    created_docs: list[Path] = []
    equipment_registry_path: Path | None = None
    if wants_order:
        rule_code = _detect_project_order_code(question_lower)
        created_docs.append(
            _create_assistant_draft(
                root=root,
                employee_root=employee_root,
                rule_code=rule_code,
                employee_name=employee_name_effective,
                employee_id=employee_id_effective,
                profession=profession,
            )
        )

    pass_rows = processed_employees if (wants_shared_pass or len(processed_employees) > 1) else [primary_employee]
    unique_equipment_rows: list[dict[str, str]] = []
    seen_equipment_keys: set[str] = set()
    for row in equipment_rows:
        dedupe_key = (
            f"{(row.get('owner_name') or '').lower()}::"
            f"{(row.get('equipment_name') or '').lower()}::"
            f"{(row.get('vehicle_model') or '').lower()}::"
            f"{(row.get('vehicle_plate') or '').lower()}"
        )
        if dedupe_key in seen_equipment_keys:
            continue
        seen_equipment_keys.add(dedupe_key)
        unique_equipment_rows.append(row)

    any_plate_detected = any((row.get("vehicle_plate") or "").strip("_") for row in unique_equipment_rows)

    if wants_vehicle_pass:
        if not unique_equipment_rows:
            fallback_plate = _extract_vehicle_plate_from_text(raw_text)
            unique_equipment_rows.append(
                {
                    "owner_name": employee_name_effective,
                    "equipment_name": "Легковой автомобиль",
                    "vehicle_model": "________________",
                    "vehicle_plate": fallback_plate or "________________",
                }
            )
            any_plate_detected = bool(fallback_plate)

        equipment_bullets_lines: list[str] = []
        for row in unique_equipment_rows:
            model = row.get("vehicle_model") or "________________"
            plate = row.get("vehicle_plate") or "________________"
            equipment_name = row.get("equipment_name") or "Техника"
            if plate == "________________":
                equipment_bullets_lines.append(f"- {equipment_name}, {model}, госномер уточнить")
            else:
                equipment_bullets_lines.append(f"- {equipment_name}, {model}, госномер {plate}")
        equipment_bullets = "\n".join(equipment_bullets_lines) if equipment_bullets_lines else "- Техника уточняется"

        workers_rows = [
            (
                str(row.get("employee_name") or "Сотрудник"),
                str(row.get("profession") or "Сотрудник"),
                str(row.get("birth_date") or "________________") or "________________",
            )
            for row in pass_rows
        ]
        equipment_table_rows = [
            (
                row.get("equipment_name") or "Техника",
                row.get("vehicle_model") or "________________",
                row.get("vehicle_plate") or "________________",
            )
            for row in unique_equipment_rows
        ]

        workers_list_table = _workers_list_table_markdown(workers_rows)
        equipment_list_table = _equipment_list_table_markdown(equipment_table_rows)

        pass_employee = pass_rows[0]
        pass_employee_root = pass_employee["employee_root"]
        pass_employee_name = str(pass_employee["employee_name"])
        pass_employee_id = str(pass_employee["employee_id"])
        pass_employee_profession = str(pass_employee["profession"])

        created_docs.append(
            _create_assistant_draft(
                root=root,
                employee_root=pass_employee_root,
                rule_code="P21_ADMISSION_LETTER",
                employee_name=pass_employee_name,
                employee_id=pass_employee_id,
                profession=pass_employee_profession,
                context_overrides={
                    "EQUIPMENT_BULLETS": equipment_bullets,
                    "WORKERS_LIST_TABLE": workers_list_table,
                    "EQUIPMENT_LIST_TABLE": equipment_list_table,
                },
            )
        )

    if wants_equipment_folder and unique_equipment_rows:
        equipment_registry_path = _create_equipment_registry_file(root, unique_equipment_rows)

    lines = ["Сценарий кадрового оформления выполнен."]
    lines.append(f"Обработано сотрудников: {len(processed_employees)}")
    lines.append("Сотрудники:")
    for row in processed_employees:
        lines.append(
            "- "
            f"{row['employee_name']} (ID: {row['employee_id'] or '-'}) - {row['status']}; "
            f"папка: {row['employee_rel_path']}"
        )

    if created_docs:
        lines.append("Созданные документы:")
        lines.extend(f"- {_to_rel_path(root, path)}" for path in created_docs)
    else:
        lines.append("Документы не запрашивались, выполнено оформление карточек сотрудников.")

    if archived_bikbulatov_requests:
        lines.append(f"Архивировано заявок по Бикбулатову: {len(archived_bikbulatov_requests)}")

    if wants_delete_bikbulatov_pass and not archived_bikbulatov_requests:
        lines.append("Для Бикбулатова активных заявок на пропуск в рабочей папке не найдено.")

    if equipment_registry_path is not None:
        lines.append(f"Создана папка с техникой: {_to_rel_path(root, equipment_registry_path.parent)}")

    if wants_vehicle_pass and len(pass_rows) > 1:
        lines.append(f"Сформирована общая заявка на пропуск: сотрудников {len(pass_rows)}, техники {len(unique_equipment_rows)}.")
    elif wants_vehicle_pass:
        lines.append(f"Заявка на пропуск сформирована для: {employee_name_effective}.")

    if wants_vehicle_pass and not any_plate_detected:
        lines.append("Госномер авто не распознан: в заявке оставлена пометка «госномер уточнить».")

    return "\n".join(lines)


def _try_execute_assistant_scenario(payload: ArmAssistRequest, db: Session) -> ArmAssistResponse | None:
    raw_question = (payload.question or "").strip()
    question = raw_question.lower()
    root = resolve_object_root()

    employee_command_response = _execute_employee_documents_command(raw_text=raw_question, root=root)
    if employee_command_response is not None:
        return ArmAssistResponse(
            model="scenario-engine",
            response=employee_command_response,
            done=True,
            used_profile=payload.profile,
            fallback_used=False,
            total_duration_sec=0.0,
            eval_tokens=None,
            eval_tokens_per_sec=None,
        )

    if any(token in question for token in ("список всех приказ", "покажи все приказы", "все приказы")):
        files = _list_order_markdown_paths(root)
        if not files:
            text = "Приказы в формате Markdown не найдены в текущем объекте."
        else:
            preview = "\n".join(f"- {path}" for path in files[:120])
            suffix = "\n..." if len(files) > 120 else ""
            text = (
                f"Найдено приказов/черновиков: {len(files)}.\n"
                "Список:\n"
                f"{preview}{suffix}"
            )

        return ArmAssistResponse(
            model="scenario-engine",
            response=text,
            done=True,
            used_profile=payload.profile,
            fallback_used=False,
            total_duration_sec=0.0,
            eval_tokens=None,
            eval_tokens_per_sec=None,
        )

    if any(token in question for token in ("выполни ocr", "запусти ocr", "выполни распознать", "выполни разлож")):
        inbox = root / "10_scan_inbox"
        results = ingest_inbox(
            object_root=root,
            inbox_folder=inbox,
            db=db,
            enable_ocr=True,
            ocr_lang="rus+eng",
            tesseract_cmd=None,
            max_pdf_pages=4,
        )
        archived = sum(1 for item in results if item.status == "archived")
        manual_review = sum(1 for item in results if item.status == "manual_review")
        response = (
            "Сценарий OCR/разбора выполнен. "
            f"Архивировано: {archived}; отправлено в ручной разбор: {manual_review}."
        )

        return ArmAssistResponse(
            model="scenario-engine",
            response=response,
            done=True,
            used_profile=payload.profile,
            fallback_used=False,
            total_duration_sec=0.0,
            eval_tokens=None,
            eval_tokens_per_sec=None,
        )

    if "manual_review" in question and any(token in question for token in ("проверь", "классифиц", "разбери", "сценар")):
        rows = _classify_manual_review_rows(root)
        if not rows:
            response = "Папка ручного разбора пуста. Дополнительная классификация не требуется."
        else:
            lines = "\n".join(
                f"- {row['rel_path']} -> {row['predicted_doc_type']} ({row['confidence']})"
                for row in rows[:80]
            )
            suffix = "\n..." if len(rows) > 80 else ""
            response = f"Классификация ручного разбора выполнена, файлов: {len(rows)}.\n{lines}{suffix}"

        return ArmAssistResponse(
            model="scenario-engine",
            response=response,
            done=True,
            used_profile=payload.profile,
            fallback_used=False,
            total_duration_sec=0.0,
            eval_tokens=None,
            eval_tokens_per_sec=None,
        )

    return None


@router.get("/checklist", response_model=list[ArmChecklistItem])
def arm_checklist() -> list[ArmChecklistItem]:
    root = resolve_object_root()
    return _build_checklist(root=root)


@router.get("/metrics", response_model=ArmDashboardResponse)
def arm_metrics(db: Session = Depends(get_db)) -> ArmDashboardResponse:
    return _build_dashboard_payload(db=db)


@router.get("/todo/today", response_model=ArmTodoResponse)
def arm_todo_today(db: Session = Depends(get_db)) -> ArmTodoResponse:
    payload = _build_dashboard_payload(db=db)
    items = _build_todos(
        checklist=payload.checklist,
        metrics=payload.metrics,
        local_llm_reachable=payload.local_llm_reachable,
    )
    return ArmTodoResponse(
        generated_at=payload.generated_at,
        object_root=payload.object_root,
        items=items,
    )


@router.post("/assist", response_model=ArmAssistResponse)
def arm_assist(payload: ArmAssistRequest, db: Session = Depends(get_db)) -> ArmAssistResponse:
    scenario_result = _try_execute_assistant_scenario(payload=payload, db=db)
    if scenario_result is not None:
        return scenario_result

    if not settings.local_llm_enabled:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Интеграция с локальной LLM отключена",
        )

    dashboard = _build_dashboard_payload(db=db)
    todos = ArmTodoResponse(
        generated_at=dashboard.generated_at,
        object_root=dashboard.object_root,
        items=_build_todos(
            checklist=dashboard.checklist,
            metrics=dashboard.metrics,
            local_llm_reachable=dashboard.local_llm_reachable,
        ),
    )
    context = _build_arm_context(payload=dashboard, todos=todos)

    try:
        result, used_profile, fallback_used = generate_with_local_llm_profile(
            prompt=payload.question,
            context=context,
            profile=payload.profile,
            model=payload.model,
            system_prompt=None,
            temperature=payload.temperature,
            num_predict=payload.num_predict,
            allow_fallback=payload.allow_fallback,
        )
    except LocalLLMConnectionError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(exc),
        ) from exc
    except LocalLLMRequestError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=str(exc),
        ) from exc

    return ArmAssistResponse(
        model=result.model,
        response=result.response,
        done=result.done,
        used_profile=used_profile,
        fallback_used=fallback_used,
        total_duration_sec=result.total_duration_sec,
        eval_tokens=result.eval_tokens,
        eval_tokens_per_sec=result.eval_tokens_per_sec,
    )


@router.get("/exports/orders-docx")
def arm_export_orders_docx(classification: str = "all") -> FileResponse:
    root = resolve_object_root()
    selected = _normalize_export_classification(classification)
    try:
        result = export_orders_docx_bundle(object_root=root, classification=selected)
    except OfficeExportDependencyError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(exc),
        ) from exc

    return FileResponse(
        path=result.bundle_path,
        media_type="application/zip",
        filename=result.bundle_path.name,
    )


@router.get("/exports/registers-xlsx")
def arm_export_registers_xlsx(db: Session = Depends(get_db)) -> FileResponse:
    root = resolve_object_root()
    checklist = _build_checklist(root=root)
    documents, journal_entries, schedules = _collect_export_data(db=db)

    try:
        result = export_registers_xlsx(
            object_root=root,
            documents=documents,
            journal_entries=journal_entries,
            schedules=schedules,
            checklist=checklist,
        )
    except OfficeExportDependencyError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(exc),
        ) from exc

    return FileResponse(
        path=result.file_path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename=result.file_path.name,
    )


@router.get("/exports/office-pack")
def arm_export_office_pack(classification: str = "all", db: Session = Depends(get_db)) -> FileResponse:
    root = resolve_object_root()
    selected = _normalize_export_classification(classification)
    checklist = _build_checklist(root=root)
    documents, journal_entries, schedules = _collect_export_data(db=db)

    try:
        docx_result = export_orders_docx_bundle(object_root=root, classification=selected)
        xlsx_result = export_registers_xlsx(
            object_root=root,
            documents=documents,
            journal_entries=journal_entries,
            schedules=schedules,
            checklist=checklist,
        )
        pack_path = build_office_pack_zip(
            object_root=root,
            docx_result=docx_result,
            xlsx_result=xlsx_result,
        )
    except OfficeExportDependencyError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(exc),
        ) from exc

    return FileResponse(
        path=pack_path,
        media_type="application/zip",
        filename=pack_path.name,
    )


@router.get("/fs/tree", response_model=ArmFsTreeResponse)
def arm_fs_tree(rel_path: str = "") -> ArmFsTreeResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)
    if not target.exists() or not target.is_dir():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Папка не найдена")

    entries: list[ArmFsEntry] = []
    for item in sorted(target.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower())):
        stat = item.stat()
        entries.append(
            ArmFsEntry(
                name=item.name,
                rel_path=_to_rel_path(root, item),
                is_dir=item.is_dir(),
                size=None if item.is_dir() else stat.st_size,
                modified_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
            )
        )

    return ArmFsTreeResponse(
        root=str(root.as_posix()),
        rel_path=_to_rel_path(root, target) if target != root else "",
        entries=entries,
    )


@router.get("/fs/file", response_model=ArmFileReadResponse)
def arm_fs_file_read(rel_path: str) -> ArmFileReadResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден")

    if target.suffix.lower() not in TEXT_PREVIEW_EXTENSIONS:
        raise HTTPException(status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE, detail="Предпросмотр поддерживает только текстовые файлы")

    size_bytes = target.stat().st_size
    if size_bytes > MAX_TEXT_PREVIEW_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=(
                f"Файл слишком большой для предпросмотра ({size_bytes} байт). "
                f"Лимит: {MAX_TEXT_PREVIEW_BYTES} байт. Используйте скачивание."
            ),
        )

    content, encoding = _read_text_preview_content(target)

    return ArmFileReadResponse(
        rel_path=_to_rel_path(root, target),
        content=content,
        encoding=encoding,
    )


@router.post("/fs/file", response_model=ArmActionResponse)
def arm_fs_file_write(payload: ArmFileWriteRequest, db: Session = Depends(get_db)) -> ArmActionResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=payload.rel_path)
    if target.suffix.lower() not in TEXT_PREVIEW_EXTENSIONS:
        raise HTTPException(status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE, detail="Редактирование поддерживает только текстовые файлы")

    content_size = len(payload.content.encode("utf-8"))
    if content_size > MAX_TEXT_PREVIEW_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=(
                f"Слишком большой объем текста ({content_size} байт). "
                f"Лимит: {MAX_TEXT_PREVIEW_BYTES} байт."
            ),
        )

    previous_exists = target.exists()
    previous_bytes = target.read_bytes() if previous_exists else None

    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(payload.content, encoding="utf-8")

    rel_path = _to_rel_path(root, target)
    try:
        snapshot = db.scalar(select(DocumentContent).where(DocumentContent.rel_path == rel_path))
        if snapshot is None:
            db.add(DocumentContent(rel_path=rel_path, content=payload.content))
        else:
            snapshot.content = payload.content
        db.commit()
    except Exception as exc:  # noqa: BLE001
        db.rollback()
        if previous_exists and previous_bytes is not None:
            target.write_bytes(previous_bytes)
        elif target.exists():
            target.unlink(missing_ok=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Не удалось сохранить изменения в БД. Изменения в файле откатены.",
        ) from exc

    return ArmActionResponse(ok=True, message=f"Сохранено: {rel_path} (файл + БД)")


@router.post("/fs/upload", response_model=ArmActionResponse)
async def arm_fs_upload(
    file: UploadFile = File(...),
    rel_dir: str = "",
) -> ArmActionResponse:
    root = resolve_object_root()
    target_dir = _resolve_safe_path(root=root, rel_path=rel_dir)

    original_name = Path(file.filename or "upload").name
    # Strip any directory component from the filename to prevent path traversal
    original_name = Path(original_name).name
    suffix = Path(original_name).suffix.lower()
    if suffix not in UPLOAD_ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Тип файла '{suffix}' не разрешён. Разрешены: {', '.join(sorted(UPLOAD_ALLOWED_EXTENSIONS))}",
        )

    data = await file.read()
    if len(data) > UPLOAD_MAX_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"Файл превышает лимит {UPLOAD_MAX_BYTES // (1024 * 1024)} МБ.",
        )

    target_dir.mkdir(parents=True, exist_ok=True)
    dest = target_dir / original_name
    dest.write_bytes(data)

    rel_path = _to_rel_path(root, dest)
    return ArmActionResponse(ok=True, message=f"Загружено: {rel_path}")


@router.post("/fs/mkdir", response_model=ArmActionResponse)
def arm_fs_mkdir(rel_path: str = "") -> ArmActionResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)
    target.mkdir(parents=True, exist_ok=True)
    normalized = _to_rel_path(root, target) if target != root else ""
    return ArmActionResponse(ok=True, message=f"Папка готова: {normalized or '/'}")


@router.get("/fs/download")
def arm_fs_download(rel_path: str) -> FileResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден")

    return FileResponse(
        path=target,
        media_type="application/octet-stream",
        filename=target.name,
    )


@router.get("/fs/view")
def arm_fs_view(rel_path: str) -> FileResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден")

    return FileResponse(
        path=target,
        media_type=_guess_media_type(target),
        filename=target.name,
        content_disposition_type="inline",
        headers={
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


@router.get("/fs/print-preview", response_class=HTMLResponse)
def arm_fs_print_preview(rel_path: str, auto_print: bool = True) -> HTMLResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден")

    rel = _to_rel_path(root, target)
    suffix = target.suffix.lower()
    is_timesheet = False

    if suffix == ".md":
        if auto_print:
            version = int(target.stat().st_mtime_ns)
            return RedirectResponse(
                url=f"/arm/fs/print-render-pdf?rel_path={quote(rel, safe='')}&v={version}",
                status_code=status.HTTP_307_TEMPORARY_REDIRECT,
            )

        print_ready_pdf = _resolve_print_ready_pdf_for_markdown(root, target)
        if print_ready_pdf is not None:
            pdf_rel = _to_rel_path(root, print_ready_pdf)
            version = int(print_ready_pdf.stat().st_mtime_ns)
            return RedirectResponse(
                url=f"/arm/fs/view?rel_path={quote(pdf_rel, safe='')}&v={version}",
                status_code=status.HTTP_307_TEMPORARY_REDIRECT,
            )

        content, _ = _read_text_preview_content(target)
        is_timesheet = _is_timesheet_document(content)
        body_html = _render_markdown_for_print(content)
    elif suffix == ".docx":
        content = _read_docx_preview(target)
        body_html = (
            "<div class=\"meta\">DOCX-предпросмотр: извлеченный текст для печати. "
            "Для точной верстки используйте печать из Word.</div>"
            f"<pre>{escape(content)}</pre>"
        )
    elif suffix in TEXT_PREVIEW_EXTENSIONS:
        content, encoding = _read_text_preview_content(target)
        body_html = (
            f"<div class=\"meta\">Режим предпросмотра: {escape(encoding)}</div>"
            f"<pre>{escape(content)}</pre>"
        )
    elif suffix == ".pdf":
        version = int(target.stat().st_mtime_ns)
        view_url = f"/arm/fs/view?rel_path={quote(rel, safe='')}&v={version}"
        body_html = f"<iframe class=\"preview-frame\" src=\"{view_url}\" title=\"PDF предпросмотр\"></iframe>"
    elif suffix in INLINE_IMAGE_EXTENSIONS:
        view_url = f"/arm/fs/view?rel_path={quote(rel, safe='')}"
        body_html = f"<img class=\"preview-image\" src=\"{view_url}\" alt=\"Предпросмотр\" />"
    else:
        body_html = (
            "<div class=\"meta\">"
            "Этот формат не поддерживает встроенный предпросмотр. "
            "Используйте скачивание оригинала и печать через локальное приложение."
            "</div>"
        )

    html = _build_print_preview_html(
        file_rel_path=rel,
        file_name=target.name,
        body_html=body_html,
        auto_print=auto_print,
        is_timesheet=is_timesheet,
    )
    return HTMLResponse(content=html)


@router.get("/fs/print-render-pdf")
def arm_fs_print_render_pdf(rel_path: str) -> Response:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден")

    if target.suffix.lower() != ".md":
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Серверная печать без колонтитулов поддерживается для Markdown-файлов",
        )

    pdf_bytes = _render_markdown_to_pdf_bytes(target)
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


@router.post("/fs/print", response_model=ArmActionResponse)
def arm_fs_print(rel_path: str) -> ArmActionResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден")

    rel = _to_rel_path(root, target)
    preview_url = f"/arm/fs/print-preview?rel_path={quote(rel, safe='')}&auto_print=1"

    # Legacy path: keep direct Windows print attempt for binary office/media files,
    # but do not fail hard if there is no shell association.
    if os.name == "nt" and target.suffix.lower() not in TEXT_PREVIEW_EXTENSIONS:
        try:
            os.startfile(str(target), "print")  # type: ignore[attr-defined]
            return ArmActionResponse(ok=True, message=f"Отправлено на печать: {rel}")
        except Exception:  # noqa: BLE001
            pass

    return ArmActionResponse(
        ok=True,
        message=(
            "Прямая печать через системную ассоциацию недоступна для этого формата. "
            f"Откройте браузерный предпросмотр печати: {preview_url}"
        ),
    )


@router.post("/fs/move", response_model=ArmActionResponse)
def arm_fs_move(source_rel_path: str, target_rel_path: str, overwrite: bool = False) -> ArmActionResponse:
    root = resolve_object_root()
    source = _resolve_safe_path(root=root, rel_path=source_rel_path)
    target = _resolve_safe_path(root=root, rel_path=target_rel_path)

    if not source.exists() or not source.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Исходный файл не найден")

    if source == target:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Исходный и целевой путь совпадают")

    if target.exists() and not overwrite:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Целевой файл уже существует")
    if target.exists() and target.is_dir():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Целевой путь указывает на папку")

    source_sidecar = source.with_suffix(f"{source.suffix}.ocr.txt")
    target_sidecar = target.with_suffix(f"{target.suffix}.ocr.txt")
    move_sidecar = source_sidecar.exists() and source_sidecar.is_file()
    if move_sidecar and target_sidecar.exists() and not overwrite:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="OCR-файл с текстом для целевого файла уже существует",
        )

    target.parent.mkdir(parents=True, exist_ok=True)

    if target.exists() and overwrite:
        target.unlink()
    if move_sidecar and target_sidecar.exists() and overwrite:
        target_sidecar.unlink()

    shutil.move(str(source), str(target))
    if move_sidecar and source_sidecar.exists():
        target_sidecar.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(source_sidecar), str(target_sidecar))

    moved_path = _to_rel_path(root, target)
    if move_sidecar:
        return ArmActionResponse(
            ok=True,
            message=(
                "Файл и OCR-файл с текстом перемещены: "
                + moved_path
            ),
        )
    return ArmActionResponse(ok=True, message="Файл перемещен: " + moved_path)


@router.post("/fs/delete", response_model=ArmActionResponse)
def arm_fs_delete(rel_path: str, with_sidecar: bool = True) -> ArmActionResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)

    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден")

    sidecar = target.with_suffix(f"{target.suffix}.ocr.txt")
    deleted_path = _to_rel_path(root, target)
    target.unlink()

    sidecar_deleted = False
    if with_sidecar and sidecar.exists() and sidecar.is_file():
        sidecar.unlink()
        sidecar_deleted = True

    if sidecar_deleted:
        return ArmActionResponse(ok=True, message="Файл и OCR-файл с текстом удалены: " + deleted_path)
    return ArmActionResponse(ok=True, message="Файл удален: " + deleted_path)


@router.get("/scanner/devices", response_model=ArmScannerDevicesResponse)
def arm_scanner_devices() -> ArmScannerDevicesResponse:
    completed = _run_scanner_command(["list"])
    if completed.returncode != 0:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=(completed.stderr.strip() or completed.stdout.strip() or "Не удалось получить список сканеров"),
        )

    devices = _parse_scanner_list_stdout(completed.stdout)
    return ArmScannerDevicesResponse(devices=devices)


@router.post("/scanner/scan-to-inbox", response_model=ArmActionResponse)
def arm_scanner_scan_to_inbox(payload: ArmScanCaptureRequest) -> ArmActionResponse:
    root = resolve_object_root()

    doc_type = (payload.doc_type or "").upper().strip()
    if doc_type not in SCANNER_DOC_TYPES:
        allowed = ", ".join(sorted(SCANNER_DOC_TYPES.keys()))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Неизвестный тип скана '{payload.doc_type}'. Разрешены: {allowed}",
        )

    if doc_type in SCANNER_DOC_TYPES_REQUIRING_EMPLOYEE and not (payload.employee_id or "").strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Для типа «Удостоверение/протокол» заполните код сотрудника.",
        )

    profile_settings = _scan_profile_settings(payload.scan_profile)
    effective_dpi = int(profile_settings.get("dpi", payload.dpi))
    effective_grayscale = bool(profile_settings.get("grayscale", payload.grayscale))

    args = [
        "scan-to-inbox",
        "--object-root",
        str(root),
        "--doc-type",
        doc_type,
        "--subject",
        payload.subject,
        "--device-index",
        str(payload.device_index),
        "--format",
        payload.image_format,
        "--dpi",
        str(effective_dpi),
    ]
    if effective_grayscale:
        args.append("--grayscale")
    if payload.employee_id:
        args.extend(["--employee-id", payload.employee_id])

    completed = _run_scanner_command(args)
    if completed.returncode != 0:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_humanize_scanner_error(completed.stderr, completed.stdout),
        )

    profile_label = str(profile_settings.get("label") or f"Профиль {int(payload.scan_profile or 1)}")
    message = completed.stdout.strip() or f"Скан добавлен во входящую папку ({profile_label})"
    return ArmActionResponse(ok=True, message=message)


@router.post("/scanner/recompress-history", response_model=ArmActionResponse)
def arm_scanner_recompress_history() -> ArmActionResponse:
    root = resolve_object_root()

    folders = [
        root / "10_scan_inbox",
        root / "08_outgoing_submissions",
        root / "02_personnel" / "employees",
    ]
    image_extensions = {".jpg", ".jpeg", ".png", ".tif", ".tiff"}

    scanned = 0
    changed = 0
    before_total = 0
    after_total = 0

    for base in folders:
        if not base.exists():
            continue
        for file in base.rglob("*"):
            if not file.is_file() or file.suffix.lower() not in image_extensions:
                continue
            scanned += 1
            profile_id = _detect_recompress_profile_for_file(file)
            ok, before_size, after_size, _ = _recompress_image_file(file, profile_id)
            before_total += before_size
            after_total += after_size
            if ok and after_size < before_size:
                changed += 1

    if scanned == 0:
        return ArmActionResponse(ok=True, message="Сканы для оптимизации не найдены.")

    saved_bytes = max(0, before_total - after_total)
    saved_mb = round(saved_bytes / (1024 * 1024), 2)
    before_mb = round(before_total / (1024 * 1024), 2)
    after_mb = round(after_total / (1024 * 1024), 2)
    message = (
        f"Оптимизация завершена: обработано {scanned} файлов, сжато {changed}. "
        f"Было {before_mb} МБ, стало {after_mb} МБ, экономия {saved_mb} МБ."
    )
    return ArmActionResponse(ok=True, message=message)


@router.post("/scan/ingest", response_model=ArmScanIngestResponse)
def arm_scan_ingest(payload: ArmScanIngestRequest, db: Session = Depends(get_db)) -> ArmScanIngestResponse:
    root = resolve_object_root()
    inbox = root / "10_scan_inbox"
    results = ingest_inbox(
        object_root=root,
        inbox_folder=inbox,
        db=db,
        enable_ocr=payload.enable_ocr,
        ocr_lang=payload.ocr_lang,
        tesseract_cmd=payload.tesseract_cmd,
        max_pdf_pages=payload.max_pdf_pages,
    )

    archived = sum(1 for item in results if item.status == "archived")
    manual_review = sum(1 for item in results if item.status == "manual_review")

    return ArmScanIngestResponse(
        archived=archived,
        manual_review=manual_review,
        items=[ArmScanIngestItem(**asdict(item)) for item in results],
    )


@router.get("/scan/manual-review")
def arm_scan_manual_review() -> list[dict[str, object]]:
    root = resolve_object_root()
    manual = root / "10_scan_inbox" / "manual_review"
    if not manual.exists():
        return []

    items: list[dict[str, object]] = []
    for file in sorted(manual.iterdir()):
        if not file.is_file():
            continue
        if file.suffix.lower() not in {".pdf", ".jpg", ".jpeg", ".png", ".tif", ".tiff"}:
            continue

        sidecar = file.with_suffix(f"{file.suffix}.ocr.txt")
        ocr_text = sidecar.read_text(encoding="utf-8", errors="ignore") if sidecar.exists() else None
        pred = classify_scan_candidate(filename=file.name, ocr_text=ocr_text)
        scan_doc_type = _detect_scan_doc_type_from_name(file.name)
        scan_subject_tag = _extract_scan_subject_tag(file.name)
        suggested_target = _suggest_manual_review_target_from_scan_name(file.name)
        if suggested_target == MANUAL_REVIEW_TARGET_BY_SCAN_TYPE["OTHER"]:
            suggested_target = _suggest_manual_review_target(pred.predicted_doc_type)
        suggested_target = f"{suggested_target.rstrip('/')}/{file.name}" if suggested_target else file.name
        items.append(
            {
                "rel_path": _to_rel_path(root, file),
                "scan_doc_type": scan_doc_type,
                "scan_subject_tag": scan_subject_tag,
                "predicted_doc_type": pred.predicted_doc_type,
                "confidence": pred.confidence,
                "source": pred.source,
                "matched_keywords": pred.matched_keywords,
                "ocr_text_rel_path": _to_rel_path(root, sidecar) if sidecar.exists() else None,
                "suggested_target_rel_path": suggested_target,
            }
        )
    return items


@router.post("/maintenance/reset-rebuild", response_model=ArmActionResponse)
def arm_maintenance_reset_rebuild(
    regenerate_project_orders: bool = True,
    overwrite_orders: bool = True,
    order_date: str = DEFAULT_ORDER_DATE,
) -> ArmActionResponse:
    root = resolve_object_root()
    if not root.exists() or not root.is_dir():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Корневая папка объекта не найдена")

    removed_items = 0
    for rel_path in MAINTENANCE_RESET_TARGETS:
        target = (root / rel_path).resolve()
        try:
            target.relative_to(root)
        except ValueError as exc:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Некорректная конфигурация reset") from exc
        removed_items += _clear_directory_contents(target)

    _ensure_object_structure_for_maintenance(root)

    employees_total = len(_iter_employee_catalog_rows(root))
    created_project_drafts = 0
    seed_employee_name: str | None = None

    if regenerate_project_orders:
        seed_employee = _pick_maintenance_seed_employee(root)
        if seed_employee is not None:
            seed_employee_name = seed_employee.employee_name
            created_project_drafts = _rebuild_project_order_drafts_from_seed_employee(
                root=root,
                seed_employee=seed_employee,
                overwrite=overwrite_orders,
                order_date=order_date,
            )

    message_parts = [
        "Сервисная очистка и пересборка структуры выполнены.",
        f"Удалено элементов: {removed_items}.",
        f"Сотрудников сохранено: {employees_total}.",
    ]

    if regenerate_project_orders:
        if seed_employee_name:
            message_parts.append(
                "Пересобрано проектных черновиков: "
                f"{created_project_drafts} (базовый сотрудник: {seed_employee_name})."
            )
        else:
            message_parts.append(
                "Проектные черновики не пересобраны: в каталоге нет сотрудников."
            )

    return ArmActionResponse(ok=True, message=" ".join(message_parts))


@router.post("/speech/google-transcribe", response_model=ArmSpeechTranscribeResponse)
def arm_speech_google_transcribe(audio: UploadFile = File(...), language: str = "ru-RU") -> ArmSpeechTranscribeResponse:
    if sr is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="SpeechRecognition не установлен в окружении сервера",
        )

    raw_audio = audio.file.read()
    if not raw_audio:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Пустой аудиофайл")

    if len(raw_audio) > 8_000_000:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="Аудиофайл слишком большой для распознавания (лимит 8 МБ)",
        )

    recognizer = sr.Recognizer()
    try:
        with sr.AudioFile(io.BytesIO(raw_audio)) as source:
            audio_data = recognizer.record(source)
        recognized_text = recognizer.recognize_google(audio_data, language=language)
    except sr.UnknownValueError:
        return ArmSpeechTranscribeResponse(
            ok=False,
            text="",
            provider="google-webspeech",
            message="Речь не распознана. Говорите чуть медленнее и ближе к микрофону.",
        )
    except sr.RequestError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=f"Сервис Google Speech недоступен: {exc}",
        ) from exc
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Не удалось обработать аудиофайл: {exc}",
        ) from exc

    clean_text = (recognized_text or "").strip()
    if not clean_text:
        return ArmSpeechTranscribeResponse(
            ok=False,
            text="",
            provider="google-webspeech",
            message="Google вернул пустой результат распознавания",
        )

    return ArmSpeechTranscribeResponse(
        ok=True,
        text=clean_text,
        provider="google-webspeech",
        message="Распознавание выполнено",
    )


@router.get("/employees/catalog", response_model=ArmEmployeeCatalogResponse)
def arm_employee_catalog(profession: str | None = None) -> ArmEmployeeCatalogResponse:
    root = resolve_object_root()
    return _build_employee_catalog(root=root, profession_filter=profession)


@router.get("/employees/checklist/overview", response_model=ArmEmployeeChecklistOverviewResponse)
def arm_employee_checklist_overview(profession: str | None = None) -> ArmEmployeeChecklistOverviewResponse:
    root = resolve_object_root()
    return _build_employee_checklist_overview(root=root, profession_filter=profession)


@router.get("/employees/checklist", response_model=ArmEmployeeChecklistResponse)
def arm_employee_checklist(employee_rel_path: str, profession: str | None = None) -> ArmEmployeeChecklistResponse:
    root = resolve_object_root()
    employee_root = _resolve_employee_root(root=root, employee_rel_path=employee_rel_path)
    return _build_employee_checklist_data(root=root, employee_root=employee_root, profession=profession)


@router.post("/employees/checklist/generate", response_model=ArmEmployeeChecklistGenerateResponse)
def arm_employee_checklist_generate(payload: ArmEmployeeChecklistGenerateRequest) -> ArmEmployeeChecklistGenerateResponse:
    root = resolve_object_root()
    employee_root = _resolve_employee_root(root=root, employee_rel_path=payload.employee_rel_path)

    analysis = _build_employee_checklist_data(
        root=root,
        employee_root=employee_root,
        profession=payload.profession,
    )

    rules_by_code = {rule.code: rule for rule in _iter_employee_tb_rules(analysis.profession)}
    checklist_by_code = {item.code: item for item in analysis.items}

    if payload.mode == "selected":
        if not payload.codes:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Для режима selected нужно передать список codes.",
            )
        selected_codes = [code.strip().upper() for code in payload.codes if code.strip()]
    elif payload.mode == "missing":
        selected_codes = [item.code for item in analysis.items if not item.ready]
    else:
        selected_codes = [item.code for item in analysis.items]

    unknown_codes = [code for code in selected_codes if code not in rules_by_code]
    if unknown_codes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                "Неизвестные коды чеклиста: "
                + ", ".join(sorted(set(unknown_codes)))
            ),
        )

    employee_output_dir = (employee_root / "07_templates_to_print").resolve()
    employee_output_dir.mkdir(parents=True, exist_ok=True)
    project_output_dir = (root / "01_orders_and_appointments" / "drafts_from_checklist").resolve()
    project_output_dir.mkdir(parents=True, exist_ok=True)

    employee_name = analysis.employee_name or employee_root.name
    created_files: list[str] = []
    skipped_files: list[str] = []
    context_overrides = (
        {
            "ORDER_DATE": payload.order_date,
            "ISSUE_DATE": payload.order_date,
        }
        if payload.order_date
        else None
    )

    for code in selected_codes:
        rule = rules_by_code[code]
        existing_item = checklist_by_code.get(code)

        # In missing mode with fresh analysis this should not happen, but keep guard explicit.
        if payload.mode == "missing" and existing_item and existing_item.ready:
            skipped_files.append(f"{code}: уже закрыто")
            continue

        file_name = ORDER_DRAFT_FILE_NAME_MAP.get(rule.code) or f"{rule.code}_{_slugify_filename(rule.title)}_draft.md"
        output_dir = project_output_dir if rule.scope == "project" else employee_output_dir
        draft_path = output_dir / file_name
        if draft_path.exists() and not payload.overwrite:
            skipped_files.append(f"{code}: {_to_rel_path(root, draft_path)}")
            continue

        draft_content = _build_draft_content(
            root=root,
            employee_root=employee_root,
            rule=rule,
            employee_rel_path=analysis.employee_rel_path,
            employee_name=employee_name,
            employee_id=analysis.employee_id,
            profession=analysis.profession,
            context_overrides=context_overrides,
        )
        draft_path.write_text(draft_content, encoding="utf-8")
        created_files.append(_to_rel_path(root, draft_path))

    if created_files:
        message = f"Создано файлов: {len(created_files)}"
    elif skipped_files:
        message = "Новые файлы не созданы: проверьте существующие черновики или режим генерации"
    else:
        message = "Все пункты чеклиста уже закрыты, генерация не требуется"

    return ArmEmployeeChecklistGenerateResponse(
        ok=bool(created_files),
        employee_rel_path=analysis.employee_rel_path,
        profession=analysis.profession,
        mode=payload.mode,
        created_files=created_files,
        skipped_files=skipped_files,
        message=message,
    )


@router.get("/", include_in_schema=False)
def arm_root_redirect() -> RedirectResponse:
    return RedirectResponse(url="/arm/dashboard", status_code=status.HTTP_307_TEMPORARY_REDIRECT)


@router.get("/research/competitors", response_class=HTMLResponse)
def arm_competitors_research_html() -> HTMLResponse:
    if not COMPETITOR_RESEARCH_HTML.exists() or not COMPETITOR_RESEARCH_HTML.is_file():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл исследования конкурентов не найден",
        )

    content: str
    try:
        content = COMPETITOR_RESEARCH_HTML.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        content = COMPETITOR_RESEARCH_HTML.read_text(encoding="cp1251", errors="ignore")

    return HTMLResponse(content=content)


@router.get("/research", include_in_schema=False)
def arm_research_redirect() -> RedirectResponse:
    return RedirectResponse(
        url="/arm/research/competitors",
        status_code=status.HTTP_307_TEMPORARY_REDIRECT,
    )


@router.get("/object-profile", response_model=ArmObjectProfileResponse)
def arm_object_profile() -> ArmObjectProfileResponse:
    root = resolve_object_root()
    metadata = _read_project_metadata(root)
    metadata_path = _project_metadata_path(root)
    ppr_sources = sorted(
        _to_rel_path(root, item)
        for item in (root / "00_incoming_requests").glob("*ППР*.pdf")
        if item.is_file()
    )
    context_path = _resolve_safe_path(root=root, rel_path=PPR_CONTEXT_REL_PATH)
    context_rel_path = _to_rel_path(root, context_path) if context_path.exists() else None

    return ArmObjectProfileResponse(
        object_name=metadata.get("object_name", ""),
        project_code=metadata.get("project_code", ""),
        organization=metadata.get("organization", ""),
        work_stage=metadata.get("work_stage", ""),
        start_date=(metadata.get("start_date") or "").strip() or None,
        metadata_rel_path=_to_rel_path(root, metadata_path),
        ppr_source_options=ppr_sources,
        ppr_context_rel_path=context_rel_path,
    )


@router.post("/object-profile", response_model=ArmActionResponse)
def arm_object_profile_save(payload: ArmObjectProfileUpdateRequest) -> ArmActionResponse:
    root = resolve_object_root()
    path = _write_project_metadata(root, payload)
    return ArmActionResponse(ok=True, message=f"Карточка объекта обновлена: {_to_rel_path(root, path)}")


@router.post("/ppr/import", response_model=ArmActionResponse)
def arm_ppr_import(payload: ArmPprImportRequest, db: Session = Depends(get_db)) -> ArmActionResponse:
    root = resolve_object_root()
    source = _resolve_safe_path(root=root, rel_path=payload.rel_path)
    if not source.exists() or not source.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл ППР не найден")

    extracted_text = _extract_source_document_text(source)
    context_markdown = _extract_ppr_context_markdown(extracted_text)
    context_path = _resolve_safe_path(root=root, rel_path=PPR_CONTEXT_REL_PATH)
    context_path.parent.mkdir(parents=True, exist_ok=True)
    context_path.write_text(context_markdown, encoding="utf-8")

    context_rel = _to_rel_path(root, context_path)
    snapshot = db.scalar(select(DocumentContent).where(DocumentContent.rel_path == context_rel))
    if snapshot is None:
        db.add(DocumentContent(rel_path=context_rel, content=context_markdown))
    else:
        snapshot.content = context_markdown
    db.commit()

    return ArmActionResponse(
        ok=True,
        message=(
            f"ППР импортирован в базу знаний: {context_rel}. "
            "Контекст будет использоваться для автозаполнения спецдокументов."
        ),
    )


def _arm_simple_nav_html(active: str) -> str:
    items = [
        ("dashboard", "/arm/dashboard", "Дашборд"),
        ("structure", "/arm/structure/view", "Структура и действия"),
        ("employees", "/arm/employees", "Сотрудники"),
        ("checklist", "/arm/checklist/view", "Чеклист"),
        ("permit-height", "/arm/permit/height", "Наряд высота"),
        ("aosr", "/arm/aosr", "АОСР"),
        ("todo", "/arm/todo/view", "План дня"),
        ("periodic", "/arm/periodic/view", "Периодические"),
        ("api", "/docs", "API"),
    ]
    links: list[str] = []
    for key, href, label in items:
        classes = "site-nav-link"
        if key == active:
            classes += " active"
        target = " target=\"_blank\" rel=\"noopener\"" if key == "api" else ""
        links.append(f"<a class=\"{classes}\" href=\"{escape(href)}\"{target}>{escape(label)}</a>")

    return (
        "<nav class=\"site-nav\"><div class=\"site-nav-inner\">"
        "<span class=\"site-nav-brand\">X5 УФА Э2</span>"
        f"{''.join(links)}"
        "</div></nav>"
    )


@router.get("/structure/view", response_class=HTMLResponse)
def arm_structure_view_html(db: Session = Depends(get_db)) -> HTMLResponse:
    root = resolve_object_root()
    documents = db.execute(select(Document).order_by(Document.created_at.desc())).scalars().all()
    rows: list[str] = []
    for doc in documents:
        rel_path = (doc.file_path or "").strip()
        order_no, order_title = _extract_order_header_info(root, rel_path)
        display_title = (doc.title or "").strip()
        if order_no or order_title:
            display_title = (("Приказ №" + order_no) if order_no else "").strip()
            if order_title:
                display_title = (display_title + " - " + order_title) if display_title else order_title
        if not display_title:
            display_title = "Документ"
        actions = "<span class=\"meta\">Путь не указан</span>"
        if rel_path:
            actions = _arm_file_actions_html(rel_path, back_href="/arm/structure/view")

        status = (doc.status or "new").strip().lower()
        if status not in DOC_STATUS_LABELS:
            status = "new"
        fix_value = escape(doc.fix_comment or "")
        status_badge = (
            '<span class="status-badge status-approved">✔ Утверждено</span>'
            if status == "approved"
            else '<span class="status-badge status-new">✔ Вновь созданные</span>'
            if status == "new"
            else '<span class="status-badge status-fix">⚙ Исправить</span>'
        )
        deletion_badge = (
            '<span class="status-badge status-delete">Помечен на удаление</span>'
            if getattr(doc, "marked_for_deletion", False)
            else '<span class="status-badge status-keep">В работе</span>'
        )

        rows.append(
            "<tr>"
            f"<td>{escape(str(doc.id))}</td>"
            f"<td>{escape(display_title)}</td>"
            f"<td>{escape(order_no or '—')}</td>"
            f"<td>{escape(order_title or '—')}</td>"
            f"<td>{escape(doc.doc_type)}</td>"
            f"<td>{status_badge}</td>"
            f"<td>{deletion_badge}</td>"
            f"<td><input class=\"fix-input\" type=\"text\" data-doc-id=\"{doc.id}\" value=\"{fix_value}\" placeholder=\"Что исправить\" /></td>"
            f"<td><select class=\"status-select\" data-doc-id=\"{doc.id}\">"
            f"<option value=\"approved\"{' selected' if status == 'approved' else ''}>Утверждено (зеленая галочка)</option>"
            f"<option value=\"new\"{' selected' if status == 'new' else ''}>Вновь созданные (синяя галочка)</option>"
            f"<option value=\"fix\"{' selected' if status == 'fix' else ''}>Исправить (желтая шестеренка)</option>"
            "</select></td>"
            f"<td><button type=\"button\" class=\"btn-inline\" data-save-id=\"{doc.id}\">Сохранить статус</button></td>"
            f"<td><div class=\"action-links\">"
            f"<button type=\"button\" class=\"btn-inline\" data-mark-delete-id=\"{doc.id}\">Пометить на удаление</button>"
            f"<button type=\"button\" class=\"btn-inline\" data-unmark-delete-id=\"{doc.id}\">Снять пометку</button>"
            f"</div></td>"
            f"<td>{actions}</td>"
            "</tr>"
        )

    rows_html = "".join(rows) if rows else "<tr><td colspan=\"12\">Документы отсутствуют</td></tr>"

    body_html = (
        "<section class=\"card\">"
        "<h1>Структура и действия</h1>"
        "<div class=\"meta\">Управление статусами документов в БД: зеленая галочка — утверждено, синяя — вновь создано, желтая шестеренка — требует исправления.</div>"
        '<div id="structureStatusMsg" class="meta" style="margin-top:8px;">Готово к изменениям.</div>'
        "<style>"
        ".structure-wrap{overflow-x:auto;padding-bottom:8px}"
        ".structure-table{width:100%;table-layout:fixed}"
        ".structure-table th,.structure-table td{padding:6px 8px;font-size:12px;line-height:1.2;vertical-align:top}"
        ".structure-table th{font-size:11px}"
        ".structure-table td:nth-child(1){width:48px}"
        ".structure-table td:nth-child(2){width:220px}"
        ".structure-table td:nth-child(3){width:72px}"
        ".structure-table td:nth-child(4){width:180px}"
        ".structure-table td:nth-child(5){width:74px}"
        ".structure-table td:nth-child(6){width:120px}"
        ".structure-table td:nth-child(7){width:112px}"
        ".structure-table td:nth-child(8){width:170px}"
        ".structure-table td:nth-child(9){width:132px}"
        ".structure-table td:nth-child(10){width:150px}"
        ".structure-table td:nth-child(11){width:240px}"
        ".structure-table td:nth-child(12){width:300px}"
        ".status-badge{display:inline-block;padding:3px 8px;border-radius:999px;font-size:11px;font-weight:700;line-height:1.15}"
        ".status-approved{background:#dcfce7;color:#166534}"
        ".status-new{background:#dbeafe;color:#1d4ed8}"
        ".status-fix{background:#fef3c7;color:#b45309}"
        ".status-delete{background:#fee2e2;color:#991b1b}"
        ".status-keep{background:#e2e8f0;color:#1e293b}"
        ".status-select,.fix-input{width:100%;min-width:120px;box-sizing:border-box;padding:6px 8px;font-size:12px}"
        ".fix-input{min-width:160px}"
        ".action-links{display:flex;gap:6px;flex-wrap:wrap}"
        ".action-links .btn-inline{min-width:150px;padding:6px 10px;font-size:12px;line-height:1.2}"
        "td .btn-inline{padding:6px 10px;font-size:12px;line-height:1.2;min-width:110px}"
        "@media (max-width: 1400px){.structure-table{table-layout:auto}}"
        "</style>"
        "<div class=\"structure-wrap\"><table class=\"structure-table\"><thead><tr>"
        "<th>ID</th><th>Документ</th><th>№ приказа</th><th>Наименование из шапки</th><th>Тип</th><th>Статус</th><th>Удаление</th><th>Что исправить</th><th>Новый статус</th><th>Действие</th><th>Управление</th><th>Файл</th>"
        "</tr></thead>"
        f"<tbody>{rows_html}</tbody></table></div>"
        "</section>"
        "<script>"
        "const structureStatusMsg = document.getElementById('structureStatusMsg');"
        "async function structureSaveStatus(docId){"
        "const statusSelect = document.querySelector('select.status-select[data-doc-id=\"' + docId + '\"]');"
        "const fixInput = document.querySelector('input.fix-input[data-doc-id=\"' + docId + '\"]');"
        "const payload = {status: statusSelect ? statusSelect.value : 'new', fix_comment: fixInput ? fixInput.value : ''};"
        "const response = await fetch('/documents/' + docId + '/status', {method:'PATCH', headers:{'Content-Type':'application/json'}, body: JSON.stringify(payload)});"
        "if (!response.ok){ const text = await response.text(); throw new Error(text || ('HTTP ' + response.status)); }"
        "return response.json();"
        "}"
        "async function structureSetDeletionMark(docId, marked){"
        "const response = await fetch('/documents/' + docId + '/deletion-mark', {method:'PATCH', headers:{'Content-Type':'application/json'}, body: JSON.stringify({marked_for_deletion: !!marked})});"
        "if (!response.ok){ const text = await response.text(); throw new Error(text || ('HTTP ' + response.status)); }"
        "return response.json();"
        "}"
        "for (const btn of document.querySelectorAll('button[data-save-id]')) {"
        "btn.addEventListener('click', async () => {"
        "const docId = btn.getAttribute('data-save-id');"
        "structureStatusMsg.textContent = 'Сохранение статуса...';"
        "try { await structureSaveStatus(docId); structureStatusMsg.textContent = 'Статус сохранен для документа #' + docId + '. Обновите страницу для обновления бейджа.'; }"
        "catch (err) { structureStatusMsg.textContent = 'Ошибка сохранения: ' + err.message; }"
        "});"
        "}"
        "for (const btn of document.querySelectorAll('button[data-mark-delete-id]')) {"
        "btn.addEventListener('click', async () => {"
        "const docId = btn.getAttribute('data-mark-delete-id');"
        "structureStatusMsg.textContent = 'Установка пометки на удаление...';"
        "try { await structureSetDeletionMark(docId, true); structureStatusMsg.textContent = 'Документ #' + docId + ' помечен на удаление. Обновите страницу для обновления бейджа.'; }"
        "catch (err) { structureStatusMsg.textContent = 'Ошибка пометки: ' + err.message; }"
        "});"
        "}"
        "for (const btn of document.querySelectorAll('button[data-unmark-delete-id]')) {"
        "btn.addEventListener('click', async () => {"
        "const docId = btn.getAttribute('data-unmark-delete-id');"
        "structureStatusMsg.textContent = 'Снятие пометки...';"
        "try { await structureSetDeletionMark(docId, false); structureStatusMsg.textContent = 'Пометка на удаление снята для документа #' + docId + '. Обновите страницу для обновления бейджа.'; }"
        "catch (err) { structureStatusMsg.textContent = 'Ошибка снятия пометки: ' + err.message; }"
        "});"
        "}"
        "</script>"
    )
    return _arm_simple_page(title="АРМ: структура и действия", active_nav="structure", body_html=body_html)


def _arm_simple_page(title: str, active_nav: str, body_html: str) -> HTMLResponse:
    html = f"""
<!doctype html>
<html lang=\"ru\">
<head>
    <meta charset=\"utf-8\" />
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
    <title>{escape(title)}</title>
    <style>
    :root {{ --bg: #f4efe7; --card: #fffaf3; --ink: #1f2a36; --accent: #0f766e; --muted: #5b6672; }}
    body {{ margin: 0; font-family: "Segoe UI", Tahoma, sans-serif; background: radial-gradient(circle at 85% 15%, #d7efe6 0%, #f4efe7 45%, #e6eff4 100%); color: var(--ink); }}
    .site-nav {{ position: sticky; top: 0; z-index: 50; background: var(--card); border-bottom: 1px solid #d8cfc2; box-shadow: 0 2px 8px rgba(0,0,0,0.06); }}
    .site-nav-inner {{ max-width: 1180px; margin: 0 auto; padding: 0 16px; display: flex; align-items: center; gap: 4px; height: 46px; flex-wrap: wrap; }}
    .site-nav-brand {{ font-weight: 700; color: var(--accent); font-size: 14px; margin-right: 12px; white-space: nowrap; }}
    .site-nav-link {{ padding: 6px 12px; border-radius: 999px; font-size: 13px; font-weight: 600; color: var(--ink); text-decoration: none; white-space: nowrap; }}
    .site-nav-link:hover {{ background: #eff6fb; }}
    .site-nav-link.active {{ background: var(--accent); color: #fff; }}
    .wrap {{ max-width: 1180px; margin: 16px auto; padding: 0 16px 24px; }}
    .card {{ background: var(--card); border: 1px solid #d8cfc2; border-radius: 16px; padding: 14px; box-shadow: 0 8px 24px rgba(0, 0, 0, 0.05); }}
    .meta {{ color: var(--muted); font-size: 14px; }}
    .controls {{ display: flex; gap: 8px; flex-wrap: wrap; margin: 10px 0; align-items: center; }}
    .action-links {{ display: flex; gap: 10px; flex-wrap: wrap; align-items: center; }}
    .btn-inline {{ display: inline-flex; align-items: center; justify-content: center; padding: 7px 12px; border-radius: 999px; border: 1px solid #b8c7c5; background: #fff; color: #0f4c5c; text-decoration: none; font-weight: 600; font-size: 13px; }}
    .btn-inline:hover {{ background: #eef7f5; }}
    .btn-inline.primary {{ background: var(--accent); border-color: var(--accent); color: #fff; }}
    table {{ width: 100%; border-collapse: collapse; background: #fff; border: 1px solid #d8e0e0; border-radius: 12px; overflow: hidden; }}
    th, td {{ padding: 8px 10px; border-bottom: 1px solid #e5e7eb; text-align: left; vertical-align: top; font-size: 14px; }}
    th {{ background: #f8fbfd; color: #334155; }}
    tr:last-child td {{ border-bottom: none; }}
    .ok {{ color: #166534; font-weight: 700; }}
    .warn {{ color: #b45309; font-weight: 700; }}
    .btn-link {{ color: #0f4c5c; text-decoration: none; font-weight: 600; }}
    .btn-link:hover {{ text-decoration: underline; }}
    input, select {{ border-radius: 10px; border: 1px solid #c5cbc9; padding: 8px 10px; font-size: 14px; background: #fff; }}
    </style>
</head>
<body>
    {_arm_simple_nav_html(active_nav)}
    <div class=\"wrap\">{body_html}</div>
</body>
</html>
"""
    return HTMLResponse(content=html)


def _arm_is_text_editable(rel_path: str) -> bool:
    suffix = PurePosixPath(rel_path).suffix.lower()
    return suffix in {".md", ".txt", ".json", ".csv", ".html", ".xml", ".yaml", ".yml"}


def _arm_action_anchor(href: str, label: str, primary: bool = False, new_tab: bool = False) -> str:
    cls = "btn-inline primary" if primary else "btn-inline"
    target = ' target="_blank" rel="noopener noreferrer"' if new_tab else ""
    return f'<a class="{cls}" href="{escape(href)}"{target}>{escape(label)}</a>'


def _arm_file_actions_html(
    rel_path: str,
    *,
    folder_rel_path: str | None = None,
    back_href: str | None = None,
) -> str:
    file_href = quote(rel_path, safe="")
    folder_rel = folder_rel_path
    if folder_rel is None:
        folder_rel = str(PurePosixPath(rel_path).parent)
        if folder_rel == ".":
            folder_rel = ""

    links = [
        _arm_action_anchor(f"/arm/fs/view?rel_path={file_href}", "Просмотр", new_tab=True),
        _arm_action_anchor(f"/arm/fs/print-preview?rel_path={file_href}&auto_print=1", "Печать", new_tab=True),
        _arm_action_anchor(f"/arm/fs/download?rel_path={file_href}", "Скачать", new_tab=True),
    ]

    if _arm_is_text_editable(rel_path):
        editor_href = f"/arm/editor?rel_path={file_href}"
        if back_href:
            editor_href += f"&back={quote(back_href, safe='')}"
        links.insert(0, _arm_action_anchor(editor_href, "Редактировать", primary=True))

    if folder_rel:
        links.append(
            _arm_action_anchor(
                f"/arm/dashboard?open_path={quote(folder_rel, safe='')}",
                "Папка",
            )
        )

    return f'<div class="action-links">{"".join(links)}</div>'


@router.get("/editor", response_class=HTMLResponse)
def arm_editor_html(rel_path: str, back: str | None = None) -> HTMLResponse:
    root = resolve_object_root()
    target = _resolve_safe_path(root=root, rel_path=rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден")
    if not _arm_is_text_editable(rel_path):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Редактор доступен только для текстовых файлов")

    rel = _to_rel_path(root, target)
    folder_rel = str(PurePosixPath(rel).parent)
    if folder_rel == ".":
        folder_rel = ""
    back_html = _arm_action_anchor(back, "Назад") if back else ""
    body_html = f"""
    <section class=\"card\">
        <h1>Редактор документа</h1>
        <div class=\"meta\">Файл: {escape(rel)}</div>
        <div class=\"controls\">
            {back_html}
            {_arm_action_anchor(f'/arm/fs/view?rel_path={quote(rel, safe="")}', 'Просмотр', new_tab=True)}
            {_arm_action_anchor(f'/arm/fs/print-preview?rel_path={quote(rel, safe="")}&auto_print=1', 'Печать', new_tab=True)}
            {_arm_action_anchor(f'/arm/fs/download?rel_path={quote(rel, safe="")}', 'Скачать', new_tab=True)}
            {_arm_action_anchor(f'/arm/dashboard?open_path={quote(folder_rel, safe="")}', 'Папка') if folder_rel else ''}
            <button type=\"button\" class=\"btn-inline\" id=\"editorSourceBtn\">Исходник (Markdown)</button>
            <button type=\"button\" class=\"btn-inline primary\" id=\"editorSaveBtn\">Сохранить</button>
            <button type=\"button\" class=\"btn-inline\" id=\"editorReloadBtn\">Перечитать</button>
        </div>
        <div id=\"editorStatus\" class=\"meta\">Загрузка...</div>
    </section>
    <section class=\"editor-main\">
        <div class=\"card\">
            <div id=\"editorPreview\" contenteditable=\"true\" spellcheck=\"false\" class=\"editor-form\"></div>
        </div>
    </section>

    <!-- Hidden modal for markdown source view -->
    <div id=\"editorSourceModal\" class=\"modal\" style=\"display: none;\">
        <div class=\"modal-content\">
            <div class=\"modal-header\">
                <h2>Markdown исходник</h2>
                <button type=\"button\" class=\"btn-close\" id=\"editorSourceClose\">✕</button>
            </div>
            <textarea id=\"editorSource\" spellcheck=\"false\" class=\"editor-source-modal\"></textarea>
            <div class=\"modal-footer\">
                <button type=\"button\" class=\"btn-inline\" id=\"editorSourceClose2\">Закрыть</button>
                <button type=\"button\" class=\"btn-inline primary\" id=\"editorSourceReload\">Перезагрузить из MD</button>
            </div>
        </div>
    </div>

    <style>
    .editor-main {{ margin-top: 16px; }}
    .editor-form {{ 
        width: 100%; 
        min-height: 72vh; 
        border: 1px solid #ccd4d7; 
        border-radius: 12px; 
        padding: 24px 32px; 
        box-sizing: border-box; 
        background: #fff; 
        color: #1f2a36; 
        font: 14px/1.6 'Segoe UI', Arial, sans-serif; 
        overflow: auto; 
        outline: none; 
        cursor: text;
    }}
    .editor-form:focus {{ border-color: #1a9c7b; box-shadow: 0 0 0 2px #1a9c7b30; }}
    .editor-form table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
    .editor-form th, .editor-form td {{ border: 1px solid #ccd4d7; padding: 8px 12px; }}
    .editor-form th {{ background: #f3f6f8; font-weight: 600; }}
    .editor-form h1 {{ font-size: 1.5em; margin: 20px 0 10px; font-weight: 700; }}
    .editor-form h2 {{ font-size: 1.25em; margin: 16px 0 8px; font-weight: 700; }}
    .editor-form h3 {{ font-size: 1.1em; margin: 12px 0 6px; font-weight: 700; }}
    .editor-form p {{ margin: 8px 0; }}
    .editor-form .fill-line {{ border-bottom: 1px solid #333; min-width: 200px; display: inline-block; text-align: center; }}
    .editor-form .fill-note {{ font-size: 11px; color: #666; text-align: center; display: block; margin-top: 2px; }}
    .editor-form .meta-row {{ display: flex; gap: 24px; margin: 6px 0; }}

    .modal {{ position: fixed; top: 0; left: 0; right: 0; bottom: 0; background: rgba(0,0,0,0.5); display: flex; align-items: center; justify-content: center; z-index: 1000; }}
    .modal-content {{ background: #fff; border-radius: 12px; width: 90%; max-width: 900px; max-height: 90vh; display: flex; flex-direction: column; box-shadow: 0 8px 32px rgba(0,0,0,0.2); }}
    .modal-header {{ padding: 16px 20px; border-bottom: 1px solid #e0e6ea; display: flex; justify-content: space-between; align-items: center; }}
    .modal-header h2 {{ margin: 0; font-size: 1.1em; }}
    .btn-close {{ background: none; border: none; font-size: 24px; cursor: pointer; color: #8a9aab; }}
    .btn-close:hover {{ color: #1f2a36; }}
    #editorSource {{ 
        flex: 1; 
        border: none; 
        padding: 12px; 
        font: 13px/1.5 Consolas, 'Courier New', monospace; 
        resize: none; 
        background: #f8fafb; 
        color: #1f2a36;
    }}
    .modal-footer {{ padding: 12px 20px; border-top: 1px solid #e0e6ea; display: flex; gap: 8px; justify-content: flex-end; }}
    </style>
    <script src=\"https://cdn.jsdelivr.net/npm/marked@9/marked.min.js\"></script>
    <script src=\"https://cdn.jsdelivr.net/npm/turndown@7/umd/turndown.min.js\"></script>
    <script>
    const editorRelPath = {rel!r};
    const editorPreview = document.getElementById('editorPreview');
    const editorSource = document.getElementById('editorSource');
    const editorStatus = document.getElementById('editorStatus');
    const editorSaveBtn = document.getElementById('editorSaveBtn');
    const editorReloadBtn = document.getElementById('editorReloadBtn');
    const editorSourceBtn = document.getElementById('editorSourceBtn');
    const editorSourceModal = document.getElementById('editorSourceModal');
    const editorSourceClose = document.getElementById('editorSourceClose');
    const editorSourceClose2 = document.getElementById('editorSourceClose2');
    const editorSourceReload = document.getElementById('editorSourceReload');

    // --- MD ↔ HTML компилятор ---
    marked.use({{ gfm: true, breaks: false }});
    const td = new TurndownService({{ headingStyle: 'atx', bulletListMarker: '-', codeBlockStyle: 'fenced' }});
    td.addRule('keep-classed-divs', {{
        filter: function(node) {{ return node.nodeName === 'DIV' && node.className; }},
        replacement: function(_c, node) {{ return '\n\n' + node.outerHTML + '\n\n'; }}
    }});
    td.addRule('keep-classed-spans', {{
        filter: function(node) {{ return node.nodeName === 'SPAN' && node.className; }},
        replacement: function(_c, node) {{ return node.outerHTML; }}
    }});

    function mdToHtml(md) {{ return marked.parse(md || ''); }}
    function htmlToMd(html) {{ return td.turndown(html || ''); }}

    let _syncLock = false;
    let _debounce = null;
    let _isModified = false;

    // Редактирование печатной формы → обновляем скрытый markdown (с дебаунсом)
    editorPreview.addEventListener('input', function() {{
        if (_syncLock) return;
        _isModified = true;
        clearTimeout(_debounce);
        _debounce = setTimeout(function() {{
            _syncLock = true;
            editorSource.value = htmlToMd(editorPreview.innerHTML);
            _syncLock = false;
        }}, 200);
    }});

            editorStatus.textContent = 'Символов: ' + editorSource.value.length + ' · не сохранено';
            _syncLock = false;
        }}, 200);
    }});

    async function editorApi(url, options) {{
        const response = await fetch(url, options || {{}});
        if (!response.ok) {{
            const text = await response.text();
            throw new Error(text || ('HTTP ' + response.status));
        }}
        const contentType = response.headers.get('content-type') || '';
        if (contentType.includes('application/json')) {{
            return response.json();
        }}
        return response.text();
    }}

    async function loadEditorFile() {{
        editorStatus.textContent = 'Чтение файла...';
        const data = await editorApi('/arm/fs/file?rel_path=' + encodeURIComponent(editorRelPath));
        const content = data.content || '';
        editorSource.value = content;
        _syncLock = true;
        editorPreview.innerHTML = mdToHtml(content);
        _syncLock = false;
        _isModified = false;
        editorStatus.textContent = 'Файл загружен. Символов: ' + content.length + '.';
    }}

    async function saveEditorFile() {{
        editorSaveBtn.disabled = true;
        editorStatus.textContent = 'Сохранение...';
        try {{
            // Перед сохранением: конвертируем текущий HTML обратно в MD
            const mdContent = htmlToMd(editorPreview.innerHTML);
            const data = await editorApi('/arm/fs/file', {{
                method: 'POST',
                headers: {{ 'Content-Type': 'application/json' }},
                body: JSON.stringify({{ rel_path: editorRelPath, content: mdContent }})
            }});
            editorStatus.textContent = data.message || 'Сохранено.';
            _isModified = false;
        }} finally {{
            editorSaveBtn.disabled = false;
        }}
    }}

    // Обработчики кнопок
    editorSaveBtn.addEventListener('click', () => saveEditorFile().catch((err) => {{
        editorStatus.textContent = 'Ошибка сохранения: ' + err.message;
    }}));
    editorReloadBtn.addEventListener('click', () => loadEditorFile().catch((err) => {{
        editorStatus.textContent = 'Ошибка чтения: ' + err.message;
    }}));

    // Модальное окно для просмотра markdown исходника
    editorSourceBtn.addEventListener('click', () => {{
        // Обновляем markdown перед открытием модала
        editorSource.value = htmlToMd(editorPreview.innerHTML);
        editorSourceModal.style.display = 'flex';
    }});
    editorSourceClose.addEventListener('click', () => {{
        editorSourceModal.style.display = 'none';
    }});
    editorSourceClose2.addEventListener('click', () => {{
        editorSourceModal.style.display = 'none';
    }});
    editorSourceModal.addEventListener('click', (e) => {{
        if (e.target === editorSourceModal) {{
            editorSourceModal.style.display = 'none';
        }}
    }});
    editorSourceReload.addEventListener('click', () => {{
        // Загружаем из markdown и обновляем форму
        _syncLock = true;
        editorPreview.innerHTML = mdToHtml(editorSource.value);
        _syncLock = false;
        editorStatus.textContent = 'Форма перезагружена из MD.';
        editorSourceModal.style.display = 'none';
    }});

    // Загружаем файл при открытии
    loadEditorFile().catch((err) => {{
        editorStatus.textContent = 'Ошибка чтения: ' + err.message;
    }});
    </script>
    """
    return _arm_simple_page(title=f"АРМ: редактор {target.name}", active_nav="checklist", body_html=body_html)


@router.get("/employees", response_class=HTMLResponse)
def arm_employees_html(profession: str | None = None) -> HTMLResponse:
    root = resolve_object_root()
    catalog = _build_employee_catalog(root=root, profession_filter=profession)

    rows: list[str] = []
    for item in catalog.items:
        open_path_url = f"/arm/dashboard?open_path={quote(item.employee_rel_path, safe='')}"
        card_url = f"/arm/employee/view?employee_rel_path={quote(item.employee_rel_path, safe='')}"
        if profession:
            card_url += f"&profession={quote(profession, safe='')}"
        rows.append(
            "<tr>"
            f"<td>{escape(item.employee_id or '')}</td>"
            f"<td>{escape(item.employee_name)}</td>"
            f"<td>{escape(item.position or '')}</td>"
            f"<td>{escape(item.profession_label)}</td>"
            f"<td><div class=\"action-links\">"
            f"<a class=\"btn-inline primary\" href=\"{escape(card_url)}\">Карточка</a>"
            f"<a class=\"btn-inline\" href=\"{escape(open_path_url)}\">Папка</a>"
            f"</div></td>"
            "</tr>"
        )

    option_html = [
        f"<option value=\"\"{' selected' if not profession else ''}>Все группы</option>"
    ]
    for opt in catalog.profession_options:
        selected = " selected" if profession == opt.key else ""
        option_html.append(
            f"<option value=\"{escape(opt.key)}\"{selected}>{escape(opt.label)}</option>"
        )

    employee_rows_html = "".join(rows) if rows else "<tr><td colspan=\"5\">Нет сотрудников по выбранному фильтру</td></tr>"

    body_html = (
        "<section class=\"card\">"
        "<h1>Каталог сотрудников</h1>"
        f"<div class=\"meta\">Всего сотрудников: {catalog.total}. Для каждого сотрудника теперь доступна отдельная карточка с чеклистом, связями и генерацией недостающих документов.</div>"
        "<form class=\"controls\" method=\"get\" action=\"/arm/employees\">"
        "<label for=\"profession\" class=\"meta\">Группа:</label>"
        f"<select id=\"profession\" name=\"profession\">{''.join(option_html)}</select>"
        "<button type=\"submit\">Применить</button>"
        "</form>"
        "<table><thead><tr><th>ID</th><th>ФИО</th><th>Должность</th><th>Группа</th><th>Действие</th></tr></thead>"
        f"<tbody>{employee_rows_html}</tbody></table>"
        "</section>"
    )
    return _arm_simple_page(title="АРМ: сотрудники", active_nav="employees", body_html=body_html)


@router.get("/employee/view", response_class=HTMLResponse)
def arm_employee_card_html(employee_rel_path: str, profession: str | None = None) -> HTMLResponse:
    root = resolve_object_root()
    employee_root = _resolve_employee_root(root=root, employee_rel_path=employee_rel_path)
    rel = _to_rel_path(root, employee_root)
    profile = _read_employee_profile(employee_root)
    employee_name = _employee_display_name(employee_root, profile)
    position = profile.get("position") or profile.get("profession") or ""
    profile_rel = _to_rel_path(root, employee_root / "employee_profile.txt")
    profession_options = _build_employee_catalog(root=root).profession_options
    option_html = [
        f'<option value=""{(" selected" if not profession else "")}>Авто</option>'
    ]
    for opt in profession_options:
        selected = " selected" if profession == opt.key else ""
        option_html.append(f'<option value="{escape(opt.key)}"{selected}>{escape(opt.label)}</option>')
    profile_link_html = ""
    if (employee_root / "employee_profile.txt").exists():
        back_href = f"/arm/employee/view?employee_rel_path={quote(rel, safe='')}"
        if profession:
            back_href += f"&profession={quote(profession, safe='')}"
        profile_link_html = _arm_action_anchor(
            f"/arm/editor?rel_path={quote(profile_rel, safe='')}&back={quote(back_href, safe='')}",
            "Профиль",
            primary=True,
        )

    body_html = f"""
    <section class=\"card\">
        <h1>Карточка сотрудника</h1>
        <div class=\"meta\">{escape(employee_name)}{(' • ' + escape(position)) if position else ''}</div>
        <div class=\"meta\">Папка: {escape(rel)}</div>
        <div class=\"controls\">
            {_arm_action_anchor('/arm/employees', 'К каталогу')}
            {_arm_action_anchor(f'/arm/dashboard?open_path={quote(rel, safe="")}', 'Папка')}
            {profile_link_html}
            <label class=\"meta\" for=\"employeeProfession\">Группа:</label>
            <select id=\"employeeProfession\">{''.join(option_html)}</select>
            <button type=\"button\" class=\"btn-inline\" id=\"employeeRefreshBtn\">Обновить</button>
            <button type=\"button\" class=\"btn-inline primary\" id=\"employeeGenerateMissingBtn\">Сгенерировать недостающее</button>
            <button type=\"button\" class=\"btn-inline\" id=\"employeeGenerateSelectedBtn\">Сгенерировать выбранное</button>
            <button type=\"button\" class=\"btn-inline\" id=\"employeeGenerateAllBtn\">Сгенерировать все</button>
        </div>
        <div id=\"employeeCardStatus\" class=\"meta\">Загрузка чеклиста...</div>
        <div id=\"employeeCardSummary\" class=\"meta\" style=\"margin-top:6px;\"></div>
    </section>
    <section class=\"card\" style=\"margin-top:16px;\">
        <div id=\"employeeChecklistItems\" class=\"employee-card-list\"></div>
    </section>
    <style>
    .employee-card-list {{ display: grid; gap: 12px; }}
    .employee-card-item {{ border: 1px solid #d8e0e0; border-radius: 14px; padding: 12px; background: #fff; }}
    .employee-card-item.missing {{ border-color: #f3d7a8; background: #fff9ef; }}
    .employee-card-head {{ display: grid; grid-template-columns: 24px minmax(0, 1fr) auto; gap: 10px; align-items: start; }}
    .employee-card-title {{ font-weight: 700; }}
    .employee-card-files {{ display: grid; gap: 8px; margin-top: 10px; }}
    .employee-card-file {{ border-top: 1px dashed #d7dfe2; padding-top: 8px; }}
    .employee-card-links {{ display: flex; gap: 8px; flex-wrap: wrap; margin-top: 6px; }}
    </style>
    <script>
    const employeeRelPath = {rel!r};
    const employeeProfessionSelect = document.getElementById('employeeProfession');
    const employeeCardStatus = document.getElementById('employeeCardStatus');
    const employeeCardSummary = document.getElementById('employeeCardSummary');
    const employeeChecklistItems = document.getElementById('employeeChecklistItems');
    const employeeRefreshBtn = document.getElementById('employeeRefreshBtn');
    const employeeGenerateMissingBtn = document.getElementById('employeeGenerateMissingBtn');
    const employeeGenerateSelectedBtn = document.getElementById('employeeGenerateSelectedBtn');
    const employeeGenerateAllBtn = document.getElementById('employeeGenerateAllBtn');

    function employeeEscapeHtml(value) {{
        return String(value || '').replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/"/g, '&quot;');
    }}

    function employeeIsEditable(relPath) {{
        return /\.(md|txt|json|csv|html|xml|ya?ml)$/i.test(relPath || '');
    }}

    function employeeFileActions(relPath) {{
        const encoded = encodeURIComponent(relPath);
        const back = encodeURIComponent(window.location.pathname + window.location.search);
        const links = [];
        if (employeeIsEditable(relPath)) {{
            links.push('<a class="btn-inline primary" href="/arm/editor?rel_path=' + encoded + '&back=' + back + '">Редактировать</a>');
        }}
        links.push('<a class="btn-inline" target="_blank" rel="noopener noreferrer" href="/arm/fs/view?rel_path=' + encoded + '">Просмотр</a>');
        links.push('<a class="btn-inline" target="_blank" rel="noopener noreferrer" href="/arm/fs/print-preview?rel_path=' + encoded + '&auto_print=1">Печать</a>');
        links.push('<a class="btn-inline" target="_blank" rel="noopener noreferrer" href="/arm/fs/download?rel_path=' + encoded + '">Скачать</a>');
        return '<div class="employee-card-links">' + links.join('') + '</div>';
    }}

    async function employeeApi(url, options) {{
        const response = await fetch(url, options || {{}});
        if (!response.ok) {{
            const text = await response.text();
            throw new Error(text || ('HTTP ' + response.status));
        }}
        return response.json();
    }}

    function employeeSelectedCodes() {{
        return Array.from(document.querySelectorAll('.employee-item-checkbox:checked')).map((node) => node.dataset.code || '').filter(Boolean);
    }}

    async function loadEmployeeChecklist() {{
        employeeCardStatus.textContent = 'Загрузка чеклиста...';
        const profession = (employeeProfessionSelect.value || '').trim();
        let url = '/arm/employees/checklist?employee_rel_path=' + encodeURIComponent(employeeRelPath);
        if (profession) {{
            url += '&profession=' + encodeURIComponent(profession);
        }}
        const data = await employeeApi(url);
        employeeCardSummary.textContent = 'Готово: ' + data.ready_count + '/' + data.total_required + ' (' + data.progress_percent + '%). Не хватает: ' + data.missing_count + '.';
        employeeChecklistItems.innerHTML = '';
        for (const item of (data.items || [])) {{
            const wrapper = document.createElement('div');
            wrapper.className = 'employee-card-item ' + (item.ready ? 'ready' : 'missing');
            let foundHtml = '';
            for (const foundFile of (item.found_files || [])) {{
                foundHtml += '<div class="employee-card-file"><div class="meta">Найдено: ' + employeeEscapeHtml(foundFile) + '</div>' + employeeFileActions(foundFile) + '</div>';
            }}
            let relatedHtml = '';
            for (const relatedFile of (item.related_files || [])) {{
                relatedHtml += '<div class="employee-card-file"><div class="meta">Связанный документ: ' + employeeEscapeHtml(relatedFile) + '</div>' + employeeFileActions(relatedFile) + '</div>';
            }}
            wrapper.innerHTML = ''
                + '<div class="employee-card-head">'
                + '<input class="employee-item-checkbox" type="checkbox" data-code="' + employeeEscapeHtml(item.code) + '"' + (item.ready ? '' : ' checked') + '>'
                + '<div><div class="employee-card-title">' + employeeEscapeHtml(item.code) + ' • ' + employeeEscapeHtml(item.title) + '</div>'
                + '<div class="meta">Папка: ' + employeeEscapeHtml(item.folder_rel_path) + '</div>'
                + '<div class="meta">Подсказка: ' + employeeEscapeHtml(item.guidance || '-') + '</div></div>'
                + '<div class="meta">' + (item.ready ? 'Готово ' : 'Не хватает ') + item.found_count + '/' + item.required_count + '</div>'
                + '</div>'
                + '<div class="employee-card-links" style="margin-top:10px;">'
                + '<a class="btn-inline" href="/arm/dashboard?open_path=' + encodeURIComponent(item.folder_rel_path || '') + '">Папка</a>'
                + '</div>'
                + ((foundHtml || relatedHtml) ? ('<div class="employee-card-files">' + foundHtml + relatedHtml + '</div>') : '');
            employeeChecklistItems.appendChild(wrapper);
        }}
        employeeCardStatus.textContent = 'Чеклист обновлен.';
    }}

    async function generateEmployeeDocs(mode) {{
        const profession = (employeeProfessionSelect.value || '').trim() || null;
        const body = {{
            employee_rel_path: employeeRelPath,
            profession: profession,
            mode: mode,
            codes: mode === 'selected' ? employeeSelectedCodes() : []
        }};
        employeeCardStatus.textContent = 'Генерация...';
        const data = await employeeApi('/arm/employees/checklist/generate', {{
            method: 'POST',
            headers: {{ 'Content-Type': 'application/json' }},
            body: JSON.stringify(body)
        }});
        employeeCardStatus.textContent = data.message || 'Готово.';
        await loadEmployeeChecklist();
    }}

    employeeRefreshBtn.addEventListener('click', () => loadEmployeeChecklist().catch((err) => {{
        employeeCardStatus.textContent = 'Ошибка: ' + err.message;
    }}));
    employeeGenerateMissingBtn.addEventListener('click', () => generateEmployeeDocs('missing').catch((err) => {{
        employeeCardStatus.textContent = 'Ошибка: ' + err.message;
    }}));
    employeeGenerateSelectedBtn.addEventListener('click', () => generateEmployeeDocs('selected').catch((err) => {{
        employeeCardStatus.textContent = 'Ошибка: ' + err.message;
    }}));
    employeeGenerateAllBtn.addEventListener('click', () => generateEmployeeDocs('all').catch((err) => {{
        employeeCardStatus.textContent = 'Ошибка: ' + err.message;
    }}));
    loadEmployeeChecklist().catch((err) => {{
        employeeCardStatus.textContent = 'Ошибка: ' + err.message;
    }});
    </script>
    """
    return _arm_simple_page(title=f"АРМ: {employee_name}", active_nav="employees", body_html=body_html)


@router.get("/checklist/view", response_class=HTMLResponse)
def arm_checklist_html() -> HTMLResponse:
    root = resolve_object_root()
    checklist = _build_checklist(root=root)
    ready_count = sum(1 for item in checklist if item.ready)

    rows: list[str] = []
    for item in checklist:
        status_html = "<span class=\"ok\">Готово</span>" if item.ready else "<span class=\"warn\">Пробел</span>"
        used_html = "<span class=\"ok\">Да</span>" if item.found > 0 else "<span class=\"warn\">Нет</span>"
        action_html = ""
        files_html = "<div class=\"meta\">Файлы не найдены</div>"
        try:
            rel_path = _to_rel_path(root, Path(item.location).resolve())
            action_parts = [
                _arm_action_anchor(f"/arm/dashboard?open_path={quote(rel_path, safe='')}", "Открыть раздел")
            ]
            if item.code == "1.3.3":
                action_parts.insert(0, _arm_action_anchor("/arm/permit/height", "Оформить наряд", primary=True))
            action_html = f'<div class="action-links">{"".join(action_parts)}</div>'
        except Exception:  # noqa: BLE001
            action_html = "-"

        if item.found_files:
            file_parts: list[str] = []
            for found_file in item.found_files[:4]:
                file_parts.append(
                    "<div class=\"checklist-file-item\">"
                    f"<div class=\"meta\">{escape(PurePosixPath(found_file).name)}</div>"
                    f"{_arm_file_actions_html(found_file, back_href='/arm/checklist/view')}"
                    "</div>"
                )
            files_html = f'<div class="checklist-files">{"".join(file_parts)}</div>'

        rows.append(
            "<tr>"
            f"<td>{escape(item.code)}</td>"
            f"<td>{escape(item.title)}</td>"
            f"<td>{item.found}/{item.required_min}</td>"
            f"<td>{used_html}</td>"
            f"<td>{status_html}</td>"
            f"<td>{files_html}</td>"
            f"<td>{action_html}</td>"
            "</tr>"
        )

    checklist_rows_html = "".join(rows) if rows else "<tr><td colspan=\"7\">Чеклист пуст</td></tr>"

    body_html = (
        "<section class=\"card\">"
        "<h1>Чеклист объекта</h1>"
        f"<div class=\"meta\">Готово: {ready_count} из {len(checklist)}. Для найденных файлов доступны прямые действия, а по наряду на высоте есть отдельная форма под реальный бланк.</div>"
        "<style>.checklist-files{display:grid;gap:8px}.checklist-file-item{padding:8px;border:1px solid #d8e0e0;border-radius:10px;background:#fff}</style>"
        "<table><thead><tr><th>Код</th><th>Пункт</th><th>Факт</th><th>Используется</th><th>Статус</th><th>Найденные файлы</th><th>Действие</th></tr></thead>"
        f"<tbody>{checklist_rows_html}</tbody></table>"
        "</section>"
    )
    return _arm_simple_page(title="АРМ: чеклист", active_nav="checklist", body_html=body_html)


@router.get("/permit/height", response_class=HTMLResponse)
def arm_height_permit_html() -> HTMLResponse:
    root = resolve_object_root()
    metadata = _read_project_metadata(root)
    catalog = _build_employee_catalog(root=root)

    employees: list[dict[str, str]] = []
    for item in catalog.items:
        if not item.employee_name:
            continue
        team_name = "Без указания бригады"
        try:
            employee_root = _resolve_safe_path(root=root, rel_path=item.employee_rel_path)
            if employee_root.exists() and employee_root.is_dir():
                profile = _read_employee_profile(employee_root)
                team_name = (profile.get("team") or team_name).strip()
        except Exception:  # noqa: BLE001
            pass

        employees.append(
            {
                "name": item.employee_name,
                "position": item.position or item.profession_label,
                "rel_path": item.employee_rel_path,
                "team": team_name,
            }
        )
    employees.sort(key=lambda item: item["name"].lower())

    supervisors = [
        item for item in employees if "прораб" in (item["position"] or "").lower() or "руковод" in (item["position"] or "").lower()
    ]
    responsible_manager = supervisors[0]["name"] if supervisors else ""
    responsible_executor = supervisors[1]["name"] if len(supervisors) > 1 else responsible_manager

    source_docx_rel = ""
    permit_dir = root / "03_hse_and_fire_safety/permits/наряды_допуски"
    if permit_dir.exists():
        matches = sorted(permit_dir.glob("*НД*высот*.docx"))
        if matches:
            source_docx_rel = _to_rel_path(root, matches[0])

    ppr_context = ""
    ppr_context_path = root / PPR_CONTEXT_REL_PATH
    if ppr_context_path.exists() and ppr_context_path.is_file():
        ppr_context = ppr_context_path.read_text(encoding="utf-8", errors="ignore").strip()
    ppr_context = ppr_context[:3000]

    now = datetime.now()
    issue_date = now.strftime("%d.%m.%Y")
    valid_until = (now + timedelta(days=14)).strftime("%d.%m.%Y")
    default_output_rel = (
        f"03_hse_and_fire_safety/permits/наряды_допуски/{now.strftime('%Y%m%d')}_НД_на_высоте_autofill_v01.md"
    )

    employee_options_html = "".join(
        f'<label class="permit-employee-option"><input type="checkbox" class="permit-employee-checkbox" '
        f'data-name="{escape(item["name"])}" data-position="{escape(item["position"] or "")}" '
        f'data-team="{escape(item["team"] or "")}" checked />'
        f'<span>{escape(item["name"])}<small>{escape(item["position"] or "")} | {escape(item["team"] or "")}</small></span></label>'
        for item in employees
    )
    crew_row_count = 12
    crew_rows_html = "".join(
        "<tr>"
        "<td><input class=\"permit-input crew-name\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "</tr>"
        for _ in range(crew_row_count)
    )
    measure_rows_html = "".join(
        "<tr>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "</tr>"
        for _ in range(3)
    )
    daily_rows_html = "".join(
        "<tr>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "</tr>"
        for _ in range(5)
    )
    brigade_change_rows_html = "".join(
        "<tr>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "<td><input class=\"permit-input\" type=\"text\" /></td>"
        "</tr>"
        for _ in range(6)
    )

    source_actions_html = (
        _arm_file_actions_html(source_docx_rel, back_href="/arm/permit/height")
        if source_docx_rel
        else '<div class="meta">DOCX-образец не найден.</div>'
    )

    body_html = f"""
    <section class=\"card\">
        <h1>Наряд-допуск на производство работ на высоте</h1>
        <div class=\"meta\">Форма собрана по реальному DOCX-шаблону из проекта, а не по упрощенному markdown-пермиту.</div>
        <div class=\"controls\">
            <input id=\"permitOutputPath\" type=\"text\" value=\"{escape(default_output_rel)}\" style=\"min-width:420px;\" />
            <button type=\"button\" class=\"btn-inline primary\" id=\"permitSaveBtn\">Сохранить как Markdown</button>
            <button type=\"button\" class=\"btn-inline\" id=\"permitOpenDraftBtn\">Открыть черновик</button>
            <button type=\"button\" class=\"btn-inline\" id=\"permitPrintBtn\">Печать страницы</button>
            {_arm_action_anchor('/arm/checklist/view', 'К чеклисту')}
        </div>
        <div id=\"permitStatus\" class=\"meta\">Черновик еще не сохранен.</div>
    </section>
    <section class=\"permit-workspace\">
        <aside class=\"card permit-side\">
            <h2>Источники</h2>
            <div class=\"meta\">Объект: {escape(metadata.get('object_name') or '')}</div>
            <div class=\"meta\">Этап: {escape(metadata.get('work_stage') or '')}</div>
            <div class=\"meta\">Дата начала работ: {escape(metadata.get('start_date') or '')}</div>
            <h3>DOCX-образец</h3>
            {source_actions_html}
            <h3>ППР-контекст</h3>
            <textarea id=\"permitPprContext\" class=\"permit-side-text\">{escape(ppr_context)}</textarea>
            <h3>Состав бригады</h3>
            <select id=\"permitForemanSelect\" class=\"permit-side-select\"></select>
            <select id=\"permitTeamSelect\" class=\"permit-side-select\"></select>
            <div class=\"permit-employee-list\">{employee_options_html}</div>
            <button type=\"button\" class=\"btn-inline\" id=\"permitApplyForemanCrewBtn\">Подставить прораба и бригаду</button>
            <button type=\"button\" class=\"btn-inline\" id=\"permitApplyCrewBtn\">Перенести в таблицу</button>
        </aside>
        <section class=\"permit-sheet\" id=\"permitPrintSheet\">
            <div class=\"permit-title\">НАРЯД-ДОПУСК № <input id=\"permitNumber\" class=\"line-input short\" type=\"text\" value=\"12/СМР-2026\" /></div>
            <div class=\"permit-subtitle\">НА ПРОИЗВОДСТВО РАБОТ НА ВЫСОТЕ</div>

            <table class=\"permit-grid permit-head-grid\">
                <tr><td class=\"label\">Организация:</td><td><input id=\"permitOrganization\" class=\"line-input\" type=\"text\" value=\"ИП Исмагилов Вадим Шакирович\" /></td></tr>
                <tr><td class=\"label\">Подразделение:</td><td><input id=\"permitSubdivision\" class=\"line-input\" type=\"text\" value=\"{escape(metadata.get('work_stage') or 'Монтаж ЖБ колонн')}\" /></td></tr>
            </table>

            <table class=\"permit-grid permit-date-grid\">
                <tr>
                    <td class=\"label\">Выдан</td>
                    <td><input id=\"permitIssueDate\" class=\"line-input\" type=\"text\" value=\"{escape(issue_date)}\" /></td>
                    <td class=\"label\">Действителен до</td>
                    <td><input id=\"permitValidUntil\" class=\"line-input\" type=\"text\" value=\"{escape(valid_until)}\" /></td>
                </tr>
            </table>

            <table class=\"permit-grid permit-head-grid\">
                <tr><td class=\"label\">Ответственному руководителю работ:</td><td><input id=\"permitManager\" class=\"line-input\" type=\"text\" value=\"{escape(responsible_manager)}\" /></td></tr>
                <tr class=\"hint-row\"><td></td><td>(фамилия, инициалы)</td></tr>
                <tr><td class=\"label\">Ответственному исполнителю работ:</td><td><input id=\"permitExecutor\" class=\"line-input\" type=\"text\" value=\"{escape(responsible_executor)}\" /></td></tr>
                <tr class=\"hint-row\"><td></td><td>(фамилия, инициалы)</td></tr>
                <tr><td class=\"label\">На выполнение работ:</td><td><input id=\"permitWorkTitle\" class=\"line-input\" type=\"text\" value=\"Работы на высоте на объекте {escape(metadata.get('object_name') or '')}\" /></td></tr>
            </table>

            <div class=\"permit-section-title\">Состав исполнителей работ (члены бригады):</div>
            <table class=\"permit-grid\" id=\"permitCrewTable\">
                <thead>
                    <tr>
                        <th>Фамилия, имя, отчество (при наличии)</th>
                        <th>С условиями работ ознакомил, инструктаж провел (подпись)</th>
                        <th>С условиями работ ознакомлен (подпись)</th>
                    </tr>
                </thead>
                <tbody>{crew_rows_html}</tbody>
            </table>

            <div class=\"permit-number-title\">1. Необходимые для производства работ:</div>
            <textarea id=\"permitMaterials\" class=\"permit-area\">Анкерные линии, страховочные системы, каски, привязи, стропы, ограждения, инструмент.</textarea>

            <div class=\"permit-number-title\">2. До начала работ следует выполнить следующие мероприятия:</div>
            <textarea id=\"permitBeforeStart\" class=\"permit-area\">Проверить рабочие места, провести целевой инструктаж, проверить исправность СИЗ и систем страховки, оформить допуск.</textarea>

            <div class=\"permit-number-title\">3. В процессе производства работ необходимо выполнить следующие мероприятия:</div>
            <textarea id=\"permitDuringWork\" class=\"permit-area\">Соблюдать ППР, использовать системы обеспечения безопасности, не допускать работы без страховки и при изменении погодных условий.</textarea>

            <div class=\"permit-number-title\">4. Особые условия проведения работ:</div>
            <textarea id=\"permitSpecialConditions\" class=\"permit-area\">{escape((metadata.get('work_stage') or '') + '. ' + (ppr_context[:400] if ppr_context else 'Особые условия уточняются по ППР.'))}</textarea>

            <div class=\"permit-number-title\">5. Разрешение на подготовку рабочих мест и на допуск к выполнению работ</div>
            <div class=\"permit-line\">Рабочие места подготовлены. <input id=\"permitPreparedBy\" class=\"line-input\" type=\"text\" value=\"{escape(responsible_executor)}\" /></div>

            <div class=\"permit-number-title\">Место выполнения работ:</div>
            <input id=\"permitWorkPlace\" class=\"line-input\" type=\"text\" value=\"Монтаж в осях 1-30, А-Ф\" />

            <div class=\"permit-number-title\">Содержание работ:</div>
            <textarea id=\"permitContent\" class=\"permit-area\">Монтажные и сопутствующие работы на высоте в соответствии с ППР и технологическими картами.</textarea>

            <div class=\"permit-number-title\">Условия проведения работ:</div>
            <textarea id=\"permitConditions\" class=\"permit-area\">Работы выполнять при исправных системах страховки, с ограждением опасной зоны и под руководством ответственных лиц.</textarea>

            <div class=\"permit-number-title\">Опасные и вредные производственные факторы:</div>
            <textarea id=\"permitHazards\" class=\"permit-area\">Падение с высоты, падение предметов, ветровая нагрузка, неудовлетворительное состояние рабочих мест, поражение электрическим током при применении инструмента.</textarea>

            <table class=\"permit-grid permit-date-grid\">
                <tr>
                    <td class=\"label\">Начало работ</td>
                    <td><input id=\"permitStartWork\" class=\"line-input\" type=\"text\" value=\"{escape(issue_date)} 08:00\" /></td>
                    <td class=\"label\">Окончание работ</td>
                    <td><input id=\"permitEndWork\" class=\"line-input\" type=\"text\" value=\"{escape(valid_until)} 20:00\" /></td>
                </tr>
            </table>

            <table class=\"permit-grid\">
                <thead><tr><th>Системы обеспечения безопасности работ на высоте</th><th>Состав системы</th></tr></thead>
                <tbody>
                    <tr><td>Удерживающие системы</td><td><input id=\"permitSystemHold\" class=\"permit-input\" type=\"text\" value=\"Привязь, строп, анкерные точки\" /></td></tr>
                    <tr><td>Системы позиционирования</td><td><input id=\"permitSystemPosition\" class=\"permit-input\" type=\"text\" value=\"Строп позиционирования\" /></td></tr>
                    <tr><td>Страховочные системы</td><td><input id=\"permitSystemSafety\" class=\"permit-input\" type=\"text\" value=\"Страховочная привязь, вертикальная/горизонтальная линия\" /></td></tr>
                    <tr><td>Эвакуационные и спасательные системы</td><td><input id=\"permitSystemEvac\" class=\"permit-input\" type=\"text\" value=\"Спаскомплект, план эвакуации\" /></td></tr>
                </tbody>
            </table>

            <table class=\"permit-grid\">
                <tr><td class=\"label\">Материалы</td><td><input id=\"permitMaterialsList\" class=\"permit-input\" type=\"text\" value=\"Металлоконструкции, крепеж, расходные материалы\" /></td></tr>
                <tr><td class=\"label\">Инструменты</td><td><input id=\"permitTools\" class=\"permit-input\" type=\"text\" value=\"Монтажный инструмент, электроинструмент, измерительный инструмент\" /></td></tr>
                <tr><td class=\"label\">Приспособления</td><td><input id=\"permitAccessories\" class=\"permit-input\" type=\"text\" value=\"Лестницы, подмости, стропы, траверсы\" /></td></tr>
            </table>

            <div class=\"permit-number-title\">Мероприятия и сроки выполнения</div>
            <table class=\"permit-grid\" id=\"permitMeasuresTable\">
                <thead><tr><th>Наименование мероприятия или ссылки на пункт ППР или технологических карт</th><th>Срок выполнения</th><th>Ответственный исполнитель</th></tr></thead>
                <tbody>{measure_rows_html}</tbody>
            </table>

            <div class=\"permit-number-title\">Мероприятия по безопасности работ на высоте</div>
            <table class=\"permit-grid\" id=\"permitSafetyMeasuresTable\">
                <thead><tr><th>Наименование мероприятия по безопасности работ на высоте</th><th>Срок выполнения</th><th>Ответственный исполнитель</th></tr></thead>
                <tbody>{measure_rows_html}</tbody>
            </table>

            <div class=\"permit-number-title\">Наименование условий</div>
            <table class=\"permit-grid\" id=\"permitConditionTable\">
                <thead><tr><th>Наименование условий</th><th>Срок выполнения</th><th>Ответственный исполнитель</th></tr></thead>
                <tbody>{measure_rows_html}</tbody>
            </table>

            <div class=\"permit-number-title\">6. Ежедневный допуск к работе и время ее окончания</div>
            <table class=\"permit-grid\" id=\"permitDailyTable\">
                <thead><tr><th>Наименование рабочего места</th><th>Дата, время</th><th>Ответственный руководитель работ</th><th>Ответственный исполнитель работ</th><th>Окончание, дата/время</th><th>Подпись ответственного исполнителя</th></tr></thead>
                <tbody>{daily_rows_html}</tbody>
            </table>

            <div class=\"permit-number-title\">7. Изменения в составе бригады</div>
            <table class=\"permit-grid\" id=\"permitBrigadeChangesTable\">
                <thead><tr><th>Введен в состав бригады</th><th>Выведен из состава бригады</th><th>Дата, время</th><th>Разрешил</th></tr></thead>
                <tbody>{brigade_change_rows_html}</tbody>
            </table>

            <div class=\"permit-number-title\">8. Регистрация целевого инструктажа при первичном допуске</div>
            <table class=\"permit-grid\">
                <tr><td class=\"label\">Лицо, выдавшее наряд</td><td><input id=\"permitIssuer\" class=\"permit-input\" type=\"text\" value=\"Исмагилов Вадим Шакирович\" /></td><td class=\"label\">Ответственный руководитель работ</td><td><input id=\"permitInstructionManager\" class=\"permit-input\" type=\"text\" value=\"{escape(responsible_manager)}\" /></td></tr>
                <tr><td class=\"label\">Ответственный исполнитель работ</td><td><input id=\"permitInstructionExecutor\" class=\"permit-input\" type=\"text\" value=\"{escape(responsible_executor)}\" /></td><td class=\"label\">Члены бригады</td><td><input id=\"permitInstructionCrew\" class=\"permit-input\" type=\"text\" value=\"\" /></td></tr>
            </table>

            <div class=\"permit-number-title\">9. Письменное разрешение (акт-допуск) действующего предприятия</div>
            <textarea id=\"permitActAdmission\" class=\"permit-area\">Не требуется, работы ведутся по условиям объекта и согласованным мероприятиям безопасности.</textarea>

            <div class=\"permit-number-title\">10. Рабочее место и условия труда проверены</div>
            <textarea id=\"permitChecked\" class=\"permit-area\">Мероприятия по безопасности, указанные в наряде-допуске, выполнены. Рабочее место проверено ответственными лицами.</textarea>

            <div class=\"permit-number-title\">11. Работа выполнена в полном объеме</div>
            <textarea id=\"permitClosed\" class=\"permit-area\">Материалы, инструмент и приспособления убраны. Члены бригады выведены. Наряд-допуск закрыт.</textarea>
        </section>
    </section>
    <style>
    .permit-workspace {{ display:grid; grid-template-columns: 320px minmax(0, 1fr); gap:16px; margin-top:16px; align-items:start; }}
    .permit-side {{ position: sticky; top: 62px; max-height: calc(100vh - 80px); overflow:auto; }}
    .permit-side h2, .permit-side h3 {{ margin: 0 0 10px; }}
    .permit-side h3 {{ margin-top: 16px; font-size: 14px; }}
    .permit-side-select {{ width:100%; box-sizing:border-box; border:1px solid #ccd4d7; border-radius:10px; padding:8px 10px; margin-bottom:8px; background:#fff; }}
    .permit-side-text {{ width:100%; min-height:180px; box-sizing:border-box; border:1px solid #ccd4d7; border-radius:10px; padding:10px; resize:vertical; font:13px/1.45 Consolas, 'Courier New', monospace; }}
    .permit-employee-list {{ display:grid; gap:8px; margin:12px 0; }}
    .permit-employee-option {{ display:flex; gap:8px; align-items:flex-start; padding:8px; border:1px solid #d8e0e0; border-radius:10px; background:#fff; }}
    .permit-employee-option small {{ display:block; color:#5b6672; font-size:12px; margin-top:2px; }}
    .permit-sheet {{ background:#fff; color:#222; border:1px solid #d7d7d7; box-shadow:0 12px 30px rgba(0,0,0,0.08); padding:28px 28px 40px; }}
    .permit-title, .permit-subtitle {{ text-align:center; font-weight:700; }}
    .permit-title {{ font-size:24px; letter-spacing:0.02em; margin-bottom:6px; }}
    .permit-subtitle {{ font-size:20px; margin-bottom:18px; }}
    .permit-grid {{ width:100%; border-collapse:collapse; margin:10px 0 16px; }}
    .permit-grid th, .permit-grid td {{ border:1px solid #222; padding:6px 8px; vertical-align:top; font-size:14px; background:#fff; }}
    .permit-grid th {{ text-align:center; font-weight:700; }}
    .permit-grid .label {{ width:28%; font-weight:600; }}
    .permit-head-grid .hint-row td {{ border-top:none; text-align:center; color:#555; font-size:12px; }}
    .permit-date-grid td {{ width:25%; }}
    .permit-number-title {{ font-weight:700; margin:14px 0 8px; }}
    .permit-line {{ margin:8px 0 14px; }}
    .line-input {{ width:100%; border:none; border-bottom:1px solid #222; padding:4px 2px; background:transparent; font-size:14px; box-sizing:border-box; }}
    .line-input.short {{ display:inline-block; width:180px; }}
    .permit-input {{ width:100%; border:none; background:transparent; font-size:14px; box-sizing:border-box; padding:4px 2px; min-height:24px; }}
    .permit-area {{ width:100%; min-height:74px; box-sizing:border-box; border:1px solid #222; padding:8px; font-size:14px; resize:vertical; }}
    @media print {{ body {{ background:#fff; }} .site-nav, .card:not(.permit-side), .permit-side {{ display:none !important; }} .wrap {{ margin:0; max-width:none; padding:0; }} .permit-workspace {{ display:block; }} .permit-sheet {{ box-shadow:none; border:none; padding:0; }} }}
    @media (max-width: 1080px) {{ .permit-workspace {{ grid-template-columns: 1fr; }} .permit-side {{ position: static; max-height:none; }} }}
    </style>
    <script>
    const permitEmployees = {json.dumps(employees, ensure_ascii=False)};
    const permitStatus = document.getElementById('permitStatus');
    const permitOutputPath = document.getElementById('permitOutputPath');
    const permitCrewTable = document.getElementById('permitCrewTable').querySelector('tbody');
    const permitForemanSelect = document.getElementById('permitForemanSelect');
    const permitTeamSelect = document.getElementById('permitTeamSelect');

    function permitIsForeman(employee) {{
        const position = String((employee && employee.position) || '').toLowerCase();
        return position.includes('прораб') || position.includes('руковод') || position.includes('мастер');
    }}

    function permitUniqueTeams() {{
        const values = new Set();
        for (const employee of permitEmployees) {{
            const team = String((employee && employee.team) || '').trim();
            if (team) {{
                values.add(team);
            }}
        }}
        return Array.from(values).sort((a, b) => a.localeCompare(b, 'ru'));
    }}

    function permitCheckedByTeam(teamName) {{
        const normalized = String(teamName || '').trim();
        const checked = [];
        const boxes = Array.from(document.querySelectorAll('.permit-employee-checkbox'));
        for (const box of boxes) {{
            const boxTeam = String(box.dataset.team || '').trim();
            const isMatch = normalized && boxTeam === normalized;
            box.checked = isMatch;
            if (isMatch) {{
                checked.push({{
                    name: box.dataset.name || '',
                    position: box.dataset.position || '',
                    team: boxTeam,
                }});
            }}
        }}
        return checked;
    }}

    function permitApplyForemanData() {{
        if (!permitForemanSelect) {{
            return;
        }}
        const foremanName = permitForemanSelect.value || '';
        if (!foremanName) {{
            return;
        }}
        const manager = document.getElementById('permitManager');
        const executor = document.getElementById('permitExecutor');
        const preparedBy = document.getElementById('permitPreparedBy');
        const instructionManager = document.getElementById('permitInstructionManager');
        const instructionExecutor = document.getElementById('permitInstructionExecutor');
        if (manager) {{ manager.value = foremanName; }}
        if (executor) {{ executor.value = foremanName; }}
        if (preparedBy) {{ preparedBy.value = foremanName; }}
        if (instructionManager) {{ instructionManager.value = foremanName; }}
        if (instructionExecutor) {{ instructionExecutor.value = foremanName; }}
    }}

    function permitInitSelectors() {{
        if (permitForemanSelect) {{
            permitForemanSelect.innerHTML = '';
            const foremen = permitEmployees.filter((item) => permitIsForeman(item));
            const placeholder = document.createElement('option');
            placeholder.value = '';
            placeholder.textContent = 'Выберите прораба';
            permitForemanSelect.appendChild(placeholder);
            for (const foreman of foremen) {{
                const option = document.createElement('option');
                option.value = foreman.name || '';
                option.textContent = (foreman.name || '') + (foreman.team ? (' | ' + foreman.team) : '');
                permitForemanSelect.appendChild(option);
            }}
            if (foremen.length) {{
                permitForemanSelect.value = foremen[0].name || '';
            }}
        }}

        if (permitTeamSelect) {{
            permitTeamSelect.innerHTML = '';
            const placeholder = document.createElement('option');
            placeholder.value = '';
            placeholder.textContent = 'Выберите бригаду';
            permitTeamSelect.appendChild(placeholder);
            for (const team of permitUniqueTeams()) {{
                const option = document.createElement('option');
                option.value = team;
                option.textContent = team;
                permitTeamSelect.appendChild(option);
            }}
            if (permitTeamSelect.options.length > 1) {{
                permitTeamSelect.selectedIndex = 1;
            }}
        }}
    }}

    function permitSelectedEmployees() {{
        return Array.from(document.querySelectorAll('.permit-employee-checkbox:checked')).map((node) => ({{
            name: node.dataset.name || '',
            position: node.dataset.position || '',
            team: node.dataset.team || ''
        }}));
    }}

    function permitApplyCrew(selectedCrew) {{
        const selected = Array.isArray(selectedCrew) && selectedCrew.length ? selectedCrew : permitSelectedEmployees();
        const rows = Array.from(permitCrewTable.querySelectorAll('tr'));
        rows.forEach((row, index) => {{
            const inputs = row.querySelectorAll('input');
            const employee = selected[index];
            inputs[0].value = employee ? employee.name + (employee.position ? ' (' + employee.position + ')' : '') : '';
            if (!inputs[1].value) {{ inputs[1].value = employee ? 'подпись' : ''; }}
            if (!inputs[2].value) {{ inputs[2].value = employee ? 'подпись' : ''; }}
        }});
        const crewNames = selected.map((item) => item.name).join(', ');
        const instructionCrew = document.getElementById('permitInstructionCrew');
        if (instructionCrew && !instructionCrew.value) {{
            instructionCrew.value = crewNames;
        }}
    }}

    function permitTableToMarkdown(tableId) {{
        const rows = Array.from(document.querySelectorAll('#' + tableId + ' tbody tr'));
        const lines = [];
        for (const row of rows) {{
            const values = Array.from(row.querySelectorAll('input, textarea')).map((node) => (node.value || '').trim()).filter(Boolean);
            if (values.length) {{
                lines.push('- ' + values.join(' | '));
            }}
        }}
        return lines.join('\n');
    }}

    function permitCrewMarkdown() {{
        const rows = Array.from(permitCrewTable.querySelectorAll('tr'));
        const lines = [];
        for (const row of rows) {{
            const name = row.querySelector('.crew-name').value.trim();
            if (name) {{
                lines.push('| ' + name + ' | подпись | подпись |');
            }}
        }}
        return lines.length ? ['| ФИО | Инструктаж провел | Ознакомлен |', '|---|---|---|'].concat(lines).join('\n') : '';
    }}

    function permitMarkdown() {{
        return [
            '# Наряд-допуск на производство работ на высоте',
            '',
            '- Номер: ' + document.getElementById('permitNumber').value.trim(),
            '- Организация: ' + document.getElementById('permitOrganization').value.trim(),
            '- Подразделение: ' + document.getElementById('permitSubdivision').value.trim(),
            '- Выдан: ' + document.getElementById('permitIssueDate').value.trim(),
            '- Действителен до: ' + document.getElementById('permitValidUntil').value.trim(),
            '- Ответственный руководитель работ: ' + document.getElementById('permitManager').value.trim(),
            '- Ответственный исполнитель работ: ' + document.getElementById('permitExecutor').value.trim(),
            '- На выполнение работ: ' + document.getElementById('permitWorkTitle').value.trim(),
            '',
            '## Состав исполнителей работ',
            permitCrewMarkdown(),
            '',
            '## 1. Необходимые для производства работ',
            document.getElementById('permitMaterials').value.trim(),
            '',
            '## 2. До начала работ',
            document.getElementById('permitBeforeStart').value.trim(),
            '',
            '## 3. В процессе производства работ',
            document.getElementById('permitDuringWork').value.trim(),
            '',
            '## 4. Особые условия проведения работ',
            document.getElementById('permitSpecialConditions').value.trim(),
            '',
            '## 5. Подготовка рабочих мест и допуск',
            '- Рабочие места подготовлены: ' + document.getElementById('permitPreparedBy').value.trim(),
            '- Место выполнения работ: ' + document.getElementById('permitWorkPlace').value.trim(),
            '- Содержание работ: ' + document.getElementById('permitContent').value.trim(),
            '- Условия проведения работ: ' + document.getElementById('permitConditions').value.trim(),
            '- Опасные и вредные факторы: ' + document.getElementById('permitHazards').value.trim(),
            '- Начало работ: ' + document.getElementById('permitStartWork').value.trim(),
            '- Окончание работ: ' + document.getElementById('permitEndWork').value.trim(),
            '',
            '## Системы обеспечения безопасности',
            '- Удерживающие системы: ' + document.getElementById('permitSystemHold').value.trim(),
            '- Системы позиционирования: ' + document.getElementById('permitSystemPosition').value.trim(),
            '- Страховочные системы: ' + document.getElementById('permitSystemSafety').value.trim(),
            '- Эвакуационные и спасательные системы: ' + document.getElementById('permitSystemEvac').value.trim(),
            '- Материалы: ' + document.getElementById('permitMaterialsList').value.trim(),
            '- Инструменты: ' + document.getElementById('permitTools').value.trim(),
            '- Приспособления: ' + document.getElementById('permitAccessories').value.trim(),
            '',
            '## Мероприятия и сроки выполнения',
            permitTableToMarkdown('permitMeasuresTable'),
            '',
            '## Мероприятия по безопасности работ на высоте',
            permitTableToMarkdown('permitSafetyMeasuresTable'),
            '',
            '## Наименование условий',
            permitTableToMarkdown('permitConditionTable'),
            '',
            '## 6. Ежедневный допуск к работе и время ее окончания',
            permitTableToMarkdown('permitDailyTable'),
            '',
            '## 7. Изменения в составе бригады',
            permitTableToMarkdown('permitBrigadeChangesTable'),
            '',
            '## 8. Регистрация целевого инструктажа',
            '- Лицо, выдавшее наряд: ' + document.getElementById('permitIssuer').value.trim(),
            '- Ответственный руководитель работ: ' + document.getElementById('permitInstructionManager').value.trim(),
            '- Ответственный исполнитель работ: ' + document.getElementById('permitInstructionExecutor').value.trim(),
            '- Члены бригады: ' + document.getElementById('permitInstructionCrew').value.trim(),
            '',
            '## 9. Письменное разрешение (акт-допуск)',
            document.getElementById('permitActAdmission').value.trim(),
            '',
            '## 10. Проверка рабочего места и условий труда',
            document.getElementById('permitChecked').value.trim(),
            '',
            '## 11. Закрытие наряда',
            document.getElementById('permitClosed').value.trim(),
            ''
        ].join('\n');
    }}

    async function permitSave() {{
        const relPath = permitOutputPath.value.trim();
        if (!relPath) {{
            throw new Error('Не указан путь для сохранения');
        }}
        const response = await fetch('/arm/fs/file', {{
            method: 'POST',
            headers: {{ 'Content-Type': 'application/json' }},
            body: JSON.stringify({{ rel_path: relPath, content: permitMarkdown() }})
        }});
        if (!response.ok) {{
            const text = await response.text();
            throw new Error(text || ('HTTP ' + response.status));
        }}
        const data = await response.json();
        permitStatus.textContent = data.message || ('Сохранено: ' + relPath);
    }}

    document.getElementById('permitApplyCrewBtn').addEventListener('click', () => permitApplyCrew());
    document.getElementById('permitApplyForemanCrewBtn').addEventListener('click', () => {{
        const selectedTeam = permitTeamSelect ? permitTeamSelect.value : '';
        if (selectedTeam) {{
            const crew = permitCheckedByTeam(selectedTeam);
            permitApplyCrew(crew);
        }} else {{
            permitApplyCrew();
        }}
        permitApplyForemanData();
    }});
    if (permitForemanSelect) {{
        permitForemanSelect.addEventListener('change', permitApplyForemanData);
    }}
    if (permitTeamSelect) {{
        permitTeamSelect.addEventListener('change', () => {{
            const selectedTeam = permitTeamSelect.value || '';
            if (!selectedTeam) {{
                return;
            }}
            const crew = permitCheckedByTeam(selectedTeam);
            permitApplyCrew(crew);
        }});
    }}
    document.getElementById('permitSaveBtn').addEventListener('click', () => {{
        permitStatus.textContent = 'Сохранение...';
        permitSave().catch((err) => {{
            permitStatus.textContent = 'Ошибка сохранения: ' + err.message;
        }});
    }});
    document.getElementById('permitOpenDraftBtn').addEventListener('click', () => {{
        const relPath = permitOutputPath.value.trim();
        if (!relPath) {{
            permitStatus.textContent = 'Сначала укажите путь черновика.';
            return;
        }}
        window.location.href = '/arm/editor?rel_path=' + encodeURIComponent(relPath) + '&back=' + encodeURIComponent('/arm/permit/height');
    }});
    document.getElementById('permitPrintBtn').addEventListener('click', () => window.print());
    permitInitSelectors();
    permitApplyForemanData();
    if (permitTeamSelect && permitTeamSelect.value) {{
        permitApplyCrew(permitCheckedByTeam(permitTeamSelect.value));
    }} else {{
        permitApplyCrew();
    }}
    </script>
    """
    return _arm_simple_page(title="АРМ: наряд на высоте", active_nav="permit-height", body_html=body_html)


@router.get("/aosr", response_class=HTMLResponse)
def arm_aosr_html() -> HTMLResponse:
    root = resolve_object_root()
    incoming = root / "00_incoming_requests"
    requested_files = [
        incoming / "АОСР(монт колонн).xlsx",
        incoming / "Реестр к акту ( монтаж колонн).xlsx",
    ]

    file_cards: list[str] = []
    for file_path in requested_files:
        if not file_path.exists() or not file_path.is_file():
            file_cards.append(
                "<section class='card'><h2>"
                + escape(file_path.name)
                + "</h2><div class='meta warn'>Файл не найден в 00_incoming_requests.</div></section>"
            )
            continue

        rel = _to_rel_path(root, file_path)
        analysis_error = ""
        try:
            analysis = _analyze_xlsx_brief(file_path)
        except HTTPException as exc:
            analysis = {"sheets": []}
            analysis_error = str(exc.detail or "Не удалось проанализировать XLSX")
        rows_html = "".join(
            "<tr>"
            f"<td>{escape(str(sheet.get('title') or 'Лист'))}</td>"
            f"<td>{escape(str(sheet.get('rows') or 0))}</td>"
            f"<td>{escape(str(sheet.get('cols') or 0))}</td>"
            f"<td>{escape(', '.join(sheet.get('headers') or [])) or '—'}</td>"
            "</tr>"
            for sheet in analysis.get("sheets", [])
        )
        if not rows_html and not analysis_error:
            rows_html = "<tr><td colspan='4'>Листы не обнаружены.</td></tr>"

        if analysis_error:
            rows_html = "<tr><td colspan='4'>" + escape(analysis_error) + "</td></tr>"

        file_cards.append(
            "<section class='card'>"
            f"<h2>{escape(file_path.name)}</h2>"
            f"<div class='meta'>Путь: {escape(rel)}</div>"
            f"<div class='action-links' style='margin-top:8px'>{_arm_file_actions_html(rel, back_href='/arm/aosr')}</div>"
            "<div class='meta' style='margin-top:10px'>Аналитика XLSX</div>"
            "<table><thead><tr><th>Лист</th><th>Строк</th><th>Колонок</th><th>Первые заголовки</th></tr></thead>"
            f"<tbody>{rows_html}</tbody></table>"
            "</section>"
        )

    body_html = (
        "<section class='card'>"
        "<h1>АОСР: реестр паспортов и входящие XLSX</h1>"
        "<div class='meta'>Раздел аналогичен «Наряд высота»: используйте готовые действия файла, "
        "предпросмотр и печать на базе шаблонов проекта.</div>"
        "<div class='controls'>"
        + _arm_action_anchor('/arm/permit/height', 'Открыть Наряд высота')
        + _arm_action_anchor('/arm/structure/view', 'Открыть Структуру и действия')
        + "</div>"
        "</section>"
        + "".join(file_cards)
    )
    return _arm_simple_page(title="АРМ: АОСР", active_nav="aosr", body_html=body_html)


@router.get("/todo/view", response_class=HTMLResponse)
def arm_todo_html(db: Session = Depends(get_db)) -> HTMLResponse:
    payload = _build_dashboard_payload(db=db)
    items = _build_todos(
        checklist=payload.checklist,
        metrics=payload.metrics,
        local_llm_reachable=payload.local_llm_reachable,
    )

    rows: list[str] = []
    for item in items:
        action_html = "-"
        if item.action_path:
            action_html = (
                f"<a class=\"btn-link\" href=\"/arm/dashboard?open_path={quote(item.action_path, safe='')}\">"
                "Открыть в дереве</a>"
            )

        rows.append(
            "<tr>"
            f"<td>{escape(item.priority)}</td>"
            f"<td>{escape(item.title)}</td>"
            f"<td>{escape(item.details or '')}</td>"
            f"<td>{action_html}</td>"
            "</tr>"
        )

    todo_rows_html = "".join(rows) if rows else "<tr><td colspan=\"4\">Задач нет</td></tr>"

    body_html = (
        "<section class=\"card\">"
        "<h1>План на день</h1>"
        f"<div class=\"meta\">Сформировано: {escape(payload.generated_at.isoformat())}</div>"
        "<table><thead><tr><th>Приоритет</th><th>Задача</th><th>Подсказка</th><th>Действие</th></tr></thead>"
        f"<tbody>{todo_rows_html}</tbody></table>"
        "</section>"
    )
    return _arm_simple_page(title="АРМ: план дня", active_nav="todo", body_html=body_html)


@router.get("/periodic/view", response_class=HTMLResponse)
def arm_periodic_html() -> HTMLResponse:
    root = resolve_object_root()
    rows_data = _collect_periodic_doc_statuses(root=root)
    metadata = _read_project_metadata(root)

    rows: list[str] = []
    for row in rows_data:
        status_html = "<span class=\"warn\">Требует обновления</span>" if row.is_due else "<span class=\"ok\">Актуально</span>"
        if row.latest_file:
            latest_html = escape(row.latest_file)
        else:
            latest_html = "—"
        if row.days_since_update is None:
            age_html = "—"
        else:
            age_html = f"{row.days_since_update} дн"

        if row.latest_file:
            action_html = _arm_file_actions_html(
                row.latest_file,
                folder_rel_path=row.action_path,
                back_href="/arm/periodic/view",
            )
        else:
            action_html = f'<div class="action-links"><a class="btn-inline" href="/arm/dashboard?open_path={quote(row.action_path, safe="")}">Открыть папку</a></div>'
        rows.append(
            "<tr>"
            f"<td>{escape(row.rule.code)}</td>"
            f"<td>{escape(row.rule.title)}</td>"
            f"<td>{escape(str(row.rule.period_days))} дн</td>"
            f"<td>{age_html}</td>"
            f"<td>{status_html}</td>"
            f"<td>{latest_html}</td>"
            f"<td>{action_html}</td>"
            "</tr>"
        )

    periodic_rows_html = "".join(rows) if rows else "<tr><td colspan=\"7\">Периодические документы не настроены</td></tr>"
    body_html = (
        "<section class=\"card\">"
        "<h1>Карточка объекта</h1>"
        "<div class=\"meta\">Дата начала работ влияет на расчет периодичности. Здесь можно править карточку объекта без перехода на дашборд.</div>"
        "<div class=\"controls\">"
        f"<input id=\"periodicObjectName\" type=\"text\" value=\"{escape(metadata.get('object_name') or '')}\" placeholder=\"Объект\" style=\"min-width:280px;\" />"
        f"<input id=\"periodicWorkStage\" type=\"text\" value=\"{escape(metadata.get('work_stage') or '')}\" placeholder=\"Этап работ\" style=\"min-width:220px;\" />"
        f"<input id=\"periodicStartDate\" type=\"text\" value=\"{escape(metadata.get('start_date') or '')}\" placeholder=\"Дата начала работ, ДД.ММ.ГГГГ\" />"
        '<button type="button" class="btn-inline primary" id="periodicSaveProfileBtn">Сохранить карточку объекта</button>'
        "</div>"
        '<div id="periodicProfileStatus" class="meta"></div>'
        "</section>"
        "<section class=\"card\">"
        "<h1>Периодические документы</h1>"
        "<div class=\"meta\">Контур для регламентных документов с дедлайнами обновления. Для каждого файла доступны просмотр, правка, печать и скачивание.</div>"
        "<table><thead><tr><th>Код</th><th>Документ</th><th>Период</th><th>Давность</th><th>Статус</th><th>Последний файл</th><th>Действие</th></tr></thead>"
        f"<tbody>{periodic_rows_html}</tbody></table>"
        "</section>"
        "<script>"
        "const periodicSaveProfileBtn = document.getElementById('periodicSaveProfileBtn');"
        "const periodicProfileStatus = document.getElementById('periodicProfileStatus');"
        "async function periodicSaveProfile(){"
        "periodicSaveProfileBtn.disabled = true;"
        "periodicProfileStatus.textContent = 'Сохранение карточки объекта...';"
        "try {"
        "const response = await fetch('/arm/object-profile', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({object_name: document.getElementById('periodicObjectName').value, work_stage: document.getElementById('periodicWorkStage').value, start_date: document.getElementById('periodicStartDate').value})});"
        "if (!response.ok) { const text = await response.text(); throw new Error(text || ('HTTP ' + response.status)); }"
        "const data = await response.json();"
        "periodicProfileStatus.textContent = data.message || 'Карточка объекта сохранена. Обновите страницу, чтобы пересчитать сроки.';"
        "} finally { periodicSaveProfileBtn.disabled = false; }"
        "}"
        "periodicSaveProfileBtn.addEventListener('click', () => periodicSaveProfile().catch((err) => { periodicProfileStatus.textContent = 'Ошибка: ' + err.message; }));"
        "</script>"
    )
    return _arm_simple_page(title="АРМ: периодические документы", active_nav="periodic", body_html=body_html)


@router.get("/dashboard", response_class=HTMLResponse)
def arm_dashboard_html(db: Session = Depends(get_db)) -> HTMLResponse:
    payload = _build_dashboard_payload(db=db)
    todos = _build_todos(
        checklist=payload.checklist,
        metrics=payload.metrics,
        local_llm_reachable=payload.local_llm_reachable,
    )

    metric_rows = [
        ("Документы в БД", payload.metrics.db_documents_total, "шт"),
        ("Записи журнала", payload.metrics.db_journal_entries_total, "шт"),
        ("Графики работ", payload.metrics.db_schedules_total, "шт"),
        ("Приказы в Markdown", payload.metrics.orders_md_total, "шт"),
        ("Приказы в PDF (готово)", payload.metrics.orders_pdf_ready_total, "шт"),
        ("Журналы производства", payload.metrics.journals_production_total, "шт"),
        ("Журналы ОТ/ПБ", payload.metrics.journals_labor_safety_total, "шт"),
        ("Сканы во входящей папке", payload.metrics.scan_inbox_pending_total, "шт"),
        ("Сканы на ручной разбор", payload.metrics.scan_manual_review_total, "шт"),
    ]
    metrics_html = "".join(
        f"<li><b>{escape(name)}:</b> {escape(str(value))} {escape(unit)}</li>"
        for name, value, unit in metric_rows
    )

    root = resolve_object_root()
    gap_items = [item for item in payload.checklist if not item.ready][:6]
    gap_lines: list[str] = []
    for item in gap_items:
        try:
            action_path = _to_rel_path(root, Path(item.location).resolve())
        except Exception:  # noqa: BLE001
            action_path = None
        title_html = f"{escape(item.code)}. {escape(item.title)}"
        if action_path:
            title_html = (
                f"<a href=\"#\" class=\"todo-link gap-link\" data-action-path=\"{escape(action_path)}\">"
                f"{title_html}</a>"
            )

        fix_hint = escape(_build_gap_fix_hint(item))
        gap_lines.append(
            f"<li>{title_html}<div class=\"system-hint\">Как исправить: {fix_hint}. "
            f"Текущее состояние: {item.found}/{item.required_min}.</div></li>"
        )
    gaps_html = "".join(gap_lines) or "<li>Критичных пробелов не найдено</li>"

    todo_items_html: list[str] = []
    for item in todos:
        title_html = escape(item.title)
        if item.action_path:
            title_html = (
                f"<a href=\"#\" class=\"todo-link\" data-action-path=\"{escape(item.action_path)}\">"
                f"{title_html}</a>"
            )
        row = f"<li>[{escape(item.priority)}] {title_html}"
        if item.details:
            row += f"<div class=\"system-hint\">Подсказка: {escape(item.details)}</div>"
        row += "</li>"
        todo_items_html.append(row)
    todo_html = "".join(todo_items_html) or "<li>Список задач пуст</li>"

    llm_badge_class = "llm-ok" if payload.local_llm_reachable else "llm-down"
    llm_badge_text = "ДОСТУПНА" if payload.local_llm_reachable else "НЕДОСТУПНА"

    script_html = r"""
<script>
    const TIMEOUTS_MS = {
        default: 20000,
        status: 12000,
        tree: 45000,
        llm: 240000,
        scanner: 240000,
        ingest: 300000,
        checklist: 45000,
        checklistGenerate: 60000
    };

    const sendBtn = document.getElementById('armSend');
    const question = document.getElementById('armQuestion');
    const profile = document.getElementById('armProfile');
    const answer = document.getElementById('armAnswer');
    const sendOnEnterCheckbox = document.getElementById('armSendOnEnter');
    const voiceBtn = document.getElementById('armVoiceBtn');
    const voiceHint = document.getElementById('armVoiceHint');

    const treePathInput = document.getElementById('treePath');
    const treeOpenBtn = document.getElementById('treeOpenBtn');
    const treeBackBtn = document.getElementById('treeBackBtn');
    const treeForwardBtn = document.getElementById('treeForwardBtn');
    const treeUpBtn = document.getElementById('treeUpBtn');
    const treeRootBtn = document.getElementById('treeRootBtn');
    const treeBreadcrumb = document.getElementById('treeBreadcrumb');
    const treeView = document.getElementById('treeView');
    const treeHint = document.getElementById('treeHint');
    const structureCard = document.getElementById('structureCard');
    const interactionHint = document.getElementById('interactionHint');
    const actionNavigatorPath = document.getElementById('actionNavigatorPath');
    const actionNavigatorSteps = document.getElementById('actionNavigatorSteps');
    const taskActionLabel = document.getElementById('taskActionLabel');
    const taskActionPath = document.getElementById('taskActionPath');
    const taskActionOpenChecklistBtn = document.getElementById('taskActionOpenChecklistBtn');
    const taskActionRunMissingBtn = document.getElementById('taskActionRunMissingBtn');
    const taskActionRunSelectedBtn = document.getElementById('taskActionRunSelectedBtn');

    const filePathLabel = document.getElementById('filePath');
    const previewCard = document.getElementById('previewCard');
    const fileRenderWrap = document.getElementById('fileRenderWrap');
    const fileRenderFrame = document.getElementById('fileRenderFrame');
    const fileEditor = document.getElementById('fileEditor');
    const fileSourceToggleBtn = document.getElementById('fileSourceToggleBtn');
    const fileMeta = document.getElementById('fileMeta');
    const fileSaveBtn = document.getElementById('fileSaveBtn');
    const fileDownloadBtn = document.getElementById('fileDownloadBtn');
    const filePrintBtn = document.getElementById('filePrintBtn');
    const fileZoomOutBtn = document.getElementById('fileZoomOutBtn');
    const fileZoomResetBtn = document.getElementById('fileZoomResetBtn');
    const fileZoomInBtn = document.getElementById('fileZoomInBtn');
    const fileMoveBtn = document.getElementById('fileMoveBtn');
    const fileDeleteBtn = document.getElementById('fileDeleteBtn');

    const scannerDevices = document.getElementById('scannerDevices');
    const scannerDocType = document.getElementById('scannerDocType');
    const scannerProfile = document.getElementById('scannerProfile');
    const scannerProfileLabel = document.getElementById('scannerProfileLabel');
    const scannerSortMode = document.getElementById('scannerSortMode');
    const scannerSubject = document.getElementById('scannerSubject');
    const scannerEmployee = document.getElementById('scannerEmployee');
    const scannerEmployeeSelect = document.getElementById('scannerEmployeeSelect');
    const scannerEmployeeSuggestions = document.getElementById('scannerEmployeeSuggestions');
    const scannerEmployeeHint = document.getElementById('scannerEmployeeHint');
    const scannerHint = document.getElementById('scannerHint');
    const scannerProgressBar = document.getElementById('scannerProgressBar');
    const scannerProgressText = document.getElementById('scannerProgressText');
    const scannerTimeline = document.getElementById('scannerTimeline');
    const scannerMsg = document.getElementById('scannerMsg');
    const scanBtn = document.getElementById('scanBtn');
    const ingestBtn = document.getElementById('ingestBtn');
    const manualClassifyBtn = document.getElementById('manualClassifyBtn');
    const recompressScansBtn = document.getElementById('recompressScansBtn');
    const manualReview = document.getElementById('manualReview');
    const manualReviewList = document.getElementById('manualReviewList');
    const manualReviewSelection = document.getElementById('manualReviewSelection');
    const manualReviewTargetType = document.getElementById('manualReviewTargetType');
    const manualReviewSuggestBtn = document.getElementById('manualReviewSuggestBtn');
    const manualReviewMkDirBtn = document.getElementById('manualReviewMkDirBtn');
    const manualReviewMovePath = document.getElementById('manualReviewMovePath');
    const manualReviewOpenBtn = document.getElementById('manualReviewOpenBtn');
    const manualReviewDownloadBtn = document.getElementById('manualReviewDownloadBtn');
    const manualReviewPrintBtn = document.getElementById('manualReviewPrintBtn');
    const manualReviewEditOcrBtn = document.getElementById('manualReviewEditOcrBtn');
    const manualReviewMoveBtn = document.getElementById('manualReviewMoveBtn');
    const manualReviewDeleteBtn = document.getElementById('manualReviewDeleteBtn');
    const maintenanceResetBtn = document.getElementById('maintenanceResetBtn');

    const employeeChecklistEmployeeSelect = document.getElementById('employeeChecklistEmployee');
    const employeeChecklistCard = document.getElementById('employeeChecklistCard');
    const employeeChecklistBody = document.getElementById('employeeChecklistBody');
    const employeeChecklistToggleBtn = document.getElementById('employeeChecklistToggleBtn');
    const employeeChecklistPath = document.getElementById('employeeChecklistPath');
    const employeeChecklistOrderDate = document.getElementById('employeeChecklistOrderDate');
    const employeeChecklistProfession = document.getElementById('employeeChecklistProfession');
    const employeeChecklistOverviewBtn = document.getElementById('employeeChecklistOverviewBtn');
    const employeeChecklistRefreshBtn = document.getElementById('employeeChecklistRefreshBtn');
    const employeeChecklistGenerateMissingBtn = document.getElementById('employeeChecklistGenerateMissingBtn');
    const employeeChecklistGenerateSelectedBtn = document.getElementById('employeeChecklistGenerateSelectedBtn');
    const employeeChecklistGenerateAllBtn = document.getElementById('employeeChecklistGenerateAllBtn');
    const employeeChecklistApplyTypeBtn = document.getElementById('employeeChecklistApplyTypeBtn');
    const employeeChecklistOverwrite = document.getElementById('employeeChecklistOverwrite');
    const employeeChecklistSummary = document.getElementById('employeeChecklistSummary');
    const employeeChecklistList = document.getElementById('employeeChecklistList');
    const employeeChecklistOverview = document.getElementById('employeeChecklistOverview');
    const employeeChecklistMsg = document.getElementById('employeeChecklistMsg');
    const employeeChecklistSelectMissing = document.getElementById('employeeChecklistSelectMissing');
    const batchGenerateMode = document.getElementById('batchGenerateMode');
    const batchEmployees = document.getElementById('batchEmployees');
    const batchDocCodes = document.getElementById('batchDocCodes');
    const batchEmployeesSelectAllBtn = document.getElementById('batchEmployeesSelectAllBtn');
    const batchEmployeesClearBtn = document.getElementById('batchEmployeesClearBtn');
    const batchUseMissingCodesBtn = document.getElementById('batchUseMissingCodesBtn');
    const batchGenerateSelectedEmployeesBtn = document.getElementById('batchGenerateSelectedEmployeesBtn');
    const batchGenerateAllFilteredBtn = document.getElementById('batchGenerateAllFilteredBtn');

    const exportClassification = document.getElementById('exportClassification');
    const exportDocxBtn = document.getElementById('exportDocxBtn');
    const exportXlsxBtn = document.getElementById('exportXlsxBtn');
    const exportPackBtn = document.getElementById('exportPackBtn');
    const objectNameInput = document.getElementById('objectNameInput');
    const projectCodeInput = document.getElementById('projectCodeInput');
    const organizationInput = document.getElementById('organizationInput');
    const workStageInput = document.getElementById('workStageInput');
    const startDateInput = document.getElementById('startDateInput');
    const objectProfileSaveBtn = document.getElementById('objectProfileSaveBtn');
    const objectProfileMsg = document.getElementById('objectProfileMsg');
    const pprSourceSelect = document.getElementById('pprSourceSelect');
    const pprImportBtn = document.getElementById('pprImportBtn');
    const uploadBtn = document.getElementById('uploadBtn');
    const uploadFile = document.getElementById('uploadFile');
    const uploadDir = document.getElementById('uploadDir');
    const uploadMsg = document.getElementById('uploadMsg');
    const uploadUseCurrentBtn = document.getElementById('uploadUseCurrentBtn');
    const uploadAutoUseTree = document.getElementById('uploadAutoUseTree');

    const llmBadge = document.getElementById('llmBadge');
    const llmDesc = document.getElementById('llmDesc');
    const backToTopBtn = document.getElementById('backToTopBtn');

    const NODE_LABELS = {
        '00_incoming_requests': 'Входящие заявки',
        '01_orders_and_appointments': 'Приказы и назначения',
        '02_personnel': 'Персонал',
        '03_hse_and_fire_safety': 'ОТ, ПБ и пожарная безопасность',
        '04_journals': 'Журналы',
        '05_execution_docs': 'Исполнительная документация',
        '06_normative_base': 'Нормативная база',
        '07_monthly_control': 'Ежемесячный контроль',
        '08_outgoing_submissions': 'Исходящие отправки',
        'исходящие_заявки': 'Исходящие заявки',
        '09_archive': 'Архив',
        '10_scan_inbox': 'Входящие сканы',
        'наряды_допуски': 'Наряды-допуски',
        'employees': 'Сотрудники',
        '01_identity_and_contract': 'Документы личности и договор',
        '02_admission_orders': 'Приказы о допуске',
        '03_briefings_and_training': 'Инструктажи и обучение',
        '04_attestation_and_certificates': 'Аттестации и удостоверения',
        '05_ppe_issue': 'Выдача СИЗ',
        '06_permits_and_work_admission': 'Наряды и допуски',
        '07_medical_and_first_aid': 'Медосмотры и первая помощь',
        '07_templates_to_print': 'Черновики к печати',
        'production': 'Производственные журналы',
        'labor_safety': 'Журналы ОТ и ПБ',
        'employee_profile.txt': 'Профиль сотрудника',
        'print_office': 'Папка офисной печати',
        'print_pdf': 'PDF к печати',
        'print_pdf_ready': 'PDF к печати (готово)',
        'заявки': 'Заявки',
        'табели': 'Табели',
        'drafts_from_checklist': 'Черновики из чеклиста',
        'drafts_from_assistant': 'Черновики от ассистента',
        'manual_review': 'Ручной разбор'
    };

    let currentFile = '';
    let knownEmployeeIds = [];
    let knownEmployees = [];
    let employeeCatalogRows = [];
    let currentEmployeeChecklist = null;
    let manualReviewRows = [];
    let selectedManualReviewPath = '';
    let currentTaskAction = { path: '', label: 'не выбрана' };
    let treeNavigationHistory = [];
    let treeNavigationIndex = -1;
    let fileSourceVisible = false;
    let filePreviewZoom = 1;
    let currentScannerProfile = 1;
    let voiceRecognition = null;
    let voiceListening = false;
    let voiceFallbackMode = false;
    let voiceFallbackStream = null;
    let voiceFallbackAudioCtx = null;
    let voiceFallbackSource = null;
    let voiceFallbackProcessor = null;
    let voiceFallbackChunks = [];
    let voiceFallbackSampleRate = 16000;
    let voiceLastTranscript = '';

    const SCANNER_DOC_TYPES_REQUIRING_EMPLOYEE = new Set(['PASSPORT']);
    const MANUAL_REVIEW_TARGET_BY_SCAN_TYPE = {
        ORDER: '01_orders_and_appointments',
        AWR: '05_execution_docs/work_reports',
        PASSPORT: '02_personnel/employees',
        INVOICE: '08_outgoing_submissions/бухгалтерия/счета',
        UPD: '08_outgoing_submissions/бухгалтерия/упд',
        TTN: '08_outgoing_submissions/логистика/ттн',
        ACT: '05_execution_docs/admission_acts',
        OTHER: '10_scan_inbox/manual_review',
    };

    const canGoogleSpeechFallback = Boolean(
        navigator.mediaDevices
        && navigator.mediaDevices.getUserMedia
        && (window.AudioContext || window.webkitAudioContext)
    );

    async function api(url, options, timeoutMs = TIMEOUTS_MS.default, timeoutMessage = '') {
        const controller = new AbortController();
        const effectiveTimeout = Number(timeoutMs) > 0 ? Number(timeoutMs) : TIMEOUTS_MS.default;
        const timer = setTimeout(() => controller.abort(), effectiveTimeout);
        try {
            const response = await fetch(url, { ...(options || {}), signal: controller.signal });
            const data = await response.json().catch(() => ({}));
            if (!response.ok) {
                throw new Error(data.detail || response.statusText || 'Ошибка запроса');
            }
            return data;
        } catch (err) {
            if (err && err.name === 'AbortError') {
                throw new Error(timeoutMessage || ('Превышено время ожидания ответа сервера (' + Math.round(effectiveTimeout / 1000) + ' сек).'));
            }
            throw err;
        } finally {
            clearTimeout(timer);
        }
    }

    function setTreeHint(text) {
        if (!treeHint) {
            return;
        }
        treeHint.textContent = text || '';
    }

    function normalizeRelPath(path) {
        return (path || '').trim().replace(/\\/g, '/').replace(/^\/+|\/+$/g, '');
    }

    function parentRelPath(path) {
        const safe = normalizeRelPath(path);
        if (!safe) {
            return '';
        }
        const parts = safe.split('/').filter(Boolean);
        if (parts.length <= 1) {
            return '';
        }
        return parts.slice(0, -1).join('/');
    }

    function pushTreeHistory(path) {
        const safe = normalizeRelPath(path);
        if (treeNavigationIndex >= 0 && treeNavigationHistory[treeNavigationIndex] === safe) {
            return;
        }
        if (treeNavigationIndex < treeNavigationHistory.length - 1) {
            treeNavigationHistory = treeNavigationHistory.slice(0, treeNavigationIndex + 1);
        }
        treeNavigationHistory.push(safe);
        treeNavigationIndex = treeNavigationHistory.length - 1;
    }

    function updateTreeNavButtons(currentPath) {
        const safe = normalizeRelPath(currentPath);
        if (treeBackBtn) {
            treeBackBtn.disabled = treeNavigationIndex <= 0;
        }
        if (treeForwardBtn) {
            treeForwardBtn.disabled = treeNavigationIndex < 0 || treeNavigationIndex >= treeNavigationHistory.length - 1;
        }
        if (treeUpBtn) {
            treeUpBtn.disabled = !safe;
        }
        if (treeRootBtn) {
            treeRootBtn.disabled = !safe;
        }
    }

    function renderTreeBreadcrumb(path) {
        if (!treeBreadcrumb) {
            return;
        }

        const safe = normalizeRelPath(path);
        treeBreadcrumb.innerHTML = '';

        function appendSeparator() {
            const sep = document.createElement('span');
            sep.className = 'tree-crumb-sep';
            sep.textContent = '›';
            treeBreadcrumb.appendChild(sep);
        }

        function appendCrumb(label, relPath) {
            const btn = document.createElement('button');
            btn.type = 'button';
            btn.className = 'tree-crumb';
            btn.dataset.path = normalizeRelPath(relPath);
            btn.textContent = label;
            treeBreadcrumb.appendChild(btn);
        }

        appendCrumb('Корень', '');
        if (!safe) {
            return;
        }

        const parts = safe.split('/').filter(Boolean);
        let assembled = '';
        for (const part of parts) {
            appendSeparator();
            assembled = assembled ? (assembled + '/' + part) : part;
            appendCrumb(prettyNodeName(part), assembled);
        }
    }

    async function navigateTreeHistory(offset) {
        const nextIndex = treeNavigationIndex + offset;
        if (nextIndex < 0 || nextIndex >= treeNavigationHistory.length) {
            return;
        }
        treeNavigationIndex = nextIndex;
        const targetPath = treeNavigationHistory[nextIndex] || '';
        await loadTree(targetPath, { rememberHistory: false });
    }

    function setInteractionHint(text, level = 'info') {
        if (!interactionHint) {
            return;
        }
        interactionHint.textContent = text || '';
        interactionHint.classList.remove('hint-error', 'hint-ok');
        if (level === 'error') {
            interactionHint.classList.add('hint-error');
        }
        if (level === 'ok') {
            interactionHint.classList.add('hint-ok');
        }
    }

    function focusStructureCard(shouldScroll = false) {
        if (!structureCard) {
            return;
        }
        structureCard.classList.add('action-focus');
        if (shouldScroll) {
            structureCard.scrollIntoView({ behavior: 'smooth', block: 'start' });
        }
        window.setTimeout(() => structureCard.classList.remove('action-focus'), 1400);
    }

    function ensurePreviewCardExpanded() {
        if (!previewCard) {
            return;
        }

        const body = previewCard.querySelector(':scope > .card-collapse-body');
        if (!(body instanceof HTMLElement)) {
            return;
        }
        if (!body.hidden) {
            return;
        }

        body.hidden = false;
        const toggle = previewCard.querySelector(':scope > .card-collapse-head button');
        if (toggle instanceof HTMLButtonElement) {
            toggle.textContent = 'Свернуть';
        }
        try {
            localStorage.setItem('arm.card.collapse.preview', '0');
        } catch (_err) {
            // ignore storage write failures
        }
    }

    function focusPreviewCard(shouldScroll = false) {
        if (!previewCard) {
            return;
        }
        ensurePreviewCardExpanded();
        previewCard.classList.add('action-focus');
        if (shouldScroll) {
            previewCard.scrollIntoView({ behavior: 'smooth', block: 'start' });
        }
        window.setTimeout(() => previewCard.classList.remove('action-focus'), 1400);
    }

    function ensureEmployeeChecklistExpanded() {
        if (!employeeChecklistCard || !employeeChecklistBody || !employeeChecklistToggleBtn) {
            return;
        }
        employeeChecklistCard.classList.remove('is-collapsed');
        employeeChecklistBody.hidden = false;
        employeeChecklistToggleBtn.textContent = 'Свернуть';
    }

    function toggleEmployeeChecklistCard() {
        if (!employeeChecklistCard || !employeeChecklistBody || !employeeChecklistToggleBtn) {
            return;
        }
        const willCollapse = !employeeChecklistCard.classList.contains('is-collapsed');
        employeeChecklistCard.classList.toggle('is-collapsed', willCollapse);
        employeeChecklistBody.hidden = willCollapse;
        employeeChecklistToggleBtn.textContent = willCollapse ? 'Развернуть' : 'Свернуть';
    }

    function focusEmployeeChecklistCard(shouldScroll = false) {
        if (!employeeChecklistCard) {
            return;
        }
        ensureEmployeeChecklistExpanded();
        employeeChecklistCard.classList.add('action-focus');
        if (shouldScroll) {
            employeeChecklistCard.scrollIntoView({ behavior: 'smooth', block: 'start' });
        }
        window.setTimeout(() => employeeChecklistCard.classList.remove('action-focus'), 1400);
    }

    function getFilteredEmployeeRows() {
        const selectedProfession = (employeeChecklistProfession && employeeChecklistProfession.value || '').trim();
        if (!selectedProfession || selectedProfession === 'all') {
            return employeeCatalogRows.slice();
        }
        return employeeCatalogRows.filter((row) => row.profession_group === selectedProfession);
    }

    function selectedValuesFromMultiSelect(selectElement) {
        if (!selectElement) {
            return [];
        }
        return Array.from(selectElement.selectedOptions || [])
            .map((item) => (item.value || '').trim())
            .filter(Boolean);
    }

    function setFileMeta(text) {
        if (!fileMeta) {
            return;
        }
        fileMeta.textContent = text || '';
    }

    function isTextFilePath(relPath) {
        return /\.(md|txt|csv|json|yml|yaml|xml|html|py|log|ini)$/i.test(relPath || '');
    }

    function updateFileSourceVisibility() {
        if (!fileEditor) {
            return;
        }
        fileEditor.classList.toggle('is-hidden', !fileSourceVisible);
        if (fileSourceToggleBtn) {
            fileSourceToggleBtn.textContent = fileSourceVisible ? 'Скрыть исходник' : 'Показать исходник';
        }
    }

    function refreshRenderedPreview(relPath) {
        if (!fileRenderFrame) {
            return;
        }
        const safePath = normalizeRelPath(relPath || currentFile || '');
        if (!safePath) {
            fileRenderFrame.src = 'about:blank';
            return;
        }
        const baseUrl = isTextFilePath(safePath)
            ? ('/arm/fs/print-preview?rel_path=' + encodeURIComponent(safePath) + '&auto_print=0')
            : ('/arm/fs/view?rel_path=' + encodeURIComponent(safePath));
        fileRenderFrame.src = baseUrl + '&v=' + Date.now();
        applyPreviewZoom();
    }

    function applyPreviewZoom() {
        if (!fileRenderFrame || !fileRenderWrap) {
            return;
        }
        const safeZoom = Math.max(0.5, Math.min(2, Number(filePreviewZoom) || 1));
        filePreviewZoom = safeZoom;
        fileRenderFrame.style.transformOrigin = '0 0';
        fileRenderFrame.style.transform = 'scale(' + safeZoom + ')';
        fileRenderFrame.style.width = (100 / safeZoom) + '%';
        fileRenderFrame.style.height = (66 / safeZoom) + 'vh';
    }

    function setPreviewZoom(nextZoom) {
        filePreviewZoom = Math.max(0.5, Math.min(2, Number(nextZoom) || 1));
        applyPreviewZoom();
        if (fileZoomResetBtn) {
            fileZoomResetBtn.textContent = Math.round(filePreviewZoom * 100) + '%';
        }
        setFileMeta('Масштаб предпросмотра: ' + Math.round(filePreviewZoom * 100) + '%');
    }

    function scannerProfileText(profile) {
        if (profile === 1) return 'Профиль 1: 300 dpi, grayscale';
        if (profile === 2) return 'Профиль 2: 400/600 dpi, grayscale';
        return 'Профиль 3: 400/600 dpi, color';
    }

    function updateScannerProfileLabel() {
        if (!scannerProfileLabel) {
            return;
        }
        scannerProfileLabel.textContent = scannerProfileText(currentScannerProfile);
    }

    async function moveCurrentFile() {
        if (!currentFile) {
            setFileMeta('Выберите файл из дерева.');
            return;
        }
        const targetPath = window.prompt('Целевой путь для перемещения файла', currentFile);
        if (!targetPath || !targetPath.trim()) {
            return;
        }
        const data = await api(
            '/arm/fs/move?source_rel_path=' + encodeURIComponent(currentFile) + '&target_rel_path=' + encodeURIComponent(targetPath.trim()),
            { method: 'POST' }
        );
        setFileMeta(data.message || 'Файл перемещен.');
        currentFile = targetPath.trim();
        refreshRenderedPreview(currentFile);
        const targetFolder = currentFile.split('/').slice(0, -1).join('/');
        if (targetFolder) {
            await loadTree(targetFolder).catch(() => {});
        }
    }

    async function deleteCurrentFile() {
        if (!currentFile) {
            setFileMeta('Выберите файл из дерева.');
            return;
        }
        const confirmed = window.confirm('Удалить файл: ' + describePath(currentFile) + '?');
        if (!confirmed) {
            return;
        }
        const data = await api('/arm/fs/delete?rel_path=' + encodeURIComponent(currentFile) + '&with_sidecar=true', { method: 'POST' });
        setFileMeta(data.message || 'Файл удален.');
        const folder = currentFile.split('/').slice(0, -1).join('/');
        currentFile = '';
        if (filePathLabel) {
            filePathLabel.textContent = 'Файл не выбран';
        }
        if (fileEditor) {
            fileEditor.value = '';
            fileEditor.readOnly = true;
        }
        refreshRenderedPreview('');
        if (folder) {
            await loadTree(folder).catch(() => {});
        }
    }

    function scannerEmployeeValue() {
        const fromInput = scannerEmployee && typeof scannerEmployee.value === 'string'
            ? scannerEmployee.value.trim()
            : '';
        const fromSelect = scannerEmployeeSelect && typeof scannerEmployeeSelect.value === 'string'
            ? scannerEmployeeSelect.value.trim()
            : '';
        return fromInput || fromSelect;
    }

    function setScannerEmployeeValue(value) {
        const safe = (value || '').trim();
        if (scannerEmployee && typeof scannerEmployee.value === 'string') {
            scannerEmployee.value = safe;
        }
        if (scannerEmployeeSelect && safe && Array.from(scannerEmployeeSelect.options).some((item) => item.value === safe)) {
            scannerEmployeeSelect.value = safe;
        }
    }

    function appendScannerTimeline(text) {
        if (!scannerTimeline || !text) {
            return;
        }
        const stamp = new Date().toLocaleTimeString('ru-RU');
        const next = scannerTimeline.textContent
            ? (scannerTimeline.textContent + '\\n[' + stamp + '] ' + text)
            : ('[' + stamp + '] ' + text);
        scannerTimeline.textContent = next;
    }

    function setScannerProgress(percent, text) {
        const safePercent = Math.max(0, Math.min(100, Number(percent) || 0));
        if (scannerProgressBar) {
            scannerProgressBar.style.width = String(safePercent) + '%';
        }
        if (scannerProgressText) {
            scannerProgressText.textContent = 'Прогресс: ' + safePercent + '%; ' + (text || 'ожидание');
        }
    }

    function resetScannerProgress(text) {
        if (scannerTimeline) {
            scannerTimeline.textContent = '';
        }
        setScannerProgress(0, text || 'ожидание');
    }

    function capitalizeCyrillic(text) {
        return text.replace(/\\b[а-яё][а-яё]+\\b/gi, (word) => word.charAt(0).toUpperCase() + word.slice(1).toLowerCase());
    }

    function localizeLatinTokens(text) {
        let output = String(text || '');
        const replacements = [
            [/\\border\\b/gi, 'приказ'],
            [/\\bpermit\\b/gi, 'наряд-допуск'],
            [/\\bauthorized\\b/gi, 'уполномоченный'],
            [/\\brepresentative\\b/gi, 'представитель'],
            [/\\badmission\\b/gi, 'допуск'],
            [/\\bcmr\\b/gi, 'СМР'],
            [/\\bresponsible\\b/gi, 'ответственные'],
            [/\\bpersons?\\b/gi, 'лица'],
            [/\\bslinger\\b/gi, 'стропальщик'],
            [/\\bassignment\\b/gi, 'назначение'],
            [/\\binternship\\b/gi, 'стажировка'],
            [/\\bindependent\\b/gi, 'самостоятельная'],
            [/\\bloading\\b/gi, 'погрузка'],
            [/\\bunloading\\b/gi, 'разгрузка'],
            [/\\bfire\\b/gi, 'пожарная'],
            [/\\bsafety\\b/gi, 'безопасность'],
            [/\\bheight\\b/gi, 'высотные'],
            [/\\bworks?\\b/gi, 'работы'],
            [/\\bregister\\b/gi, 'реестр'],
            [/\\bpdf\\b/gi, 'PDF'],
            [/\\bready\\b/gi, 'готово'],
            [/\\bprint\\b/gi, 'печать'],
            [/\\boffice\\b/gi, 'офис'],
        ];

        for (const [pattern, replacement] of replacements) {
            output = output.replace(pattern, replacement);
        }
        return output;
    }

    function prettifyUnknownName(name) {
        const raw = String(name || '').trim();
        const noExt = raw.replace(/\.[A-Za-z0-9]{1,8}$/u, '');
        const normalizedTokens = noExt.replace(/[_-]+/g, ' ').toLowerCase();
        const isOrderFile = /\b(order|permit)\b/u.test(normalizedTokens);

        let cleaned = noExt;
        if (isOrderFile) {
            cleaned = cleaned.replace(/^\d{8}[_-]*/u, '');
            cleaned = cleaned.replace(/^order[_-]?register[_-]*/iu, 'register_');
            cleaned = cleaned.replace(/^(order|permit)[_-]?\d{1,2}[_-]*/iu, '');
        }

        cleaned = cleaned.replace(/[_-]?v\d{2}$/iu, '');
        const normalized = cleaned.replace(/[_-]+/g, ' ').replace(/\s+/g, ' ').trim();
        return capitalizeCyrillic(localizeLatinTokens(normalized));
    }

    function parseEmployeeIdFromName(name) {
        const raw = (name || '').trim();
        // Employee IDs are short numeric codes; avoid treating YYYYMMDD order dates as employee IDs.
        const match = /^(\d{2,5})(?:[_-].*)?$/u.exec(raw);
        return match ? match[1] : '';
    }

    function extractEmployeeIdFromPath(path) {
        const normalized = (path || '').replace(/\\\\/g, '/');
        const match = /02_personnel\/employees\/([^/]+)/i.exec(normalized);
        if (!match) {
            return '';
        }
        return parseEmployeeIdFromName(match[1]);
    }

    function extractEmployeeRelPath(path) {
        const normalized = (path || '').replace(/\\\\/g, '/').replace(/^\/+/, '');
        const match = /(^|\/)(02_personnel\/employees\/[^/]+)/i.exec(normalized);
        if (!match) {
            return '';
        }
        return match[2] || '';
    }

    function getEmployeeDisplayName(employeeRelPath) {
        const normalized = normalizeRelPath(employeeRelPath);
        if (!normalized || !employeeCatalogRows || !employeeCatalogRows.length) {
            return '';
        }
        const row = employeeCatalogRows.find((item) => normalizeRelPath(item.employee_rel_path || '') === normalized);
        return row && row.employee_name ? String(row.employee_name).trim() : '';
    }

    function syncEmployeeChecklistPath(path) {
        if (!employeeChecklistPath) {
            return;
        }
        const employeePath = extractEmployeeRelPath(path);
        if (!employeePath) {
            return;
        }
        if (!employeeChecklistPath.value.trim()) {
            employeeChecklistPath.value = employeePath;
            if (employeeChecklistMsg) {
                employeeChecklistMsg.textContent = 'Путь сотрудника подставлен из текущей навигации: ' + describePath(employeePath);
            }
        }
    }

    function updateEmployeeIdSuggestions(ids) {
        const unique = Array.from(new Set((ids || []).filter((item) => Boolean((item || '').trim()))));
        unique.sort((a, b) => String(a).localeCompare(String(b), 'ru'));
        knownEmployeeIds = unique;

        if (!scannerEmployeeSuggestions) {
            return;
        }

        scannerEmployeeSuggestions.innerHTML = '';
        for (const id of knownEmployeeIds) {
            const option = document.createElement('option');
            option.value = id;
            scannerEmployeeSuggestions.appendChild(option);
        }
    }

    function updateEmployeeSelectOptions(employees) {
        knownEmployees = Array.isArray(employees) ? employees : [];
        if (!scannerEmployeeSelect) {
            return;
        }

        const previousValue = scannerEmployeeSelect.value || '';
        scannerEmployeeSelect.innerHTML = '';

        const placeholder = document.createElement('option');
        placeholder.value = '';
        placeholder.textContent = 'Выберите сотрудника (для удостоверения/протокола)';
        scannerEmployeeSelect.appendChild(placeholder);

        for (const employee of knownEmployees) {
            const option = document.createElement('option');
            option.value = employee.value;
            option.textContent = employee.label;
            scannerEmployeeSelect.appendChild(option);
        }

        if (previousValue && knownEmployees.some((item) => item.value === previousValue)) {
            scannerEmployeeSelect.value = previousValue;
        }
    }

    async function refreshEmployeeIdSuggestions() {
        try {
            const data = await api(
                '/arm/fs/tree?rel_path=' + encodeURIComponent('02_personnel/employees'),
                undefined,
                TIMEOUTS_MS.tree,
                'Каталог сотрудников загружается слишком долго. Повторите попытку.'
            );
            const employees = (data.entries || [])
                .filter((item) => item && item.is_dir)
                .map((item) => {
                    const folderName = item.name || '';
                    const parsedId = parseEmployeeIdFromName(folderName);
                    const value = parsedId || folderName;
                    const prettyName = prettifyUnknownName(folderName.replace(/^\d{2,10}[_-]?/u, '')) || folderName;
                    const label = parsedId ? ('[' + parsedId + '] ' + prettyName) : prettyName;
                    return { value, label };
                })
                .filter((item) => Boolean((item.value || '').trim()));

            updateEmployeeIdSuggestions(employees.map((item) => item.value));
            updateEmployeeSelectOptions(employees);

            if (scannerEmployeeHint && employees.length) {
                scannerEmployeeHint.textContent = 'Доступно сотрудников: ' + employees.length + '. Можно выбрать из списка или ввести вручную.';
            }
        } catch (err) {
            if (scannerEmployeeHint) {
                scannerEmployeeHint.textContent = 'Не удалось загрузить список ID автоматически: ' + err;
            }
        }
    }

    function inferProfessionByGroupKey(groupKey) {
        if (groupKey === 'electric') {
            return 'электротехнический персонал';
        }
        if (groupKey === 'supervisor') {
            return 'прораб';
        }
        if (groupKey === 'itr') {
            return 'инженерно-технический работник';
        }
        return 'общий персонал';
    }

    function fileExtensionFromName(name) {
        const match = /\.([A-Za-z0-9]{1,8})$/u.exec(String(name || '').trim());
        return match ? String(match[1]).toLowerCase() : '';
    }

    function parseOrderLabelFromPath(relPath, fallbackName) {
        const safeRel = String(relPath || '').trim();
        const safeName = String(fallbackName || '').trim();
        const sourceName = safeName || safeRel.split('/').pop() || '';
        if (!sourceName) {
            return '';
        }
        const stem = sourceName.replace(/\.[^.]+$/u, '');
        const match = /ORDER[_-]?([0-9]{1,3})(?:[_-](.*))?/iu.exec(stem);
        if (!match) {
            return '';
        }
        let orderNo = String(match[1] || '').trim();
        if (/^\d+$/u.test(orderNo)) {
            orderNo = orderNo.padStart(2, '0');
        }
        let title = String(match[2] || '').replace(/[_-]?v\d+$/iu, '').trim();
        title = title.replace(/[_-]+/gu, ' ').replace(/\s+/gu, ' ').trim();
        if (title) {
            title = localizeLatinTokens(title);
            title = capitalizeCyrillic(title);
        }
        return title ? ('Приказ №' + orderNo + ' - ' + title) : ('Приказ №' + orderNo);
    }

    function renderEmployeeChecklistEmployeeOptions() {
        if (!employeeChecklistEmployeeSelect) {
            return;
        }

        const selectedProfession = (employeeChecklistProfession && employeeChecklistProfession.value || '').trim();
        const filtered = selectedProfession && selectedProfession !== 'all'
            ? employeeCatalogRows.filter((row) => row.profession_group === selectedProfession)
            : employeeCatalogRows;

        const oldValue = employeeChecklistEmployeeSelect.value || '';
        employeeChecklistEmployeeSelect.innerHTML = '';

        const first = document.createElement('option');
        first.value = '';
        first.textContent = 'Не выбран (анализ по типам сотрудников)';
        employeeChecklistEmployeeSelect.appendChild(first);

        const groups = new Map();
        for (const row of filtered) {
            const key = row.profession_group || 'default';
            if (!groups.has(key)) {
                groups.set(key, []);
            }
            groups.get(key).push(row);
        }

        for (const [groupKey, rows] of groups.entries()) {
            const optgroup = document.createElement('optgroup');
            optgroup.label = (rows[0] && rows[0].profession_label) || groupKey;
            rows
                .slice()
                .sort((a, b) => String(a.employee_name || '').localeCompare(String(b.employee_name || ''), 'ru'))
                .forEach((row) => {
                    const option = document.createElement('option');
                    option.value = row.employee_rel_path;
                    const idPart = row.employee_id ? ('[' + row.employee_id + '] ') : '';
                    const posPart = row.position ? (' - ' + row.position) : '';
                    option.textContent = idPart + row.employee_name + posPart;
                    optgroup.appendChild(option);
                });
            employeeChecklistEmployeeSelect.appendChild(optgroup);
        }

        if (oldValue && filtered.some((row) => row.employee_rel_path === oldValue)) {
            employeeChecklistEmployeeSelect.value = oldValue;
        }
    }

    function renderBatchEmployeeOptions() {
        if (!batchEmployees) {
            return;
        }

        const previousSelected = selectedValuesFromMultiSelect(batchEmployees);
        const rows = getFilteredEmployeeRows();

        batchEmployees.innerHTML = '';
        for (const row of rows) {
            const option = document.createElement('option');
            option.value = row.employee_rel_path;
            const idPart = row.employee_id ? ('[' + row.employee_id + '] ') : '';
            const posPart = row.position ? (' - ' + row.position) : '';
            option.textContent = idPart + row.employee_name + posPart;
            option.selected = previousSelected.includes(row.employee_rel_path);
            batchEmployees.appendChild(option);
        }
    }

    function renderBatchDocCodesFromChecklist(checklistData) {
        if (!batchDocCodes) {
            return;
        }

        batchDocCodes.innerHTML = '';
        const items = (checklistData && checklistData.items) || [];
        if (!items.length) {
            batchDocCodes.textContent = 'Сначала откройте чеклист сотрудника. Тогда здесь появятся коды документов для выборочной генерации.';
            return;
        }

        for (const item of items) {
            const label = document.createElement('label');
            label.className = 'batch-code-item';

            const checkbox = document.createElement('input');
            checkbox.type = 'checkbox';
            checkbox.dataset.code = item.code;
            checkbox.checked = !item.ready;

            const text = document.createElement('span');
            const state = item.ready ? 'готово' : 'не хватает';
            text.textContent = item.code + ' - ' + item.title + ' (' + state + ')';

            label.appendChild(checkbox);
            label.appendChild(text);
            batchDocCodes.appendChild(label);
        }
    }

    function collectBatchSelectedCodes() {
        if (!batchDocCodes) {
            return [];
        }
        return Array.from(batchDocCodes.querySelectorAll('input[type="checkbox"][data-code]'))
            .filter((item) => item.checked)
            .map((item) => item.dataset.code || '')
            .filter(Boolean);
    }

    function formatGenerateMode(mode) {
        const safeMode = (mode || '').trim();
        if (safeMode === 'selected') {
            return 'выбранные коды';
        }
        if (safeMode === 'all') {
            return 'полный комплект';
        }
        return 'только недостающие';
    }

    function getChecklistOrderDate() {
        const value = (employeeChecklistOrderDate && employeeChecklistOrderDate.value || '').trim() || '01.03.2026';
        if (!/^\d{2}\.\d{2}\.\d{4}$/.test(value)) {
            return null;
        }
        return value;
    }

    function getCatalogRowByEmployeePath(employeeRelPath) {
        return employeeCatalogRows.find((row) => row.employee_rel_path === employeeRelPath) || null;
    }

    function setMultiSelectState(selectElement, shouldSelect) {
        if (!selectElement) {
            return;
        }
        for (const option of Array.from(selectElement.options || [])) {
            option.selected = shouldSelect;
        }
    }

    async function generateChecklistForEmployeePath(employeeRelPath, mode, codes, orderDate) {
        const row = getCatalogRowByEmployeePath(employeeRelPath);
        const payload = {
            employee_rel_path: employeeRelPath,
            profession: row ? inferProfessionByGroupKey(row.profession_group || 'default') : null,
            order_date: orderDate,
            mode: mode,
            codes: mode === 'selected' ? codes : [],
            overwrite: Boolean(employeeChecklistOverwrite && employeeChecklistOverwrite.checked)
        };
        return api('/arm/employees/checklist/generate', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify(payload)
        }, TIMEOUTS_MS.checklistGenerate, 'Пакетная генерация заняла слишком много времени.');
    }

    async function runBatchChecklistGeneration(generateAllFiltered) {
        const mode = (batchGenerateMode && batchGenerateMode.value) || 'missing';
        const targetRows = generateAllFiltered
            ? getFilteredEmployeeRows()
            : selectedValuesFromMultiSelect(batchEmployees).map((path) => getCatalogRowByEmployeePath(path)).filter(Boolean);

        if (!targetRows.length) {
            setEmployeeChecklistMsg(
                generateAllFiltered
                    ? 'Нет сотрудников в выбранном фильтре для пакетной генерации.'
                    : 'Выберите хотя бы одного сотрудника для пакетной генерации.',
                'error'
            );
            return;
        }

        const selectedCodes = mode === 'selected' ? collectBatchSelectedCodes() : [];
        if (mode === 'selected' && !selectedCodes.length) {
            setEmployeeChecklistMsg('Для выборочного режима отметьте хотя бы один код документа.', 'error');
            return;
        }

        const orderDate = getChecklistOrderDate();
        if (!orderDate) {
            setEmployeeChecklistMsg('Дата приказа должна быть в формате ДД.ММ.ГГГГ.', 'error');
            if (employeeChecklistOrderDate) {
                employeeChecklistOrderDate.focus();
            }
            return;
        }

        setEmployeeChecklistMsg('Пакетная генерация: сотрудников ' + targetRows.length + ', режим ' + formatGenerateMode(mode) + '.');

        let createdTotal = 0;
        let skippedTotal = 0;
        const errors = [];

        for (const row of targetRows) {
            try {
                const result = await generateChecklistForEmployeePath(row.employee_rel_path, mode, selectedCodes, orderDate);
                createdTotal += (result.created_files || []).length;
                skippedTotal += (result.skipped_files || []).length;
            } catch (err) {
                errors.push((row.employee_name || row.employee_rel_path) + ': ' + String(err || 'ошибка'));
            }
        }

        const level = errors.length && createdTotal === 0 ? 'error' : 'ok';
        const parts = [
            'Пакетная генерация завершена.',
            'Сотрудников: ' + targetRows.length + '.',
            'Создано файлов: ' + createdTotal + '.',
            'Пропущено: ' + skippedTotal + '.'
        ];
        if (errors.length) {
            parts.push('Ошибок: ' + errors.length + '.');
        }
        setEmployeeChecklistMsg(parts.join(' '), level);
        if (errors.length) {
            setInteractionHint('Часть сотрудников не обработана: ' + errors.slice(0, 2).join(' | '), 'error');
        } else {
            setInteractionHint('Пакетная генерация завершена без ошибок.', 'ok');
        }

        const currentEmployeePath = (employeeChecklistPath && employeeChecklistPath.value || '').trim();
        if (currentEmployeePath && targetRows.some((row) => row.employee_rel_path === currentEmployeePath)) {
            await loadEmployeeChecklist().catch(() => {});
        }
    }

    async function applyChecklistByProfession() {
        const selectedProfession = (employeeChecklistProfession && employeeChecklistProfession.value || '').trim();
        if (!selectedProfession || selectedProfession === 'all') {
            setEmployeeChecklistMsg('Выберите конкретный тип сотрудников перед применением формы.', 'error');
            if (employeeChecklistProfession) {
                employeeChecklistProfession.focus();
            }
            return;
        }

        if (batchGenerateMode && !batchGenerateMode.value) {
            batchGenerateMode.value = 'missing';
        }

        await runBatchChecklistGeneration(true);
    }

    async function refreshEmployeeChecklistCatalog() {
        const data = await api('/arm/employees/catalog', undefined, TIMEOUTS_MS.checklist);
        employeeCatalogRows = data.items || [];

        if (employeeChecklistProfession) {
            const current = employeeChecklistProfession.value || 'all';
            employeeChecklistProfession.innerHTML = '';

            const anyOption = document.createElement('option');
            anyOption.value = 'all';
            anyOption.textContent = 'Все типы сотрудников';
            employeeChecklistProfession.appendChild(anyOption);

            for (const optionData of (data.profession_options || [])) {
                const option = document.createElement('option');
                option.value = optionData.key;
                option.textContent = optionData.label;
                employeeChecklistProfession.appendChild(option);
            }

            employeeChecklistProfession.value = current && Array.from(employeeChecklistProfession.options).some((item) => item.value === current)
                ? current
                : 'all';
        }

        renderEmployeeChecklistEmployeeOptions();
        renderBatchEmployeeOptions();
    }

    function renderEmployeeChecklistOverview(data) {
        if (!employeeChecklistOverview) {
            return;
        }

        employeeChecklistOverview.innerHTML = '';
        const groups = data.groups || [];
        if (!groups.length) {
            employeeChecklistOverview.textContent = 'Нет сотрудников для выбранного фильтра.';
            return;
        }

        for (const group of groups) {
            const wrap = document.createElement('div');
            wrap.className = 'employee-overview-group';

            const title = document.createElement('div');
            title.className = 'employee-overview-title';
            title.textContent = group.profession_label
                + ': сотрудников ' + group.employees_total
                + '; готово ' + group.ready_employees
                + '; средний прогресс ' + group.average_progress_percent + '%';
            wrap.appendChild(title);

            const actionsCaption = document.createElement('div');
            actionsCaption.className = 'meta';
            actionsCaption.textContent = 'Перечень мероприятий по типу сотрудников:';
            wrap.appendChild(actionsCaption);

            const actionsList = document.createElement('ul');
            for (const action of (group.missing_actions || [])) {
                const li = document.createElement('li');
                const prefix = action.missing_employees > 0
                    ? ('[не хватает у ' + action.missing_employees + ' чел.] ')
                    : '[готово по группе] ';
                li.textContent = prefix + action.code + ' - ' + action.title + '. ' + action.guidance;
                actionsList.appendChild(li);
            }
            wrap.appendChild(actionsList);

            const employeesCaption = document.createElement('div');
            employeesCaption.className = 'meta';
            employeesCaption.textContent = 'Сотрудники группы:';
            wrap.appendChild(employeesCaption);

            const employeesList = document.createElement('ul');
            for (const employee of (group.employees || [])) {
                const li = document.createElement('li');
                const idPart = employee.employee_id ? ('[' + employee.employee_id + '] ') : '';
                const topMissing = (employee.top_missing_codes || []).join(', ');

                const link = document.createElement('a');
                link.href = '#';
                link.className = 'todo-link';
                link.textContent = idPart + employee.employee_name + ' - ' + employee.progress_percent + '%';
                link.addEventListener('click', async (event) => {
                    event.preventDefault();
                    if (employeeChecklistEmployeeSelect) {
                        employeeChecklistEmployeeSelect.value = employee.employee_rel_path;
                    }
                    if (employeeChecklistPath) {
                        employeeChecklistPath.value = employee.employee_rel_path;
                    }
                    if (employeeChecklistProfession) {
                        employeeChecklistProfession.value = group.profession_group;
                    }
                    await loadEmployeeChecklist();
                });

                li.appendChild(link);
                if (topMissing) {
                    const hint = document.createElement('div');
                    hint.className = 'system-hint';
                    hint.textContent = 'Основные недостающие документы: ' + topMissing;
                    li.appendChild(hint);
                }
                employeesList.appendChild(li);
            }
            wrap.appendChild(employeesList);

            employeeChecklistOverview.appendChild(wrap);
        }
    }

    async function loadEmployeeChecklistOverview() {
        const professionFilter = (employeeChecklistProfession && employeeChecklistProfession.value || '').trim();
        const query = new URLSearchParams();
        if (professionFilter && professionFilter !== 'all') {
            query.set('profession', professionFilter);
        }

        setEmployeeChecklistMsg('Выполняется комплексный анализ по сотрудникам и типам ролей...');
        const data = await api(
            '/arm/employees/checklist/overview' + (query.toString() ? ('?' + query.toString()) : ''),
            undefined,
            TIMEOUTS_MS.checklistGenerate,
            'Комплексный анализ чеклиста выполняется слишком долго.'
        );

        if (employeeChecklistList) {
            employeeChecklistList.innerHTML = '';
        }
        renderBatchDocCodesFromChecklist(null);
        if (employeeChecklistSummary) {
            employeeChecklistSummary.textContent = 'Комплексный анализ завершен: групп ' + (data.groups || []).length + '.';
        }
        renderEmployeeChecklistOverview(data);
        setEmployeeChecklistMsg('Показан групповой анализ по профессиям и сотрудникам.', 'ok');
    }

    function applyEmployeeIdFromPath(path) {
        const employeeId = extractEmployeeIdFromPath(path);
        if (!employeeId) {
            if (scannerEmployeeHint) {
                const sample = knownEmployeeIds.slice(0, 8).join(', ');
                scannerEmployeeHint.textContent = sample
                    ? ('Выбор кода сотрудника из списка: ' + sample + (knownEmployeeIds.length > 8 ? '...' : ''))
                    : 'Выберите папку сотрудника, чтобы подставить код автоматически.';
            }
            return;
        }

        if (!knownEmployeeIds.includes(employeeId)) {
            updateEmployeeIdSuggestions([...knownEmployeeIds, employeeId]);
            updateEmployeeSelectOptions([
                ...knownEmployees,
                {
                    value: employeeId,
                    label: '[Авто] ' + employeeId,
                },
            ]);
        }

        if (!scannerEmployeeValue()) {
            setScannerEmployeeValue(employeeId);
        }

        if (scannerEmployeeSelect && knownEmployees.some((item) => item.value === employeeId)) {
            scannerEmployeeSelect.value = employeeId;
        }

        if (scannerEmployeeHint) {
            scannerEmployeeHint.textContent = 'Найден код сотрудника из выбранной папки: ' + employeeId + '.';
        }
    }

    function humanLabel(name) {
        const rawName = (name || '').trim();
        const key = rawName.toLowerCase();
        if (NODE_LABELS[key]) {
            return NODE_LABELS[key];
        }

        const isFileName = /\.[A-Za-z0-9]{1,8}$/u.test(rawName);
        const employeeId = isFileName ? '' : parseEmployeeIdFromName(rawName);
        if (employeeId && /^\d{2,5}[_-]/u.test(rawName)) {
            const personName = prettifyUnknownName(rawName.replace(/^\d{2,10}[_-]?/u, ''));
            return personName ? ('[' + employeeId + '] ' + personName) : ('[' + employeeId + ']');
        }

        return '';
    }

    function describePath(path) {
        const safePath = (path || '').trim().replace(/^\/+|\/+$/g, '');
        if (!safePath) {
            return 'Корень объекта';
        }
        return safePath
            .split('/')
            .map((part) => humanLabel(part) || prettifyUnknownName(part) || part)
            .join(' / ');
    }

    function prettyNodeName(name) {
        const label = humanLabel(name || '');
        return label || prettifyUnknownName(name || '') || (name || '');
    }

    function buildActionSteps(path, label) {
        const safePath = (path || '').toLowerCase();
        const safeLabel = (label || 'задача').trim();

        if (safePath.includes('10_scan_inbox/manual_review')) {
            return [
                'Откройте файл из списка ручного разбора и проверьте содержимое/скан.',
                'Нажмите «Классифицировать ручной разбор» и проверьте прогноз типа документа.',
                'После проверки переместите документ в целевую папку и обновите дерево.'
            ];
        }

        if (safePath.includes('02_personnel/employees')) {
            return [
                'Откройте папку сотрудника и выберите нужный раздел 01-07.',
                'Проверьте наличие обязательных сканов/файлов в выбранном разделе.',
                'Если документа нет: выполните сканирование, затем «Распознать и разложить».'
            ];
        }

        if (safePath.includes('01_orders_and_appointments')) {
            return [
                'Откройте папку приказов и проверьте наличие требуемых файлов приказов/нарядов-допусков.',
                'При необходимости скорректируйте файл через «Предпросмотр, правка, печать».',
                'После обновления выгрузите DOCX/XLSX в блоке «Экспорт и печать офисных документов».'
            ];
        }

        if (safePath.includes('04_journals')) {
            return [
                'Откройте нужный журнал и проверьте полноту записей.',
                'Убедитесь, что файл оформлен и сохранен в правильной подпапке.',
                'Если не хватает данных, добавьте записи и обновите метрики на дашборде.'
            ];
        }

        return [
            'Открыт маршрут по задаче: ' + safeLabel + '.',
            'Проверьте в этой папке обязательные файлы и подпапки.',
            'Дальше выберите файл: открыть, изменить, скачать или отправить на печать.'
        ];
    }

    function renderActionNavigator(path, label) {
        if (!actionNavigatorPath || !actionNavigatorSteps) {
            return;
        }
        const rel = (path || '').trim();
        actionNavigatorPath.textContent = 'Маршрут: ' + (rel ? describePath(rel) : '/');
        actionNavigatorSteps.innerHTML = '';
        for (const step of buildActionSteps(rel, label)) {
            const li = document.createElement('li');
            li.textContent = step;
            actionNavigatorSteps.appendChild(li);
        }
    }

    function updateTaskActionContext(path, label) {
        currentTaskAction = {
            path: (path || '').trim(),
            label: (label || 'задача').trim() || 'задача'
        };

        if (taskActionLabel) {
            taskActionLabel.textContent = 'Задача: ' + currentTaskAction.label;
        }
        if (taskActionPath) {
            taskActionPath.textContent = 'Путь: ' + (currentTaskAction.path ? describePath(currentTaskAction.path) : '/');
        }

        if (batchGenerateMode) {
            if (currentTaskAction.path.includes('01_orders_and_appointments')) {
                batchGenerateMode.value = 'selected';
            } else if (currentTaskAction.path.includes('02_personnel/employees')) {
                batchGenerateMode.value = 'missing';
            }
        }
    }

    async function openChecklistByTaskAction() {
        const path = (currentTaskAction.path || '').trim();
        if (!path) {
            setEmployeeChecklistMsg('Сначала выберите задачу из списка задач/пробелов.', 'error');
            return;
        }

        const employeePath = extractEmployeeRelPath(path);
        focusEmployeeChecklistCard();

        if (employeePath) {
            if (employeeChecklistPath) {
                employeeChecklistPath.value = employeePath;
            }
            if (employeeChecklistEmployeeSelect) {
                employeeChecklistEmployeeSelect.value = employeePath;
            }
            await loadEmployeeChecklist();
            return;
        }

        if (path.includes('02_personnel/employees')) {
            if (employeeChecklistPath) {
                employeeChecklistPath.value = '';
            }
            await loadEmployeeChecklistOverview();
            return;
        }

        setEmployeeChecklistMsg('Для этой задачи откройте нужного сотрудника в чеклисте и запустите генерацию.', 'ok');
    }

    async function runTaskActionGeneration(mode) {
        await openChecklistByTaskAction();
        const employeePath = (employeeChecklistPath && employeeChecklistPath.value || '').trim();
        if (!employeePath) {
            setEmployeeChecklistMsg('Для запуска генерации выберите сотрудника в блоке ТБ-чеклиста.', 'error');
            return;
        }
        await generateEmployeeChecklist(mode);
    }

    function syncScannerRequirements() {
        const selectedDocType = scannerDocType ? scannerDocType.value : 'ORDER';
        const requiresEmployee = SCANNER_DOC_TYPES_REQUIRING_EMPLOYEE.has(selectedDocType);
        if (scannerEmployee) {
            scannerEmployee.required = requiresEmployee;
        }
        if (requiresEmployee) {
            if (scannerEmployee) {
                scannerEmployee.classList.add('required-input');
                scannerEmployee.placeholder = 'Код сотрудника обязателен для удостоверения/протокола';
            }
            if (!scannerEmployeeValue()) {
                scannerMsg.textContent = 'Для удостоверения/протокола укажите код сотрудника, затем запустите сканирование.';
            }
            if (scannerHint) {
                scannerHint.textContent = 'Режим удостоверения/протокола: код сотрудника обязателен.';
            }
            if (scannerEmployeeHint) {
                scannerEmployeeHint.textContent = 'Можно выбрать код из списка или открыть папку сотрудника для автоподстановки.';
            }
            return;
        }

        if (scannerEmployee) {
            scannerEmployee.classList.remove('required-input');
            scannerEmployee.placeholder = 'Код сотрудника (нужен только для удостоверения/протокола)';
        }
        if (scannerHint) {
            if (selectedDocType === 'INVOICE' || selectedDocType === 'UPD') {
                scannerHint.textContent = 'Режим бухгалтерских сканов: код сотрудника не обязателен, проверьте тег в теме (например: счет, упд).';
            } else {
                scannerHint.textContent = 'Код сотрудника можно не заполнять для текущего типа скана.';
            }
        }
        if (scannerEmployeeHint) {
            const sample = knownEmployeeIds.slice(0, 8).join(', ');
            scannerEmployeeHint.textContent = sample
                ? ('Выбор кода сотрудника из списка: ' + sample + (knownEmployeeIds.length > 8 ? '...' : ''))
                : 'Код сотрудника нужен только для удостоверения/протокола.';
        }
    }

    function setEmployeeChecklistMsg(text, level = 'info') {
        if (!employeeChecklistMsg) {
            return;
        }
        employeeChecklistMsg.textContent = text;
        employeeChecklistMsg.classList.remove('hint-error', 'hint-ok');
        if (level === 'error') {
            employeeChecklistMsg.classList.add('hint-error');
        }
        if (level === 'ok') {
            employeeChecklistMsg.classList.add('hint-ok');
        }
    }

    function collectSelectedChecklistCodes() {
        if (!employeeChecklistList) {
            return [];
        }
        return Array.from(employeeChecklistList.querySelectorAll('input[type="checkbox"][data-code]'))
            .filter((item) => item.checked)
            .map((item) => item.dataset.code || '')
            .filter(Boolean);
    }

    async function openChecklistPath(relPath, label) {
        const safePath = (relPath || '').trim();
        if (!safePath) {
            return;
        }
        const fileLike = /\.[a-z0-9]{1,8}$/i.test(safePath);
        if (fileLike) {
            try {
                await loadFile(safePath);
                renderActionNavigator(safePath, label || 'Файл из чеклиста');
                setInteractionHint('Открыт файл из чеклиста: ' + describePath(safePath), 'ok');
                return;
            } catch (_err) {
                const opened = openPrintPreview(safePath, false);
                if (opened) {
                    setInteractionHint('Файл из чеклиста открыт в браузерном просмотре.', 'ok');
                    return;
                }
            }
        }
        await openActionPath(safePath, label || 'Пункт чеклиста');
    }

    function renderEmployeeChecklist(data) {
        currentEmployeeChecklist = data;
        renderBatchDocCodesFromChecklist(data);
        if (!employeeChecklistList || !employeeChecklistSummary) {
            return;
        }

        if (employeeChecklistOverview) {
            employeeChecklistOverview.innerHTML = '';
        }

        const employeeName = data.employee_name || data.employee_id || 'Сотрудник';
        employeeChecklistSummary.textContent =
            'Сотрудник: ' + employeeName
            + '; роль: ' + (data.profession || '-')
            + '; готово: ' + data.ready_count + '/' + data.total_required
            + ' (' + data.progress_percent + '%).'
            + ' Не хватает: ' + data.missing_count + '.';

        employeeChecklistList.innerHTML = '';
        for (const item of (data.items || [])) {
            const row = document.createElement('div');
            row.className = 'employee-checklist-item ' + (item.ready ? 'is-ready' : 'is-missing');

            const top = document.createElement('div');
            top.className = 'employee-checklist-top';

            const checkbox = document.createElement('input');
            checkbox.type = 'checkbox';
            checkbox.dataset.code = item.code;
            checkbox.dataset.ready = item.ready ? '1' : '0';
            checkbox.checked = !item.ready;

            const title = document.createElement('div');
            title.className = 'employee-checklist-title';
            title.textContent = item.code + ' - ' + item.title;

            const status = document.createElement('span');
            status.className = 'employee-checklist-state ' + (item.ready ? 'ok' : 'warn');
            status.textContent = item.ready
                ? ('Готово ' + item.found_count + '/' + item.required_count)
                : ('Не хватает ' + item.found_count + '/' + item.required_count);

            const openBtn = document.createElement('button');
            openBtn.type = 'button';
            openBtn.className = 'btn mini secondary';
            openBtn.textContent = 'Открыть папку';
            openBtn.addEventListener('click', () => {
                openChecklistPath(item.folder_rel_path, item.title).catch((err) => {
                    setEmployeeChecklistMsg('Не удалось открыть путь чеклиста: ' + err, 'error');
                });
            });

            top.appendChild(checkbox);
            top.appendChild(title);
            top.appendChild(status);
            top.appendChild(openBtn);
            row.appendChild(top);

            const hint = document.createElement('div');
            hint.className = 'system-hint';
            hint.textContent = 'Подсказка: ' + (item.guidance || '-');
            row.appendChild(hint);

            const folderMeta = document.createElement('div');
            folderMeta.className = 'meta';
            folderMeta.textContent = 'Папка: ' + describePath(item.folder_rel_path || '');
            row.appendChild(folderMeta);

            if ((item.found_files || []).length) {
                const foundWrap = document.createElement('div');
                foundWrap.className = 'employee-found-files';
                const caption = document.createElement('div');
                caption.className = 'meta';
                caption.textContent = 'Найдено файлов: ' + item.found_files.length;
                foundWrap.appendChild(caption);

                for (const foundFile of item.found_files) {
                    const link = document.createElement('a');
                    link.href = '#';
                    link.className = 'todo-link';
                    link.textContent = describePath(foundFile);
                    link.addEventListener('click', (event) => {
                        event.preventDefault();
                        openChecklistPath(foundFile, item.title).catch((err) => {
                            setEmployeeChecklistMsg('Не удалось открыть файл из чеклиста: ' + err, 'error');
                        });
                    });
                    foundWrap.appendChild(link);
                }
                row.appendChild(foundWrap);
            }

            if ((item.related_files || []).length) {
                const relatedWrap = document.createElement('div');
                relatedWrap.className = 'employee-found-files';
                const relatedCaption = document.createElement('div');
                relatedCaption.className = 'meta';
                relatedCaption.textContent = 'Связанные документы по участию: ' + (item.related_count || item.related_files.length);
                relatedWrap.appendChild(relatedCaption);

                for (const relatedFile of item.related_files) {
                    const relatedLink = document.createElement('a');
                    relatedLink.href = '#';
                    relatedLink.className = 'todo-link';
                    relatedLink.textContent = 'Связь: ' + describePath(relatedFile);
                    relatedLink.addEventListener('click', (event) => {
                        event.preventDefault();
                        openChecklistPath(relatedFile, item.title).catch((err) => {
                            setEmployeeChecklistMsg('Не удалось открыть связанный документ: ' + err, 'error');
                        });
                    });
                    relatedWrap.appendChild(relatedLink);
                }
                row.appendChild(relatedWrap);
            }

            employeeChecklistList.appendChild(row);
        }
    }

    async function loadEmployeeChecklist() {
        if (!employeeChecklistPath) {
            return;
        }
        const selectedEmployeePath = (employeeChecklistEmployeeSelect && employeeChecklistEmployeeSelect.value || '').trim();
        const employeePath = (employeeChecklistPath.value || '').trim() || selectedEmployeePath;
        if (!employeePath) {
            await loadEmployeeChecklistOverview();
            return;
        }

        if (employeeChecklistPath.value !== employeePath) {
            employeeChecklistPath.value = employeePath;
        }
        if (employeeChecklistEmployeeSelect && employeeChecklistEmployeeSelect.value !== employeePath) {
            employeeChecklistEmployeeSelect.value = employeePath;
        }

        setEmployeeChecklistMsg('Проверка комплектности документов сотрудника...');
        const query = new URLSearchParams();
        query.set('employee_rel_path', employeePath);
        const selectedProfession = (employeeChecklistProfession && employeeChecklistProfession.value || '').trim();
        const professionText = selectedProfession && selectedProfession !== 'all'
            ? inferProfessionByGroupKey(selectedProfession)
            : '';
        if (professionText) {
            query.set('profession', professionText);
        }

        const data = await api(
            '/arm/employees/checklist?' + query.toString(),
            undefined,
            TIMEOUTS_MS.checklist,
            'Проверка чеклиста сотрудника заняла слишком много времени.'
        );

        renderEmployeeChecklist(data);
        setEmployeeChecklistMsg('Чеклист обновлен: готово ' + data.ready_count + '/' + data.total_required + '.', 'ok');
    }

    async function generateEmployeeChecklist(mode) {
        if (!employeeChecklistPath) {
            return;
        }
        const selectedEmployeePath = (employeeChecklistEmployeeSelect && employeeChecklistEmployeeSelect.value || '').trim();
        const employeePath = (employeeChecklistPath.value || '').trim() || selectedEmployeePath;
        if (!employeePath) {
            setEmployeeChecklistMsg('Укажите путь сотрудника перед генерацией документов.', 'error');
            if (employeeChecklistEmployeeSelect) {
                employeeChecklistEmployeeSelect.focus();
            } else {
                employeeChecklistPath.focus();
            }
            return;
        }

        if (employeeChecklistPath.value !== employeePath) {
            employeeChecklistPath.value = employeePath;
        }

        const selectedProfession = (employeeChecklistProfession && employeeChecklistProfession.value || '').trim();
        const professionText = selectedProfession && selectedProfession !== 'all'
            ? inferProfessionByGroupKey(selectedProfession)
            : null;
        const orderDate = getChecklistOrderDate();
        if (!orderDate) {
            setEmployeeChecklistMsg('Дата приказа должна быть в формате ДД.ММ.ГГГГ.', 'error');
            if (employeeChecklistOrderDate) {
                employeeChecklistOrderDate.focus();
            }
            return;
        }

        const payload = {
            employee_rel_path: employeePath,
            profession: professionText,
            order_date: orderDate,
            mode: mode,
            codes: mode === 'selected' ? collectSelectedChecklistCodes() : [],
            overwrite: Boolean(employeeChecklistOverwrite && employeeChecklistOverwrite.checked)
        };

        if (mode === 'selected' && !payload.codes.length) {
            setEmployeeChecklistMsg('Выберите хотя бы один пункт чеклиста для точечной генерации.', 'error');
            return;
        }

        setEmployeeChecklistMsg('Генерация черновиков по ТБ-чеклисту...');
        const result = await api(
            '/arm/employees/checklist/generate',
            {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify(payload)
            },
            TIMEOUTS_MS.checklistGenerate,
            'Генерация черновиков заняла слишком много времени.'
        );

        const createdCount = (result.created_files || []).length;
        const skippedCount = (result.skipped_files || []).length;
        setEmployeeChecklistMsg(
            (result.message || 'Генерация завершена')
            + ' Создано: ' + createdCount
            + '; пропущено: ' + skippedCount + '.',
            createdCount > 0 ? 'ok' : 'error'
        );

        if (createdCount > 0) {
            setInteractionHint('Черновики созданы. Проверьте папку templates_to_print сотрудника и дооформите документы.', 'ok');
        }

        await loadEmployeeChecklist();
    }

    function selectOnlyMissingInChecklist() {
        if (!employeeChecklistList) {
            return;
        }
        const checkboxes = employeeChecklistList.querySelectorAll('input[type="checkbox"][data-code]');
        for (const checkbox of checkboxes) {
            checkbox.checked = checkbox.dataset.ready !== '1';
        }
        setEmployeeChecklistMsg('Отмечены только отсутствующие документы.', 'ok');
    }

    async function withBusy(button, work) {
        button.disabled = true;
        button.classList.add('is-busy');
        try {
            return await work();
        } finally {
            button.disabled = false;
            button.classList.remove('is-busy');
        }
    }

    function initCollapsibleCards() {
        const configs = [
            { id: 'metricsCard', key: 'metrics', defaultCollapsed: false },
            { id: 'gapsCard', key: 'gaps', defaultCollapsed: false },
            { id: 'todoCard', key: 'todo', defaultCollapsed: false },
            { id: 'structureCard', key: 'structure', defaultCollapsed: false },
            { id: 'previewCard', key: 'preview', defaultCollapsed: false },
            { id: 'scannerCard', key: 'scanner', defaultCollapsed: false },
            { id: 'exportCard', key: 'export', defaultCollapsed: true },
            { id: 'assistantCard', key: 'assistant', defaultCollapsed: true },
        ];

        for (const config of configs) {
            const card = document.getElementById(config.id);
            if (!card || card.dataset.collapseReady === '1') {
                continue;
            }

            const heading = card.querySelector(':scope > h2');
            if (!heading) {
                continue;
            }

            const head = document.createElement('div');
            head.className = 'card-collapse-head';
            card.insertBefore(head, heading);
            head.appendChild(heading);

            const toggle = document.createElement('button');
            toggle.type = 'button';
            toggle.className = 'btn secondary mini';
            head.appendChild(toggle);

            const body = document.createElement('div');
            body.className = 'card-collapse-body';
            const children = Array.from(card.children);
            let move = false;
            for (const child of children) {
                if (child === head) {
                    move = true;
                    continue;
                }
                if (move) {
                    body.appendChild(child);
                }
            }
            card.appendChild(body);

            const storageKey = 'arm.card.collapse.' + config.key;
            let collapsed = Boolean(config.defaultCollapsed);
            try {
                const stored = localStorage.getItem(storageKey);
                if (stored !== null) {
                    collapsed = stored === '1';
                }
            } catch (_err) {
                // localStorage may be unavailable in restricted environments
            }

            const apply = (isCollapsed) => {
                body.hidden = isCollapsed;
                toggle.textContent = isCollapsed ? 'Развернуть' : 'Свернуть';
            };

            apply(collapsed);
            toggle.addEventListener('click', () => {
                collapsed = !collapsed;
                apply(collapsed);
                try {
                    localStorage.setItem(storageKey, collapsed ? '1' : '0');
                } catch (_err) {
                    // ignore storage write failures
                }
            });

            card.dataset.collapseReady = '1';
        }
    }

    async function refreshLLMStatus() {
        const hasBadge = llmBadge instanceof HTMLElement;
        const hasDesc = llmDesc instanceof HTMLElement;

        try {
            const statusData = await api('/local-llm/status', undefined, TIMEOUTS_MS.status);
            const runtimeData = await api('/local-llm/runtime', undefined, TIMEOUTS_MS.status);
            if (statusData.is_reachable) {
                if (hasBadge) {
                    llmBadge.textContent = 'ДОСТУПНА';
                    llmBadge.classList.remove('llm-down');
                    llmBadge.classList.add('llm-ok');
                }
                if (hasDesc) {
                    llmDesc.textContent = 'Модель: ' + (statusData.default_model || '-') + '; ускорение: ' + (runtimeData.acceleration || '-') + '; активных моделей: ' + (runtimeData.running_models_count ?? 0);
                }
            } else {
                if (hasBadge) {
                    llmBadge.textContent = 'НЕДОСТУПНА';
                    llmBadge.classList.remove('llm-ok');
                    llmBadge.classList.add('llm-down');
                }
                if (hasDesc) {
                    llmDesc.textContent = 'Локальная LLM недоступна. Проверьте процесс Ollama.';
                }
            }
        } catch (err) {
            if (hasBadge) {
                llmBadge.textContent = 'ОШИБКА';
                llmBadge.classList.remove('llm-ok');
                llmBadge.classList.add('llm-down');
            }
            if (hasDesc) {
                llmDesc.textContent = 'Ошибка проверки LLM: ' + err;
            }
        }
    }

    const PROFILE_LABELS = {
        fast: 'Быстрый',
        balanced: 'Сбалансированный',
        quality: 'Качество',
    };

    function profileLabel(profileId) {
        const key = String(profileId || '').trim();
        return PROFILE_LABELS[key] || key || '-';
    }

    async function requestAssist(questionText, profileId) {
        return api(
            '/arm/assist',
            {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ question: questionText, profile: profileId, allow_fallback: true })
            },
            TIMEOUTS_MS.llm,
            'Локальная LLM не успела ответить за отведенное время. Попробуйте профиль «Быстрый» или повторите запрос позже.'
        );
    }

    function formatAssistMeta(data) {
        return 'модель=' + (data.model || '-') + '; профиль=' + profileLabel(data.used_profile) + '; резерв=' + String(data.fallback_used) + '; токенов/сек=' + (data.eval_tokens_per_sec ?? '-');
    }

    function setVoiceHint(text) {
        if (!voiceHint) {
            return;
        }
        voiceHint.textContent = text || '';
    }

    function mergeVoiceChunks(chunks) {
        const totalLength = chunks.reduce((sum, chunk) => sum + chunk.length, 0);
        const result = new Float32Array(totalLength);
        let offset = 0;
        for (const chunk of chunks) {
            result.set(chunk, offset);
            offset += chunk.length;
        }
        return result;
    }

    function encodeWav(samples, sampleRate) {
        const bytesPerSample = 2;
        const blockAlign = bytesPerSample;
        const buffer = new ArrayBuffer(44 + samples.length * bytesPerSample);
        const view = new DataView(buffer);

        const writeString = (offset, value) => {
            for (let i = 0; i < value.length; i += 1) {
                view.setUint8(offset + i, value.charCodeAt(i));
            }
        };

        writeString(0, 'RIFF');
        view.setUint32(4, 36 + samples.length * bytesPerSample, true);
        writeString(8, 'WAVE');
        writeString(12, 'fmt ');
        view.setUint32(16, 16, true);
        view.setUint16(20, 1, true);
        view.setUint16(22, 1, true);
        view.setUint32(24, sampleRate, true);
        view.setUint32(28, sampleRate * blockAlign, true);
        view.setUint16(32, blockAlign, true);
        view.setUint16(34, 16, true);
        writeString(36, 'data');
        view.setUint32(40, samples.length * bytesPerSample, true);

        let offset = 44;
        for (let i = 0; i < samples.length; i += 1) {
            const sample = Math.max(-1, Math.min(1, samples[i]));
            view.setInt16(offset, sample < 0 ? sample * 0x8000 : sample * 0x7FFF, true);
            offset += 2;
        }

        return new Blob([view], { type: 'audio/wav' });
    }

    function cleanupVoiceFallbackCapture() {
        if (voiceFallbackProcessor) {
            try {
                voiceFallbackProcessor.disconnect();
            } catch (_err) {
                // ignore node cleanup errors
            }
            voiceFallbackProcessor.onaudioprocess = null;
            voiceFallbackProcessor = null;
        }

        if (voiceFallbackSource) {
            try {
                voiceFallbackSource.disconnect();
            } catch (_err) {
                // ignore node cleanup errors
            }
            voiceFallbackSource = null;
        }

        if (voiceFallbackStream) {
            for (const track of voiceFallbackStream.getTracks()) {
                track.stop();
            }
            voiceFallbackStream = null;
        }

        if (voiceFallbackAudioCtx) {
            voiceFallbackAudioCtx.close().catch(() => {});
            voiceFallbackAudioCtx = null;
        }
    }

    async function transcribeVoiceWithGoogle(wavBlob) {
        const form = new FormData();
        form.append('audio', wavBlob, 'voice.wav');

        const response = await fetch('/arm/speech/google-transcribe?language=ru-RU', {
            method: 'POST',
            body: form,
        });
        const data = await response.json().catch(() => ({}));
        if (!response.ok) {
            throw new Error(data.detail || response.statusText || 'Ошибка сервиса распознавания');
        }
        return data;
    }

    async function startGoogleVoiceFallbackCapture() {
        if (!canGoogleSpeechFallback) {
            setVoiceHint('Google fallback недоступен: браузер не поддерживает захват аудио с микрофона.');
            return;
        }

        try {
            const AudioCtor = window.AudioContext || window.webkitAudioContext;
            voiceFallbackStream = await navigator.mediaDevices.getUserMedia({ audio: true });
            voiceFallbackAudioCtx = new AudioCtor({ sampleRate: 16000 });
            voiceFallbackSampleRate = Number(voiceFallbackAudioCtx.sampleRate || 16000);
            voiceFallbackChunks = [];

            voiceFallbackSource = voiceFallbackAudioCtx.createMediaStreamSource(voiceFallbackStream);
            voiceFallbackProcessor = voiceFallbackAudioCtx.createScriptProcessor(4096, 1, 1);
            voiceFallbackProcessor.onaudioprocess = (event) => {
                const input = event.inputBuffer.getChannelData(0);
                voiceFallbackChunks.push(new Float32Array(input));
            };

            voiceFallbackSource.connect(voiceFallbackProcessor);
            voiceFallbackProcessor.connect(voiceFallbackAudioCtx.destination);

            voiceFallbackMode = true;
            voiceListening = true;
            updateVoiceButton();
            setVoiceHint('Слушаю через Google fallback... Нажмите кнопку еще раз, чтобы остановить и распознать.');
        } catch (err) {
            cleanupVoiceFallbackCapture();
            voiceFallbackMode = false;
            voiceListening = false;
            updateVoiceButton();
            setVoiceHint('Не удалось начать запись через Google fallback: ' + err);
        }
    }

    async function stopGoogleVoiceFallbackCapture(transcribe = true, message = '') {
        const chunks = voiceFallbackChunks.slice();
        const sampleRate = voiceFallbackSampleRate;
        voiceFallbackChunks = [];

        voiceFallbackMode = false;
        voiceListening = false;
        cleanupVoiceFallbackCapture();
        updateVoiceButton();

        if (!transcribe) {
            if (message) {
                setVoiceHint(message);
            }
            return;
        }

        if (!chunks.length) {
            setVoiceHint('Запись слишком короткая. Попробуйте еще раз и говорите дольше.');
            return;
        }

        try {
            setVoiceHint('Обработка записи и отправка на распознавание Google...');
            const merged = mergeVoiceChunks(chunks);
            const wavBlob = encodeWav(merged, sampleRate || 16000);
            const data = await transcribeVoiceWithGoogle(wavBlob);

            if (data.ok && data.text) {
                appendVoiceToQuestion(data.text);
                setVoiceHint('Речь распознана через Google fallback и добавлена в поле вопроса.');
                return;
            }
            setVoiceHint(data.message || 'Распознавание не дало результата. Повторите запись.');
        } catch (err) {
            setVoiceHint('Ошибка Google fallback: ' + err);
        }
    }

    function updateVoiceButton() {
        if (!voiceBtn) {
            return;
        }

        const useGoogleMode = canGoogleSpeechFallback;

        if (!voiceRecognition && !useGoogleMode) {
            voiceBtn.disabled = true;
            voiceBtn.textContent = 'Голос недоступен';
            return;
        }

        voiceBtn.disabled = false;
        if (voiceListening) {
            voiceBtn.textContent = 'Остановить запись';
            return;
        }

        voiceBtn.textContent = useGoogleMode ? 'Голосовой ввод (Google)' : 'Голосовой ввод';
    }

    function appendVoiceToQuestion(transcript) {
        const safe = (transcript || '').trim();
        if (!safe || !question) {
            return;
        }
        const current = (question.value || '').trim();
        question.value = current ? (current + '\n' + safe) : safe;
        question.focus();
    }

    async function stopVoiceInput(message) {
        if (voiceFallbackMode) {
            await stopGoogleVoiceFallbackCapture(false, message || 'Голосовой ввод остановлен.');
            return;
        }

        voiceListening = false;
        if (voiceRecognition) {
            try {
                voiceRecognition.stop();
            } catch (_err) {
                // ignore stop race errors from browser speech API
            }
        }
        updateVoiceButton();
        if (message) {
            setVoiceHint(message);
        }
    }

    function initVoiceInput() {
        if (canGoogleSpeechFallback) {
            voiceRecognition = null;
            updateVoiceButton();
            setVoiceHint('Включен серверный Google-режим голосового ввода. Нажмите кнопку, скажите фразу и нажмите кнопку повторно для распознавания.');
            return;
        }

        const SpeechCtor = window.SpeechRecognition || window.webkitSpeechRecognition;
        if (!SpeechCtor) {
            voiceRecognition = null;
            updateVoiceButton();
            setVoiceHint('Голосовой ввод не поддерживается в этом браузере. Откройте страницу в Chrome или Edge.');
            return;
        }

        voiceRecognition = new SpeechCtor();
        voiceRecognition.lang = 'ru-RU';
        voiceRecognition.interimResults = false;
        voiceRecognition.continuous = false;
        voiceRecognition.maxAlternatives = 1;

        voiceRecognition.onresult = (event) => {
            let recognized = '';
            for (let i = event.resultIndex; i < event.results.length; i += 1) {
                const result = event.results[i];
                if (!result || !result[0]) {
                    continue;
                }
                recognized += result[0].transcript + ' ';
            }
            const safeRecognized = recognized.trim();
            if (!safeRecognized) {
                return;
            }
            voiceLastTranscript = safeRecognized;
            appendVoiceToQuestion(safeRecognized);
            voiceListening = false;
            updateVoiceButton();
            setVoiceHint('Речь распознана и добавлена в поле вопроса.');
        };

        voiceRecognition.onerror = (event) => {
            const code = String(event && event.error ? event.error : 'unknown');
            voiceListening = false;
            updateVoiceButton();
            if (code === 'not-allowed' || code === 'service-not-allowed') {
                setVoiceHint('Нет доступа к микрофону. Разрешите доступ к микрофону для сайта.');
                return;
            }
            if (code === 'no-speech') {
                setVoiceHint('Речь не распознана. Повторите попытку.');
                return;
            }
            setVoiceHint('Ошибка голосового ввода: ' + code + '.');
        };

        voiceRecognition.onend = () => {
            if (voiceListening) {
                voiceListening = false;
                updateVoiceButton();
                if (!voiceLastTranscript) {
                    setVoiceHint('Речь не распознана. Попробуйте еще раз и говорите чуть медленнее.');
                }
                return;
            }
            updateVoiceButton();
        };

        updateVoiceButton();
        setVoiceHint('Голосовой ввод готов. Нажмите кнопку и говорите.');
    }

    async function toggleVoiceInput() {
        if (canGoogleSpeechFallback) {
            if (voiceListening) {
                await stopGoogleVoiceFallbackCapture(true);
            } else {
                await startGoogleVoiceFallbackCapture();
            }
            return;
        }

        if (!voiceRecognition) {
            setVoiceHint('Голосовой ввод недоступен в текущем браузере.');
            return;
        }

        if (voiceListening) {
            await stopVoiceInput('Голосовой ввод остановлен.');
            return;
        }
        try {
            voiceLastTranscript = '';
            voiceRecognition.start();
            voiceListening = true;
            updateVoiceButton();
            setVoiceHint('Слушаю... Говорите, текст добавится в поле вопроса.');
        } catch (err) {
            voiceListening = false;
            updateVoiceButton();
            setVoiceHint('Не удалось запустить голосовой ввод: ' + err);
        }
    }

    async function sendAssist() {
        if (sendBtn.disabled) {
            return;
        }

        if (voiceListening) {
            await stopVoiceInput('Голосовой ввод остановлен перед отправкой запроса.');
        }

        const text = (question.value || '').trim();
        if (!text) {
            answer.textContent = 'Введите вопрос для ассистента.';
            return;
        }

        await withBusy(sendBtn, async () => {
            answer.textContent = 'Выполняется запрос к локальной LLM...';
            try {
                let requestedProfile = profile.value || 'fast';
                let data;
                try {
                    data = await requestAssist(text, requestedProfile);
                } catch (firstError) {
                    const message = String(firstError || '');
                    if (requestedProfile !== 'fast' && message.toLowerCase().includes('не успела')) {
                        answer.textContent = 'Профиль ' + profileLabel(requestedProfile) + ' отвечает слишком долго. Пробую повторно в режиме «Быстрый»...';
                        requestedProfile = 'fast';
                        profile.value = 'fast';
                        data = await requestAssist(text, 'fast');
                    } else {
                        throw firstError;
                    }
                }

                answer.textContent = formatAssistMeta(data) + '\\n\\n' + (data.response || '');
                setInteractionHint('Ответ ассистента получен. Режим: ' + profileLabel(data.used_profile || requestedProfile) + '.', 'ok');
            } catch (err) {
                answer.textContent = 'Ошибка LLM: ' + err;
                setInteractionHint('Запрос к ассистенту завершился ошибкой: ' + err, 'error');
            }
        });
    }

    async function loadTree(relPath, options = {}) {
        const rememberHistory = options.rememberHistory !== false;
        const requestedPath = normalizeRelPath(relPath);
        const data = await api(
            '/arm/fs/tree?rel_path=' + encodeURIComponent(requestedPath),
            undefined,
            TIMEOUTS_MS.tree,
            'Загрузка структуры занимает слишком много времени. Повторите попытку.'
        );

        const rel = normalizeRelPath(data.rel_path || requestedPath);
        if (treePathInput) {
            treePathInput.dataset.currentPath = rel;
            if (!(document.activeElement === treePathInput)) {
                treePathInput.value = rel;
            }
        }
        if (uploadDir && uploadAutoUseTree && uploadAutoUseTree.checked) {
            uploadDir.value = rel;
        }

        treeView.innerHTML = '';
        const ul = document.createElement('ul');
        const entries = data.entries || [];
        const isEmployeeCardPath = /^02_personnel\/employees\/[^/]+$/i.test(rel);
        const requiredEmployeeFolders = [
            '01_identity_and_contract',
            '02_admission_orders',
            '03_briefings_and_training',
            '04_attestation_and_certificates',
            '05_ppe_issue',
            '06_permits_and_work_admission',
            '07_medical_and_first_aid',
            '07_templates_to_print'
        ];
        const existingEmployeeFolders = new Set(
            entries
                .filter((entry) => entry && entry.is_dir)
                .map((entry) => String(entry.name || '').trim())
        );
        let missingEmployeeFolders = [];

        if (rel) {
            const parentLi = document.createElement('li');
            const parentLink = document.createElement('a');
            parentLink.href = '#';
            parentLink.dataset.path = parentRelPath(rel);
            parentLink.dataset.isDir = '1';
            parentLink.textContent = '[..] На уровень выше';
            parentLi.appendChild(parentLink);
            ul.appendChild(parentLi);
        }

        for (const entry of entries) {
            const li = document.createElement('li');
            const a = document.createElement('a');
            a.href = '#';
            a.dataset.path = entry.rel_path;
            a.dataset.isDir = entry.is_dir ? '1' : '0';
            let prettyName = prettyNodeName(entry.name);
            if (entry.is_dir && /^02_personnel\/employees$/i.test(rel)) {
                const displayName = getEmployeeDisplayName(entry.rel_path);
                if (displayName) {
                    const employeeId = parseEmployeeIdFromName(entry.name || '');
                    prettyName = employeeId ? ('[' + employeeId + '] ' + displayName) : displayName;
                }
            }
            if (entry.is_dir) {
                const isRequiredEmployeeFolder = isEmployeeCardPath && requiredEmployeeFolders.includes(String(entry.name || '').trim());
                if (isRequiredEmployeeFolder) {
                    li.classList.add('tree-item-ok');
                }
                a.textContent = '[Папка] ' + prettyName;
            } else {
                const orderFileLabel = parseOrderLabelFromPath(entry.rel_path, entry.name);
                if (orderFileLabel) {
                    prettyName = orderFileLabel;
                }
                const ext = fileExtensionFromName(entry.name);
                const typeLabel = ext ? ext.toUpperCase() : 'FILE';
                if (ext === 'md' || ext === 'txt' || ext === 'csv' || ext === 'json' || ext === 'yml' || ext === 'yaml' || ext === 'py' || ext === 'log' || ext === 'ini') {
                    li.classList.add('tree-item-editable');
                }
                a.textContent = '[Файл ' + typeLabel + '] ' + prettyName;
            }
            li.appendChild(a);
            ul.appendChild(li);
        }

        if (isEmployeeCardPath) {
            missingEmployeeFolders = requiredEmployeeFolders.filter((folderName) => !existingEmployeeFolders.has(folderName));
            for (const folderName of missingEmployeeFolders) {
                const li = document.createElement('li');
                li.className = 'tree-item-missing';
                li.textContent = '[Не хватает] ' + (humanLabel(folderName) || prettifyUnknownName(folderName) || folderName);
                ul.appendChild(li);
            }
        }

        if (!entries.length) {
            const li = document.createElement('li');
            li.textContent = '[Пусто] В этой папке пока нет файлов или подпапок.';
            ul.appendChild(li);
        }

        treeView.appendChild(ul);
        const localizedPath = describePath(rel);
        if (treePathInput && !(document.activeElement === treePathInput)) {
            treePathInput.placeholder = rel
                ? ('Открыт раздел: ' + localizedPath)
                : 'Относительный путь (если нужен ручной переход)';
        }

        if (isEmployeeCardPath && missingEmployeeFolders.length) {
            setTreeHint(
                'Открыта папка: ' + localizedPath
                + '. Не хватает обязательных разделов: ' + missingEmployeeFolders.length
                + '. Красным отмечены отсутствующие разделы.'
            );
        } else if (isEmployeeCardPath) {
            setTreeHint('Открыта папка: ' + localizedPath + '. Обязательные разделы сотрудника по структуре на месте.');
        } else if (entries.length) {
            setTreeHint('Открыта папка: ' + localizedPath + '. Элементов: ' + entries.length + '. Следующий шаг: выберите папку или файл из списка ниже.');
        } else if (rel.toLowerCase().includes('04_journals')) {
            setTreeHint('Открыта папка журнала: ' + localizedPath + '. Папка пока пуста, поэтому визуально список не меняется.');
            setInteractionHint('Журнал открыт, но в разделе пока нет файлов. Добавьте файл или выберите соседнюю подпапку.', 'error');
        } else {
            setTreeHint('Открыта папка: ' + localizedPath + '. Внутри пока нет файлов.');
        }

        if (rememberHistory) {
            pushTreeHistory(rel);
        }
        renderTreeBreadcrumb(rel);
        updateTreeNavButtons(rel);
        applyEmployeeIdFromPath(rel);
        syncEmployeeChecklistPath(rel);
        renderActionNavigator(rel, 'Открытие пути');
    }

    async function loadFile(relPath) {
        const data = await api('/arm/fs/file?rel_path=' + encodeURIComponent(relPath));
        currentFile = data.rel_path;
        filePathLabel.textContent = describePath(data.rel_path);
        fileEditor.value = data.content;
        fileEditor.readOnly = false;
        updateFileSourceVisibility();
        refreshRenderedPreview(data.rel_path);
        if (fileSaveBtn) {
            fileSaveBtn.disabled = false;
        }
        applyEmployeeIdFromPath(data.rel_path);
        syncEmployeeChecklistPath(data.rel_path);
        setFileMeta('Кодировка: ' + data.encoding + '; длина: ' + String(data.content.length));
        focusPreviewCard(true);
        setInteractionHint('Открыт текстовый файл: можно редактировать и сохранять в блоке «Предпросмотр, правка, печать».', 'ok');
    }

    function selectBinaryFile(relPath, sourceLabel) {
        const safePath = normalizeRelPath(relPath);
        if (!safePath) {
            return;
        }
        currentFile = safePath;
        if (filePathLabel) {
            filePathLabel.textContent = describePath(safePath);
        }
        if (fileEditor) {
            fileEditor.value = '';
            fileEditor.readOnly = true;
        }
        fileSourceVisible = false;
        updateFileSourceVisibility();
        refreshRenderedPreview(safePath);
        if (fileSaveBtn) {
            fileSaveBtn.disabled = true;
        }
        setFileMeta('Выбран бинарный файл. Используйте «Скачать» или «Печать».');
        updateTaskActionContext(safePath, sourceLabel || 'Файл');
        renderActionNavigator(safePath, sourceLabel || 'Файл');
        applyEmployeeIdFromPath(safePath);
        syncEmployeeChecklistPath(safePath);
        focusPreviewCard(true);
    }

    async function saveCurrentFile() {
        if (!currentFile) {
            setFileMeta('Выберите файл из дерева.');
            return;
        }
        const data = await api('/arm/fs/file', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify({ rel_path: currentFile, content: fileEditor.value })
        });
        setFileMeta(data.message || 'Сохранено.');
        refreshRenderedPreview(currentFile);
    }

    function downloadCurrentFile() {
        if (!currentFile) {
            setFileMeta('Выберите файл из дерева.');
            return;
        }
        window.open('/arm/fs/download?rel_path=' + encodeURIComponent(currentFile), '_blank');
    }

    function openPrintPreview(relPath, autoPrint) {
        const safePath = normalizeRelPath(relPath);
        if (!safePath) {
            return false;
        }
        const url = '/arm/fs/print-preview?rel_path=' + encodeURIComponent(safePath)
            + '&auto_print=' + (autoPrint ? '1' : '0');
        const popup = window.open(url, '_blank');
        return Boolean(popup);
    }

    async function printCurrentFile() {
        if (!currentFile) {
            setFileMeta('Выберите файл из дерева.');
            return;
        }
        const opened = openPrintPreview(currentFile, true);
        if (!opened) {
            setFileMeta('Браузер заблокировал всплывающее окно печати. Разрешите всплывающие окна для этого сайта.');
            return;
        }
        setFileMeta('Открыт браузерный предпросмотр печати. Если диалог печати не открылся автоматически, нажмите «Печать» в новом окне.');
    }

    async function refreshScannerDevices() {
        try {
            const data = await api(
                '/arm/scanner/devices',
                undefined,
                TIMEOUTS_MS.tree,
                'Список сканеров загружается слишком долго. Проверьте подключение устройства.'
            );
            scannerDevices.innerHTML = '';
            if (!data.devices.length) {
                const option = document.createElement('option');
                option.value = '1';
                option.textContent = 'Сканеры не найдены';
                scannerDevices.appendChild(option);
                return;
            }
            for (const device of data.devices) {
                const option = document.createElement('option');
                option.value = String(device.index);
                option.textContent = '[' + device.index + '] ' + device.name;
                scannerDevices.appendChild(option);
            }
        } catch (err) {
            scannerMsg.textContent = 'Ошибка списка сканеров: ' + err;
        }
    }

    async function scanToInbox() {
        resetScannerProgress('подготовка к сканированию');
        appendScannerTimeline('Проверка параметров сканирования.');
        setScannerProgress(10, 'подготовка параметров');

        const selectedDocType = scannerDocType ? scannerDocType.value : 'ORDER';
        const employeeId = scannerEmployeeValue();
        if (SCANNER_DOC_TYPES_REQUIRING_EMPLOYEE.has(selectedDocType) && !employeeId) {
            scannerMsg.textContent = 'Сканирование остановлено: для удостоверения/протокола требуется код сотрудника.';
            setInteractionHint('Перед сканированием удостоверения/протокола заполните поле кода сотрудника.', 'error');
            if (scannerEmployee) {
                scannerEmployee.focus();
            }
            setScannerProgress(0, 'ожидание');
            return;
        }

        const currentInputValue = scannerEmployee && typeof scannerEmployee.value === 'string'
            ? scannerEmployee.value.trim()
            : '';
        if (employeeId && !currentInputValue) {
            setScannerEmployeeValue(employeeId);
        }

        scannerMsg.textContent = 'Сканирование...';
        appendScannerTimeline('Отправка задания на устройство сканера.');
        setScannerProgress(35, 'ожидание ответа сканера');

        const data = await api('/arm/scanner/scan-to-inbox', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify({
                doc_type: selectedDocType,
                subject: (scannerSubject && scannerSubject.value) || 'scan_from_arm',
                employee_id: employeeId || null,
                device_index: Number((scannerDevices && scannerDevices.value) || 1),
                image_format: 'jpg',
                scan_profile: currentScannerProfile,
                dpi: 300,
                grayscale: false
            })
        }, TIMEOUTS_MS.scanner, 'Сканирование заняло слишком много времени. Скан мог завершиться, но сервер не успел вернуть ответ.');
        scannerMsg.textContent = data.message || 'Скан добавлен во входящую папку.';
        setInteractionHint('Сканирование выполнено. Следующий шаг: нажмите «Распознать и разложить».', 'ok');
        setScannerProgress(100, 'сканирование завершено');
        appendScannerTimeline('Скан помещен во входящую папку.');
    }

    async function ingestInbox() {
        resetScannerProgress('подготовка OCR/разбора');
        appendScannerTimeline('Запуск OCR и разложения входящих сканов.');
        setScannerProgress(20, 'OCR: запуск');

        scannerMsg.textContent = 'OCR и разбор входящих сканов...';
        const data = await api('/arm/scan/ingest', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify({ enable_ocr: true, ocr_lang: 'rus+eng', max_pdf_pages: 4 })
        }, TIMEOUTS_MS.ingest, 'OCR/разбор выполняется дольше обычного. Подождите и повторите проверку ручного разбора.');
        scannerMsg.textContent = 'Обработано: архив=' + data.archived + ', ручной разбор=' + data.manual_review;
        setInteractionHint('OCR завершен: архив=' + data.archived + ', ручной разбор=' + data.manual_review + '.', 'ok');
        setScannerProgress(85, 'OCR завершен, формируется сводка');
        appendScannerTimeline('Архивировано: ' + data.archived + '; ручной разбор: ' + data.manual_review + '.');
        setScannerProgress(100, 'разбор завершен');
    }

    function formatManualReviewConfidence(value) {
        const numeric = Number(value);
        if (!Number.isFinite(numeric)) {
            return '-';
        }
        return numeric.toFixed(2);
    }

    function getSelectedManualReviewRow() {
        return manualReviewRows.find((row) => row.rel_path === selectedManualReviewPath) || null;
    }

    function updateManualReviewSelectionInfo(forceSuggestedPath = false) {
        const selectedRow = getSelectedManualReviewRow();
        if (!selectedRow) {
            if (manualReviewSelection) {
                manualReviewSelection.textContent = 'Файл из ручного разбора не выбран.';
            }
            if (manualReviewMovePath && forceSuggestedPath) {
                manualReviewMovePath.value = '';
            }
            return;
        }

        if (manualReviewSelection) {
            manualReviewSelection.textContent = 'Выбран: '
                + describePath(selectedRow.rel_path)
                + '; тип=' + (selectedRow.predicted_doc_type || '-')
                + '; уверенность=' + formatManualReviewConfidence(selectedRow.confidence) + '.';
        }
        if (manualReviewMovePath && (forceSuggestedPath || !manualReviewMovePath.value.trim())) {
            manualReviewMovePath.value = (selectedRow.suggested_target_rel_path || '').trim();
        }
    }

    function isManualReviewUnresolved(row) {
        const confidence = Number(row && row.confidence);
        const predictedType = String((row && row.predicted_doc_type) || '').toLowerCase();
        return !Number.isFinite(confidence) || confidence < 0.7 || predictedType === 'unknown';
    }

    function sortManualReviewRows(rows) {
        const mode = scannerSortMode ? scannerSortMode.value : 'unresolved';
        const next = Array.isArray(rows) ? rows.slice() : [];
        if (mode === 'name') {
            return next.sort((left, right) => String(left.rel_path || '').localeCompare(String(right.rel_path || ''), 'ru'));
        }
        return next.sort((left, right) => {
            const unresolvedDelta = Number(isManualReviewUnresolved(right)) - Number(isManualReviewUnresolved(left));
            if (unresolvedDelta !== 0) {
                return unresolvedDelta;
            }
            const leftConfidence = Number(left.confidence);
            const rightConfidence = Number(right.confidence);
            if (Number.isFinite(leftConfidence) && Number.isFinite(rightConfidence)) {
                if (leftConfidence !== rightConfidence) {
                    return leftConfidence - rightConfidence;
                }
            }
            return String(left.rel_path || '').localeCompare(String(right.rel_path || ''), 'ru');
        });
    }

    function deriveManualReviewRoute(typeCode, row) {
        const normalizedType = String(typeCode || '').toUpperCase();
        const base = MANUAL_REVIEW_TARGET_BY_SCAN_TYPE[normalizedType] || MANUAL_REVIEW_TARGET_BY_SCAN_TYPE.OTHER;
        const fileName = String(((row && row.rel_path) || '').split('/').pop() || '').trim();
        return base ? (base + (fileName ? ('/' + fileName) : '')) : fileName;
    }

    function renderManualReviewRows(rows) {
        if (!manualReviewList) {
            return;
        }

        manualReviewList.innerHTML = '';
        if (!rows.length) {
            selectedManualReviewPath = '';
            manualReviewList.textContent = 'Папка ручного разбора пуста.';
            updateManualReviewSelectionInfo(true);
            return;
        }

        if (!rows.some((row) => row.rel_path === selectedManualReviewPath)) {
            selectedManualReviewPath = rows[0].rel_path;
        }

        for (const row of rows) {
            const selected = row.rel_path === selectedManualReviewPath;
            const item = document.createElement('div');
            item.className = 'manual-review-item' + (selected ? ' is-selected' : '');

            const head = document.createElement('div');
            head.className = 'manual-review-head';

            const pick = document.createElement('input');
            pick.type = 'radio';
            pick.name = 'manualReviewPick';
            pick.checked = selected;
            pick.addEventListener('change', () => {
                selectedManualReviewPath = row.rel_path;
                renderManualReviewRows(manualReviewRows);
            });

            const title = document.createElement('div');
            title.className = 'manual-review-title';
            title.textContent = describePath(row.rel_path || '');

            const score = document.createElement('span');
            score.className = 'manual-review-score';
            score.textContent = 'увер.=' + formatManualReviewConfidence(row.confidence);

            head.appendChild(pick);
            head.appendChild(title);
            head.appendChild(score);
            item.appendChild(head);

            const meta = document.createElement('div');
            meta.className = 'meta';
            const unresolvedLabel = isManualReviewUnresolved(row) ? 'да' : 'нет';
            meta.textContent = 'Тип: ' + (row.predicted_doc_type || '-')
                + '; тип скана: ' + (row.scan_doc_type || '-')
                + '; тег: ' + (row.scan_subject_tag || '-')
                + '; источник: ' + (row.source || '-')
                + '; неразобран: ' + unresolvedLabel
                + '; OCR: ' + (row.ocr_text_rel_path ? 'есть' : 'нет');
            item.appendChild(meta);

            if (row.suggested_target_rel_path) {
                const hint = document.createElement('div');
                hint.className = 'system-hint';
                hint.textContent = 'Рекомендуемый путь: ' + describePath(row.suggested_target_rel_path);
                item.appendChild(hint);
            }

            item.addEventListener('click', (event) => {
                if (event.target instanceof HTMLInputElement) {
                    return;
                }
                selectedManualReviewPath = row.rel_path;
                renderManualReviewRows(manualReviewRows);
            });

            manualReviewList.appendChild(item);
        }

        updateManualReviewSelectionInfo(true);
    }

    async function refreshManualReviewRows() {
        const rows = await api('/arm/scan/manual-review');
        manualReviewRows = sortManualReviewRows(Array.isArray(rows) ? rows : []);
        if (manualReview) {
            manualReview.textContent = manualReviewRows.length
                ? ('Ручной разбор: найдено файлов ' + manualReviewRows.length + '. Выберите файл в списке ниже.')
                : 'Папка ручного разбора пуста.';
        }
        renderManualReviewRows(manualReviewRows);
        return manualReviewRows;
    }

    function requireSelectedManualReviewRow() {
        const selectedRow = getSelectedManualReviewRow();
        if (!selectedRow) {
            setInteractionHint('Выберите файл из списка ручного разбора перед выполнением действия.', 'error');
            if (scannerMsg) {
                scannerMsg.textContent = 'Файл из ручного разбора не выбран.';
            }
            return null;
        }
        return selectedRow;
    }

    async function openSelectedManualReview() {
        const selectedRow = requireSelectedManualReviewRow();
        if (!selectedRow) {
            return;
        }
        try {
            await loadFile(selectedRow.rel_path);
            updateTaskActionContext(selectedRow.rel_path, 'ручной разбор: просмотр файла');
            renderActionNavigator(selectedRow.rel_path, 'ручной разбор: просмотр файла');
            setInteractionHint('Открыт файл ручного разбора: ' + describePath(selectedRow.rel_path), 'ok');
        } catch (_err) {
            const opened = openPrintPreview(selectedRow.rel_path, false);
            if (opened) {
                updateTaskActionContext(selectedRow.rel_path, 'ручной разбор: браузерный просмотр');
                renderActionNavigator(selectedRow.rel_path, 'ручной разбор: браузерный просмотр');
                setInteractionHint('Файл ручного разбора открыт в браузерном просмотре.', 'ok');
                return;
            }
            const folderPath = selectedRow.rel_path.split('/').slice(0, -1).join('/');
            await loadTree(folderPath);
            updateTaskActionContext(folderPath, 'ручной разбор: папка файла');
            renderActionNavigator(folderPath, 'ручной разбор: папка файла');
            setInteractionHint('Файл не поддерживает текстовый предпросмотр. Открыта его папка.', 'ok');
        }
    }

    function downloadSelectedManualReview() {
        const selectedRow = requireSelectedManualReviewRow();
        if (!selectedRow) {
            return;
        }
        window.open('/arm/fs/download?rel_path=' + encodeURIComponent(selectedRow.rel_path), '_blank');
        setInteractionHint('Начато скачивание файла из ручного разбора.', 'ok');
    }

    async function printSelectedManualReview() {
        const selectedRow = requireSelectedManualReviewRow();
        if (!selectedRow) {
            return;
        }
        const opened = openPrintPreview(selectedRow.rel_path, true);
        if (!opened) {
            if (scannerMsg) {
                scannerMsg.textContent = 'Браузер заблокировал окно печати. Разрешите всплывающие окна.';
            }
            setInteractionHint('Не удалось открыть окно печати: всплывающее окно заблокировано браузером.', 'error');
            return;
        }
        if (scannerMsg) {
            scannerMsg.textContent = 'Открыт браузерный предпросмотр печати для файла ручного разбора.';
        }
        setInteractionHint('Файл из ручного разбора открыт в браузерном окне печати.', 'ok');
    }

    async function openSelectedManualReviewOcr() {
        const selectedRow = requireSelectedManualReviewRow();
        if (!selectedRow) {
            return;
        }
        if (!selectedRow.ocr_text_rel_path) {
            setInteractionHint('У выбранного файла нет OCR-файла с текстом. Сначала выполните OCR/разбор.', 'error');
            return;
        }
        await loadFile(selectedRow.ocr_text_rel_path);
        updateTaskActionContext(selectedRow.ocr_text_rel_path, 'ручной разбор: OCR-текст');
        renderActionNavigator(selectedRow.ocr_text_rel_path, 'ручной разбор: OCR-текст');
        setInteractionHint('Открыт OCR-файл с текстом для редактирования.', 'ok');
    }

    function suggestPathForSelectedManualReview() {
        const selectedRow = requireSelectedManualReviewRow();
        if (!selectedRow) {
            return '';
        }
        const selectedType = manualReviewTargetType && manualReviewTargetType.value
            ? manualReviewTargetType.value
            : (selectedRow.scan_doc_type || 'OTHER');
        const nextPath = deriveManualReviewRoute(selectedType, selectedRow);
        if (manualReviewMovePath) {
            manualReviewMovePath.value = nextPath;
        }
        return nextPath;
    }

    async function ensureManualReviewMoveFolder() {
        const targetPath = (manualReviewMovePath && manualReviewMovePath.value || '').trim();
        const folderPath = targetPath.split('/').slice(0, -1).join('/');
        if (!folderPath) {
            throw new Error('Сначала укажите путь файла, чтобы создать целевую папку.');
        }
        const data = await api('/arm/fs/mkdir?rel_path=' + encodeURIComponent(folderPath), { method: 'POST' });
        if (scannerMsg) {
            scannerMsg.textContent = data.message || ('Папка готова: ' + folderPath);
        }
        return folderPath;
    }

    async function moveSelectedManualReview() {
        const selectedRow = requireSelectedManualReviewRow();
        if (!selectedRow) {
            return;
        }
        const targetPath = (manualReviewMovePath && manualReviewMovePath.value || '').trim();
        if (!targetPath) {
            setInteractionHint('Укажите целевой путь для перемещения файла из ручного разбора.', 'error');
            return;
        }

        const targetFolder = targetPath.split('/').slice(0, -1).join('/');
        if (targetFolder) {
            await api('/arm/fs/mkdir?rel_path=' + encodeURIComponent(targetFolder), { method: 'POST' });
        }

        const data = await api(
            '/arm/fs/move?source_rel_path=' + encodeURIComponent(selectedRow.rel_path) + '&target_rel_path=' + encodeURIComponent(targetPath),
            { method: 'POST' }
        );
        if (scannerMsg) {
            scannerMsg.textContent = data.message || 'Файл перемещен.';
        }
        await refreshManualReviewRows();
        if (targetFolder) {
            await loadTree(targetFolder).catch(() => {});
            updateTaskActionContext(targetFolder, 'ручной разбор: перенос');
            renderActionNavigator(targetFolder, 'ручной разбор: перенос');
        }
        setInteractionHint('Файл из ручного разбора перемещен.', 'ok');
    }

    async function deleteSelectedManualReview() {
        const selectedRow = requireSelectedManualReviewRow();
        if (!selectedRow) {
            return;
        }

        const data = await api(
            '/arm/fs/delete?rel_path=' + encodeURIComponent(selectedRow.rel_path) + '&with_sidecar=true',
            { method: 'POST' }
        );
        if (scannerMsg) {
            scannerMsg.textContent = data.message || 'Файл удален.';
        }
        await refreshManualReviewRows();
        setInteractionHint('Файл из ручного разбора удален.', 'ok');
    }

    async function runMaintenanceResetAndRebuild() {
        const orderDate = getChecklistOrderDate();
        if (!orderDate) {
            setInteractionHint('Дата приказа должна быть в формате ДД.ММ.ГГГГ.', 'error');
            if (employeeChecklistOrderDate) {
                employeeChecklistOrderDate.focus();
            }
            return;
        }

        const confirmed = window.confirm(
            'Очистить накопленные черновики/сканы и пересобрать структуру? '
            + 'Папки сотрудников и их профили сохранятся.'
        );
        if (!confirmed) {
            return;
        }

        setInteractionHint('Запущен сервисный reset/rebuild. Ожидайте завершения операции...');
        const data = await api(
            '/arm/maintenance/reset-rebuild?regenerate_project_orders=1&overwrite_orders=1&order_date=' + encodeURIComponent(orderDate),
            { method: 'POST' },
            TIMEOUTS_MS.ingest,
            'Сервисный reset/rebuild выполняется слишком долго.'
        );

        setInteractionHint(data.message || 'Сервисный reset/rebuild завершен.', 'ok');
        if (scannerMsg) {
            scannerMsg.textContent = data.message || 'Сервисный reset/rebuild завершен.';
        }

        await loadTree('');
        await refreshManualReviewRows().catch(() => {});
        await refreshEmployeeChecklistCatalog().catch(() => {});
    }

    async function classifyManualReview() {
        resetScannerProgress('подготовка классификации ручного разбора');
        appendScannerTimeline('Старт классификации ручного разбора.');
        setScannerProgress(25, 'классификация: запуск');

        manualReview.textContent = 'Классификация ручного разбора...';
        const rows = await refreshManualReviewRows();
        if (!rows.length) {
            setInteractionHint('Папка ручного разбора пуста: дополнительных действий не требуется.', 'ok');
            setScannerProgress(100, 'ручной разбор пуст');
            appendScannerTimeline('Папка ручного разбора пуста.');
            return;
        }
        manualReview.textContent = 'Ручной разбор: найдено файлов ' + rows.length + '. Проверьте список и выполните действие.';
        setInteractionHint('Классификация ручного разбора выполнена: проверьте список и переместите документы.', 'ok');
        setScannerProgress(100, 'классификация завершена');
        appendScannerTimeline('Классифицировано файлов: ' + rows.length + '.');
    }

    async function openActionPath(path, label) {
        if (!path) {
            setInteractionHint('Для выбранной позиции не указан путь действия. Откройте нужный раздел вручную через поле пути.', 'error');
            return;
        }
        focusStructureCard();
        const safeLabel = (label || 'задача').trim();
        updateTaskActionContext(path, safeLabel);
        setInteractionHint('Переход по задаче: ' + safeLabel + '. Открываю путь: ' + describePath(path));
        const looksLikeFile = /\.[a-z0-9]{1,8}$/i.test(path);
        if (looksLikeFile) {
            try {
                await loadFile(path);
                renderActionNavigator(path, safeLabel);
                setInteractionHint('Готово: открыт файл ' + describePath(path) + '.', 'ok');
                return;
            } catch (_err) {
                selectBinaryFile(path, safeLabel);
                const opened = openPrintPreview(path, false);
                if (opened) {
                    setInteractionHint('Файл выбран без текстового режима: открыт браузерный предпросмотр.', 'ok');
                } else {
                    setInteractionHint('Файл выбран без предпросмотра: доступно скачать или отправить на печать.', 'ok');
                }
                return;
            }
        }
        await loadTree(path);
        renderActionNavigator(path, safeLabel);
        if (path.includes('02_personnel/employees')) {
            await openChecklistByTaskAction();
        }
        setInteractionHint('Готово: открыта папка ' + describePath(path) + '. Следующий шаг: выберите файл/подпапку в блоке «Структура и действия».', 'ok');
    }

    treeView.addEventListener('click', (event) => {
        const rawTarget = event.target;
        const target = rawTarget instanceof Element
            ? rawTarget
            : (rawTarget && rawTarget.parentElement ? rawTarget.parentElement : null);
        if (!target) {
            return;
        }
        const link = target.closest('a[data-path]');
        if (!(link instanceof HTMLAnchorElement)) {
            return;
        }
        event.preventDefault();
        const relPath = link.dataset.path || '';
        const isDir = link.dataset.isDir === '1';
        if (isDir) {
            loadTree(relPath)
                .then(() => {
                    updateTaskActionContext(relPath, 'Навигация по дереву');
                    setInteractionHint('Открыта подпапка: ' + describePath(relPath), 'ok');
                    renderActionNavigator(relPath, 'Навигация по дереву');
                })
                .catch((err) => {
                    setTreeHint('Ошибка дерева: ' + err);
                    setInteractionHint('Не удалось открыть подпапку: ' + err, 'error');
                });
        } else {
            loadFile(relPath)
                .then(() => {
                    updateTaskActionContext(relPath, 'Просмотр файла');
                    setInteractionHint('Открыт файл для просмотра: ' + describePath(relPath), 'ok');
                    renderActionNavigator(relPath, 'Просмотр файла');
                })
                .catch((_err) => {
                    selectBinaryFile(relPath, 'Выбор файла');
                    setInteractionHint('Файл выбран без текстового предпросмотра. Доступны действия: скачать или печать.', 'ok');
                });
        }
    });

    document.addEventListener('click', (event) => {
        const rawTarget = event.target;
        const target = rawTarget instanceof Element
            ? rawTarget
            : (rawTarget && rawTarget.parentElement ? rawTarget.parentElement : null);
        if (!target) {
            return;
        }
        const link = target.closest('a.todo-link');
        if (!(link instanceof HTMLAnchorElement)) {
            return;
        }
        event.preventDefault();
        event.stopPropagation();
        withBusy(treeOpenBtn, async () => {
            await openActionPath(link.dataset.actionPath || '', link.textContent || '');
        }).catch((err) => {
            setTreeHint('Не удалось открыть путь: ' + err);
            setInteractionHint('Ошибка перехода по интерактивной ссылке: ' + err, 'error');
        });
    });

    sendBtn.addEventListener('click', sendAssist);
    if (voiceBtn) {
        voiceBtn.addEventListener('click', toggleVoiceInput);
    }
    question.addEventListener('keydown', (event) => {
        if (event.key !== 'Enter' || event.isComposing) {
            return;
        }
        const sendOnEnter = Boolean(sendOnEnterCheckbox && sendOnEnterCheckbox.checked);
        const shouldSend = event.ctrlKey || (sendOnEnter && !event.shiftKey);
        if (!shouldSend) {
            return;
        }
        event.preventDefault();
        sendAssist();
    });
    if (sendOnEnterCheckbox) {
        sendOnEnterCheckbox.addEventListener('change', () => {
            if (sendOnEnterCheckbox.checked) {
                setInteractionHint('Режим ввода ассистента: Enter отправляет, Shift+Enter перенос строки.', 'ok');
            } else {
                setInteractionHint('Режим ввода ассистента: Ctrl+Enter отправляет, Enter перенос строки.', 'ok');
            }
        });
    }
    if (scannerDocType) {
        scannerDocType.addEventListener('change', syncScannerRequirements);
    }
    if (scannerSortMode) {
        scannerSortMode.addEventListener('change', () => {
            manualReviewRows = sortManualReviewRows(manualReviewRows);
            renderManualReviewRows(manualReviewRows);
        });
    }
    if (scannerEmployee) {
        scannerEmployee.addEventListener('input', syncScannerRequirements);
    }
    if (scannerEmployeeSelect) {
        scannerEmployeeSelect.addEventListener('change', () => {
            const value = (scannerEmployeeSelect.value || '').trim();
            if (value) {
                setScannerEmployeeValue(value);
                if (scannerEmployeeHint) {
                    const selected = scannerEmployeeSelect.options[scannerEmployeeSelect.selectedIndex];
                    scannerEmployeeHint.textContent = 'Выбран сотрудник: ' + (selected ? selected.textContent : value) + '.';
                }
            }
            syncScannerRequirements();
        });
    }
    if (treePathInput && treeOpenBtn) {
        treePathInput.addEventListener('keydown', (event) => {
            if (event.key !== 'Enter' || event.isComposing) {
                return;
            }
            event.preventDefault();
            treeOpenBtn.click();
        });
    }
    treeOpenBtn.addEventListener('click', () => withBusy(treeOpenBtn, async () => {
        const manualPath = normalizeRelPath(treePathInput && treePathInput.value || '');
        await loadTree(manualPath);
        const openedPath = normalizeRelPath(treePathInput && treePathInput.dataset.currentPath || manualPath);
        updateTaskActionContext(openedPath, 'Открыто вручную');
        renderActionNavigator(openedPath, 'Открыто вручную');
        setInteractionHint('Открыт путь из поля: ' + describePath(openedPath), 'ok');
    }).catch((err) => {
        setTreeHint('Ошибка дерева: ' + err);
        setInteractionHint('Не удалось открыть путь из поля: ' + err, 'error');
    }));
    if (treeBackBtn) {
        treeBackBtn.addEventListener('click', () => withBusy(treeBackBtn, () => navigateTreeHistory(-1)).catch((err) => {
            setTreeHint('Ошибка перехода назад: ' + err);
            setInteractionHint('Не удалось вернуться к предыдущей папке: ' + err, 'error');
        }));
    }
    if (treeForwardBtn) {
        treeForwardBtn.addEventListener('click', () => withBusy(treeForwardBtn, () => navigateTreeHistory(1)).catch((err) => {
            setTreeHint('Ошибка перехода вперед: ' + err);
            setInteractionHint('Не удалось перейти вперед: ' + err, 'error');
        }));
    }
    if (treeUpBtn) {
        treeUpBtn.addEventListener('click', () => withBusy(treeUpBtn, async () => {
            const currentPath = normalizeRelPath(treePathInput && treePathInput.dataset.currentPath || '');
            const nextPath = parentRelPath(currentPath);
            await loadTree(nextPath);
            setInteractionHint('Переход на уровень выше выполнен: ' + describePath(nextPath), 'ok');
        }).catch((err) => {
            setTreeHint('Ошибка перехода на уровень выше: ' + err);
            setInteractionHint('Не удалось перейти на уровень выше: ' + err, 'error');
        }));
    }
    if (treeRootBtn) {
        treeRootBtn.addEventListener('click', () => withBusy(treeRootBtn, async () => {
            await loadTree('');
            setInteractionHint('Открыт корень объекта.', 'ok');
        }).catch((err) => {
            setTreeHint('Ошибка перехода в корень: ' + err);
            setInteractionHint('Не удалось открыть корень объекта: ' + err, 'error');
        }));
    }
    if (treeBreadcrumb) {
        treeBreadcrumb.addEventListener('click', (event) => {
            const rawTarget = event.target;
            const target = rawTarget instanceof Element
                ? rawTarget
                : (rawTarget && rawTarget.parentElement ? rawTarget.parentElement : null);
            if (!target) {
                return;
            }
            const crumb = target.closest('button.tree-crumb[data-path]');
            if (!(crumb instanceof HTMLButtonElement)) {
                return;
            }
            event.preventDefault();
            const relPath = crumb.dataset.path || '';
            withBusy(treeOpenBtn, async () => {
                await loadTree(relPath);
                setInteractionHint('Открыт раздел по хлебным крошкам: ' + describePath(relPath), 'ok');
            }).catch((err) => {
                setTreeHint('Ошибка перехода по крошкам: ' + err);
                setInteractionHint('Не удалось открыть раздел по хлебным крошкам: ' + err, 'error');
            });
        });
    }
    fileSaveBtn.addEventListener('click', () => withBusy(fileSaveBtn, saveCurrentFile).catch((err) => setFileMeta('Ошибка сохранения: ' + err)));
    fileDownloadBtn.addEventListener('click', downloadCurrentFile);
    filePrintBtn.addEventListener('click', () => withBusy(filePrintBtn, printCurrentFile).catch((err) => setFileMeta('Ошибка печати: ' + err)));
    if (fileZoomOutBtn) {
        fileZoomOutBtn.addEventListener('click', () => setPreviewZoom(filePreviewZoom - 0.1));
    }
    if (fileZoomResetBtn) {
        fileZoomResetBtn.addEventListener('click', () => setPreviewZoom(1));
    }
    if (fileZoomInBtn) {
        fileZoomInBtn.addEventListener('click', () => setPreviewZoom(filePreviewZoom + 0.1));
    }
    if (fileMoveBtn) {
        fileMoveBtn.addEventListener('click', () => withBusy(fileMoveBtn, moveCurrentFile).catch((err) => setFileMeta('Ошибка перемещения: ' + err)));
    }
    if (fileDeleteBtn) {
        fileDeleteBtn.addEventListener('click', () => withBusy(fileDeleteBtn, deleteCurrentFile).catch((err) => setFileMeta('Ошибка удаления: ' + err)));
    }
    if (fileSourceToggleBtn) {
        fileSourceToggleBtn.addEventListener('click', () => {
            fileSourceVisible = !fileSourceVisible;
            updateFileSourceVisibility();
        });
    }
    if (scannerProfile) {
        scannerProfile.addEventListener('input', () => {
            currentScannerProfile = Math.max(1, Math.min(3, Number(scannerProfile.value) || 1));
            updateScannerProfileLabel();
        });
    }
    scanBtn.addEventListener('click', () => withBusy(scanBtn, scanToInbox).catch((err) => {
        const message = String(err || '');
        if (message.toLowerCase().includes('заняло слишком много времени')) {
            scannerMsg.textContent = 'Сканер мог завершить работу, но ответ сервера не успел прийти. Проверьте папку входящих сканов и запустите «Распознать и разложить».';
            setInteractionHint('Сканирование завершилось с задержкой ответа. Проверьте входящую папку и запустите OCR.', 'error');
            return;
        }
        scannerMsg.textContent = 'Ошибка сканирования: ' + err;
        setInteractionHint('Ошибка сканирования: ' + err, 'error');
    }));
    ingestBtn.addEventListener('click', () => withBusy(ingestBtn, ingestInbox).catch((err) => {
        scannerMsg.textContent = 'Ошибка OCR/разбора: ' + err;
        setInteractionHint('Ошибка OCR/разбора: ' + err, 'error');
    }));
    manualClassifyBtn.addEventListener('click', () => withBusy(manualClassifyBtn, classifyManualReview).catch((err) => {
        manualReview.textContent = 'Ошибка классификации: ' + err;
        setInteractionHint('Ошибка классификации ручного разбора: ' + err, 'error');
    }));
    if (recompressScansBtn) {
        recompressScansBtn.addEventListener('click', () => withBusy(recompressScansBtn, async () => {
            const data = await api('/arm/scanner/recompress-history', { method: 'POST' }, TIMEOUTS_MS.scanner);
            scannerMsg.textContent = data.message || 'Оптимизация сканов завершена.';
            appendScannerTimeline(data.message || 'Оптимизация исторических сканов завершена.');
            setInteractionHint('Оптимизация сканов завершена.', 'ok');
        }).catch((err) => {
            scannerMsg.textContent = 'Ошибка оптимизации сканов: ' + err;
            setInteractionHint('Не удалось оптимизировать исторические сканы: ' + err, 'error');
        }));
    }
    if (manualReviewOpenBtn) {
        manualReviewOpenBtn.addEventListener('click', () => withBusy(manualReviewOpenBtn, openSelectedManualReview).catch((err) => {
            setInteractionHint('Не удалось открыть файл ручного разбора: ' + err, 'error');
        }));
    }
    if (manualReviewDownloadBtn) {
        manualReviewDownloadBtn.addEventListener('click', downloadSelectedManualReview);
    }
    if (manualReviewPrintBtn) {
        manualReviewPrintBtn.addEventListener('click', () => withBusy(manualReviewPrintBtn, printSelectedManualReview).catch((err) => {
            setInteractionHint('Не удалось отправить на печать: ' + err, 'error');
        }));
    }
    if (manualReviewEditOcrBtn) {
        manualReviewEditOcrBtn.addEventListener('click', () => withBusy(manualReviewEditOcrBtn, openSelectedManualReviewOcr).catch((err) => {
            setInteractionHint('Не удалось открыть OCR-текст: ' + err, 'error');
        }));
    }
    if (manualReviewSuggestBtn) {
        manualReviewSuggestBtn.addEventListener('click', () => {
            const path = suggestPathForSelectedManualReview();
            if (!path) {
                return;
            }
            setInteractionHint('Подставлен маршрут переноса: ' + describePath(path), 'ok');
        });
    }
    if (manualReviewMkDirBtn) {
        manualReviewMkDirBtn.addEventListener('click', () => withBusy(manualReviewMkDirBtn, ensureManualReviewMoveFolder).then((folderPath) => {
            setInteractionHint('Папка маршрута подготовлена: ' + describePath(folderPath), 'ok');
        }).catch((err) => {
            setInteractionHint('Не удалось создать папку маршрута: ' + err, 'error');
        }));
    }
    if (manualReviewTargetType) {
        manualReviewTargetType.addEventListener('change', () => {
            suggestPathForSelectedManualReview();
        });
    }
    if (manualReviewMoveBtn) {
        manualReviewMoveBtn.addEventListener('click', () => withBusy(manualReviewMoveBtn, moveSelectedManualReview).catch((err) => {
            setInteractionHint('Ошибка перемещения файла из ручного разбора: ' + err, 'error');
        }));
    }
    if (manualReviewDeleteBtn) {
        manualReviewDeleteBtn.addEventListener('click', () => withBusy(manualReviewDeleteBtn, deleteSelectedManualReview).catch((err) => {
            setInteractionHint('Ошибка удаления файла из ручного разбора: ' + err, 'error');
        }));
    }
    if (maintenanceResetBtn) {
        maintenanceResetBtn.addEventListener('click', () => withBusy(maintenanceResetBtn, runMaintenanceResetAndRebuild).catch((err) => {
            setInteractionHint('Ошибка сервисного reset/rebuild: ' + err, 'error');
        }));
    }
    if (employeeChecklistToggleBtn) {
        employeeChecklistToggleBtn.addEventListener('click', toggleEmployeeChecklistCard);
    }
    if (taskActionOpenChecklistBtn) {
        taskActionOpenChecklistBtn.addEventListener('click', () => withBusy(taskActionOpenChecklistBtn, openChecklistByTaskAction).catch((err) => {
            setEmployeeChecklistMsg('Не удалось открыть ТБ-чеклист: ' + err, 'error');
        }));
    }
    if (taskActionRunMissingBtn) {
        taskActionRunMissingBtn.addEventListener('click', () => withBusy(taskActionRunMissingBtn, () => runTaskActionGeneration('missing')).catch((err) => {
            setEmployeeChecklistMsg('Ошибка генерации недостающих документов из панели задачи: ' + err, 'error');
        }));
    }
    if (taskActionRunSelectedBtn) {
        taskActionRunSelectedBtn.addEventListener('click', () => withBusy(taskActionRunSelectedBtn, () => runTaskActionGeneration('selected')).catch((err) => {
            setEmployeeChecklistMsg('Ошибка генерации выбранных документов из панели задачи: ' + err, 'error');
        }));
    }
    if (employeeChecklistSelectMissing) {
        employeeChecklistSelectMissing.addEventListener('click', selectOnlyMissingInChecklist);
    }
    if (employeeChecklistProfession) {
        employeeChecklistProfession.addEventListener('change', () => {
            renderEmployeeChecklistEmployeeOptions();
            renderBatchEmployeeOptions();
            if (employeeChecklistEmployeeSelect && !employeeChecklistEmployeeSelect.value && employeeChecklistPath) {
                employeeChecklistPath.value = '';
            }
        });
    }
    if (employeeChecklistEmployeeSelect) {
        employeeChecklistEmployeeSelect.addEventListener('change', () => {
            const relPath = (employeeChecklistEmployeeSelect.value || '').trim();
            if (employeeChecklistPath) {
                employeeChecklistPath.value = relPath;
            }

            const selected = employeeCatalogRows.find((row) => row.employee_rel_path === relPath);
            if (selected && employeeChecklistProfession) {
                employeeChecklistProfession.value = selected.profession_group || 'all';
            }
            renderBatchEmployeeOptions();
        });
    }
    if (employeeChecklistOverviewBtn) {
        employeeChecklistOverviewBtn.addEventListener('click', () => withBusy(employeeChecklistOverviewBtn, loadEmployeeChecklistOverview).catch((err) => {
            setEmployeeChecklistMsg('Ошибка комплексного анализа: ' + err, 'error');
        }));
    }
    if (employeeChecklistRefreshBtn) {
        employeeChecklistRefreshBtn.addEventListener('click', () => withBusy(employeeChecklistRefreshBtn, loadEmployeeChecklist).catch((err) => {
            setEmployeeChecklistMsg('Ошибка проверки чеклиста: ' + err, 'error');
        }));
    }
    if (employeeChecklistGenerateMissingBtn) {
        employeeChecklistGenerateMissingBtn.addEventListener('click', () => withBusy(employeeChecklistGenerateMissingBtn, () => generateEmployeeChecklist('missing')).catch((err) => {
            setEmployeeChecklistMsg('Ошибка генерации недостающих документов: ' + err, 'error');
        }));
    }
    if (employeeChecklistGenerateSelectedBtn) {
        employeeChecklistGenerateSelectedBtn.addEventListener('click', () => withBusy(employeeChecklistGenerateSelectedBtn, () => generateEmployeeChecklist('selected')).catch((err) => {
            setEmployeeChecklistMsg('Ошибка точечной генерации: ' + err, 'error');
        }));
    }
    if (employeeChecklistGenerateAllBtn) {
        employeeChecklistGenerateAllBtn.addEventListener('click', () => withBusy(employeeChecklistGenerateAllBtn, () => generateEmployeeChecklist('all')).catch((err) => {
            setEmployeeChecklistMsg('Ошибка полной генерации: ' + err, 'error');
        }));
    }
    if (employeeChecklistApplyTypeBtn) {
        employeeChecklistApplyTypeBtn.addEventListener('click', () => withBusy(employeeChecklistApplyTypeBtn, applyChecklistByProfession).catch((err) => {
            setEmployeeChecklistMsg('Ошибка применения формы к типу сотрудников: ' + err, 'error');
        }));
    }
    if (batchGenerateMode) {
        batchGenerateMode.addEventListener('change', () => {
            const mode = batchGenerateMode.value || 'missing';
            setEmployeeChecklistMsg('Пакетный режим: ' + formatGenerateMode(mode) + '.', 'ok');
        });
    }
    if (batchEmployeesSelectAllBtn) {
        batchEmployeesSelectAllBtn.addEventListener('click', () => {
            setMultiSelectState(batchEmployees, true);
            setEmployeeChecklistMsg('Выбраны все сотрудники из текущего фильтра.', 'ok');
        });
    }
    if (batchEmployeesClearBtn) {
        batchEmployeesClearBtn.addEventListener('click', () => {
            setMultiSelectState(batchEmployees, false);
            setEmployeeChecklistMsg('Выбор сотрудников очищен.', 'ok');
        });
    }
    if (batchUseMissingCodesBtn) {
        batchUseMissingCodesBtn.addEventListener('click', () => {
            if (!currentEmployeeChecklist) {
                setEmployeeChecklistMsg('Сначала откройте чеклист конкретного сотрудника, чтобы взять недостающие коды.', 'error');
                return;
            }
            renderBatchDocCodesFromChecklist(currentEmployeeChecklist);
            setEmployeeChecklistMsg('В пакетный блок подставлены коды из текущего чеклиста.', 'ok');
        });
    }
    if (batchGenerateSelectedEmployeesBtn) {
        batchGenerateSelectedEmployeesBtn.addEventListener('click', () => withBusy(batchGenerateSelectedEmployeesBtn, () => runBatchChecklistGeneration(false)).catch((err) => {
            setEmployeeChecklistMsg('Ошибка пакетной генерации (выбранные сотрудники): ' + err, 'error');
        }));
    }
    if (batchGenerateAllFilteredBtn) {
        batchGenerateAllFilteredBtn.addEventListener('click', () => withBusy(batchGenerateAllFilteredBtn, () => runBatchChecklistGeneration(true)).catch((err) => {
            setEmployeeChecklistMsg('Ошибка пакетной генерации (все сотрудники фильтра): ' + err, 'error');
        }));
    }
    if (employeeChecklistPath && employeeChecklistRefreshBtn) {
        employeeChecklistPath.addEventListener('keydown', (event) => {
            if (event.key === 'Enter' && !event.shiftKey && !event.ctrlKey) {
                event.preventDefault();
                withBusy(employeeChecklistRefreshBtn, loadEmployeeChecklist).catch((err) => {
                    setEmployeeChecklistMsg('Ошибка проверки чеклиста: ' + err, 'error');
                });
            }
        });
    }
    if (exportDocxBtn) {
        exportDocxBtn.addEventListener('click', () => {
            const classification = (exportClassification && exportClassification.value) || 'all';
            window.open('/arm/exports/orders-docx?classification=' + encodeURIComponent(classification), '_blank');
        });
    }
    if (exportXlsxBtn) {
        exportXlsxBtn.addEventListener('click', () => {
            window.open('/arm/exports/registers-xlsx', '_blank');
        });
    }
    if (exportPackBtn) {
        exportPackBtn.addEventListener('click', () => {
            const classification = (exportClassification && exportClassification.value) || 'all';
            window.open('/arm/exports/office-pack?classification=' + encodeURIComponent(classification), '_blank');
        });
    }
    if (backToTopBtn) {
        const updateBackToTopVisibility = () => {
            backToTopBtn.hidden = window.scrollY < 320;
        };
        backToTopBtn.addEventListener('click', () => {
            window.scrollTo({ top: 0, behavior: 'smooth' });
        });
        window.addEventListener('scroll', updateBackToTopVisibility, { passive: true });
        updateBackToTopVisibility();
    }

    async function loadObjectProfile() {
        const data = await api('/arm/object-profile', undefined, TIMEOUTS_MS.default);
        if (objectNameInput) objectNameInput.value = data.object_name || '';
        if (projectCodeInput) projectCodeInput.value = data.project_code || '';
        if (organizationInput) organizationInput.value = data.organization || '';
        if (workStageInput) workStageInput.value = data.work_stage || '';
        if (startDateInput) startDateInput.value = data.start_date || '';
        if (pprSourceSelect) {
            pprSourceSelect.innerHTML = '';
            const blank = document.createElement('option');
            blank.value = '';
            blank.textContent = 'Выберите ППР из входящих';
            pprSourceSelect.appendChild(blank);
            for (const relPath of (data.ppr_source_options || [])) {
                const option = document.createElement('option');
                option.value = relPath;
                option.textContent = describePath(relPath);
                pprSourceSelect.appendChild(option);
            }
        }
        if (objectProfileMsg) {
            objectProfileMsg.textContent = data.ppr_context_rel_path
                ? ('Карточка объекта загружена. Контекст ППР: ' + describePath(data.ppr_context_rel_path))
                : 'Карточка объекта загружена.';
        }
    }

    if (objectProfileSaveBtn) {
        objectProfileSaveBtn.addEventListener('click', () => withBusy(objectProfileSaveBtn, async () => {
            const data = await api('/arm/object-profile', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({
                    object_name: objectNameInput ? objectNameInput.value.trim() : '',
                    project_code: projectCodeInput ? projectCodeInput.value.trim() : '',
                    organization: organizationInput ? organizationInput.value.trim() : '',
                    work_stage: workStageInput ? workStageInput.value.trim() : '',
                    start_date: startDateInput ? startDateInput.value.trim() : '',
                }),
            }, TIMEOUTS_MS.default);
            if (objectProfileMsg) {
                objectProfileMsg.textContent = data.message || 'Карточка объекта обновлена.';
            }
        }).catch((err) => {
            if (objectProfileMsg) {
                objectProfileMsg.textContent = 'Ошибка сохранения карточки объекта: ' + err;
            }
        }));
    }

    if (pprImportBtn) {
        pprImportBtn.addEventListener('click', () => withBusy(pprImportBtn, async () => {
            const relPath = pprSourceSelect ? (pprSourceSelect.value || '').trim() : '';
            if (!relPath) {
                throw new Error('Выберите ППР для импорта.');
            }
            const data = await api('/arm/ppr/import', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ rel_path: relPath }),
            }, TIMEOUTS_MS.default);
            if (objectProfileMsg) {
                objectProfileMsg.textContent = data.message || 'ППР импортирован в базу знаний.';
            }
            await loadObjectProfile();
        }).catch((err) => {
            if (objectProfileMsg) {
                objectProfileMsg.textContent = 'Ошибка импорта ППР: ' + err;
            }
        }));
    }

    initCollapsibleCards();
    loadObjectProfile().catch((err) => {
        if (objectProfileMsg) {
            objectProfileMsg.textContent = 'Не удалось загрузить карточку объекта: ' + err;
        }
    });
    refreshLLMStatus();
    setInterval(refreshLLMStatus, 10000);
    resetScannerProgress('ожидание действий');
    refreshScannerDevices();
    refreshManualReviewRows().catch((err) => {
        if (manualReview) {
            manualReview.textContent = 'Ошибка загрузки ручного разбора: ' + err;
        }
    });
    refreshEmployeeIdSuggestions();
    refreshEmployeeChecklistCatalog().catch((err) => {
        setEmployeeChecklistMsg('Не удалось загрузить список сотрудников: ' + err, 'error');
    });
    syncScannerRequirements();
    currentScannerProfile = Math.max(1, Math.min(3, Number(scannerProfile && scannerProfile.value) || 1));
    updateScannerProfileLabel();
    setPreviewZoom(1);
    initVoiceInput();
    window.addEventListener('beforeunload', () => {
        if (voiceListening) {
            stopVoiceInput('');
        }
    });
    renderActionNavigator('', 'Старт');
    setInteractionHint('Подсказка: кликните пункт в «Критичные пробелы» или «Задачи на сегодня», система откроет нужную папку и подсветит рабочий блок.');
    setEmployeeChecklistMsg('Выберите сотрудника в дереве или задайте путь вручную, затем нажмите «Проверить комплект».');
    updateFileSourceVisibility();
    refreshRenderedPreview('');
    const startupPath = normalizeRelPath(new URLSearchParams(window.location.search).get('open_path') || '');
    loadTree(startupPath);

    // ---------- FILE UPLOAD ----------
    if (uploadUseCurrentBtn) {
        uploadUseCurrentBtn.addEventListener('click', () => {
            const currentPath = normalizeRelPath((treePathInput && treePathInput.dataset.currentPath) || '');
            if (uploadDir) {
                uploadDir.value = currentPath;
            }
            if (uploadMsg) {
                uploadMsg.textContent = currentPath
                    ? ('Папка для загрузки выбрана: ' + describePath(currentPath))
                    : 'Выбран корень объекта для загрузки.';
            }
        });
    }
    if (uploadBtn) {
        uploadBtn.addEventListener('click', async () => {
            const file = uploadFile && uploadFile.files[0];
            if (!file) { uploadMsg.textContent = 'Не выбран файл.'; return; }
            const dir = (uploadDir && uploadDir.value.trim()) || '';
            const fd = new FormData();
            fd.append('file', file);
            uploadBtn.classList.add('is-busy');
            uploadMsg.textContent = 'Загрузка...';
            try {
                const res = await fetch('/arm/fs/upload?rel_dir=' + encodeURIComponent(dir), {
                    method: 'POST',
                    body: fd,
                    signal: AbortSignal.timeout(30000),
                });
                const data = await res.json();
                uploadMsg.textContent = data.message || (data.detail ? '\u26d4 ' + data.detail : 'Готово');
                if (data.ok) { loadTree(dir); if (uploadFile) uploadFile.value = ''; }
            } catch (err) {
                uploadMsg.textContent = '\u26d4 Ошибка: ' + err;
            } finally {
                uploadBtn.classList.remove('is-busy');
            }
        });
    }
</script>
"""

    html = f"""
<!doctype html>
<html lang=\"ru\">
<head>
    <meta charset=\"utf-8\" />
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
    <title>АРМ объекта: дашборд</title>
    <style>
    :root {{
        --bg: #f4efe7;
        --card: #fffaf3;
        --ink: #1f2a36;
        --accent: #0f766e;
        --muted: #5b6672;
    }}
    body {{
        margin: 0;
        font-family: "Segoe UI", Tahoma, sans-serif;
        background: radial-gradient(circle at 85% 15%, #d7efe6 0%, #f4efe7 45%, #e6eff4 100%);
        color: var(--ink);
    }}
    .wrap {{ max-width: 1180px; margin: 24px auto; padding: 0 16px 24px; }}
    .hero {{ background: var(--card); border: 1px solid #d8cfc2; border-radius: 16px; padding: 18px; box-shadow: 0 8px 24px rgba(0, 0, 0, 0.06); }}
    .tabs {{ display: inline-flex; gap: 8px; margin-bottom: 12px; flex-wrap: wrap; }}
    .tab {{ display: inline-block; text-decoration: none; padding: 7px 12px; border-radius: 999px; border: 1px solid #c8d2d8; background: #f8fbfd; color: #334155; font-size: 13px; font-weight: 600; }}
    .tab:hover {{ background: #eff6fb; }}
    .tab.active {{ background: var(--accent); border-color: var(--accent); color: #fff; }}
    .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(340px, 1fr)); gap: 16px; margin-top: 16px; }}
    .dashboard-main-grid {{ grid-template-columns: repeat(2, minmax(420px, 1fr)); }}
    .card {{ background: var(--card); border: 1px solid #d8cfc2; border-radius: 16px; padding: 14px; box-shadow: 0 8px 24px rgba(0, 0, 0, 0.05); }}
    .card-collapse-head {{ display: flex; justify-content: space-between; align-items: center; gap: 10px; margin-bottom: 6px; }}
    .card-collapse-head h2 {{ margin: 0; }}
    .card-collapse-body[hidden] {{ display: none !important; }}
    .actions {{ display: flex; flex-wrap: wrap; gap: 8px; margin-top: 8px; }}
    .btn {{ background: var(--accent); color: #fff; text-decoration: none; border-radius: 10px; padding: 8px 12px; font-size: 14px; border: none; cursor: pointer; }}
    .btn.secondary {{ background: #334155; }}
    .btn.is-busy {{ opacity: 0.65; pointer-events: none; }}
    h1 {{ margin: 0 0 8px; font-size: 24px; }}
    h2 {{ margin: 0 0 8px; font-size: 18px; color: var(--accent); }}
    .meta {{ color: var(--muted); font-size: 14px; }}
    .kpi {{ font-size: 34px; font-weight: 700; color: var(--accent); margin: 8px 0; }}
    ul {{ margin: 8px 0 0; padding-left: 18px; }}
    li {{ margin: 6px 0; }}
    .todo-link {{ color: #0f4c5c; font-weight: 600; text-decoration: none; }}
    .todo-link:hover {{ text-decoration: underline; }}
    .system-hint {{ margin-top: 4px; color: #475569; font-size: 13px; line-height: 1.35; }}
    .action-note {{ margin-top: 10px; padding: 10px 12px; border-radius: 10px; border: 1px solid #c9dbe4; background: #eef6f9; color: #1e3a4a; font-size: 13px; line-height: 1.35; }}
    .action-note.hint-error {{ background: #fef2f2; border-color: #fecaca; color: #991b1b; }}
    .action-note.hint-ok {{ background: #ecfdf5; border-color: #bbf7d0; color: #065f46; }}
    .llm-badge {{ display: inline-block; margin-left: 6px; padding: 2px 8px; border-radius: 999px; font-size: 12px; font-weight: 700; animation: pulse 1.2s infinite; }}
    .llm-ok {{ background: #d1fae5; color: #065f46; }}
    .llm-down {{ background: #fee2e2; color: #991b1b; }}
    @keyframes pulse {{ 0% {{ opacity: 1; }} 50% {{ opacity: 0.35; }} 100% {{ opacity: 1; }} }}
    .card.action-focus {{ box-shadow: 0 0 0 3px rgba(15, 118, 110, 0.25), 0 8px 24px rgba(0, 0, 0, 0.08); transition: box-shadow 0.2s ease; }}
    .tree {{ background: #fff; border: 1px solid #d8e0e0; border-radius: 10px; max-height: 280px; overflow: auto; padding: 8px; }}
    .tree ul {{ margin: 0; padding-left: 16px; }}
    .tree li {{ margin: 4px 0; }}
    .tree a {{ color: #0f4c5c; text-decoration: none; }}
    .tree a:hover {{ text-decoration: underline; }}
    .tree li.tree-item-ok > a {{ color: #166534; font-weight: 600; }}
    .tree li.tree-item-editable > a {{ color: #0f766e; }}
    .tree li.tree-item-missing {{
        color: #991b1b;
        background: #fef2f2;
        border: 1px dashed #fecaca;
        border-radius: 8px;
        padding: 4px 8px;
        list-style-position: inside;
    }}
    .tree-toolbar {{ margin-top: 8px; }}
    .tree-toolbar .btn[disabled] {{ opacity: 0.45; cursor: not-allowed; }}
    .tree-breadcrumb {{ margin-top: 8px; display: flex; flex-wrap: wrap; gap: 6px; align-items: center; }}
    .tree-crumb {{ border: 1px solid #c8d2d8; border-radius: 999px; background: #ffffff; color: #0f4c5c; padding: 4px 9px; font-size: 12px; cursor: pointer; }}
    .tree-crumb:hover {{ background: #eef6fb; }}
    .tree-crumb-sep {{ color: #64748b; font-size: 12px; }}
    input, select, textarea {{ width: 100%; box-sizing: border-box; border-radius: 10px; border: 1px solid #c5cbc9; padding: 10px; font-size: 14px; background: #fff; }}
    .required-input {{ border-color: #f59e0b; box-shadow: 0 0 0 2px rgba(245, 158, 11, 0.15); }}
    textarea {{ min-height: 220px; resize: vertical; font-family: Consolas, "Courier New", monospace; }}
    .is-hidden {{ display: none !important; }}
    #fileRenderWrap {{ width: 100%; overflow: auto; border-radius: 10px; }}
    #fileRenderFrame {{ width: 100%; min-height: 66vh; border: 1px solid #d8e0e0; border-radius: 10px; background: #fff; margin-top: 10px; }}
    pre {{ background: #f8fbfb; border: 1px solid #d8e0e0; border-radius: 10px; padding: 10px; white-space: pre-wrap; word-break: break-word; margin: 8px 0 0; font-size: 13px; }}
    .inline-grid {{ display: grid; grid-template-columns: repeat(2, minmax(120px, 1fr)); gap: 8px; margin-top: 8px; }}
    .action-nav {{ margin-top: 10px; padding: 10px 12px; border-radius: 10px; border: 1px dashed #b9ced8; background: #f6fbfd; }}
    .action-nav ol {{ margin: 8px 0 0; padding-left: 18px; }}
    .action-nav li {{ margin: 4px 0; color: #334155; font-size: 13px; }}
    .btn.mini {{ padding: 5px 8px; font-size: 12px; border-radius: 8px; }}
    .employee-checklist-tools {{ display: grid; grid-template-columns: 1fr; gap: 8px; margin-top: 8px; }}
    .employee-checklist-list {{ margin-top: 10px; max-height: 360px; overflow: auto; border: 1px solid #d8e0e0; border-radius: 10px; background: #ffffff; padding: 8px; }}
    .employee-checklist-item {{ border: 1px solid #e2e8f0; border-radius: 10px; padding: 8px; margin-bottom: 8px; background: #f9fbfc; }}
    .employee-checklist-item.is-ready {{ border-color: #bbf7d0; background: #f0fdf4; }}
    .employee-checklist-item.is-missing {{ border-color: #fde68a; background: #fffbeb; }}
    .employee-checklist-top {{ display: grid; grid-template-columns: auto 1fr auto auto; gap: 8px; align-items: center; }}
    .employee-checklist-title {{ font-weight: 600; color: #0f172a; }}
    .employee-checklist-state {{ font-size: 12px; border-radius: 999px; padding: 3px 8px; white-space: nowrap; }}
    .employee-checklist-state.ok {{ background: #dcfce7; color: #166534; }}
    .employee-checklist-state.warn {{ background: #fef3c7; color: #92400e; }}
    .employee-found-files {{ margin-top: 6px; display: grid; gap: 4px; }}
    .employee-checklist-msg {{ margin-top: 8px; padding: 8px 10px; border: 1px solid #d1d5db; border-radius: 10px; background: #f8fafc; font-size: 13px; color: #334155; }}
    .employee-overview-group {{ border: 1px solid #d7e1e8; border-radius: 10px; background: #f8fbff; padding: 8px 10px; margin-top: 8px; }}
    .employee-overview-title {{ font-weight: 700; color: #0f4c5c; margin-bottom: 6px; }}
    .scanner-progress {{ margin-top: 8px; width: 100%; height: 10px; border-radius: 999px; background: #dfe7eb; overflow: hidden; border: 1px solid #cbd5e1; }}
    .scanner-progress-bar {{ width: 0%; height: 100%; background: linear-gradient(90deg, #0f766e, #16a34a); transition: width 0.25s ease; }}
    .task-action-panel {{ margin-top: 10px; padding: 10px 12px; border-radius: 10px; border: 1px solid #d8e5ec; background: #f8fcff; }}
    .manual-review-list {{ margin-top: 8px; max-height: 220px; overflow: auto; border: 1px solid #d8e0e0; border-radius: 10px; background: #fff; padding: 8px; display: grid; gap: 6px; }}
    .manual-review-item {{ border: 1px solid #e2e8f0; border-radius: 10px; padding: 8px; background: #f8fbff; cursor: pointer; }}
    .manual-review-item:hover {{ border-color: #94a3b8; }}
    .manual-review-item.is-selected {{ border-color: #0f766e; box-shadow: 0 0 0 2px rgba(15, 118, 110, 0.15); background: #ecfeff; }}
    .manual-review-head {{ display: grid; grid-template-columns: auto 1fr auto; align-items: center; gap: 8px; }}
    .manual-review-title {{ font-weight: 600; color: #0f172a; }}
    .manual-review-score {{ font-size: 12px; color: #0f766e; font-weight: 700; }}
    .employee-checklist-head {{ display: flex; justify-content: space-between; align-items: center; gap: 8px; }}
    #todoMainCard {{ grid-column: 1 / -1; }}
    #structureCard {{ grid-column: 1; }}
    #uploadCard {{ grid-column: 1; }}
    #previewCard {{ grid-column: 2; }}
    #employeeChecklistCard {{ grid-column: 1 / -1; }}
    .employee-checklist-head h2 {{ margin: 0; }}
    .batch-panel {{ margin-top: 10px; border: 1px solid #d7e1e8; border-radius: 10px; padding: 10px; background: #f8fbff; }}
    .batch-codes {{ margin-top: 8px; max-height: 180px; overflow: auto; border: 1px solid #d8e0e0; border-radius: 10px; background: #fff; padding: 8px; display: grid; gap: 6px; }}
    .batch-code-item {{ display: grid; grid-template-columns: auto 1fr; gap: 8px; align-items: start; font-size: 13px; color: #0f172a; }}
    .multi-select {{ min-height: 140px; }}
    .back-to-top {{ position: fixed; right: 18px; bottom: 18px; z-index: 30; border-radius: 999px; padding: 10px 14px; }}
    @media (max-width: 1120px) {{
        .dashboard-main-grid {{ grid-template-columns: 1fr; }}
    }}
    @media (max-width: 860px) {{
        .employee-checklist-top {{ grid-template-columns: auto 1fr; }}
        .manual-review-head {{ grid-template-columns: auto 1fr; }}
        .employee-checklist-head {{ flex-direction: column; align-items: flex-start; }}
    }}
    .site-nav {{ position: sticky; top: 0; z-index: 50; background: var(--card); border-bottom: 1px solid #d8cfc2; box-shadow: 0 2px 8px rgba(0,0,0,0.06); }}
    .site-nav-inner {{ max-width: 1180px; margin: 0 auto; padding: 0 16px; display: flex; align-items: center; gap: 4px; height: 46px; flex-wrap: wrap; }}
    .site-nav-brand {{ font-weight: 700; color: var(--accent); font-size: 14px; margin-right: 12px; white-space: nowrap; }}
    .site-nav-link {{ padding: 6px 12px; border-radius: 999px; font-size: 13px; font-weight: 600; color: var(--ink); text-decoration: none; white-space: nowrap; }}
    .site-nav-link:hover {{ background: #eff6fb; }}
    .site-nav-link.active {{ background: var(--accent); color: #fff; }}
    </style>
</head>
<body>
    <nav class=\"site-nav\">
        <div class=\"site-nav-inner\">
            <span class=\"site-nav-brand\">X5 УФА Э2</span>
            <a class=\"site-nav-link active\" href=\"/arm/dashboard\">Дашборд</a>
            <a class=\"site-nav-link\" href=\"/arm/structure/view\">Структура и действия</a>
            <a class=\"site-nav-link\" href=\"/arm/employees\">Сотрудники</a>
            <a class=\"site-nav-link\" href=\"/arm/checklist/view\">Чеклист</a>
            <a class=\"site-nav-link\" href=\"/arm/permit/height\">Наряд высота</a>
            <a class=\"site-nav-link\" href=\"/arm/aosr\">АОСР</a>
            <a class=\"site-nav-link\" href=\"/arm/todo/view\">План дня</a>
            <a class=\"site-nav-link\" href=\"/arm/periodic/view\">Периодические</a>
            <a class=\"site-nav-link\" href=\"/docs\" target=\"_blank\" rel=\"noopener\">API</a>
        </div>
    </nav>
    <div class=\"wrap\">
    <section class="grid">
        <article class="card" id="metricsCard">
            <a class=\"tab active\" href=\"/arm/dashboard\">АРМ</a>
            <a class=\"tab\" href=\"/arm/research\">Исследование</a>
        </nav>
        <h1>АРМ объекта: X5 UFA E2</h1>
        <article class="card" id="gapsCard">
        <div class=\"meta\">Сформировано (UTC): {escape(payload.generated_at.isoformat())}</div>
        <div class=\"kpi\">{payload.checklist_progress_percent}%</div>
        <div>Комплектность: {payload.checklist_ready}/{payload.checklist_total}</div>
        <div class=\"meta\">LLM статус: <span id=\"llmBadge\" class=\"llm-badge {llm_badge_class}\">{llm_badge_text}</span></div>
        <div class="meta" id="llmDesc">Проверка локальной LLM...</div>
        <article class="card" id="todoCard"> 
        <div class=\"action-note\" id=\"interactionHint\">Подсказка навигации появится здесь после действий.</div>
    </section>

    <section class=\"grid dashboard-main-grid\">
        <article class=\"card\">
            <h2>Метрики</h2>
            <div class=\"system-hint\">Системная подсказка: следите за ростом сканов на ручной разбор, это главный индикатор блокеров.</div>
            <ul>{metrics_html}</ul>
        </article>
        <article class=\"card\">
            <h2>Критичные пробелы</h2>
            <div class=\"system-hint\">Системная подсказка: клик по позиции откроет нужную папку; под каждой позицией указано действие по устранению.</div>
            <ul>{gaps_html}</ul>
        </article>
        <article class=\"card\" id=\"todoMainCard\"> 
            <h2>Задачи на сегодня</h2>
            <div class=\"system-hint\">Системная подсказка: начните с пунктов приоритета high.</div>
            <ul>{todo_html}</ul>
        </article>
        <article class=\"card\" id=\"structureCard\">
            <h2>Структура и действия</h2>
            <div class=\"system-hint\">Системная подсказка: фамилии отображаются с заглавной буквы, фактические пути не изменяются.</div>
            <div class=\"actions\" style=\"margin-top:0;\">
                <input id=\"treePath\" placeholder=\"Относительный путь, например 02_personnel/employees\" />
                <button class=\"btn secondary\" id=\"treeOpenBtn\">Открыть</button>
            </div>
            <div class=\"actions tree-toolbar\">
                <button class=\"btn secondary mini\" id=\"treeBackBtn\" type=\"button\" disabled>Назад</button>
                <button class=\"btn secondary mini\" id=\"treeForwardBtn\" type=\"button\" disabled>Вперед</button>
                <button class=\"btn secondary mini\" id=\"treeUpBtn\" type=\"button\" disabled>На уровень выше</button>
                <button class=\"btn secondary mini\" id=\"treeRootBtn\" type=\"button\" disabled>Корень</button>
            </div>
            <div id=\"treeBreadcrumb\" class=\"tree-breadcrumb\">Корень</div>
            <div id=\"treeView\" class=\"tree\"></div>
            <div class=\"meta\" id=\"treeHint\">Загрузка дерева...</div>
            <div class=\"action-nav\" id=\"actionNavigator\">
                <div><b>Навигатор действий (предпросмотр)</b></div>
                <div class=\"meta\" id=\"actionNavigatorPath\">Маршрут: /</div>
                <ol id=\"actionNavigatorSteps\">
                    <li>Выберите задачу в «Критичные пробелы» или «Задачи на сегодня».</li>
                    <li>Проверьте обязательные документы в открытой папке.</li>
                    <li>Выполните действие: правка, скачивание, печать или сканирование.</li>
                </ol>
            </div>
            <div class="task-action-panel">
                <div class="meta" id="taskActionLabel">Задача: не выбрана</div>
                <div class="meta" id="taskActionPath">Путь: /</div>
                <div class="actions">
                    <button class="btn secondary" id="taskActionOpenChecklistBtn" type="button">Открыть ТБ-чеклист</button>
                    <button class="btn secondary" id="taskActionRunMissingBtn" type="button">Сгенерировать недостающие</button>
                    <button class="btn secondary" id="taskActionRunSelectedBtn" type="button">Сгенерировать выбранные</button>
                </div>
            </div>
        </article>
        <article class=\"card\" id=\"uploadCard\">
            <h2>Загрузить файл на объект</h2>
            <div class=\"system-hint\">Загрузите DOCX/PDF или скан в любую папку объекта. Загруженные DOCX-шаблоны сразу доступны в файловом дереве.</div>
            <div class=\"actions\" style=\"margin-top:8px; flex-wrap:wrap;\">
                <input id=\"uploadDir\" placeholder=\"Папка назначения, напр. 06_normative_base\" style=\"flex:2; min-width:200px;\" />
                <input type=\"file\" id=\"uploadFile\" accept=\".docx,.pdf,.png,.jpg,.jpeg,.tiff,.tif,.xlsx\" style=\"flex:2; min-width:180px;\" />
                <button class=\"btn\" id=\"uploadBtn\" type=\"button\">Загрузить</button>
            </div>
            <div class=\"actions\" style=\"margin-top:8px;\">
                <button class=\"btn secondary\" id=\"uploadUseCurrentBtn\" type=\"button\">Взять текущую папку из дерева</button>
                <label class=\"meta\" style=\"display:flex; gap:6px; align-items:center;\"><input id=\"uploadAutoUseTree\" type=\"checkbox\" checked />Авто-синхронизация папки с деревом</label>
            </div>
            <div class=\"meta\" id=\"uploadMsg\">Выберите файл и папку назначения.</div>
        </article>
        <article class=\"card\" id=\"previewCard\">
            <h2>Предпросмотр, правка, печать</h2>
            <div class=\"system-hint\">Системная подсказка: для крупных файлов используйте «Скачать», а не онлайн-редактирование.</div>
            <div class=\"meta\" id=\"filePath\">Файл не выбран</div>
            <div class="actions" style="margin-top:8px;">
                <button class="btn secondary mini" id="fileZoomOutBtn" type="button">- Зум</button>
                <button class="btn secondary mini" id="fileZoomResetBtn" type="button">100%</button>
                <button class="btn secondary mini" id="fileZoomInBtn" type="button">+ Зум</button>
            </div>
            <div id="fileRenderWrap"><iframe id="fileRenderFrame" title="Рендер файла" src="about:blank"></iframe></div>
            <textarea id=\"fileEditor\" placeholder=\"Содержимое текстового файла\"></textarea>
            <div class=\"actions\">
                <button class=\"btn secondary\" id=\"fileSourceToggleBtn\" type=\"button\">Показать исходник</button>
                <button class=\"btn secondary\" id=\"fileSaveBtn\">Сохранить</button>
                <button class=\"btn secondary\" id=\"fileDownloadBtn\">Скачать</button>
                <button class=\"btn secondary\" id=\"filePrintBtn\">Печать</button>
                <button class="btn secondary" id="fileMoveBtn" type="button">Переместить</button>
                <button class="btn secondary" id="fileDeleteBtn" type="button">Удалить</button>
            </div>
            <div class=\"meta\" id=\"fileMeta\">Ожидание</div>
        </article>
        <article class=\"card\" id=\"employeeChecklistCard\">
            <div class=\"employee-checklist-head\">
                <h2>ТБ-чеклист сотрудника</h2>
                <button class=\"btn secondary mini\" id=\"employeeChecklistToggleBtn\" type=\"button\">Свернуть</button>
            </div>
            <div class=\"system-hint\">Системная подсказка: можно анализировать конкретного сотрудника или сразу тип сотрудников (по профессии).</div>
            <div id=\"employeeChecklistBody\">
            <div class=\"employee-checklist-tools\">
                <select id=\"employeeChecklistProfession\">
                    <option value=\"all\">Все типы сотрудников</option>
                </select>
                <select id=\"employeeChecklistEmployee\">
                    <option value=\"\">Не выбран (анализ по типам сотрудников)</option>
                </select>
                <input id=\"employeeChecklistPath\" placeholder=\"Путь сотрудника, например 02_personnel/employees/001_ivanov_ivan_ivanovich\" />
                <input id=\"employeeChecklistOrderDate\" value=\"01.03.2026\" placeholder=\"Дата приказа (ДД.ММ.ГГГГ)\" />
                <label class=\"meta\" style=\"display:flex; gap:6px; align-items:center;\">
                    <input type=\"checkbox\" id=\"employeeChecklistOverwrite\" />
                    Перезаписывать существующие черновики
                </label>
                <div class=\"actions\">
                    <button class=\"btn secondary\" id=\"employeeChecklistOverviewBtn\">Комплексный анализ</button>
                    <button class=\"btn secondary\" id=\"employeeChecklistRefreshBtn\">Проверить комплект</button>
                    <button class=\"btn secondary\" id=\"employeeChecklistSelectMissing\">Выбрать только недостающие</button>
                    <button class=\"btn secondary\" id=\"employeeChecklistGenerateSelectedBtn\">Сгенерировать выбранные</button>
                    <button class=\"btn secondary\" id=\"employeeChecklistGenerateMissingBtn\">Сгенерировать недостающие</button>
                    <button class=\"btn secondary\" id=\"employeeChecklistGenerateAllBtn\">Сгенерировать полный комплект</button>
                    <button class=\"btn secondary\" id=\"employeeChecklistApplyTypeBtn\">Применить форму к типу сотрудников</button>
                </div>
            </div>
            <div class=\"batch-panel\">
                <div class=\"meta\"><b>Пакетная генерация</b>: отмечайте сотрудников и коды документов, затем запускайте в 1 клик.</div>
                <div class=\"employee-checklist-tools\">
                    <select id=\"batchGenerateMode\">
                        <option value=\"missing\">Режим: только недостающие</option>
                        <option value=\"selected\">Режим: только выбранные коды</option>
                        <option value=\"all\">Режим: полный комплект</option>
                    </select>
                    <select id=\"batchEmployees\" class=\"multi-select\" multiple></select>
                    <div class=\"actions\">
                        <button class=\"btn secondary mini\" id=\"batchEmployeesSelectAllBtn\" type=\"button\">Выбрать всех в фильтре</button>
                        <button class=\"btn secondary mini\" id=\"batchEmployeesClearBtn\" type=\"button\">Снять выбор</button>
                    </div>
                    <div class=\"actions\">
                        <button class=\"btn secondary mini\" id=\"batchUseMissingCodesBtn\" type=\"button\">Взять недостающие коды из текущего чеклиста</button>
                    </div>
                    <div id=\"batchDocCodes\" class=\"batch-codes\">Сначала откройте чеклист сотрудника. Тогда здесь появятся коды документов для выборочной генерации.</div>
                    <div class=\"actions\">
                        <button class=\"btn secondary\" id=\"batchGenerateSelectedEmployeesBtn\" type=\"button\">Сгенерировать для выбранных сотрудников</button>
                        <button class=\"btn secondary\" id=\"batchGenerateAllFilteredBtn\" type=\"button\">Сгенерировать для всех сотрудников фильтра</button>
                    </div>
                </div>
            </div>
            <div class=\"meta\" id=\"employeeChecklistSummary\">Статус: ожидание проверки.</div>
            <div id=\"employeeChecklistOverview\" class=\"employee-overview\"></div>
            <div id=\"employeeChecklistList\" class=\"employee-checklist-list\"></div>
            <div id=\"employeeChecklistMsg\" class=\"employee-checklist-msg\">Ожидание действий.</div>
            </div>
        </article>
        <article class=\"card\" id=\"scannerCard\">
            <h2>Сканирование и OCR</h2>
            <div class=\"system-hint\">Системная подсказка: шаги процесса: сканирование → OCR-разбор → классификация ручного разбора.</div>
            <div class=\"inline-grid\">
                <select id=\"scannerDevices\"></select>
                <select id=\"scannerDocType\">
                    <option value=\"ORDER\">Приказ</option>
                    <option value=\"AWR\">Акт выполненных работ</option>
                    <option value=\"PASSPORT\">Удостоверение/протокол</option>
                    <option value=\"INVOICE\">Счет</option>
                    <option value=\"UPD\">УПД</option>
                    <option value=\"TTN\">ТТН</option>
                    <option value=\"ACT\">Акт</option>
                    <option value=\"OTHER\">Другое</option>
                </select>
                <select id=\"scannerSortMode\">
                    <option value=\"unresolved\">Сортировка: сначала неразобранные</option>
                    <option value=\"name\">Сортировка: по имени файла</option>
                </select>
                <input id=\"scannerSubject\" placeholder=\"Тема: например приказ_допуск_20260311\" />
                <select id=\"scannerEmployeeSelect\">
                    <option value=\"\">Выберите сотрудника</option>
                </select>
                <input id=\"scannerEmployee\" list=\"scannerEmployeeSuggestions\" placeholder=\"Код сотрудника (для удостоверения/протокола)\" />
                <datalist id=\"scannerEmployeeSuggestions\"></datalist>
            </div>
            <div class=\"actions\" style=\"margin-top:8px; align-items:center; gap:10px; flex-wrap:wrap;\">
                <label class=\"meta\" for=\"scannerProfile\" style=\"min-width:120px;\">Профиль скана:</label>
                <input id=\"scannerProfile\" type=\"range\" min=\"1\" max=\"3\" step=\"1\" value=\"1\" style=\"max-width:280px;\" />
                <span class=\"meta\" id=\"scannerProfileLabel\">Профиль 1: 300 dpi, grayscale</span>
            </div>
            <div class=\"meta\" id=\"scannerEmployeeHint\">Выберите папку сотрудника, чтобы подставить код автоматически.</div>
            <div class=\"system-hint\" id=\"scannerHint\">Режим приказа/акта: код сотрудника можно не заполнять.</div>
            <div class=\"actions\">
                <button class=\"btn secondary\" id=\"scanBtn\">Сканировать во входящую папку</button>
                <button class=\"btn secondary\" id=\"ingestBtn\">Распознать и разложить</button>
                <button class=\"btn secondary\" id=\"manualClassifyBtn\">Классифицировать ручной разбор</button>
                <button class=\"btn secondary\" id=\"recompressScansBtn\" type=\"button\">Сжать исторические сканы</button>
                <button class=\"btn secondary\" id=\"maintenanceResetBtn\" type=\"button\">Service: reset + rebuild</button>
            </div>
            <div class=\"scanner-progress\"><div class=\"scanner-progress-bar\" id=\"scannerProgressBar\"></div></div>
            <div class=\"meta\" id=\"scannerProgressText\">Прогресс: 0%; ожидание.</div>
            <pre id=\"scannerTimeline\">Журнал этапов: ожидание</pre>
            <pre id=\"manualReview\">Ручной разбор: ожидание</pre>
            <div id=\"manualReviewList\" class=\"manual-review-list\">Список ручного разбора появится после классификации.</div>
            <div class=\"meta\" id=\"manualReviewSelection\">Файл из ручного разбора не выбран.</div>
            <div class=\"inline-grid\">
                <select id=\"manualReviewTargetType\">
                    <option value=\"\">Маршрут вручную (как в поле пути ниже)</option>
                    <option value=\"ORDER\">Маршрут: Приказ</option>
                    <option value=\"AWR\">Маршрут: АВР</option>
                    <option value=\"PASSPORT\">Маршрут: Удостоверение/протокол сотрудника</option>
                    <option value=\"INVOICE\">Маршрут: Счет</option>
                    <option value=\"UPD\">Маршрут: УПД</option>
                    <option value=\"TTN\">Маршрут: ТТН</option>
                    <option value=\"ACT\">Маршрут: Акт</option>
                    <option value=\"OTHER\">Маршрут: Прочее/ручной разбор</option>
                </select>
                <button class=\"btn secondary\" id=\"manualReviewSuggestBtn\" type=\"button\">Подставить маршрут</button>
                <button class=\"btn secondary\" id=\"manualReviewMkDirBtn\" type=\"button\">Создать папку маршрута</button>
            </div>
            <input id=\"manualReviewMovePath\" placeholder=\"Целевой путь для переноса, например 02_personnel/employees/001_ivanov/06_permits_and_work_admission/file.jpg\" />
            <div class=\"actions\">
                <button class=\"btn secondary\" id=\"manualReviewOpenBtn\" type=\"button\">Открыть</button>
                <button class=\"btn secondary\" id=\"manualReviewDownloadBtn\" type=\"button\">Скачать</button>
                <button class=\"btn secondary\" id=\"manualReviewPrintBtn\" type=\"button\">Печать</button>
                <button class=\"btn secondary\" id=\"manualReviewEditOcrBtn\" type=\"button\">Открыть OCR-текст</button>
                <button class=\"btn secondary\" id=\"manualReviewMoveBtn\" type=\"button\">Переместить</button>
                <button class=\"btn secondary\" id=\"manualReviewDeleteBtn\" type=\"button\">Удалить</button>
            </div>
            <div class=\"meta\" id=\"scannerMsg\">Скан-центр готов.</div>
        </article>
        <article class=\"card\" id=\"exportCard\">
            <h2>Экспорт и печать офисных документов</h2>
            <div class=\"system-hint\">Системная подсказка: выберите классификацию источников перед выгрузкой пакета DOCX/ZIP.</div>
            <select id=\"exportClassification\">
                <option value=\"all\">Все источники (база + черновики чеклиста + черновики сотрудников)</option>
                <option value=\"base-orders\">Только базовые приказы объекта</option>
                <option value=\"checklist-drafts\">Только черновики из чеклиста</option>
                <option value=\"employee-drafts\">Только черновики сотрудников</option>
            </select>
            <div class=\"actions\">
                <button class=\"btn\" id=\"exportDocxBtn\">Скачать DOCX по классификации</button>
                <button class=\"btn\" id=\"exportPackBtn\">Скачать офисный пакет (ZIP) по классификации</button>
                <button class=\"btn secondary\" id=\"exportXlsxBtn\">Скачать реестры XLSX</button>
                <a class=\"btn secondary\" href=\"/arm/research/competitors\" target=\"_blank\" rel=\"noopener\">Открыть исследование конкурентов (HTML)</a>
            </div>
            <p class=\"meta\">Файлы сохраняются в папку офисной печати внутри раздела «Приказы и назначения».</p>
        </article>
        <article class=\"card\" id=\"objectProfileCard\">
            <h2>Карточка объекта и ППР</h2>
            <div class=\"system-hint\">Дата начала работ управляет периодичностью документов. ППР импортируется в базу знаний проекта для автозаполнения форм без повторного ввода.</div>
            <div class=\"inline-grid\">
                <input id=\"objectNameInput\" placeholder=\"Объект\" />
                <input id=\"workStageInput\" placeholder=\"Этап работ\" />
                <input id=\"projectCodeInput\" placeholder=\"Шифр проекта\" />
                <input id=\"organizationInput\" placeholder=\"Организация\" />
            </div>
            <div class=\"actions\" style=\"margin-top:8px; align-items:center;\">
                <input id=\"startDateInput\" placeholder=\"Дата начала работ, напр. 01.03.2026\" style=\"max-width:260px;\" />
                <button class=\"btn\" id=\"objectProfileSaveBtn\" type=\"button\">Сохранить карточку объекта</button>
            </div>
            <div class=\"actions\" style=\"margin-top:8px; align-items:center;\">
                <select id=\"pprSourceSelect\" style=\"min-width:280px;\"></select>
                <button class=\"btn secondary\" id=\"pprImportBtn\" type=\"button\">Импортировать ППР в базу</button>
            </div>
            <div class=\"meta\" id=\"objectProfileMsg\">Загрузка карточки объекта...</div>
        </article>
        <article class=\"card\" id=\"assistantCard\">
            <h2>Ассистент (локальная LLM и Copilot)</h2>
            <div class=\"system-hint\">Системная подсказка: если ответ медленный, выберите профиль «Быстрый».</div>
            <div style=\"display:flex; gap:8px; align-items:center; margin:8px 0;\">
                <label for=\"armProfile\" class=\"meta\">Профиль:</label>
                <select id=\"armProfile\">
                    <option value=\"fast\" selected>Быстрый</option>
                    <option value=\"balanced\">Сбалансированный</option>
                    <option value=\"quality\">Качество</option>
                </select>
                <button class=\"btn secondary\" id=\"armSend\">Отправить</button>
                <button class=\"btn secondary\" id=\"armVoiceBtn\" type=\"button\">Голосовой ввод</button>
            </div>
            <label class=\"meta\" style=\"display:flex; gap:6px; align-items:center; margin:4px 0 8px;\">
                <input type=\"checkbox\" id=\"armSendOnEnter\" />
                Enter отправляет сообщение (если выключено, отправка по Ctrl+Enter)
            </label>
            <div class=\"meta\" id=\"armVoiceHint\">Голосовой ввод: нажмите кнопку, чтобы начать запись с микрофона.</div>
            <textarea id=\"armQuestion\" placeholder=\"Например: сформируй задачи на день и предложи порядок устранения пробелов\"></textarea>
            <pre id=\"armAnswer\">Ожидание запроса...</pre>
        </article>
    </section>
    </div>
    <button class=\"btn secondary back-to-top\" id=\"backToTopBtn\" type=\"button\" hidden>Наверх</button>
    {script_html}
</body>
</html>
"""
    return HTMLResponse(content=html)




