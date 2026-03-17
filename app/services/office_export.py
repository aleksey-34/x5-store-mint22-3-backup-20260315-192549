from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from app.models.document import Document
from app.models.journal_entry import JournalEntry
from app.models.work_schedule import WorkSchedule
from app.schemas.arm_admin import ArmChecklistItem


ORDER_DOC_NAME_MAP = {
    "ORDER_01": "Приказ_01_Уполномоченный_представитель",
    "ORDER_02": "Приказ_02_Допуск_к_СМР",
    "ORDER_03": "Приказ_03_Ответственные_по_направлениям",
    "ORDER_04": "Приказ_04_Стропальщики",
    "ORDER_05": "Приказ_05_Электрохозяйство",
    "ORDER_06": "Приказ_06_Стажировка_рабочих",
    "ORDER_07": "Приказ_07_Допуск_по_профессиям",
    "ORDER_08": "Приказ_08_Допуск_монтажников_СМР",
    "ORDER_09": "Приказ_09_Распределение_по_прорабам",
    "ORDER_10": "Приказ_10_ТБ_по_прорабам",
    "ORDER_11": "Приказ_11_Ответственные_лица_по_ПС",
    "ORDER_12": "Приказ_12_Ответственные_лица_за_наряды-допуски",
    "ORDER_13": "Приказ_13_Ответственные_лица_и_меры_безопасности_на_высоте",
    "ORDER_14": "Приказ_14_Ответственные_лица_за_пожарную_безопасность",
    "ORDER_15": "Приказ_15_Ответственные_за_погрузочно_разгрузочные_работы",
    "ORDER_16": "Приказ_16_Ответственные_лица_за_сосуды_под_давлением",
    "ORDER_17": "Приказ_17_Закрытие_смены",
    "ORDER_18": "Приказ_18_Допуск_к_стажировке",
    "ORDER_19": "Приказ_19_Допуск_к_самостоятельной_работе",
    "ORDER_20": "Приказ_20_Ответственные_лица_по_электропрогреву_бетона_зима",
    "LETTER_ADMISSION": "Письмо_допуск_персонал_техника",
    "PERMIT_12": "Наряд_допуск_12_Опасные_работы",
}

VALID_EXPORT_CLASSIFICATIONS: tuple[str, ...] = (
    "all",
    "base-orders",
    "checklist-drafts",
    "employee-drafts",
)

SCAN_EXTENSIONS: tuple[str, ...] = (".pdf", ".jpg", ".jpeg", ".png", ".tif", ".tiff")


class OfficeExportDependencyError(RuntimeError):
    """Raised when optional office export dependencies are missing."""


@dataclass(frozen=True)
class OfficeDocxExportResult:
    output_dir: Path
    bundle_path: Path
    files_count: int


@dataclass(frozen=True)
class OfficeXlsxExportResult:
    file_path: Path


def _ensure_office_dependencies() -> tuple[object, object]:
    try:
        from docx import Document as DocxDocument
        from openpyxl import Workbook
    except Exception as exc:  # noqa: BLE001
        raise OfficeExportDependencyError(
            "Для офисного экспорта установите python-docx и openpyxl"
        ) from exc
    return DocxDocument, Workbook


def _apply_docx_defaults(doc: object) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    if style._element is not None:  # pyright: ignore[reportAttributeAccessIssue]
        rpr = style._element.get_or_add_rPr()  # pyright: ignore[reportAttributeAccessIssue]
        rfonts = rpr.get_or_add_rFonts()  # pyright: ignore[reportAttributeAccessIssue]
        rfonts.set(qn("w:eastAsia"), "Times New Roman")  # pyright: ignore[reportCallIssue]


def _set_table_borders(table: object) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")
        borders.append(node)
    tbl_pr.append(borders)


def _is_markdown_separator(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell or "") is not None for cell in cells)


def _extract_markdown_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start_index
    while index < len(lines):
        raw = lines[index].strip()
        if not (raw.startswith("|") and raw.endswith("|")):
            break
        cells = [part.strip() for part in raw.strip("|").split("|")]
        if _is_markdown_separator(cells):
            index += 1
            continue
        rows.append(cells)
        index += 1
    return rows, index


def _build_table_from_markdown_rows(doc: object, rows: list[list[str]]) -> None:
    if not rows:
        return

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    _set_table_borders(table)

    for row_index, row_values in enumerate(rows):
        for col_index in range(column_count):
            value = row_values[col_index] if col_index < len(row_values) else ""
            cell = table.cell(row_index, col_index)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            if row_index == 0:
                run.bold = True


def _collect_employee_scan_files(object_root: Path) -> list[Path]:
    employees_root = object_root / "02_personnel" / "employees"
    if not employees_root.exists() or not employees_root.is_dir():
        return []

    seen: set[Path] = set()
    files: list[Path] = []
    for file in employees_root.rglob("*"):
        if not file.is_file():
            continue
        if file.suffix.lower() not in SCAN_EXTENSIONS:
            continue
        resolved = file.resolve()
        if resolved in seen:
            continue
        seen.add(resolved)
        files.append(resolved)
    return sorted(files, key=lambda p: p.as_posix().lower())


def _office_root(object_root: Path) -> Path:
    out = object_root / "01_orders_and_appointments" / "print_office"
    out.mkdir(parents=True, exist_ok=True)
    return out


def _write_md_to_docx(md_text: str, doc: object) -> None:
    lines = md_text.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            index += 1
            continue

        if line.startswith("#"):
            level = min(3, max(1, len(line) - len(line.lstrip("#"))))
            doc.add_heading(line[level:].strip(), level=level)
            index += 1
            continue

        if re.match(r"^[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
            index += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            index += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            rows, next_index = _extract_markdown_table(lines, index)
            _build_table_from_markdown_rows(doc, rows)
            index = next_index
            continue

        doc.add_paragraph(line)
        index += 1


def _russian_safe_filename(stem: str) -> str:
    upper_stem = stem.upper()
    for token, name in ORDER_DOC_NAME_MAP.items():
        if token in upper_stem:
            return name

    normalized = stem.replace("-", "_").replace(" ", "_")
    normalized = re.sub(r"[^0-9A-Za-zА-Яа-яЁё_]+", "", normalized)
    return normalized or "Документ"


def _collect_order_markdown_files(object_root: Path, classification: str) -> list[Path]:
    selected = (classification or "all").strip().lower()
    if selected not in VALID_EXPORT_CLASSIFICATIONS:
        selected = "all"

    source_dir = object_root / "01_orders_and_appointments"
    files: list[Path] = []

    if selected in {"all", "base-orders"}:
        files.extend(sorted(source_dir.glob("*.md")))

    if selected in {"all", "checklist-drafts"}:
        drafts_dir = source_dir / "drafts_from_checklist"
        if drafts_dir.exists() and drafts_dir.is_dir():
            files.extend(sorted(drafts_dir.glob("*.md")))

    if selected in {"all", "employee-drafts"}:
        employees_root = object_root / "02_personnel" / "employees"
        if employees_root.exists() and employees_root.is_dir():
            files.extend(sorted(employees_root.glob("*/07_templates_to_print/*.md")))

    seen: set[Path] = set()
    unique: list[Path] = []
    for file in files:
        resolved = file.resolve()
        if resolved in seen:
            continue
        seen.add(resolved)
        unique.append(resolved)

    return sorted(unique, key=lambda p: (p.parent.as_posix().lower(), p.name.lower()))


def export_orders_docx_bundle(object_root: Path, classification: str = "all") -> OfficeDocxExportResult:
    DocxDocument, _ = _ensure_office_dependencies()

    selected = (classification or "all").strip().lower()
    if selected not in VALID_EXPORT_CLASSIFICATIONS:
        selected = "all"

    office_root = _office_root(object_root)
    docx_dir = office_root / "docx"
    docx_dir.mkdir(parents=True, exist_ok=True)

    md_files = _collect_order_markdown_files(object_root=object_root, classification=selected)

    generated: list[Path] = []
    for md_file in md_files:
        doc = DocxDocument()
        _apply_docx_defaults(doc)
        title = _russian_safe_filename(md_file.stem)
        doc.add_heading(title.replace("_", " "), level=1)
        text = md_file.read_text(encoding="utf-8")
        _write_md_to_docx(text, doc)

        target = docx_dir / f"{title}.docx"
        doc.save(target)
        generated.append(target)

    if not generated:
        doc = DocxDocument()
        _apply_docx_defaults(doc)
        doc.add_heading("Офисный экспорт", level=1)
        doc.add_paragraph(
            "По выбранной классификации не найдено файлов Markdown для экспорта. "
            "Проверьте фильтр и наличие документов в целевых папках."
        )
        placeholder = docx_dir / "README_приказы_не_найдены.docx"
        doc.save(placeholder)
        generated.append(placeholder)

    bundle_name = "пакет_приказов_docx.zip" if selected == "all" else f"пакет_приказов_docx_{selected}.zip"
    bundle = office_root / bundle_name
    with ZipFile(bundle, mode="w", compression=ZIP_DEFLATED) as archive:
        for file in generated:
            archive.write(file, arcname=file.name)

    return OfficeDocxExportResult(output_dir=docx_dir, bundle_path=bundle, files_count=len(generated))


def _append_sheet_header(ws: object, header: list[str]) -> None:
    ws.append(header)


def export_registers_xlsx(
    object_root: Path,
    documents: list[Document],
    journal_entries: list[JournalEntry],
    schedules: list[WorkSchedule],
    checklist: list[ArmChecklistItem],
) -> OfficeXlsxExportResult:
    _, Workbook = _ensure_office_dependencies()

    office_root = _office_root(object_root)
    output_path = office_root / "арм_реестры.xlsx"

    wb = Workbook()
    ws_docs = wb.active
    ws_docs.title = "документы"
    _append_sheet_header(ws_docs, ["ID", "Название", "Тип документа", "Статус", "Путь к файлу", "Создан"])
    for doc in documents:
        ws_docs.append(
            [
                doc.id,
                doc.title,
                doc.doc_type,
                doc.status,
                doc.file_path or "",
                doc.created_at.isoformat() if doc.created_at else "",
            ]
        )

    ws_journal = wb.create_sheet("журнал")
    _append_sheet_header(ws_journal, ["ID", "Категория", "Содержание", "Источник", "Создано"])
    for entry in journal_entries:
        ws_journal.append(
            [
                entry.id,
                entry.category,
                entry.content,
                entry.source,
                entry.created_at.isoformat() if entry.created_at else "",
            ]
        )

    ws_schedules = wb.create_sheet("графики")
    _append_sheet_header(
        ws_schedules,
        [
            "ID",
            "Наименование",
            "План старт",
            "План финиш",
            "Факт старт",
            "Факт финиш",
            "Прогресс, %",
            "Статус",
            "Примечания",
        ],
    )
    for schedule in schedules:
        ws_schedules.append(
            [
                schedule.id,
                schedule.title,
                schedule.planned_start.isoformat() if schedule.planned_start else "",
                schedule.planned_end.isoformat() if schedule.planned_end else "",
                schedule.actual_start.isoformat() if schedule.actual_start else "",
                schedule.actual_end.isoformat() if schedule.actual_end else "",
                schedule.progress_percent,
                schedule.status,
                schedule.notes or "",
            ]
        )

    ws_check = wb.create_sheet("чеклист")
    _append_sheet_header(ws_check, ["Код", "Позиция", "Расположение", "Требуется минимум", "Найдено", "Готово"])
    for item in checklist:
        ws_check.append([item.code, item.title, item.location, item.required_min, item.found, item.ready])

    ws_meta = wb.create_sheet("мета")
    ws_meta.append(["Сформировано UTC", datetime.utcnow().isoformat()])
    ws_meta.append(["Корень объекта", str(object_root)])

    wb.save(output_path)
    return OfficeXlsxExportResult(file_path=output_path)


def build_office_pack_zip(
    object_root: Path,
    docx_result: OfficeDocxExportResult,
    xlsx_result: OfficeXlsxExportResult,
) -> Path:
    office_root = _office_root(object_root)
    pack_path = office_root / "офисный_пакет_печати.zip"

    with ZipFile(pack_path, mode="w", compression=ZIP_DEFLATED) as archive:
        archive.write(xlsx_result.file_path, arcname=xlsx_result.file_path.name)
        archive.write(docx_result.bundle_path, arcname=docx_result.bundle_path.name)
        for docx_file in sorted(docx_result.output_dir.glob("*.docx")):
            archive.write(docx_file, arcname=f"docx/{docx_file.name}")
        for scan_file in _collect_employee_scan_files(object_root):
            rel_scan = scan_file.relative_to(object_root)
            archive.write(scan_file, arcname=f"scans/{rel_scan.as_posix()}")

    return pack_path
