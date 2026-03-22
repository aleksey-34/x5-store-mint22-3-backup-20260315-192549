#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Конвертация инструкций по охране труда (ИОТ) из Markdown в DOCX.

Объект: Логистический парк, г. Уфа, Этап 2
Организация: ИП Исмагилов Вадим Шакирович

Использование:
    python convert_iot_to_docx.py

Требования:
    - pandoc (установить с https://pandoc.org/installing.html)
      Windows: winget install --id JohnMacFarlane.Pandoc
    ИЛИ
    - python-docx (pip install python-docx)
      для упрощённой конвертации без pandoc

Результат:
    Папка docx_export/ рядом с папкой instructions/
    с готовыми .docx файлами для ТБшницы
"""

import os
import subprocess
import shutil
import sys
from pathlib import Path

# --- Пути ---
SCRIPT_DIR = Path(__file__).parent
REPO_ROOT = SCRIPT_DIR.parent
OBJECT_DIR = REPO_ROOT / "docflow" / "objects" / "x5-ufa-e2_logistics_park"
INSTRUCTIONS_DIR = OBJECT_DIR / "03_hse_and_fire_safety" / "instructions"
OUTPUT_DIR = OBJECT_DIR / "03_hse_and_fire_safety" / "docx_export"

# Дополнительно конвертировать Приказ № 23
EXTRA_FILES = [
    OBJECT_DIR / "01_orders_and_appointments" / "20260322_ORDER_23_утверждение_ИОТ_v01.md",
]

# Шаблон DOCX (если есть — подставить путь; если нет — оставить None)
DOCX_REFERENCE = None  # или Path("reference.docx") для кастомного стиля Word

# --- Вспомогательные функции ---

def check_pandoc() -> bool:
    """Проверяет наличие pandoc в PATH."""
    return shutil.which("pandoc") is not None


def convert_with_pandoc(md_path: Path, out_path: Path) -> bool:
    """Конвертирует один MD файл в DOCX через pandoc."""
    cmd = [
        "pandoc",
        str(md_path),
        "-o", str(out_path),
        "--from", "markdown+raw_html",
        "--to", "docx",
        "-s",                             # standalone document
        "--wrap=none",                    # не разбивать длинные строки
    ]
    if DOCX_REFERENCE and DOCX_REFERENCE.exists():
        cmd += ["--reference-doc", str(DOCX_REFERENCE)]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8")
        if result.returncode == 0:
            return True
        else:
            print(f"  [ОШИБКА pandoc] {md_path.name}:\n{result.stderr.strip()}")
            return False
    except Exception as e:
        print(f"  [ИСКЛЮЧЕНИЕ] {md_path.name}: {e}")
        return False


def convert_with_python_docx(md_path: Path, out_path: Path) -> bool:
    """
    Упрощённая конвертация через python-docx без pandoc.
    Создаёт .docx с содержимым MD как plain text (без форматирования).
    Для полноценной конвертации рекомендуется pandoc.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  [ОШИБКА] python-docx не установлен. Выполните: pip install python-docx")
        return False

    doc = Document()

    # Стиль страницы: A4, поля 20/20/20/20 мм
    section = doc.sections[0]
    section.page_width  = int(21.0 * 360000)
    section.page_height = int(29.7 * 360000)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    text = md_path.read_text(encoding="utf-8")

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Заголовки
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        # HTML-теги — убрать (грубо)
        elif stripped.startswith("<") and stripped.endswith(">"):
            # Пропускаем HTML-теги (они будут потеряны в этом режиме)
            continue
        else:
            # Убираем markdown-символы bold/italic для plain text
            clean = stripped.replace("**", "").replace("*", "").replace("`", "")
            doc.add_paragraph(clean)

    doc.save(str(out_path))
    return True


def main():
    print("=" * 60)
    print("Конвертация ИОТ: Markdown → DOCX")
    print(f"Источник: {INSTRUCTIONS_DIR}")
    print(f"Результат: {OUTPUT_DIR}")
    print("=" * 60)

    # Создать папку вывода
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Определить метод конвертации
    use_pandoc = check_pandoc()
    if use_pandoc:
        print("✓ pandoc обнаружен — используем pandoc (полная конвертация с таблицами и HTML)")
    else:
        print("⚠ pandoc не найден — используем python-docx (упрощённая конвертация)")
        print("  Для лучшего результата установите pandoc:")
        print("    winget install --id JohnMacFarlane.Pandoc")
        print()

    # Собрать файлы для конвертации
    md_files = sorted(INSTRUCTIONS_DIR.glob("*.md"))
    md_files += [f for f in EXTRA_FILES if f.exists()]

    if not md_files:
        print("Файлы .md не найдены!")
        sys.exit(1)

    print(f"\nНашено файлов: {len(md_files)}\n")

    success_count = 0
    fail_count = 0

    for md_file in md_files:
        out_name = md_file.stem + ".docx"
        out_path = OUTPUT_DIR / out_name
        print(f"  → {md_file.name}")

        if use_pandoc:
            ok = convert_with_pandoc(md_file, out_path)
        else:
            ok = convert_with_python_docx(md_file, out_path)

        if ok:
            size_kb = out_path.stat().st_size // 1024
            print(f"    ✓ Сохранён: {out_path.name} ({size_kb} КБ)")
            success_count += 1
        else:
            print(f"    ✗ ОШИБКА конвертации")
            fail_count += 1

    print()
    print("=" * 60)
    print(f"Готово: {success_count} файлов конвертировано, {fail_count} ошибок")
    print(f"Папка DOCX: {OUTPUT_DIR}")
    print("=" * 60)

    if success_count > 0:
        print()
        print("Следующие шаги:")
        print(f"  1. Откройте папку: {OUTPUT_DIR}")
        print("  2. Проверьте документы в Word (особенно таблицы)")
        print("  3. При необходимости скорректируйте форматирование")
        print("  4. Отправьте ТБшнице (все .docx или ZIP-архив)")


if __name__ == "__main__":
    main()
